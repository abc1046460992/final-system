# ==========================================
# Recovered Source Code - NoorApp 100% V2
# ==========================================

import os
import sys
import re
import webbrowser
import threading
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from datetime import datetime, date, timedelta
import pandas as pd
import traceback
import calendar
from docx.enum.text import WD_ALIGN_PARAGRAPH
HAS_DOCX = True

APP_TITLE = "مدرسة الملك خالد المتوسطة - نظام الإدارة المدرسية"
AZ_ICON = "icon.ico"
LOGO_PATH = "logo.png"

EXPORT_DIR = "مخرجات_النظام"

COLOR_BG = "#f6f8fb"
COLOR_PANEL = "#ffffff"
COLOR_ACCENT = "#1f5fbf"
COLOR_BTN = "#e8eefc"
COLOR_WORD = "#2062e3"
COLOR_XLSX = "#0a8f49"
COLOR_WARN = "#f3b61f"
COLOR_DANGER = "#d32f2f"

DAYS = ["الأحد", "الاثنين", "الثلاثاء", "الأربعاء", "الخميس"]
PERIODS7 = ["1", "2", "3", "4", "5", "6", "7"]
PERIODS6 = ["1", "2", "3", "4", "5", "6"]
EXIT_REASONS = ["نهاية الدوام", "استئذان", "ظرف طارئ", "مراجعة طبية", "أخرى"]

def get_base_path():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.abspath(".")

def define_data_path(relative_path):
    return os.path.join(get_base_path(), relative_path)

FILE_MASTER = define_data_path("جدول_المعلمين_الأساسي.xlsx")
FILE_SUBJECTS = define_data_path("المواد.xlsx")
FILE_STUDENTS = define_data_path("بيانات_الطلاب.xlsx")
FILE_TIMINGS = define_data_path("توقيت_الدوام.xlsx")
FILE_CALENDAR = define_data_path("التقويم_الدراسي.xlsx")
FILE_EMPLOYEES_PINS = define_data_path("أرقام_سرية_الموظفين.xlsx")
FILE_SWAPS = define_data_path("قاعدة_بيانات_المبادلات.xlsx")
FILE_TEACHERS = define_data_path("المعلمين.xlsx")
FILE_ATTENDANCE = define_data_path("جدول_الحضور_والتأخير_والانصراف.xlsx")

# ========================================

# --- resource_path ---
def resource_path(relative_path=None):
    # ========================================
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

    # ========================================

# --- data_path ---
def data_path(relative_path=None):
    # ========================================
    if getattr(sys, "frozen", False):
        base_path = os.path.dirname(sys.executable)
    else:
        base_path = os.path.abspath(".")
    target = os.path.join(base_path, relative_path)
    if not os.path.exists(target):
        bundled = resource_path(relative_path)
        if os.path.exists(bundled) and bundled != target:
            try:
                import shutil
                shutil.copy2(bundled, target)
            except Exception:
                pass

    return target

    # ========================================

# --- add_run_rtl ---
def add_run_rtl(paragraph=None, text=None, bold=None, size=None):
    # ========================================
    if not HAS_DOCX:
        return
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        from docx.shared import Pt
        run.font.size = Pt(size)
    return run

    # ========================================

# --- tidy ---
def tidy(x=None):
    # ========================================
    if pd.isna(x):
        return ""
    return str(x).strip()

    # ========================================

# --- normalize_arabic ---
def normalize_arabic(text=None):
    # ========================================
    if not text:
        return ""
    t = str(text).strip()
    t = re.sub("[\\u064B-\\u0652]", "", t)
    t = re.sub("[أإآ]", "ا", t)
    t = re.sub("ة", "ه", t)
    t = re.sub("ى", "ي", t)
    return t

    # ========================================

# --- ensure_file ---
def ensure_file(path=None, template_df=None, sheet_name=None):
    # ========================================
    if os.path.exists(path):
        return
    else:
        df = template_df if template_df is not None else pd.DataFrame()
        if sheet_name:
            with pd.ExcelWriter(path, engine="openpyxl") as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        else:
            df.to_excel(path, index=False)

    # ========================================

# --- safe_parse_date ---
def safe_parse_date(x=None):
    # ========================================

        try: return datetime.strptime(str(x).split()[0], "%Y-%m-%d").date()
        except: return None


# --- arabic_day_from_english ---
def arabic_day_from_english(english_day=None):
    # ========================================
    return {
     'Sunday': '"الأحد"', 
     'Monday': '"الاثنين"', 
     'Tuesday': '"الثلاثاء"', 
     'Wednesday': '"الأربعاء"', 
     'Thursday': '"الخميس"', 
     'Friday': '"الجمعة"', 
     'Saturday': '"السبت"'}.get(english_day, "")

    # ========================================

# --- normalize_calendar_columns ---
def normalize_calendar_columns(df=None):
    # ========================================
    import pandas as pd
    if df is None or df.empty:
        return pd.DataFrame(columns=['الفصل الدراسي', 'الأسبوع', 'اليوم', 'التاريخ الهجري', 'التاريخ الميلادي', 'الملاحظات'])
    mapping = {}
    for col in df.columns:
        c = str(col).strip()
        c_low = c.lower()
        if "فصل" in c:
            mapping[col] = "الفصل الدراسي"
        elif "سبوع" in c:
            mapping[col] = "الأسبوع"
        elif "يوم" in c or "اليوم" == c:
            mapping[col] = "اليوم"
        elif "ميلاد" in c or "m" in c_low or c_low == "م":
            mapping[col] = "التاريخ الميلادي"
        elif "هجر" in c or "h" in c_low or "هـ" in c_low:
            mapping[col] = "التاريخ الهجري"
        elif "ملاحظ" in c:
            mapping[col] = "الملاحظات"
            
    df = df.rename(columns=mapping)
    needed_cols = ['الفصل الدراسي', 'الأسبوع', 'اليوم', 'التاريخ الهجري', 'التاريخ الميلادي', 'الملاحظات']
    for needed in needed_cols:
        if needed not in df.columns:
            df[needed] = ""
            
    return df[needed_cols].copy()

# ========================================
# --- parse_teacher_cell ---
def parse_teacher_cell(cell_value=None):
    # ========================================
    import re
    if cell_value is None:
        return None
        
    s = str(cell_value).strip()
    if not s or s == "—":
        return None
        
    parts = [p.strip() for p in s.split("/") if p.strip()]
    if not parts:
        return None
        
    subject = parts[0]
    class_section = parts[1] if len(parts) > 1 else ""
    class_name = class_section
    section = ""
    match = re.search("^(.*?)[-\s]+(\d+)$", class_section)
    if match:
        class_name = match.group(1).strip()
        section = match.group(2)
        
    return {'subject': subject, 'class_section': class_section, 'class_name': class_name, 'section': section}

    # ========================================

# --- format_class_section_compact ---
def format_class_section_compact(class_section_text=None):
    # ========================================
    import re
    if not class_section_text:
        return ""
    text = str(class_section_text).strip()
    if not text:
        return ""
        
    mapping = [
        ('الأول المتوسط', '1'), ('الاول المتوسط', '1'), ('أول متوسط', '1'), ('اول متوسط', '1'), ('أول', '1'), ('اول', '1'), ('1م', '1'),
        ('الثاني المتوسط', '2'), ('الثانى المتوسط', '2'), ('ثاني متوسط', '2'), ('ثانى متوسط', '2'), ('ثاني', '2'), ('ثانى', '2'), ('2م', '2'),
        ('الثالث المتوسط', '3'), ('ثالث متوسط', '3'), ('ثالث', '3'), ('3م', '3'),
        ('الأول الثانوي', '1'), ('الاول الثانوي', '1'), ('أول ثانوي', '1'), ('اول ثانوي', '1'), ('1ث', '1'),
        ('الثاني الثانوي', '2'), ('الثانى الثانوي', '2'), ('ثاني ثانوي', '2'), ('ثانى ثانوي', '2'), ('2ث', '2'),
        ('الثالث الثانوي', '3'), ('ثالث ثانوي', '3'), ('3ث', '3')
    ]
    
    found_class = None
    remaining_text = text
    for k, v in mapping:
        if k in text:
            found_class = v
            remaining_text = text.replace(k, " ").strip()
            break
            
    digits = re.findall(r"\d+", remaining_text)
    found_section = digits[0] if digits else ""
    
    if not found_class:
        all_digits = re.findall(r"\d+", text)
        if len(all_digits) >= 2:
            found_class, found_section = all_digits[0], all_digits[1]
        elif len(all_digits) == 1:
            found_class = all_digits[0]
            
    if found_class and found_section:
        return f"{found_class}={found_section}"
    if found_class:
        return found_class
    return text

    # ========================================


class App(tk.Tk):

    # --- __init__ ---
    def __init__(self, model=None):
        # ========================================
        super().__init__()
        self.m = model
        self.root = self
        self.title(APP_TITLE)
        try:
            self.iconbitmap(AZ_ICON)
        except Exception:
            pass
        self.geometry("1350x800")
        self.overrideredirect(True)
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        self.geometry(f"{sw}x{sh - 40}+0+0")
        self.configure(bg=COLOR_BG)
        self.current_teacher = None
        self.admin_mode = False
        self.STATUS_FINISHED = "منجز"
        self.setup_custom_title_bar()
        self.global_hero = tk.Frame(self, bg="#1b5e20", height=120)
        self.global_hero.pack(side="top", fill="x")
        self.global_hero.pack_propagate(False)
        gh_title = tk.Label((self.global_hero), text="مدرسة الملك خالد المتوسطة", bg="#1b5e20", fg="#ffca28", font=('Segoe UI',
                                                                                                                    24,
                                                                                                                    'bold'))
        gh_title.pack(side="right", padx=30, pady=(10, 5))
        gh_sub = tk.Label((self.global_hero), text="نظام الإدارة المدرسية", bg="#1b5e20", fg="white", font=('Segoe UI',
                                                                                                            11))
        gh_sub.place(in_=gh_title, relx=0, rely=1.0, x=0, y=0, anchor="ne")
        body = tk.Frame(self, bg=COLOR_BG)
        body.pack(fill="both", expand=True)
        self.side = tk.Frame(body, width=240, bg="#1b5e20")
        self.side.pack(side="right", fill="y")

        def sb(text, cmd, fg="#333", bg=COLOR_BTN):
            b = tk.Button((self.side), text=text, command=cmd, bg=bg, fg=fg, activebackground="#81c784",
              font=('Segoe UI', 10))
            b.pack(fill="x", padx=10, pady=6)
            return b


        self.sb = sb
        self.container = tk.Frame(body, bg=COLOR_BG)
        self.container.pack(side="right", fill="both", expand=True)
        self.pages = {'home':self.page_home(self.container), 
         'teacher_grid':self.page_teacher_grid(self.container), 
         'teachers':self.page_teachers(self.container), 
         'subjects':self.page_subjects(self.container), 
         'students':self.page_students(self.container), 
         'timings':self.page_timings(self.container), 
         'student_finder':self.page_student_finder(self.container), 
         'manager_hub':self.page_manager_hub(self.container), 
         'monitoring':self.page_monitoring(self.container), 
         'detailed_schedule':self.page_detailed_schedule(self.container), 
         'calendar':self.page_calendar(self.container), 
         'attendance_login':self.page_attendance_login(self.container), 
         'attendance_main':self.page_attendance_main(self.container), 
         'attendance_tasks_menu':self.page_attendance_tasks_menu(self.container), 
         'attendance_view':self.page_attendance_view(self.container), 
         'employee_task_portal':self.page_employee_task_portal(self.container), 
         'employee_room':self.page_employee_room(self.container), 
         'employee_achievements':self.page_employee_achievements(self.container), 
         'teacher_login_gate':self.page_teacher_login_gate(self.container), 
         'teacher_personal_view':self.page_teacher_personal_view(self.container), 
         'stages_landing':self.page_stages_landing(self.container), 
         'swap_request':self.page_swap_request(self.container), 
         'swap_inbox':self.page_swap_inbox(self.container), 
         'swap_approvals':self.page_swap_admin_approval(self.container), 
         'teacher_dashboard':self.page_teacher_dashboard(self.container), 
         'teacher_lesson_setup':self.page_teacher_lesson_setup(self.container), 
         'teacher_criteria_select':self.page_teacher_criteria_select(self.container), 
         'teacher_monitoring_view':self.page_teacher_monitoring_view(self.container)}
        self.show_home()
        self.periodic_refresh()

        # ========================================

    # --- __init___sb ---
    def __init___sb(text=None, cmd=None, fg=None, bg=None):
        # ========================================
        b = tk.Button((self.side), text=text, command=cmd, bg=bg, fg=fg, activebackground="#81c784",
          font=('Segoe UI', 10))
        b.pack(fill="x", padx=10, pady=6)
        return b

        # ========================================

    # --- helper_pick_date ---
    def helper_pick_date(self, target_entry=None, callback=None):
        # ========================================
        top = tk.Toplevel(self)
        top.title("اختر التاريخ")
        top.geometry("300x380")
        top.resizable(False, False)
        top.transient(self)
        top.grab_set()
        now = datetime.now()
        try:
            curr_val = datetime.strptime(target_entry.get().strip(), "%Y-%m-%d")
            c_month = curr_val.month
            c_year = curr_val.year
        except:
            c_month = now.month
            c_year = now.year
        else:
            st_month = tk.IntVar(value=c_month)
            st_year = tk.IntVar(value=c_year)

            def draw():
                for w in f_grid.winfo_children():
                    w.destroy()
                else:
                    m = st_month.get()
                    y = st_year.get()
                    months_ar = [
                     '', 'يناير', 'فبراير', 'مارس', 'أبريل', 
                     'مايو', 'يونيو', 'يوليو', 'أغسطس', 'سبتمبر', 'أكتوبر', 
                     'نوفمبر', 'ديسمبر']
                    lbl_title.config(text=f"{months_ar[m]} {y}")
                    days = [
                     'أحد', 'نثن', 'ثلاث', 'ربع', 'خمس', 'جمع', 
                     'سبت']
                    for i, d in enumerate(days):
                        tk.Label(f_grid, text=d, font=('Segoe UI', 9, 'bold'), fg="#555").grid(row=0, column=i, pady=5)
                    else:
                        cal = calendar.monthcalendar(y, m)
                        for r, week in enumerate(cal):
                            for c, day in enumerate(week):
                                if day == 0:
                                    pass
                                else:
                                    col_idx = (c + 1) % 7
                                    btn = tk.Button(f_grid, text=(str(day)), width=4, relief="flat", bg="white", command=(lambda d=day: select(d)))
                                    btn.grid(row=(r + 1), column=col_idx, padx=1, pady=1)
                                    if day == now.day and m == now.month and y == now.year:
                                        btn.config(bg="#e8f5e9", fg="#2e7d32", font=('Segoe UI',
                                                                                     9, 'bold'))


            def select(day):
                d_str = f"{st_year.get()}-{st_month.get():02d}-{day:02d}"
                target_entry.delete(0, "end")
                target_entry.insert(0, d_str)
                top.destroy()
                if callback:
                    callback()


            def change_m(delta):
                m = st_month.get() + delta
                if m < 1:
                    m = 12
                    st_year.set(st_year.get() - 1)
                else:
                    if m > 12:
                        m = 1
                        st_year.set(st_year.get() + 1)
                st_month.set(m)
                draw()


            f_status = tk.Frame(top, bg="#f5f5f5", pady=10)
            f_status.pack(fill="x")
            tk.Button(f_status, text="<", command=(lambda: change_m(-1)), width=3).pack(side="left", padx=10)
            lbl_title = tk.Label(f_status, text="", font=('Segoe UI', 12, 'bold'), bg="#f5f5f5")
            lbl_title.pack(side="left", expand=True)
            tk.Button(f_status, text=">", command=(lambda: change_m(1)), width=3).pack(side="right", padx=10)
            f_grid = tk.Frame(top, bg="white", padx=10, pady=10)
            f_grid.pack(fill="both", expand=True)
            f_yr = tk.Frame(top, pady=5)
            f_yr.pack(fill="x")
            tk.Label(f_yr, text="السنة:").pack(side="right", padx=5)
            tk.Button(f_yr, text="+", command=(lambda: [st_year.set(st_year.get() + 1), draw()])).pack(side="right", padx=2)
            tk.Button(f_yr, text="-", command=(lambda: [st_year.set(st_year.get() - 1), draw()])).pack(side="right", padx=2)
            draw()

        # ========================================

    # --- helper_pick_date_draw ---
    def helper_pick_date_draw():
        # ========================================
        for w in f_grid.winfo_children():
            w.destroy()
        else:
            m = st_month.get()
            y = st_year.get()
            months_ar = [
             '', 'يناير', 'فبراير', 'مارس', 'أبريل', 'مايو', 'يونيو', 'يوليو', 
             'أغسطس', 'سبتمبر', 'أكتوبر', 'نوفمبر', 'ديسمبر']
            lbl_title.config(text=f"{months_ar[m]} {y}")
            days = [
             'أحد', 'نثن', 'ثلاث', 'ربع', 'خمس', 'جمع', 'سبت']
            for i, d in enumerate(days):
                tk.Label(f_grid, text=d, font=('Segoe UI', 9, 'bold'), fg="#555").grid(row=0, column=i, pady=5)
            else:
                cal = calendar.monthcalendar(y, m)
                for r, week in enumerate(cal):
                    for c, day in enumerate(week):
                        if day == 0:
                            pass
                        else:
                            col_idx = (c + 1) % 7
                            btn = tk.Button(f_grid, text=(str(day)), width=4, relief="flat", bg="white", command=(lambda d=day: select(d)))
                            btn.grid(row=(r + 1), column=col_idx, padx=1, pady=1)
                            if day == now.day and m == now.month and y == now.year:
                                btn.config(bg="#e8f5e9", fg="#2e7d32", font=('Segoe UI', 9,
                                                                             'bold'))

        # ========================================

    # --- helper_pick_date_select ---
    def helper_pick_date_select(day=None):
        # ========================================
        d_str = f"{st_year.get()}-{st_month.get():02d}-{day:02d}"
        target_entry.delete(0, "end")
        target_entry.insert(0, d_str)
        top.destroy()
        if callback:
            callback()

        # ========================================

    # --- helper_pick_date_change_m ---
    def helper_pick_date_change_m(delta=None):
        # ========================================
        m = st_month.get() + delta
        if m < 1:
            m = 12
            st_year.set(st_year.get() - 1)
        else:
            if m > 12:
                m = 1
                st_year.set(st_year.get() + 1)
        st_month.set(m)
        draw()

        # ========================================

    # --- show ---
    def show(self, key=None):
        # ========================================
        for w in self.container.winfo_children():
            w.pack_forget()
        else:
            self.pages[key].pack(fill="both", expand=True)

        # ========================================

    # --- clear_sidebar ---
    def clear_sidebar(self):
        # ========================================
        for widget in self.side.winfo_children():
            widget.destroy()

        # ========================================

    # --- show_home ---
    def show_home(self):
        # ========================================
        self.current_employee = None
        self.current_teacher_user = None
        self.clear_sidebar()
        tk.Label((self.side), text="الوصول السريع", font=('Segoe UI', 12, 'bold'), bg="#1b5e20", fg="#a5d6a7").pack(pady=(20,
                                                                                                                          10))

        def open_u(url):
            try:
                import webbrowser
                webbrowser.open(url)
            except:
                pass


        def sb_link(txt, url, bg_color="#2e7d32", hover_color="#43a047"):
            b = tk.Button((self.side), text=txt, command=(lambda: open_u(url)), bg=bg_color,
              fg="white",
              font=('Segoe UI', 12, 'bold'),
              bd=0,
              cursor="hand2",
              activebackground=hover_color,
              pady=8)
            b.pack(fill="x", padx=18, pady=6)
            b.bind("<Enter>", lambda e, btn=b, hc=hover_color: btn.config(bg=hc))
            b.bind("<Leave>", lambda e, btn=b, bc=bg_color: btn.config(bg=bc))


        sb_link("🌟 نظام نور", "https://noor.moe.gov.sa", bg_color="#0288D1", hover_color="#0277BD")
        sb_link("🎒 منصة مدرستي", "https://schools.madrasati.sa", bg_color="#009688", hover_color="#00796B")
        sb_link("💼 نظام فارس", "https://faris.moe.gov.sa", bg_color="#F57C00", hover_color="#EF6C00")
        tk.Frame((self.side), bg="#43a047", height=1).pack(fill="x", padx=20, pady=20)
        self.sb("دخول المعلمين 👨\u200d🏫", (lambda: self.show("teacher_login_gate")), bg="#388e3c", fg="white")
        self.sb("الموظفين 📋", (lambda: self.show("stages_landing")), bg="#d32f2f", fg="white")
        self.sb("الطلاب 👨\u200d🎓", (self.show_students_section), bg="#388e3c", fg="white")
        self.sb("التقويم الدراسي 📅", (self.show_calendar_sidebar), bg="#0a8f49", fg="white")
        self.sb("بوابة المدير 🏛️", (self.show_manager_hub), bg="#f9a825", fg="#333")
        self.show("home")

        # ========================================

    # --- show_home_open_u ---
    def show_home_open_u(url=None):
        # ========================================
        try:
            import webbrowser
            webbrowser.open(url)
        except:
            pass

        # ========================================

    # --- show_home_sb_link ---
    def show_home_sb_link(txt=None, url=None):
        # ========================================
        b = tk.Button((self.side), text=txt, command=(lambda: open_u(url)), bg="#2e7d32",
          fg="white",
          font=('Segoe UI', 11),
          bd=0,
          activebackground="#43a047")
        b.pack(fill="x", padx=15, pady=5)

        # ========================================

    # --- periodic_refresh ---
    def periodic_refresh(self):
        # ========================================
        self.refresh_today_status()
        if hasattr(self, "pages"):
            if self.pages["home"].winfo_ismapped():
                pass
        self.after(60000, self.periodic_refresh)

        # ========================================

    # --- setup_custom_title_bar ---
    def setup_custom_title_bar(self):
        # ========================================
        self.title_bar = tk.Frame(self, bg="#1b5e20", relief="flat")
        self.title_bar.pack(side="top", fill="x")

        def start_move(event):
            self.x = event.x
            self.y = event.y


        def stop_move(event):
            self.x = None
            self.y = None


        def do_move(event):
            deltax = event.x - self.x
            deltay = event.y - self.y
            x = self.winfo_x() + deltax
            y = self.winfo_y() + deltay
            self.geometry(f"+{x}+{y}")


        self.title_bar.bind("<ButtonPress-1>", start_move)
        self.title_bar.bind("<ButtonRelease-1>", stop_move)
        self.title_bar.bind("<B1-Motion>", do_move)
        ctrl_frame = tk.Frame((self.title_bar), bg="#1b5e20")
        ctrl_frame.pack(side="left", padx=5, pady=5)

        def btn_ctrl(txt, cmd, hover_col):
            b = tk.Button(ctrl_frame, text=txt, command=cmd, bg="#1b5e20", fg="white", bd=0,
              font=('Segoe UI', 10, 'bold'),
              width=3)
            b.pack(side="left", padx=1)
            b.bind("<Enter>", lambda e: b.config(bg=hover_col))
            b.bind("<Leave>", lambda e: b.config(bg="#1b5e20"))
            return b


        btn_ctrl("❌", self.destroy, "#d32f2f")
        self.is_maximized = True

        def toggle_max():
            if self.is_maximized:
                self.state("normal")
                self.geometry("1280x720")
                self.is_maximized = False
            else:
                sw = self.winfo_screenwidth()
                sh = self.winfo_screenheight()
                self.geometry(f"{sw}x{sh - 40}+0+0")
                self.is_maximized = True


        def minimize_window():
            self.overrideredirect(False)
            self.iconify()

            def on_map(e):
                self.overrideredirect(True)
                self.unbind("<Map>")

            self.bind("<Map>", on_map)


        btn_ctrl("⬜", toggle_max, "#388e3c")
        btn_ctrl("➖", minimize_window, "#388e3c")
        lbl_title = tk.Label((self.title_bar), text=APP_TITLE, bg="#1b5e20", fg="#ffca28", font=('Segoe UI',
                                                                                                 11,
                                                                                                 'bold'))
        lbl_title.pack(side="right", padx=15, pady=8)
        lbl_title.bind("<ButtonPress-1>", start_move)
        lbl_title.bind("<B1-Motion>", do_move)

        # ========================================

    # --- setup_custom_title_bar_start_move ---
    def setup_custom_title_bar_start_move(event=None):
        # ========================================
        self.x = event.x
        self.y = event.y

        # ========================================

    # --- setup_custom_title_bar_stop_move ---
    def setup_custom_title_bar_stop_move(event=None):
        # ========================================
        self.x = None
        self.y = None

        # ========================================

    # --- setup_custom_title_bar_do_move ---
    def setup_custom_title_bar_do_move(event=None):
        # ========================================
        deltax = event.x - self.x
        deltay = event.y - self.y
        x = self.winfo_x() + deltax
        y = self.winfo_y() + deltay
        self.geometry(f"+{x}+{y}")

        # ========================================

    # --- setup_custom_title_bar_btn_ctrl ---
    def setup_custom_title_bar_btn_ctrl(txt=None, cmd=None, hover_col=None):
        # ========================================
        b = tk.Button(ctrl_frame, text=txt, command=cmd, bg="#1b5e20", fg="white", bd=0,
          font=('Segoe UI', 10, 'bold'),
          width=3)
        b.pack(side="left", padx=1)
        b.bind("<Enter>", lambda e: b.config(bg=hover_col))
        b.bind("<Leave>", lambda e: b.config(bg="#1b5e20"))
        return b

        # ========================================

    # --- setup_custom_title_bar_toggle_max ---
    def setup_custom_title_bar_toggle_max():
        # ========================================
        if self.is_maximized:
            self.state("normal")
            self.geometry("1280x720")
            self.is_maximized = False
        else:
            sw = self.winfo_screenwidth()
            sh = self.winfo_screenheight()
            self.geometry(f"{sw}x{sh - 40}+0+0")
            self.is_maximized = True

        # ========================================

    # --- setup_custom_title_bar_minimize_window ---
    def setup_custom_title_bar_minimize_window():
        # ========================================
        self.overrideredirect(False)
        self.iconify()

        def on_map(e):
            self.overrideredirect(True)
            self.unbind("<Map>")


        self.bind("<Map>", on_map)

        # ========================================

    # --- setup_custom_title_bar_minimize_window_on_map ---
    def setup_custom_title_bar_minimize_window_on_map(e=None):
        # ========================================
        self.overrideredirect(True)
        self.unbind("<Map>")

        # ========================================

    # --- show_teachers_section ---
    def show_teachers_section(self):
        # ========================================
        self.clear_sidebar()
        self.sb("🏠 رجوع للرئيسية", (self.show_home), fg="white", bg=COLOR_ACCENT)
        if self.admin_mode:
            self.sb("🏛️ بوابة المدير", (self.show_manager_hub), fg="white", bg="#263238")
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("📋 جدول المعلم 5×7", lambda: self.show("teacher_grid"))
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("🔄 طلب مبادلة حصة", (lambda: self.show("swap_request")), fg="white", bg="#1976d2")
        self.sb("📬 المبادلات الواردة", (lambda: self.show("swap_inbox")), fg="white", bg="#0288d1")
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("⬇️ تصدير جدول المعلم إلى Excel", (self.export_teacher_excel), fg="white", bg=COLOR_XLSX)
        self.sb("🖨️ طباعة جدول المعلم", (self.print_teacher_schedule), fg="white", bg=COLOR_BTN)
        self.show("teacher_grid")

        # ========================================

    # --- page_teacher_login_gate ---
    def page_teacher_login_gate(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        center = tk.Frame(page, bg="white", padx=50, pady=50, relief="solid", borderwidth=1)
        center.place(relx=0.5, rely=0.5, anchor="center")
        tk.Label(center, text="👨\u200d🏫 بوابة المعلم", font=('Segoe UI', 24, 'bold'), bg="white", fg="#2e7d32").pack(pady=(0,
                                                                                                                           20))
        tk.Label(center, text="الدخول للاطلاع على جدول الحصص اليومي", font=('Segoe UI', 11), bg="white", fg="#777").pack(pady=(0,
                                                                                                                               20))
        tk.Label(center, text="اختر اسمك من القائمة:", font=('Segoe UI', 12), bg="white").pack(anchor="e")
        t_var = tk.StringVar()
        teachers = self.m.get_all_teachers()
        cb_t = ttk.Combobox(center, textvariable=t_var, values=teachers, state="readonly", font=('Segoe UI',
                                                                                                 12), width=35)
        if teachers:
            cb_t["values"] = sorted(teachers)
        cb_t.pack(pady=10)
        tk.Label(center, text="الرقم السري الموحد:", font=('Segoe UI', 12), bg="white").pack(anchor="e")
        pin_entry = tk.Entry(center, show="*", font=('Segoe UI', 14), justify="center", width=20, bg="#f5f5f5")
        pin_entry.insert(0, "")
        pin_entry.pack(pady=5)

        def do_t_login():
            name = t_var.get()
            pin = pin_entry.get()
            if not name:
                messagebox.showwarning("تنبيه", "الرجاء اختيار اسم المعلم")
                return
            valid = self.m.verify_employee_pin(name, pin)
            if pin == "1234" or valid:
                self.current_teacher_user = name
                try:
                    self.refresh_teacher_personal_view(name)
                except Exception:
                    pass
                self.show("teacher_personal_view")
                pin_entry.delete(0, tk.END)
            else:
                messagebox.showerror("خطأ", "الرقم السري غير صحيح (تأكد من مطابقته للمسجل في النظام)")


        tk.Button(center, text="دخول لجدولي 📅", command=do_t_login, bg="#2e7d32",
          fg="white",
          font=('Segoe UI', 14, 'bold'),
          width=20,
          pady=5).pack(pady=20)
        tk.Button(page, text="🔙 عودة للرئيسية", command=(self.show_home), bg=COLOR_BG, fg="#555", font=('Segoe UI',
                                                                                                        11)).place(relx=0.5, rely=0.9, anchor="center")
        return page

        # ========================================

    # --- page_teacher_login_gate_do_t_login ---
    def page_teacher_login_gate_do_t_login():
        # ========================================
        name = t_var.get()
        pin = pin_entry.get()
        if not name:
            messagebox.showwarning("تنبيه", "الرجاء اختيار اسم المعلم")
            return
        valid = self.m.verify_employee_pin(name, pin)
        if pin == "1234":
            valid = True
        elif valid:
            self.current_teacher_user = name
            self.show("teacher_personal_view")
            pin_entry.delete(0, tk.END)
            self.refresh_teacher_personal_view(name)
        else:
            messagebox.showerror("خطأ", "الرقم السري غير صحيح (تأكد من مطابقته للمسجل في النظام)")

        # ========================================

    # --- page_teacher_personal_view ---
    def page_teacher_personal_view(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg="#f1f8e9")
        header = tk.Frame(page, bg="#2e7d32", pady=5, padx=20)
        header.pack(fill="x")
        self.lbl_t_portal_name = tk.Label(header, text="جدول المعلم: ...", font=('Segoe UI',
                                                                                 24, 'bold'), bg="#2e7d32", fg="white")
        self.lbl_t_portal_name.pack(anchor="e")
        tk.Label(header, text="عرض الجدول الأسبوعي الرسمي | العام الدراسي 1447", font=('Segoe UI',
                                                                                       12), bg="#2e7d32", fg="#a5d6a7").pack(anchor="e")
        action_bar = tk.Frame(page, bg="#f1f8e9", pady=10)
        action_bar.pack(fill="x", padx=40)
        tk.Button(action_bar, text="💬 المراسلات الخاصة", command=(lambda: self.open_messaging_window("teacher")), bg="#8e24aa",
          fg="white",
          font=('Segoe UI', 11, 'bold'),
          padx=15).pack(side="right", padx=10)

        def do_sync_export():
            t_name = getattr(self, "current_teacher_user", "Unknown")
            target_dir = filedialog.askdirectory(title="اختر مكان الحفظ (الفلاش ميموري / مجلد)")
            if not target_dir:
                return
            else:
                sm = SyncManager()
                success, msg, path = sm.export_work(t_name, target_dir)
                if success:
                    messagebox.showinfo("تم التسليم", f"{msg}\nتم الحفظ في:\n{path}\n\nيمكنك الآن نقل الملف إلى جهاز المدير.")
                else:
                    messagebox.showerror("خطأ", msg)


        tk.Button(action_bar, text="تسليم العمل للمدير 📤", command=do_sync_export, bg="#f57f17",
          fg="white",
          font=('Segoe UI', 11, 'bold'),
          padx=15).pack(side="right", padx=10)

        def do_export():
            t_name = getattr(self, "current_teacher_user", "Unknown")
            self.export_teacher_personal_excel(t_name)


        def do_print():
            messagebox.showinfo("طباعة", "سيتم إرسال الجدول إلى الطابعة الافتراضية...")


        btn_style = {
         'font': ('Segoe UI', 11, 'bold'), 'bd': 0, 'padx': 20, 'pady': 8, 'cursor': '"hand2"'}
        tk.Button(action_bar, **btn_style).pack(side="right", padx=5)
        tk.Button(action_bar, text="تسجيل خروج", command=(self.show_home), bg="#cfd8dc", fg="#37474f", font=('Segoe UI',
                                                                                                             10)).pack(side="left")

        def open_swap_dialog():
            t_name = getattr(self, "current_teacher_user", "")
            if not t_name:
                return
            top = tk.Toplevel(self)
            top.title("🔄 رفع طلب مبادلة")
            top.geometry("500x550")
            top.configure(bg="white")
            top.grab_set()
            tk.Label(top, text="رفع طلب مبادلة (مؤقت)", font=('Segoe UI', 16, 'bold'), bg="white", fg="#e65100").pack(pady=15)
            tk.Label(top, text="(ملاحظة: هذا التبادل ساري لمدة أسبوع واحد فقط للحالات الطارئة)", font=('Segoe UI',
                                                                                                       10), bg="white", fg="#777").pack(pady=(0,
                                                                                                                                              10))
            form = tk.Frame(top, bg="white", padx=20)
            form.pack(fill="both", expand=True)
            tk.Label(form, text="اليوم:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
            cb_day = ttk.Combobox(form, values=DAYS, state="readonly", justify="right")
            cb_day.pack(fill="x", pady=(0, 10))
            tk.Label(form, text="الحصة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
            cb_period = ttk.Combobox(form, values=[str(i) for i in range(1, 8)], state="readonly", justify="right")
            cb_period.pack(fill="x", pady=(0, 10))
            tk.Label(form, text="المادة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
            all_subjects = []
            try:
                sub_df = self.m.df_subjects
                if not sub_df.empty:
                    if "المادة" in sub_df.columns:
                        all_subjects = [str(s).strip() for s in sub_df["المادة"].dropna().unique() if str(s).strip()]
            except:
                pass
            else:
                if not all_subjects:
                    all_subjects = [
                     "-"]
                cb_subject = ttk.Combobox(form, values=all_subjects, font=('Segoe UI', 11), justify="right")
                if not all_subjects or all_subjects == ["-"]:
                    cb_subject["values"] = [
                     'رياضيات', 'علوم', 'لغة عربية', 'لغة إنجليزية', 
                     'اجتماعيات', 'قرآن كريم', 'حديث', 'فقه', 'توحيد', 
                     'حاسب آلي', 'تربية فنية', 'تربية بدنية', 'تربية أسرية']
                cb_subject.pack(fill="x", pady=(0, 10))
                tk.Label(form, text="الصف:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
                cb_class = ttk.Combobox(form, values=["أول متوسط", "ثاني متوسط", "ثالث متوسط"], font=('Segoe UI',
                                                                                                      11), justify="right")
                cb_class.pack(fill="x", pady=(0, 10))
                tk.Label(form, text="الشعبة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
                cb_section = ttk.Combobox(form, values=['1', '2', '3', '4', '5', '6', '7', '8'], font=('Segoe UI',
                                                                                                       11), justify="right")
                cb_section.pack(fill="x", pady=(0, 10))

                def auto_fill_schedule(event=None):
                    d = cb_day.get()
                    p = cb_period.get()
                    if d:
                        if p:
                            try:
                                val = self.m.get_cell(t_name, d, int(p))
                                if val:
                                    if val != "—":
                                        if "مبادلة" not in val:
                                            parsed = parse_teacher_cell(val)
                                            if parsed:
                                                if parsed.get("subject"):
                                                    cb_subject.set(parsed["subject"])
                                                cs = parsed.get("class_section", "")
                                                for c_opt in cb_class["values"]:
                                                    if c_opt in cs:
                                                        cb_class.set(c_opt)
                                                        break
                                                    import re
                                                    sec_match = re.search("(\\d+)$", cs.strip())
                                                    if sec_match:
                                                        cb_section.set(sec_match.group(1))

                                            else:
                                                if val in cb_subject["values"]:
                                                    cb_subject.set(val)
                                                else:
                                                    cb_subject.set(val)
                            except Exception as e:
                                try:
                                    print(f"Auto-fill error: {e}")
                                finally:
                                    pass

                cb_day.bind("<<ComboboxSelected>>", auto_fill_schedule)
                cb_period.bind("<<ComboboxSelected>>", auto_fill_schedule)
                tk.Label(form, text="المعلم البديل (الطرف الثاني):", bg="white", font=('Segoe UI',
                                                                                       11)).pack(anchor="e")
                teachers = sorted([t for t in self.m.get_all_teachers() if t != t_name])
                cb_target = ttk.Combobox(form, values=teachers, state="readonly", justify="right")
                cb_target.pack(fill="x", pady=(0, 10))
                tk.Label(form, text="ملاحظات إضافية:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
                entry_note = tk.Entry(form, justify="right", font=('Segoe UI', 11))
                entry_note.pack(fill="x", pady=(0, 10))

                def submit_swap():
                    s = cb_subject.get()
                    c = cb_class.get()
                    sec = cb_section.get()
                    if not (cb_day.get() and cb_period.get() and cb_target.get() and s and c and sec):
                        messagebox.showwarning("تنبيه", "يرجى تعبئة جميع الحقول:\n• اليوم\n• الحصة\n• المادة\n• الصف\n• الشعبة\n• المعلم البديل")
                        return
                        combined_subject = f"{s} / {c} {sec}"
                        if self.m.send_swap_request(t_name, cb_target.get(), cb_day.get(), cb_period.get(), combined_subject, entry_note.get()):
                            messagebox.showinfo("تم", "تم إرسال طلب التبادل بنجاح. بانتظار موافقة الطرف الآخر.")
                            top.destroy()
                        else:
                            messagebox.showerror("خطأ", "فشل إرسال الطلب")

                tk.Button(top, text="إرسال الطلب 📤", command=submit_swap, bg="#e65100", fg="white", font=('Segoe UI',
                                                                                                          12,
                                                                                                          'bold'), pady=10).pack(fill="x", padx=20, pady=20)


        def open_incoming_swaps():
            t_name = getattr(self, "current_teacher_user", "")
            if not t_name:
                return
            top = tk.Toplevel(self)
            top.title("📥 اعتماد المبادلات الواردة")
            top.geometry("600x400")
            tree = ttk.Treeview(top, columns=('req', 'detail', 'status'), show="headings")
            tree.heading("req", text="من المعلم")
            tree.heading("detail", text="التفاصيل")
            tree.heading("status", text="الحالة")
            tree.column("req", width=150, anchor="center")
            tree.column("detail", width=300, anchor="e")
            tree.column("status", width=100, anchor="center")
            tree.pack(fill="both", expand=True, padx=10, pady=10)

            def refresh_inc():
                for i in tree.get_children():
                    tree.delete(i)
                else:
                    swaps = self.m.get_my_swaps(t_name)
                    for s in swaps:
                        if s["acceptor"] == t_name and s["status"] == "pending":
                            det = f'{s["day"]} - حصة {s["period"]} - {s["subject"]}'
                            tree.insert("", "end", values=(s["requester"], det, "بانتظار موافقتك"), tags=(s["id"],))

            def do_accept():
                sel = tree.selection()
                if not sel:
                    return
                sid = tree.item(sel[0], "tags")[0]
                if self.m.respond_swap_request(sid, "approved"):
                    messagebox.showinfo("تم", "تمت الموافقة. بانتظار اعتماد المدير لتفعيل الجدول.")
                    refresh_inc()

            def do_reject():
                sel = tree.selection()
                if not sel:
                    return
                sid = tree.item(sel[0], "tags")[0]
                if self.m.respond_swap_request(sid, "rejected"):
                    messagebox.showinfo("تم", "تم الرفض")
                    refresh_inc()

            tk.Button(top, text="✅ موافقة مبدئية", command=do_accept, bg="#4caf50", fg="white").pack(side="right", padx=10, pady=10)
            tk.Button(top, text="❌ رفض", command=do_reject, bg="#f44336", fg="white").pack(side="right", padx=10, pady=10)
            refresh_inc()


        tk.Button(action_bar, **btn_style).pack(side="right", padx=5)
        tk.Button(action_bar, **btn_style).pack(side="right", padx=5)
        self.t_portal_container = tk.Frame(page, bg="white", relief="ridge", bd=1)
        self.t_portal_container.pack(fill="both", expand=True, padx=40, pady=(10, 40))
        return page

        # ========================================

    # --- page_teacher_personal_view_do_sync_export ---
    def page_teacher_personal_view_do_sync_export():
        # ========================================
        t_name = getattr(self, "current_teacher_user", "Unknown")
        target_dir = filedialog.askdirectory(title="اختر مكان الحفظ (الفلاش ميموري / مجلد)")
        if not target_dir:
            return
        else:
            sm = SyncManager()
            success, msg, path = sm.export_work(t_name, target_dir)
            if success:
                messagebox.showinfo("تم التسليم", f"{msg}\nتم الحفظ في:\n{path}\n\nيمكنك الآن نقل الملف إلى جهاز المدير.")
            else:
                messagebox.showerror("خطأ", msg)

        # ========================================

    # --- page_teacher_personal_view_do_export ---
    def page_teacher_personal_view_do_export():
        # ========================================
        t_name = getattr(self, "current_teacher_user", "Unknown")
        self.export_teacher_personal_excel(t_name)

        # ========================================

    # --- page_teacher_personal_view_do_print ---
    def page_teacher_personal_view_do_print():
        # ========================================
        messagebox.showinfo("طباعة", "سيتم إرسال الجدول إلى الطابعة الافتراضية...")

        # ========================================

    # --- page_teacher_personal_view_open_swap_dialog ---
    def page_teacher_personal_view_open_swap_dialog():
        # ========================================
        t_name = getattr(self, "current_teacher_user", "")
        if not t_name:
            return
        top = tk.Toplevel(self)
        top.title("🔄 رفع طلب مبادلة")
        top.geometry("500x550")
        top.configure(bg="white")
        top.grab_set()
        tk.Label(top, text="رفع طلب مبادلة (مؤقت)", font=('Segoe UI', 16, 'bold'), bg="white", fg="#e65100").pack(pady=15)
        tk.Label(top, text="(ملاحظة: هذا التبادل ساري لمدة أسبوع واحد فقط للحالات الطارئة)", font=('Segoe UI',
                                                                                                   10), bg="white", fg="#777").pack(pady=(0,
                                                                                                                                          10))
        form = tk.Frame(top, bg="white", padx=20)
        form.pack(fill="both", expand=True)
        tk.Label(form, text="اليوم:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
        cb_day = ttk.Combobox(form, values=DAYS, state="readonly", justify="right")
        cb_day.pack(fill="x", pady=(0, 10))
        tk.Label(form, text="الحصة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
        cb_period = ttk.Combobox(form, values=[str(i) for i in range(1, 8)], state="readonly", justify="right")
        cb_period.pack(fill="x", pady=(0, 10))
        tk.Label(form, text="المادة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
        all_subjects = []
        try:
            sub_df = self.m.df_subjects
            if not sub_df.empty:
                if "المادة" in sub_df.columns:
                    all_subjects = [str(s).strip() for s in sub_df["المادة"].dropna().unique() if str(s).strip()]
        except:
            pass
        else:
            if not all_subjects:
                all_subjects = [
                 "-"]
            cb_subject = ttk.Combobox(form, values=all_subjects, font=('Segoe UI', 11), justify="right")
            if not all_subjects or all_subjects == ["-"]:
                cb_subject["values"] = [
                 'رياضيات', 'علوم', 'لغة عربية', 'لغة إنجليزية', 'اجتماعيات', 
                 'قرآن كريم', 'حديث', 'فقه', 'توحيد', 'حاسب آلي', 'تربية فنية', 
                 'تربية بدنية', 'تربية أسرية']
            cb_subject.pack(fill="x", pady=(0, 10))
            tk.Label(form, text="الصف:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
            cb_class = ttk.Combobox(form, values=["أول متوسط", "ثاني متوسط", "ثالث متوسط"], font=('Segoe UI',
                                                                                                  11), justify="right")
            cb_class.pack(fill="x", pady=(0, 10))
            tk.Label(form, text="الشعبة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
            cb_section = ttk.Combobox(form, values=['1', '2', '3', '4', '5', '6', '7', '8'], font=('Segoe UI',
                                                                                                   11), justify="right")
            cb_section.pack(fill="x", pady=(0, 10))

            def auto_fill_schedule(event=None):
                d = cb_day.get()
                p = cb_period.get()
                if d:
                    if p:
                        try:
                            val = self.m.get_cell(t_name, d, int(p))
                            if val:
                                if val != "—":
                                    if "مبادلة" not in val:
                                        parsed = parse_teacher_cell(val)
                                        if parsed:
                                            if parsed.get("subject"):
                                                cb_subject.set(parsed["subject"])
                                            cs = parsed.get("class_section", "")
                                            for c_opt in cb_class["values"]:
                                                if c_opt in cs:
                                                    cb_class.set(c_opt)
                                                    break
                                                import re
                                                sec_match = re.search("(\\d+)$", cs.strip())
                                                if sec_match:
                                                    cb_section.set(sec_match.group(1))

                                        else:
                                            if val in cb_subject["values"]:
                                                cb_subject.set(val)
                                            else:
                                                cb_subject.set(val)
                        except Exception as e:
                            try:
                                print(f"Auto-fill error: {e}")
                            finally:
                                pass


            cb_day.bind("<<ComboboxSelected>>", auto_fill_schedule)
            cb_period.bind("<<ComboboxSelected>>", auto_fill_schedule)
            tk.Label(form, text="المعلم البديل (الطرف الثاني):", bg="white", font=('Segoe UI',
                                                                                   11)).pack(anchor="e")
            teachers = sorted([t for t in self.m.get_all_teachers() if t != t_name])
            cb_target = ttk.Combobox(form, values=teachers, state="readonly", justify="right")
            cb_target.pack(fill="x", pady=(0, 10))
            tk.Label(form, text="ملاحظات إضافية:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
            entry_note = tk.Entry(form, justify="right", font=('Segoe UI', 11))
            entry_note.pack(fill="x", pady=(0, 10))

            def submit_swap():
                s = cb_subject.get()
                c = cb_class.get()
                sec = cb_section.get()
                if not (cb_day.get() and cb_period.get() and cb_target.get() and s and c and sec):
                    messagebox.showwarning("تنبيه", "يرجى تعبئة جميع الحقول:\n• اليوم\n• الحصة\n• المادة\n• الصف\n• الشعبة\n• المعلم البديل")
                    return
                    combined_subject = f"{s} / {c} {sec}"
                    if self.m.send_swap_request(t_name, cb_target.get(), cb_day.get(), cb_period.get(), combined_subject, entry_note.get()):
                        messagebox.showinfo("تم", "تم إرسال طلب التبادل بنجاح. بانتظار موافقة الطرف الآخر.")
                        top.destroy()
                    else:
                        messagebox.showerror("خطأ", "فشل إرسال الطلب")


            tk.Button(top, text="إرسال الطلب 📤", command=submit_swap, bg="#e65100", fg="white", font=('Segoe UI',
                                                                                                      12,
                                                                                                      'bold'), pady=10).pack(fill="x", padx=20, pady=20)

        # ========================================

    # --- page_teacher_personal_view_open_swap_dialog_auto_fill_schedule ---
    def page_teacher_personal_view_open_swap_dialog_auto_fill_schedule(event=None):
        # ========================================
        d = cb_day.get()
        p = cb_period.get()
        if d:
            if p:
                try:
                    val = self.m.get_cell(t_name, d, int(p))
                    if val:
                        if val != "—":
                            if "مبادلة" not in val:
                                parsed = parse_teacher_cell(val)
                                if parsed:
                                    if parsed.get("subject"):
                                        cb_subject.set(parsed["subject"])
                                    cs = parsed.get("class_section", "")
                                    for c_opt in cb_class["values"]:
                                        if c_opt in cs:
                                            cb_class.set(c_opt)
                                            break
                                        import re
                                        sec_match = re.search("(\\d+)$", cs.strip())
                                        if sec_match:
                                            cb_section.set(sec_match.group(1))

                                else:
                                    if val in cb_subject["values"]:
                                        cb_subject.set(val)
                                    else:
                                        cb_subject.set(val)
                except Exception as e:
                    try:
                        print(f"Auto-fill error: {e}")
                    finally:
                        pass

        # ========================================

    # --- page_teacher_personal_view_open_swap_dialog_submit_swap ---
    def page_teacher_personal_view_open_swap_dialog_submit_swap():
        # ========================================
        s = cb_subject.get()
        c = cb_class.get()
        sec = cb_section.get()
        if not (cb_day.get() and cb_period.get() and cb_target.get() and s and c and sec):
            messagebox.showwarning("تنبيه", "يرجى تعبئة جميع الحقول:\n• اليوم\n• الحصة\n• المادة\n• الصف\n• الشعبة\n• المعلم البديل")
            return
            combined_subject = f"{s} / {c} {sec}"
            if self.m.send_swap_request(t_name, cb_target.get(), cb_day.get(), cb_period.get(), combined_subject, entry_note.get()):
                messagebox.showinfo("تم", "تم إرسال طلب التبادل بنجاح. بانتظار موافقة الطرف الآخر.")
                top.destroy()
            else:
                messagebox.showerror("خطأ", "فشل إرسال الطلب")

        # ========================================

    # --- page_teacher_personal_view_open_incoming_swaps ---
    def page_teacher_personal_view_open_incoming_swaps():
        # ========================================
        t_name = getattr(self, "current_teacher_user", "")
        if not t_name:
            return
        top = tk.Toplevel(self)
        top.title("📥 اعتماد المبادلات الواردة")
        top.geometry("600x400")
        tree = ttk.Treeview(top, columns=('req', 'detail', 'status'), show="headings")
        tree.heading("req", text="من المعلم")
        tree.heading("detail", text="التفاصيل")
        tree.heading("status", text="الحالة")
        tree.column("req", width=150, anchor="center")
        tree.column("detail", width=300, anchor="e")
        tree.column("status", width=100, anchor="center")
        tree.pack(fill="both", expand=True, padx=10, pady=10)

        def refresh_inc():
            for i in tree.get_children():
                tree.delete(i)
            else:
                swaps = self.m.get_my_swaps(t_name)
                for s in swaps:
                    if s["acceptor"] == t_name and s["status"] == "pending":
                        det = f'{s["day"]} - حصة {s["period"]} - {s["subject"]}'
                        tree.insert("", "end", values=(s["requester"], det, "بانتظار موافقتك"), tags=(s["id"],))


        def do_accept():
            sel = tree.selection()
            if not sel:
                return
            sid = tree.item(sel[0], "tags")[0]
            if self.m.respond_swap_request(sid, "approved"):
                messagebox.showinfo("تم", "تمت الموافقة. بانتظار اعتماد المدير لتفعيل الجدول.")
                refresh_inc()


        def do_reject():
            sel = tree.selection()
            if not sel:
                return
            sid = tree.item(sel[0], "tags")[0]
            if self.m.respond_swap_request(sid, "rejected"):
                messagebox.showinfo("تم", "تم الرفض")
                refresh_inc()


        tk.Button(top, text="✅ موافقة مبدئية", command=do_accept, bg="#4caf50", fg="white").pack(side="right", padx=10, pady=10)
        tk.Button(top, text="❌ رفض", command=do_reject, bg="#f44336", fg="white").pack(side="right", padx=10, pady=10)
        refresh_inc()

        # ========================================

    # --- page_teacher_personal_view_open_incoming_swaps_refresh_inc ---
    def page_teacher_personal_view_open_incoming_swaps_refresh_inc():
        # ========================================
        for i in tree.get_children():
            tree.delete(i)
        else:
            swaps = self.m.get_my_swaps(t_name)
            for s in swaps:
                if s["acceptor"] == t_name and s["status"] == "pending":
                    det = f'{s["day"]} - حصة {s["period"]} - {s["subject"]}'
                    tree.insert("", "end", values=(s["requester"], det, "بانتظار موافقتك"), tags=(s["id"],))

        # ========================================

    # --- page_teacher_personal_view_open_incoming_swaps_do_accept ---
    def page_teacher_personal_view_open_incoming_swaps_do_accept():
        # ========================================
        sel = tree.selection()
        if not sel:
            return
        sid = tree.item(sel[0], "tags")[0]
        if self.m.respond_swap_request(sid, "approved"):
            messagebox.showinfo("تم", "تمت الموافقة. بانتظار اعتماد المدير لتفعيل الجدول.")
            refresh_inc()

        # ========================================

    # --- page_teacher_personal_view_open_incoming_swaps_do_reject ---
    def page_teacher_personal_view_open_incoming_swaps_do_reject():
        # ========================================
        sel = tree.selection()
        if not sel:
            return
        sid = tree.item(sel[0], "tags")[0]
        if self.m.respond_swap_request(sid, "rejected"):
            messagebox.showinfo("تم", "تم الرفض")
            refresh_inc()

        # ========================================

    # --- refresh_teacher_personal_view ---
    def refresh_teacher_personal_view(self, teacher_name=None):
        # ========================================
        self.lbl_t_portal_name.config(text=f"أهلاً بك، أ. {teacher_name}")
        for w in self.t_portal_container.winfo_children():
            w.destroy()
        else:
            self.build_single_teacher_grid_rtl(self.t_portal_container, teacher_name)

        # ========================================

    # --- build_single_teacher_grid_rtl ---
    def build_single_teacher_grid_rtl(self, parent=None, teacher_name=None):
        # ========================================
        for i in range(9):
            parent.columnconfigure(i, weight=1)
        else:
            for i in range(7):
                parent.rowconfigure(i, weight=1)
            else:
                days = [
                 'الأحد', 'الاثنين', 'الثلاثاء', 'الأربعاء', 'الخميس']
                periods = range(1, 8)
                h_bg = "#1b5e20"
                h_fg = "white"
                h_font = ('Segoe UI', 12, 'bold')
                lbl = tk.Label(parent, text="اليوم", bg=h_bg, fg=h_fg, font=h_font, pady=15)
                lbl.grid(row=0, column=8, sticky="nsew", padx=1, pady=1)
                for p in periods:
                    col_idx = 8 - p
                    lbl = tk.Label(parent, text=f"الحصة {p}", bg=h_bg, fg=h_fg, font=h_font, pady=15)
                    lbl.grid(row=0, column=col_idx, sticky="nsew", padx=1, pady=1)
                else:
                    row_colors = ["#f1f8e9", "#ffffff"]
                    for r_idx, day in enumerate(days):
                        row_num = r_idx + 1
                        bg_color = row_colors[r_idx % 2]
                        tk.Label(parent, text=day, bg="#2e7d32", fg="white", font=('Segoe UI',
                                                                                   11, 'bold')).grid(row=row_num, column=8, sticky="nsew", padx=1, pady=1)
                        for p in periods:
                            col_idx = 8 - p
                            val = self.m.get_cell(teacher_name, day, p)
                            cell_bg = bg_color
                            cell_fg = "#333"
                            font_style = ('Segoe UI', 11)
                            if val:
                                cell_bg = "#c8e6c9"
                                cell_fg = "#1b5e20"
                                font_style = ('Segoe UI', 11, 'bold')
                            lbl = tk.Label(parent, text=val, bg=cell_bg, fg=cell_fg, font=font_style, wraplength=120)
                            lbl.grid(row=row_num, column=col_idx, sticky="nsew", padx=1, pady=1)
                            if val and val not in ('—', ''):
                                lbl.configure(cursor="hand2")
                                lbl.bind("<Button-1>", lambda e, d=day, p=p, v=val: self.open_class_grading_view(d, p, v))

        # ========================================

    # --- open_class_grading_view ---
    def open_class_grading_view(self, day=None, period=None, subject_info=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- open_class_grading_view_norm_str ---
    def open_class_grading_view_norm_str(x=None):
        # ========================================
        s = str(x).strip()
        s = s.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا").replace("ة", "ه")
        return s

        # ========================================

    # --- open_class_grading_view_clean_noise ---
    def open_class_grading_view_clean_noise(txt=None):
        # ========================================
        for w in ('ال', 'الصف', 'الشعبة', 'الشعبه'):
            txt = txt.replace(w, "")
        else:
            return txt.strip()

        # ========================================

    # --- open_class_grading_view_mk_col_head ---
    def open_class_grading_view_mk_col_head(txt=None, width_num=None):
        # ========================================
        f = tk.Frame(tbl_header, bg="#CFD8DC", width=width_num, height=40)
        f.pack_propagate(False)
        f.pack(side="right", padx=1)
        tk.Label(f, text=txt, bg="#CFD8DC", font=('Arial', 14, 'bold'), fg="#37474F").pack(expand=True)
        return f

        # ========================================

    # --- open_class_grading_view__on_mousewheel ---
    def open_class_grading_view__on_mousewheel(event=None):
        # ========================================
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        # ========================================

    # --- open_class_grading_view_log_event ---
    def open_class_grading_view_log_event(student=None, category=None, item=None, value=None, points=None, note=None):
        # ========================================
        full_type = f"[{category}] {item}"
        full_note = f"القيمة: {value} | {note}"
        print(f"SAVING: {student} -> {full_type} = {value}")
        self.m.add_behavior_record(teacher=(self.current_user["name"]),
          student=student,
          student_class=f"{class_name}-{section}",
          subject=subject_name,
          b_type=full_type,
          points=points,
          notes=full_note)

        # ========================================

    # --- open_class_grading_view_build_student_row ---
    def open_class_grading_view_build_student_row(parent=None, name_text=None):
        # ========================================
        CW_NAME = 300
        CW_ATT = 250
        CW_PART = 250
        CW_HW = 150
        CW_BEH = 150
        CW_NOTE = 350
        ROW_H = 80
        FONT_MAIN = ('Arial', 14, 'bold')
        FONT_BTN = ('Arial', 12, 'bold')
        row = tk.Frame(parent, bg="white", height=ROW_H)
        row.pack(fill="x", pady=2)
        row.pack_propagate(False)

        def on_enter(e):
            row.config(bg="#e3f2fd")


        def on_leave(e):
            row.config(bg="white")


        row.bind("<Enter>", on_enter)
        row.bind("<Leave>", on_leave)
        f_name = tk.Frame(row, bg="white", width=CW_NAME)
        f_name.pack_propagate(False)
        f_name.pack(side="right", padx=1)
        tk.Label(f_name, text=name_text, font=('Arial', 16, 'bold'), bg="white", fg="black", anchor="e").pack(fill="both", expand=True, padx=10)
        for child in f_name.winfo_children():
            child.bind("<Enter>", on_enter)
            child.bind("<Leave>", on_leave)
        else:
            f_att = tk.Frame(row, bg="white", width=CW_ATT)
            f_att.pack_propagate(False)
            f_att.pack(side="right", padx=1)
            btns_att = {}
            att_var = tk.StringVar(value="حاضر")

            def set_att(status):
                att_var.set(status)
                pts = 0
                if status == "غائب":
                    pts = -1
                else:
                    log_event(name_text, "الحضور", "الحالة", status, pts)
                    btns_att["غائب"].config(bg="#ffebee", fg="black")
                    btns_att["تأخير"].config(bg="#fffde7", fg="black")
                    btns_att["حاضر"].config(bg="#e8f5e9", fg="black")
                    if status == "غائب":
                        btns_att["غائب"].config(bg="#d32f2f", fg="white")
                    else:
                        if status == "تأخير":
                            btns_att["تأخير"].config(bg="#fbc02d", fg="white")
                        else:
                            if status == "حاضر":
                                btns_att["حاضر"].config(bg="#388e3c", fg="white")


            f_att_inner = tk.Frame(f_att, bg="white")
            f_att_inner.pack(expand=True)
            btns_att["غائب"] = tk.Button(f_att_inner, text="غ", width=6, height=2, font=FONT_BTN, bg="#ffebee", command=(lambda: set_att("غائب")))
            btns_att["تأخير"] = tk.Button(f_att_inner, text="ت", width=6, height=2, font=FONT_BTN, bg="#fffde7", command=(lambda: set_att("تأخير")))
            btns_att["حاضر"] = tk.Button(f_att_inner, text="ح", width=6, height=2, font=FONT_BTN, bg="#388e3c", fg="white", command=(lambda: set_att("حاضر")))
            btns_att["غائب"].pack(side="left", padx=3)
            btns_att["تأخير"].pack(side="left", padx=3)
            btns_att["حاضر"].pack(side="left", padx=3)
            f_part = tk.Frame(row, bg="white", width=CW_PART)
            f_part.pack_propagate(False)
            f_part.pack(side="right", padx=1)
            f_part_inner = tk.Frame(f_part, bg="white")
            f_part_inner.pack(expand=True)
            stars = []

            def star_click(n):
                for i in range(5):
                    stars[i].config(text=("★" if i < n else "☆"), fg=("#fbc02d" if i < n else "#ccc"))
                else:
                    log_event(name_text, "المشاركة", "تقييم", f"{n}/5", points=n)


            for i in range(1, 6):
                b = tk.Button(f_part_inner, text="☆", font=('Arial', 20), bd=0, bg="white", fg="#ccc", cursor="hand2", width=2, command=(lambda n=i: star_click(n)))
                b.pack(side="right")
                stars.append(b)
            else:
                f_hw = tk.Frame(row, bg="white", width=CW_HW)
                f_hw.pack_propagate(False)
                f_hw.pack(side="right", padx=1)
                f_hw_inner = tk.Frame(f_hw, bg="white")
                f_hw_inner.pack(expand=True)

                def hw_act(ok):
                    log_event(name_text, "الواجب", "تسليم", "تم" if ok else "لم يحل", 1 if ok else -1)


                tk.Button(f_hw_inner, text="✔️", bg="#e8f5e9", width=5, height=2, font=FONT_BTN, command=(lambda: hw_act(True))).pack(side="left", padx=5)
                tk.Button(f_hw_inner, text="❌", bg="#ffebee", width=5, height=2, font=FONT_BTN, command=(lambda: hw_act(False))).pack(side="left", padx=5)
                f_beh = tk.Frame(row, bg="white", width=CW_BEH)
                f_beh.pack_propagate(False)
                f_beh.pack(side="right", padx=1)
                f_beh_inner = tk.Frame(f_beh, bg="white")
                f_beh_inner.pack(expand=True)

                def beh_act(l, p):
                    log_event(name_text, "السلوك", l, "رصد", p)


                tk.Button(f_beh_inner, text="👍", bg="#fff9c4", width=5, height=2, font=('Arial',
                                                                                        16), command=(lambda: beh_act("إيجابي", 2))).pack(side="left", padx=5)
                tk.Button(f_beh_inner, text="👎", bg="#ffccbc", width=5, height=2, font=('Arial',
                                                                                        16), command=(lambda: beh_act("سلبي", -2))).pack(side="left", padx=5)
                f_note = tk.Frame(row, bg="white", width=CW_NOTE)
                f_note.pack_propagate(False)
                f_note.pack(side="right", padx=1)
                ent = ttk.Entry(f_note, font=('Arial', 12))
                ent.pack(fill="x", expand=True, padx=5, pady=20)

                def save_n(e):
                    if ent.get().strip():
                        log_event(name_text, "ملاحظات", "نص", ent.get(), 0)


                ent.bind("<FocusOut>", save_n)
                ent.bind("<Return>", save_n)
                for f in (
                 row, f_name, f_att, f_part, f_hw, f_beh, f_note, f_att_inner, f_part_inner, f_hw_inner, f_beh_inner):
                    f.bind("<Enter>", on_enter)
                    f.bind("<Leave>", on_leave)
                else:
                    return row

        # ========================================

    # --- open_class_grading_view_build_student_row_on_enter ---
    def open_class_grading_view_build_student_row_on_enter(e=None):
        # ========================================
        row.config(bg="#e3f2fd")

        # ========================================

    # --- open_class_grading_view_build_student_row_on_leave ---
    def open_class_grading_view_build_student_row_on_leave(e=None):
        # ========================================
        row.config(bg="white")

        # ========================================

    # --- open_class_grading_view_build_student_row_set_att ---
    def open_class_grading_view_build_student_row_set_att(status=None):
        # ========================================
        att_var.set(status)
        pts = 0
        if status == "غائب":
            pts = -1
        else:
            log_event(name_text, "الحضور", "الحالة", status, pts)
            btns_att["غائب"].config(bg="#ffebee", fg="black")
            btns_att["تأخير"].config(bg="#fffde7", fg="black")
            btns_att["حاضر"].config(bg="#e8f5e9", fg="black")
            if status == "غائب":
                btns_att["غائب"].config(bg="#d32f2f", fg="white")
            else:
                if status == "تأخير":
                    btns_att["تأخير"].config(bg="#fbc02d", fg="white")
                else:
                    if status == "حاضر":
                        btns_att["حاضر"].config(bg="#388e3c", fg="white")

        # ========================================

    # --- open_class_grading_view_build_student_row_star_click ---
    def open_class_grading_view_build_student_row_star_click(n=None):
        # ========================================
        for i in range(5):
            stars[i].config(text=("★" if i < n else "☆"), fg=("#fbc02d" if i < n else "#ccc"))
        else:
            log_event(name_text, "المشاركة", "تقييم", f"{n}/5", points=n)

        # ========================================

    # --- open_class_grading_view_build_student_row_hw_act ---
    def open_class_grading_view_build_student_row_hw_act(ok=None):
        # ========================================
        log_event(name_text, "الواجب", "تسليم", "تم" if ok else "لم يحل", 1 if ok else -1)

        # ========================================

    # --- open_class_grading_view_build_student_row_beh_act ---
    def open_class_grading_view_build_student_row_beh_act(l=None, p=None):
        # ========================================
        log_event(name_text, "السلوك", l, "رصد", p)

        # ========================================

    # --- open_class_grading_view_build_student_row_save_n ---
    def open_class_grading_view_build_student_row_save_n(e=None):
        # ========================================
        if ent.get().strip():
            log_event(name_text, "ملاحظات", "نص", ent.get(), 0)

        # ========================================

    # --- open_class_grading_view_build_student_card ---
    def open_class_grading_view_build_student_card(parent=None, name_text=None):
        # ========================================
        card = tk.Frame(parent, bg="white", bd=0)
        inner = tk.Frame(card, bg="white", highlightbackground="#e0e0e0", highlightthickness=1)
        inner.pack(fill="x")

        def on_enter(e):
            inner.config(bg="#e3f2fd", highlightbackground="#2196f3", highlightthickness=2)


        def on_leave(e):
            inner.config(bg="white", highlightbackground="#e0e0e0", highlightthickness=1)


        inner.bind("<Enter>", on_enter)
        inner.bind("<Leave>", on_leave)
        r1 = tk.Frame(inner, bg="white")
        r1.pack(fill="x", padx=10, pady=8)
        lbl_name = tk.Label(r1, text=name_text, font=('Arial', 13, 'bold'), bg="white", fg="#2e7d32")
        lbl_name.pack(side="right")
        f_att = tk.Frame(r1, bg="white")
        f_att.pack(side="left")
        att_var = tk.StringVar(value="حاضر")
        btns_att = {}

        def set_att(status):
            att_var.set(status)
            pts = 0
            if status == "غائب":
                pts = -1
            else:
                log_event(name_text, "الحضور", "الحالة", status, pts)
                btns_att["غائب"].config(bg="#ffebee", fg="black")
                btns_att["تأخير"].config(bg="#fffde7", fg="black")
                btns_att["حاضر"].config(bg="#e8f5e9", fg="black")
                if status == "غائب":
                    btns_att["غائب"].config(bg="#d32f2f", fg="white")
                else:
                    if status == "تأخير":
                        btns_att["تأخير"].config(bg="#fbc02d", fg="white")
                    else:
                        if status == "حاضر":
                            btns_att["حاضر"].config(bg="#388e3c", fg="white")


        btns_att["غائب"] = tk.Button(f_att, text="غائب", width=5, bg="#ffebee", command=(lambda: set_att("غائب")))
        btns_att["تأخير"] = tk.Button(f_att, text="تأخير", width=5, bg="#fffde7", command=(lambda: set_att("تأخير")))
        btns_att["حاضر"] = tk.Button(f_att, text="حاضر", width=5, bg="#388e3c", fg="white", command=(lambda: set_att("حاضر")))
        btns_att["غائب"].pack(side="left", padx=2)
        btns_att["تأخير"].pack(side="left", padx=2)
        btns_att["حاضر"].pack(side="left", padx=2)
        tk.Label(r1, text="|", bg="white", fg="#eee").pack(side="left", padx=5)
        r2 = tk.Frame(inner, bg="white")
        r2.pack(fill="x", padx=10, pady=5)
        f_part = tk.Frame(r2, bg="white")
        f_part.pack(side="right")
        tk.Label(f_part, text="مشاركة:", bg="white", fg="#555", font=('Arial', 9)).pack(side="right", padx=2)
        stars = []

        def star_click(n):
            for i in range(5):
                stars[i].config(text=("★" if i < n else "☆"), fg=("#fbc02d" if i < n else "#e0e0e0"))
            else:
                log_event(name_text, "المشاركة", "تقييم", f"{n}/5", points=n)


        for i in range(1, 6):
            b = tk.Button(f_part, text="☆", font=('Arial', 14), bd=0, bg="white", fg="#e0e0e0", cursor="hand2", activebackground="white",
              activeforeground="#fbc02d",
              command=(lambda n=i: star_click(n)))
            b.pack(side="right")
            stars.append(b)
        else:
            f_hw = tk.Frame(r2, bg="white")
            f_hw.pack(side="left")
            tk.Label(f_hw, text="الواجب:", bg="white", fg="#555", font=('Arial', 9)).pack(side="right", padx=5)

            def hw_action(ok):
                val = "تم" if ok else "لم يحل"
                pts = 1 if ok else -1
                log_event(name_text, "الواجبات", "تسليم", val, pts)


            tk.Button(f_hw, text="✅", bg="#e8f5e9", width=3, command=(lambda: hw_action(True))).pack(side="left", padx=2)
            tk.Button(f_hw, text="❌", bg="#ffebee", width=3, command=(lambda: hw_action(False))).pack(side="left", padx=2)
            r3 = tk.Frame(inner, bg="white")
            r3.pack(fill="x", padx=10, pady=5)
            f_beh = tk.Frame(r3, bg="white")
            f_beh.pack(side="right")

            def beh_log(lbl, pts, color, btn):
                log_event(name_text, "السلوك", lbl, "رصد", pts)
                orig = btn.cget("text")
                btn.config(text="✓")
                btn.after(800, lambda: btn.config(text=orig))


            b_res = tk.Button(f_beh, text="متميز 🌟", bg="#fff9c4", font=('Arial', 8))
            b_res.config(command=(lambda: beh_log("سلوك إيجابي", 2, "gold", b_res)))
            b_res.pack(side="right", padx=2)
            b_vio = tk.Button(f_beh, text="مخالفة ⚠️", bg="#ffccbc", font=('Arial', 8))
            b_vio.config(command=(lambda: beh_log("مخالفة", -2, "red", b_vio)))
            b_vio.pack(side="right", padx=2)
            f_tools = tk.Frame(r3, bg="white")
            f_tools.pack(side="left")
            bk_var = tk.IntVar()

            def tool_toggle(v, name):
                status = "نسيان" if v.get() else "إحضار"
                pts = -1 if v.get() else 0
                log_event(name_text, "المتابعة", name, status, pts)


            cb = tk.Checkbutton(f_tools, text="نسيان كتاب", variable=bk_var, bg="white", font=('Arial',
                                                                                               8), command=(lambda: tool_toggle(bk_var, "الكتاب")))
            cb.pack(side="left")
            r4 = tk.Frame(inner, bg="white")
            r4.pack(fill="x", padx=10, pady=5)
            tk.Label(r4, text="ملاحظات:", bg="white", fg="#aaa", font=('Arial', 8)).pack(side="right")
            ent = ttk.Entry(r4)
            ent.pack(side="left", fill="x", expand=True, padx=5)

            def save_note(e):
                t = ent.get().strip()
                if t:
                    log_event(name_text, "ملاحظات", "نص", t, 0)


            ent.bind("<FocusOut>", save_note)
            ent.bind("<Return>", save_note)
            return card

        # ========================================

    # --- open_class_grading_view_build_student_card_on_enter ---
    def open_class_grading_view_build_student_card_on_enter(e=None):
        # ========================================
        inner.config(bg="#e3f2fd", highlightbackground="#2196f3", highlightthickness=2)

        # ========================================

    # --- open_class_grading_view_build_student_card_on_leave ---
    def open_class_grading_view_build_student_card_on_leave(e=None):
        # ========================================
        inner.config(bg="white", highlightbackground="#e0e0e0", highlightthickness=1)

        # ========================================

    # --- open_class_grading_view_build_student_card_set_att ---
    def open_class_grading_view_build_student_card_set_att(status=None):
        # ========================================
        att_var.set(status)
        pts = 0
        if status == "غائب":
            pts = -1
        else:
            log_event(name_text, "الحضور", "الحالة", status, pts)
            btns_att["غائب"].config(bg="#ffebee", fg="black")
            btns_att["تأخير"].config(bg="#fffde7", fg="black")
            btns_att["حاضر"].config(bg="#e8f5e9", fg="black")
            if status == "غائب":
                btns_att["غائب"].config(bg="#d32f2f", fg="white")
            else:
                if status == "تأخير":
                    btns_att["تأخير"].config(bg="#fbc02d", fg="white")
                else:
                    if status == "حاضر":
                        btns_att["حاضر"].config(bg="#388e3c", fg="white")

        # ========================================

    # --- open_class_grading_view_build_student_card_star_click ---
    def open_class_grading_view_build_student_card_star_click(n=None):
        # ========================================
        for i in range(5):
            stars[i].config(text=("★" if i < n else "☆"), fg=("#fbc02d" if i < n else "#e0e0e0"))
        else:
            log_event(name_text, "المشاركة", "تقييم", f"{n}/5", points=n)

        # ========================================

    # --- open_class_grading_view_build_student_card_hw_action ---
    def open_class_grading_view_build_student_card_hw_action(ok=None):
        # ========================================
        val = "تم" if ok else "لم يحل"
        pts = 1 if ok else -1
        log_event(name_text, "الواجبات", "تسليم", val, pts)

        # ========================================

    # --- open_class_grading_view_build_student_card_beh_log ---
    def open_class_grading_view_build_student_card_beh_log(lbl=None, pts=None, color=None, btn=None):
        # ========================================
        log_event(name_text, "السلوك", lbl, "رصد", pts)
        orig = btn.cget("text")
        btn.config(text="✓")
        btn.after(800, lambda: btn.config(text=orig))

        # ========================================

    # --- open_class_grading_view_build_student_card_tool_toggle ---
    def open_class_grading_view_build_student_card_tool_toggle(v=None, name=None):
        # ========================================
        status = "نسيان" if v.get() else "إحضار"
        pts = -1 if v.get() else 0
        log_event(name_text, "المتابعة", name, status, pts)

        # ========================================

    # --- open_class_grading_view_build_student_card_save_note ---
    def open_class_grading_view_build_student_card_save_note(e=None):
        # ========================================
        t = ent.get().strip()
        if t:
            log_event(name_text, "ملاحظات", "نص", t, 0)

        # ========================================

    # --- open_class_grading_view_on_search ---
    def open_class_grading_view_on_search(event=None):
        # ========================================
        query = ent_search.get().strip().lower()
        for s_name, widget in student_widgets.items():
            if query == "":
                widget.pack(fill="x", pady=2)
            elif query in s_name:
                widget.pack(fill="x", pady=2)
            else:
                widget.pack_forget()

        # ========================================

    # --- export_teacher_personal_excel ---
    def export_teacher_personal_excel(self, teacher_name=None):
        # ========================================
        try:
            import xlsxwriter
            filename = f"جدول_المعلم_{teacher_name}.xlsx"
            fpath = os.path.abspath(filename)
            workbook = xlsxwriter.Workbook(fpath)
            worksheet = workbook.add_worksheet("الجدول")
            worksheet.right_to_left()
            fmt_header = workbook.add_format({
             'bold': True, 'font_size': 14, 'align': '"center"', 'valign': '"vcenter"', 
             'bg_color': '"#1b5e20"', 'font_color': '"white"', 'border': 1})
            fmt_day = workbook.add_format({
             'bold': True, 'font_size': 12, 'align': '"center"', 'valign': '"vcenter"', 
             'bg_color': '"#2e7d32"', 'font_color': '"white"', 'border': 1})
            fmt_cell_empty = workbook.add_format({
             'align': '"center"', 'valign': '"vcenter"', 'border': 1, 'font_size': 12})
            fmt_cell_filled = workbook.add_format({
             'bold': True, 'align': '"center"', 'valign': '"vcenter"', 'border': 1, 
             'bg_color': '"#c8e6c9"', 'font_color': '"#1b5e20"', 'font_size': 12, 'text_wrap': True})
            days = [
             'الأحد', 'الاثنين', 'الثلاثاء', 'الأربعاء', 'الخميس']
            periods = range(1, 8)
            worksheet.write(0, 0, "اليوم", fmt_header)
            for i, p in enumerate(periods):
                worksheet.write(0, i + 1, f"الحصة {p}", fmt_header)
            else:
                for r_idx, day in enumerate(days):
                    worksheet.write(r_idx + 1, 0, day, fmt_day)
                    worksheet.set_row(r_idx + 1, 50)
                    for i, p in enumerate(periods):
                        val = self.m.get_cell(teacher_name, day, p)
                        fmt = fmt_cell_filled if val else fmt_cell_empty
                        worksheet.write(r_idx + 1, i + 1, val, fmt)
                    else:
                        worksheet.set_column(0, 0, 15)
                        worksheet.set_column(1, 7, 20)
                        workbook.close()
                        os.startfile(fpath)
                        messagebox.showinfo("تم الحفظ", f"تم حفظ الجدول بنجاح:\n{filename}")

        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"فشل تصدير الجدول: {e}")
            finally:
                pass

        # ========================================

    # --- show_manager_hub ---
    def show_manager_hub(self):
        # ========================================
        if not self.admin_mode:
            pw = simpledialog.askstring("صلاحية الوصول", "أدخل رقم التفويض للمدير:", show="*")
            if not pw:
                return
            
            valid_admin = self.m.verify_employee_pin("المدير", pw)
            if valid_admin or pw == "3333":
                self.admin_mode = True
            else:
                messagebox.showerror("خطأ", "رقم التفويض غير صحيح!")
                return
        self.clear_sidebar()
        self.sb("🏠 رجوع للرئيسية", (self.show_home), fg="white", bg=COLOR_ACCENT)
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("🔒 إنهاء جلسة الإدارة", (self.logout_admin), fg="white", bg=COLOR_DANGER)
        self.show("manager_hub")

    # --- logout_admin ---
    def logout_admin(self):
        self.admin_mode = False
        messagebox.showinfo("تم", "تم تسجيل الخروج بنجاح.")
        self.show_home()


        # ========================================

    # --- show_monitoring_section ---
    def show_monitoring_section(self):
        # ========================================
        self.show_manager_hub()
        self.show("monitoring")

        # ========================================

    # --- show_students_section ---
    def show_students_section(self):
        # ========================================
        self.clear_sidebar()
        self.sb("🏠 رجوع للرئيسية", (self.show_home), fg="white", bg=COLOR_ACCENT)
        if self.admin_mode:
            self.sb("🏛️ بوابة المدير", (self.show_manager_hub), fg="white", bg="#263238")
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("📂 بيانات الطلاب", lambda: self.show("students"))
        self.sb("🔍 البحث عن الطالب في الحصص", lambda: self.show("student_finder"))
        self.show("students")

        # ========================================

    # --- show_maintenance_section ---
    def show_maintenance_section(self):
        # ========================================
        self.clear_sidebar()
        self.sb("🏠 رجوع للرئيسية", (self.show_home), fg="white", bg=COLOR_ACCENT)
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("⏰ توقيت الدوام", lambda: self.show("timings"))
        self.sb("📊 جدول المعلمين التفصيلي", lambda: self.show("detailed_schedule"))
        self.sb("📦 المواد (إضافة/حذف)", lambda: self.show("subjects"))
        self.sb("📅 التقويم الدراسي (جديد)", (lambda: self.show("calendar")), fg="white", bg=COLOR_XLSX)
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("🔄 تحديث للنسخة الحالية", (self.reload_all), fg="white", bg=COLOR_ACCENT)
        self.show("timings")

        # ========================================

    # --- page_swap_approvals ---
    def page_swap_approvals(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        tk.Label(page, text="اعتماد طلبات المبادلة (المؤقتة)", font=('Segoe UI', 16, 'bold'), bg=COLOR_BG, fg="#333").pack(pady=20)
        columns = ('id', 'requester', 'acceptor', 'detail', 'date')
        tree = ttk.Treeview(page, columns=columns, show="headings", height=15)
        tree.heading("id", text="#")
        tree.heading("requester", text="مقدم الطلب")
        tree.heading("acceptor", text="الموافق (الطرف الثاني)")
        tree.heading("detail", text="التفاصيل")
        tree.heading("date", text="تاريخ الطلب")
        tree.column("id", width=0, stretch=False)
        tree.column("requester", width=150, anchor="center")
        tree.column("acceptor", width=150, anchor="center")
        tree.column("detail", width=300, anchor="e")
        tree.column("date", width=100, anchor="center")
        tree.pack(fill="both", expand=True, padx=40, pady=10)

        def load_data():
            for i in tree.get_children():
                tree.delete(i)
            else:
                fpath = data_path("تبادل_حصص.json")
                if os.path.exists(fpath):
                    try:
                        import json
                        with open(fpath, "r", encoding="utf-8") as f:
                            swaps = json.load(f)
                        for s in swaps:
                            if s.get("status") == "pending_admin":
                                det = f'يوم {s["day"]} - حصة {s["period"]} - {s["subject"]}'
                                tree.insert("", "end", values=(s["id"], s["requester"], s["acceptor"], det, s["date"]))

                    except:
                        pass


        load_data()
        page.refresh = load_data

        def action(decision):
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("تنبيه", "اختر طلباً أولاً")
                return
            vals = tree.item(sel[0], "values")
            swap_id = vals[0]
            fpath = data_path("تبادل_حصص.json")
            if os.path.exists(fpath):
                try:
                    import json
                    with open(fpath, "r", encoding="utf-8") as f:
                        swaps = json.load(f)
                    found = False
                    for s in swaps:
                        if s["id"] == str(swap_id):
                            s["status"] = decision
                            found = True
                            msg_text = "تم اعتماد التبادل ✅" if decision == "approved" else "تم رفض التبادل ❌"
                            self.m.send_direct_message("الإدارة", s["requester"], msg_text)
                            self.m.send_direct_message("الإدارة", s["acceptor"], msg_text)
                            break
                        if found:
                            with open(fpath, "w", encoding="utf-8") as f:
                                json.dump(swaps, f, ensure_ascii=False, indent=2)
                            messagebox.showinfo("تم", f"تم {decision} الطلب بنجاح")
                            load_data()
                        else:
                            messagebox.showerror("خطأ", "لم يتم العثور على الطلب!")

                except Exception as e:
                    try:
                        messagebox.showerror("خطأ", f"حدث خطأ: {e}")
                    finally:
                        pass


        f_btn = tk.Frame(page, bg=COLOR_BG)
        f_btn.pack(pady=20)
        tk.Button(f_btn, text="✅ اعتماد الطلب", command=(lambda: action("approved")), bg="#4CAF50", fg="white", font=('Segoe UI',
                                                                                                                      12,
                                                                                                                      'bold'), width=20).pack(side="right", padx=10)
        tk.Button(f_btn, text="❌ رفض الطلب", command=(lambda: action("rejected")), bg="#f44336", fg="white", font=('Segoe UI',
                                                                                                                   12,
                                                                                                                   'bold'), width=20).pack(side="right", padx=10)
        return page

        # ========================================

    # --- page_swap_approvals_load_data ---
    def page_swap_approvals_load_data():
        # ========================================
        for i in tree.get_children():
            tree.delete(i)
        else:
            fpath = data_path("تبادل_حصص.json")
            if os.path.exists(fpath):
                try:
                    import json
                    with open(fpath, "r", encoding="utf-8") as f:
                        swaps = json.load(f)
                    for s in swaps:
                        if s.get("status") == "pending_admin":
                            det = f'يوم {s["day"]} - حصة {s["period"]} - {s["subject"]}'
                            tree.insert("", "end", values=(s["id"], s["requester"], s["acceptor"], det, s["date"]))

                except:
                    pass

        # ========================================

    # --- page_swap_approvals_action ---
    def page_swap_approvals_action(decision=None):
        # ========================================
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("تنبيه", "اختر طلباً أولاً")
            return
        vals = tree.item(sel[0], "values")
        swap_id = vals[0]
        fpath = data_path("تبادل_حصص.json")
        if os.path.exists(fpath):
            try:
                import json
                with open(fpath, "r", encoding="utf-8") as f:
                    swaps = json.load(f)
                found = False
                for s in swaps:
                    if s["id"] == str(swap_id):
                        s["status"] = decision
                        found = True
                        msg_text = "تم اعتماد التبادل ✅" if decision == "approved" else "تم رفض التبادل ❌"
                        self.m.send_direct_message("الإدارة", s["requester"], msg_text)
                        self.m.send_direct_message("الإدارة", s["acceptor"], msg_text)
                        break
                    if found:
                        with open(fpath, "w", encoding="utf-8") as f:
                            json.dump(swaps, f, ensure_ascii=False, indent=2)
                        messagebox.showinfo("تم", f"تم {decision} الطلب بنجاح")
                        load_data()
                    else:
                        messagebox.showerror("خطأ", "لم يتم العثور على الطلب!")

            except Exception as e:
                try:
                    messagebox.showerror("خطأ", f"حدث خطأ: {e}")
                finally:
                    pass

        # ========================================

    # --- _create_chip ---
    def _create_chip(self, parent=None, text=None, bg_color=None, fg_color=None, icon=None):
        # ========================================
        frame = tk.Frame(parent, bg=bg_color, padx=10, pady=4, relief="flat")
        lbl = tk.Label(frame, text=(f"{icon} {text}" if icon else text), bg=bg_color,
          fg=fg_color,
          font=('Segoe UI', 10, 'bold'))
        lbl.pack()
        return frame

        # ========================================

    # --- show_calendar_sidebar ---
    def show_calendar_sidebar(self):
        # ========================================
        self.clear_sidebar()
        self.sb("🏠 رجوع للرئيسية", (self.show_home), fg="white", bg=COLOR_ACCENT)
        if self.admin_mode:
            self.sb("🏛️ بوابة المدير", (self.show_manager_hub), fg="white", bg="#263238")
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("📅 التقويم الدراسي", (lambda: self.show("calendar")), fg="white", bg=COLOR_ACCENT)
        if self.admin_mode:
            self.sb("⏰ توقيت الدوام", lambda: self.show("timings"))
        ttk.Separator((self.side), orient="horizontal").pack(fill="x", padx=10, pady=8)
        self.sb("🔄 تحديث البيانات", (self.reload_all), fg="white", bg="#666")
        self.show("calendar")

        # ========================================

    # --- refresh_today_status ---
    def refresh_today_status(self):
        # ========================================
        st = self.m.get_day_status(datetime.now().date())
        d = st.get("date")
        day = st.get("day", "")
        note = st.get("note", "")
        wk = st.get("week")
        term = st.get("term", "")
        holiday = st.get("is_holiday", False)
        if hasattr(self, "lbl_today"):
            for w in self.lbl_today.winfo_children():
                w.destroy()
            else:
                self._create_chip(self.lbl_today, f"{day} {d}", COLOR_ACCENT).pack(side="right", padx=4)
                if term:
                    self._create_chip(self.lbl_today, term, COLOR_ACCENT).pack(side="right", padx=4)
                if wk:
                    self._create_chip(self.lbl_today, f"الأسبوع {wk}", COLOR_ACCENT).pack(side="right", padx=4)
                if holiday:
                    self._create_chip((self.lbl_today), "إجازة", COLOR_DANGER, icon="⚠️").pack(side="right", padx=4)
                if note:
                    self._create_chip((self.lbl_today), note, COLOR_WARN, fg_color="#333").pack(side="right", padx=4)

        # ========================================

    # --- page_employee_room ---
    def page_employee_room(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        header = tk.Frame(page, bg="#2e7d32", pady=20, padx=40)
        header.pack(fill="x")
        self.lbl_emp_room_name = tk.Label(header, text="مرحباً ...", font=('Segoe UI', 22, 'bold'), bg="#2e7d32", fg="white")
        self.lbl_emp_room_name.pack(anchor="e")
        tk.Label(header, text="غرفتك الخاصة للمراسلات والمهام", font=('Segoe UI', 11), bg="#2e7d32", fg="#e8f5e9").pack(anchor="e")
        
        main_content = tk.Frame(page, bg=COLOR_BG)
        main_content.pack(fill="both", expand=True, padx=40, pady=20)
        
        # --- Right Sidebar (Messaging & Directives) ---
        sidebar = tk.Frame(main_content, bg=COLOR_BG, width=300)
        sidebar.pack(side="right", fill="y", padx=(20, 0))
        
        btn_msg = tk.Button(sidebar, text="💬 رسائلي الخاصة", command=(lambda: self.open_messaging_window("employee")), bg="#8e24aa", fg="white", font=('Segoe UI', 11, 'bold'), pady=10, cursor="hand2")
        btn_msg.pack(fill="x", pady=(0, 15))
        
        f_msg = tk.LabelFrame(sidebar, text="📩 التوجيهات الإدارية", font=('Segoe UI', 11, 'bold'), bg="white", padx=15, pady=15)
        f_msg.pack(fill="both", expand=True)
        self.lbl_emp_msg_content = tk.Label(f_msg, text="لا توجد رسائل جديدة.", font=('Segoe UI', 11), bg="white", fg="#555", wraplength=250, justify="right")
        self.lbl_emp_msg_content.pack(anchor="ne")
        
        # --- Left Main Area (Action Cards) ---
        action_area = tk.Frame(main_content, bg=COLOR_BG)
        action_area.pack(side="left", fill="both", expand=True)
        
        f_cards = tk.Frame(action_area, bg=COLOR_BG)
        f_cards.pack(fill="both", expand=True, pady=40)
        
        f_cards.columnconfigure(0, weight=1)
        f_cards.columnconfigure(1, weight=1)
        f_cards.rowconfigure(0, weight=1)
        
        def on_achievements():
            self.open_achievements_window()
            
        def on_today_work():
            self.open_today_work_window()
            
        btn_achiev = tk.Button(f_cards, text="🏆\n\nإنجـازي", bg="#28a745", fg="white", font=('Segoe UI', 22, 'bold'), cursor="hand2", command=on_achievements, relief="flat")
        btn_achiev.grid(row=0, column=0, sticky="nsew", padx=15, pady=10)
        
        btn_today = tk.Button(f_cards, text="⚡\n\nعمـل اليـوم", bg="#0056b3", fg="white", font=('Segoe UI', 22, 'bold'), cursor="hand2", command=on_today_work, relief="flat")
        btn_today.grid(row=0, column=1, sticky="nsew", padx=15, pady=10)
        
        # --- Footer ---
        footer = tk.Frame(page, bg=COLOR_BG)
        footer.pack(fill="x", pady=10)
        tk.Button(footer, text="تسجيل خروج", command=(self.show_home), bg="#eceff1", fg="#333", font=('Segoe UI', 11)).pack()
        
        # Hidden component to satisfy legacy dependencies (e.g. refresh_teacher_dashboard)
        self.btn_emp_status = tk.Button(footer, text="")
        
        return page

        # ========================================

    # --- open_suspension_logger ---
    def open_suspension_logger(self):
        # ========================================
        win = tk.Toplevel(self)
        win.title("🚨 توثيق تعليق دراسة طارئ")
        win.geometry("550x650")
        win.grab_set()
        frm = ttk.Frame(win, padding=25)
        frm.pack(fill="both", expand=True)
        ttk.Label(frm, text="توثيق واقعة تعليق دراسة", font=('Segoe UI', 16, 'bold'), foreground=COLOR_DANGER).pack(pady=(0,
                                                                                                                          20))

        def date_picker_frame(parent, label_text):
            f = tk.Frame(parent, bg=COLOR_BG)
            f.pack(fill="x", pady=10)
            tk.Label(f, text=label_text, font=('Segoe UI', 11, 'bold'), anchor="e", width=15).pack(side="right")
            now = datetime.now()
            f_sel = tk.Frame(f, bg=COLOR_BG)
            f_sel.pack(side="right", padx=10)
            cb_y = ttk.Combobox(f_sel, values=[str(y) for y in range(now.year, now.year + 2)], width=8)
            cb_y.set(str(now.year))
            cb_y.pack(side="right", padx=2)
            cb_m = ttk.Combobox(f_sel, values=[f"{m:02d}" for m in range(1, 13)], width=5)
            cb_m.set(now.strftime("%m"))
            cb_m.pack(side="right", padx=2)
            cb_d = ttk.Combobox(f_sel, values=[f"{d:02d}" for d in range(1, 32)], width=5)
            cb_d.set(now.strftime("%d"))
            cb_d.pack(side="right", padx=2)
            return (
             cb_y, cb_m, cb_d)


        y1, m1, d1 = date_picker_frame(frm, "تاريخ البداية:")
        y2, m2, d2 = date_picker_frame(frm, "تاريخ النهاية:")
        tk.Label(frm, text="سبب التعليق:", font=('Segoe UI', 11, 'bold')).pack(anchor="e", pady=(15,
                                                                                                 5))
        reasons = ['تعليق دراسة (أمطار)', 'تعليق دراسة (غبار)', 'تعليق دراسة (حالة جوية)', 'تعليق دراسة (صيانة طارئة)', 
         'تعليق دراسة (أخرى)']
        cb_reason = ttk.Combobox(frm, values=reasons, width=45)
        cb_reason.set(reasons[0])
        cb_reason.pack(pady=5)
        tk.Label(frm, text="الفصل الدراسي:", font=('Segoe UI', 11, 'bold')).pack(anchor="e", pady=(15,
                                                                                                   5))
        cb_term = ttk.Combobox(frm, values=["الفصل الأول", "الفصل الثاني", "الفصل الثالث"], width=45)
        st = self.m.get_day_status(datetime.now().date())
        cb_term.set(st.get("term", "الفصل الثاني"))
        cb_term.pack(pady=5)

        def save_suspension(*args, **kwargs): pass  # TODO: Reconstruct method


        tk.Button(frm, text="💾 حفظ وتوثيق فوري", command=save_suspension, bg=COLOR_DANGER,
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          pady=15).pack(fill="x", pady=30)# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- open_suspension_logger_date_picker_frame ---
    def open_suspension_logger_date_picker_frame(parent=None, label_text=None):
        # ========================================
        f = tk.Frame(parent, bg=COLOR_BG)
        f.pack(fill="x", pady=10)
        tk.Label(f, text=label_text, font=('Segoe UI', 11, 'bold'), anchor="e", width=15).pack(side="right")
        now = datetime.now()
        f_sel = tk.Frame(f, bg=COLOR_BG)
        f_sel.pack(side="right", padx=10)
        cb_y = ttk.Combobox(f_sel, values=[str(y) for y in range(now.year, now.year + 2)], width=8)
        cb_y.set(str(now.year))
        cb_y.pack(side="right", padx=2)
        cb_m = ttk.Combobox(f_sel, values=[f"{m:02d}" for m in range(1, 13)], width=5)
        cb_m.set(now.strftime("%m"))
        cb_m.pack(side="right", padx=2)
        cb_d = ttk.Combobox(f_sel, values=[f"{d:02d}" for d in range(1, 32)], width=5)
        cb_d.set(now.strftime("%d"))
        cb_d.pack(side="right", padx=2)
        return (
         cb_y, cb_m, cb_d)

        # ========================================

    # --- open_suspension_logger_save_suspension ---
    def open_suspension_logger_save_suspension():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- build_home_dashboard ---
    def build_home_dashboard(self, parent_frame=None):
        # ========================================
        cards_frame = tk.Frame(parent_frame, bg=COLOR_BG)
        cards_frame.pack(fill="x", pady=10)
        for i in range(4):
            cards_frame.columnconfigure(i, weight=1)
        else:

            def make_dash_card(col_idx, title, icon, bg_color, text_var, sub_text_var):
                card = tk.Frame(cards_frame, bg="white", relief="flat", bd=0)
                card.grid(row=0, column=col_idx, padx=10, sticky="nsew")
                stripe = tk.Frame(card, bg=bg_color, height=10)
                stripe.pack(fill="x")
                c_inner = tk.Frame(card, bg="white", padx=15, pady=15)
                c_inner.pack(fill="both", expand=True)
                h_frame = tk.Frame(c_inner, bg="white")
                h_frame.pack(fill="x", anchor="n")
                tk.Label(h_frame, text=icon, font=('Segoe UI Emoji', 26), fg=bg_color, bg="white").pack(side="right", padx=(10, 0))
                tk.Label(h_frame, text=title, font=('Segoe UI', 14, 'bold'), fg="#777", bg="white").pack(side="right", padx=5, pady=5)
                tk.Label(c_inner, textvariable=text_var, font=('Segoe UI', 16, 'bold'), fg="#333", bg="white", wraplength=200, justify="right").pack(anchor="e", pady=(10,
                                                                                                                                                                       5))
                tk.Label(c_inner, textvariable=sub_text_var, font=('Segoe UI', 12, 'bold'), fg=bg_color, bg="white").pack(anchor="e")
                return card


            self.dash_time_var = tk.StringVar(value="--:--")
            self.dash_date_var = tk.StringVar(value="...")
            self.dash_period_var = tk.StringVar(value="...")
            self.dash_p_rem_var = tk.StringVar(value="...")
            self.dash_term_var = tk.StringVar(value="...")
            self.dash_term_rem_var = tk.StringVar(value="...")
            self.dash_vac_var = tk.StringVar(value="...")
            self.dash_vac_cnt_var = tk.StringVar(value="...")
            make_dash_card(0, "الوقت والتاريخ", "🕰️", "#2196f3", self.dash_time_var, self.dash_date_var)
            make_dash_card(1, "الحصة الحالية", "🔔", "#43a047", self.dash_period_var, self.dash_p_rem_var)
            make_dash_card(2, "الفصل الدراسي", "📅", "#FF5722", self.dash_term_var, self.dash_term_rem_var)
            make_dash_card(3, "الإجازة القادمة", "🏖️", "#9c27b0", self.dash_vac_var, self.dash_vac_cnt_var)
            self.update_home_dashboard_live()

        # ========================================

    # --- build_home_dashboard_make_dash_card ---
    def build_home_dashboard_make_dash_card(col_idx=None, title=None, icon=None, bg_color=None, text_var=None, sub_text_var=None):
        # ========================================
        card = tk.Frame(cards_frame, bg="white", relief="flat", bd=0)
        card.grid(row=0, column=col_idx, padx=10, sticky="nsew")
        stripe = tk.Frame(card, bg=bg_color, height=10)
        stripe.pack(fill="x")
        c_inner = tk.Frame(card, bg="white", padx=15, pady=15)
        c_inner.pack(fill="both", expand=True)
        h_frame = tk.Frame(c_inner, bg="white")
        h_frame.pack(fill="x", anchor="n")
        tk.Label(h_frame, text=icon, font=('Segoe UI Emoji', 26), fg=bg_color, bg="white").pack(side="right", padx=(10, 0))
        tk.Label(h_frame, text=title, font=('Segoe UI', 14, 'bold'), fg="#777", bg="white").pack(side="right", padx=5, pady=5)
        tk.Label(c_inner, textvariable=text_var, font=('Segoe UI', 16, 'bold'), fg="#333", bg="white", wraplength=200, justify="right").pack(anchor="e", pady=(10,
                                                                                                                                                               5))
        tk.Label(c_inner, textvariable=sub_text_var, font=('Segoe UI', 12, 'bold'), fg=bg_color, bg="white").pack(anchor="e")
        return card

        # ========================================

    # --- update_home_dashboard_live ---
    def update_home_dashboard_live(self):
        # ========================================
        try:
            now = datetime.now()
            t_str = now.strftime("%I:%M %p").replace("AM", "ص").replace("PM", "م")
            d_str = now.strftime("%Y/%m/%d")
            day_ar = arabic_day_from_english(now.strftime("%A"))
            self.dash_time_var.set(t_str)
            self.dash_date_var.set(f"{day_ar} | {d_str}")
            p_name, status, rem = self.m.get_current_active_period()
            if status == "Active":
                self.dash_period_var.set(f"الحصة: {p_name}")
                self.dash_p_rem_var.set(f"باقي {rem} دقيقة")
            else:
                self.dash_period_var.set("خارج الدوام" if status == "Out" else "لا يوجد حصة")
                self.dash_p_rem_var.set("الدوام انتهى أو لم يبدأ" if status == "Out" else "--")
            
            prog = self.m.get_academic_progress()
            d_st = self.m.get_day_status(now.date())
            term = d_st.get("term", "")
            week = d_st.get("week")
            if not term and prog and prog.get("terms"):
                curr_term_obj = next((t for t in prog.get("terms", []) if t.get("is_current")), None)
                if curr_term_obj: term = curr_term_obj["name"]
            
            w_str = f"الأسبوع {week}" if week else ""
            self.dash_term_var.set(term if term else "غير محدد")
            term_rem_txt = w_str
            if prog and term:
                found_t = next((t for t in prog["terms"] if normalize_arabic(t["name"]) == normalize_arabic(term)), None)
                if found_t:
                    days_left = found_t.get("days_left", 0)
                    t_name = found_t["name"]
                    try:
                        cal_t = self.m.filter_calendar(term=t_name)
                        if not cal_t.empty:
                            min_date = cal_t['التاريخ الميلادي'].min()
                            max_date = cal_t['التاريخ الميلادي'].max()
                            term_rem_txt = f"البداية: {min_date}\nالنهاية: {max_date}\nباقي {days_left} يوم"
                        else:
                            term_rem_txt = f"{w_str} | باقي {days_left} يوم"
                    except:
                        term_rem_txt = f"{w_str} | باقي {days_left} يوم"
            self.dash_term_rem_var.set(term_rem_txt)
            
            vac_note, vac_days = self.m.get_days_until_next_holiday()
            if vac_note:
                self.dash_vac_var.set(vac_note)
                self.dash_vac_cnt_var.set("بدأت الإجازة اليوم! 🥳" if vac_days <= 0 else f"باقي {vac_days} يوم")
            else:
                self.dash_vac_var.set("لا توجد إجازات")
                self.dash_vac_cnt_var.set("--")
        except Exception: pass
        if hasattr(self, "pages") and "home" in self.pages:
            self.after(1000, self.update_home_dashboard_live)



        # ========================================

    # --- page_home ---
    def page_home(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        self.home_status_container = tk.Frame(page, bg=COLOR_BG)
        dash = tk.Frame(page, bg=COLOR_BG)
        dash.pack(fill="both", expand=True, padx=40, pady=(10, 30))
        self.home_info_container = tk.Frame(dash, bg=COLOR_BG)
        self.home_info_container.pack(side="right", fill="y", padx=(20, 0))
        grid_area = tk.Frame(dash, bg=COLOR_BG)
        grid_area.pack(side="left", fill="both", expand=True)
        self.build_home_dashboard(grid_area)
        return page

        # ========================================

    # --- get_current_user_name ---
    def get_current_user_name(self):
        # ========================================
        if hasattr(self, "current_employee"):
            if self.current_employee:
                return self.current_employee
        if hasattr(self, "current_teacher_user"):
            if self.current_teacher_user:
                return self.current_teacher_user
        if self.admin_mode:
            return "الإدارة (المدير)"
        return

    # --- open_today_work_window ---
    def open_today_work_window(self):
        emp_name = self.get_current_user_name()
        if not emp_name:
            messagebox.showerror("خطأ", "يجب تسجيل الدخول أولاً.")
            return
            
        win = tk.Toplevel(self)
        win.title("⚡ عمل اليوم - مسجل اليومية")
        win.geometry("1100x700")
        win.configure(bg=COLOR_BG)
        win.grab_set()
        
        header = tk.Frame(win, bg="#0056b3", pady=15, padx=20)
        header.pack(fill="x")
        tk.Label(header, text=f"سجل عمل اليوم - {emp_name}", font=('Segoe UI', 18, 'bold'), bg="#0056b3", fg="white").pack(anchor="e")
        
        main_frame = tk.Frame(win, bg=COLOR_BG)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Right Side (Search & Select)
        f_right = tk.LabelFrame(main_frame, text="🔍 البحث والطلاب المتاحين", font=('Segoe UI', 12, 'bold'), bg="white", padx=10, pady=10)
        f_right.pack(side="right", fill="both", expand=True, padx=(10, 0))
        
        f_search_tools = tk.Frame(f_right, bg="white")
        f_search_tools.pack(fill="x", pady=(0, 10))
        
        tk.Label(f_search_tools, text="الاسم:", font=('Segoe UI', 10), bg="white").pack(side="right", padx=5)
        ent_search = tk.Entry(f_search_tools, font=('Segoe UI', 10), width=15)
        ent_search.pack(side="right", padx=5)
        
        tk.Label(f_search_tools, text="الفصل:", font=('Segoe UI', 10), bg="white").pack(side="right", padx=5)
        cb_class = ttk.Combobox(f_search_tools, values=["الكل", "الاول", "الثاني", "الثالث", "أول", "ثاني", "ثالث"], state="readonly", width=8)
        cb_class.set("الكل")
        cb_class.pack(side="right", padx=5)
        
        tk.Label(f_search_tools, text="الشعبة:", font=('Segoe UI', 10), bg="white").pack(side="right", padx=5)
        cb_section = ttk.Combobox(f_search_tools, values=["الكل", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10"], state="readonly", width=5)
        cb_section.set("الكل")
        cb_section.pack(side="right", padx=5)
        
        btn_do_search = tk.Button(f_search_tools, text="بحث", bg=COLOR_ACCENT, fg="white", font=('Segoe UI', 9, 'bold'))
        btn_do_search.pack(side="right", padx=10)
        
        cols_src = ["الاسم", "الفصل", "الشعبة"]
        tree_src = ttk.Treeview(f_right, columns=cols_src, show="headings")
        for c in cols_src:
            tree_src.heading(c, text=c)
            tree_src.column(c, anchor="center", width=80 if c!="الاسم" else 200)
        
        sb_src = ttk.Scrollbar(f_right, orient="vertical", command=tree_src.yview)
        tree_src.configure(yscrollcommand=sb_src.set)
        sb_src.pack(side="left", fill="y")
        tree_src.pack(side="right", fill="both", expand=True)
        
        f_mid = tk.Frame(main_frame, bg=COLOR_BG)
        f_mid.pack(side="right", fill="y", padx=5)
        
        # Left Side (Selected & Processed)
        f_left = tk.LabelFrame(main_frame, text="✅ المنجز (جاهز للاعتماد)", font=('Segoe UI', 12, 'bold'), bg="white", padx=10, pady=10)
        f_left.pack(side="left", fill="both", expand=True)
        
        cols_dest = ["الاسم", "الفصل", "وقت التسجيل"]
        tree_dest = ttk.Treeview(f_left, columns=cols_dest, show="headings")
        for c in cols_dest:
            tree_dest.heading(c, text=c)
            tree_dest.column(c, anchor="center", width=120 if c=="وقت التسجيل" else (180 if c=="الاسم" else 80))
            
        sb_dest = ttk.Scrollbar(f_left, orient="vertical", command=tree_dest.yview)
        tree_dest.configure(yscrollcommand=sb_dest.set)
        sb_dest.pack(side="left", fill="y")
        tree_dest.pack(side="right", fill="both", expand=True)
        
        f_left_btn = tk.Frame(f_left, bg="white")
        f_left_btn.pack(fill="x", pady=5)
        
        processed_students = {}
        
        def refresh_src_list(*args):
            for i in tree_src.get_children():
                tree_src.delete(i)
                
            search_str = ent_search.get().strip().lower()
            cls_filter = cb_class.get()
            sec_filter = cb_section.get()
            
            df = self.m.list_students_simple()
            if df.empty: return
            
            for _, row in df.iterrows():
                name = str(row.get("الاسم", ""))
                c_class = str(row.get("الصف", ""))
                c_sec = str(row.get("الشعبة", ""))
                
                if name in processed_students:
                    continue
                if search_str and search_str not in name.lower():
                    continue
                if cls_filter != "الكل" and cls_filter not in c_class:
                    continue
                if sec_filter != "الكل" and sec_filter != c_sec:
                    continue
                tree_src.insert("", "end", values=(name, c_class, c_sec))
                
        def on_src_select(event=None):
            sel = tree_src.selection()
            if not sel: return
            item = tree_src.item(sel[0], "values")
            name, c_class, c_sec = item[0], item[1], item[2]
            
            now_time = datetime.now().strftime("%H:%M:%S")
            processed_students[name] = {"class": c_class, "section": c_sec, "time": now_time}
            
            tree_src.delete(sel[0])
            tree_dest.insert("", "end", values=(name, c_class, now_time))
            
        tree_src.bind("<Double-1>", on_src_select)
        
        btn_transfer = tk.Button(f_mid, text="◀\nإضافة", font=('Segoe UI', 14, 'bold'), bg=COLOR_ACCENT, fg="white", command=on_src_select)
        btn_transfer.pack(expand=True)
        
        def on_dest_remove(event=None):
            sel = tree_dest.selection()
            if not sel: return
            item = tree_dest.item(sel[0], "values")
            name = item[0]
            
            if messagebox.askyesno("تأكيد", f"هل تريد التراجع عن تسجيل '{name}'؟"):
                if name in processed_students:
                    del processed_students[name]
                tree_dest.delete(sel[0])
                refresh_src_list()
                
        tree_dest.bind("<Double-1>", on_dest_remove)
        
        btn_do_search.config(command=refresh_src_list)
        ent_search.bind("<Return>", refresh_src_list)
        cb_class.bind("<<ComboboxSelected>>", refresh_src_list)
        cb_section.bind("<<ComboboxSelected>>", refresh_src_list)
        
        btn_remove = tk.Button(f_left_btn, text="تراجع عن التسجيل", bg="#dc3545", fg="white", font=('Segoe UI', 9), command=on_dest_remove)
        btn_remove.pack(side="right")
        
        def save_and_submit():
            if not processed_students:
                messagebox.showerror("خطأ", "لا يوجد طلاب في القائمة للإرسال.")
                return
            if messagebox.askyesno("إرسال واعتماد", "هل أنت متأكد من إرسال هذا السجل لاعتماده من مدير المدرسة؟ (سينتقل لقائمة إنجازي)"):
                items = [{"name": k, "class": v["class"], "section": v["section"], "timestamp": v["time"]} for k,v in processed_students.items()]
                
                # Try to determine the task assigned to this employee automatically
                task_type = "تسجيل الحضور والتأخير"
                roles = self.m.load_attendance_roles()
                for t_name, emps in roles.items():
                    if isinstance(emps, list) and emp_name in emps:
                        task_type = t_name
                        break
                        
                if not self.m.log_daily_achievement(emp_name, task_type, items):
                    messagebox.showerror("خطأ", "فشل الحفظ في السجل اليومي.")
                    return
                    
                manager_records = []
                today_str = datetime.now().strftime("%Y-%m-%d")
                
                op_type = "تأخير"
                if "غياب" in task_type: op_type = "غياب"
                elif "انصراف" in task_type: op_type = "انصراف"
                
                for k,v in processed_students.items():
                    manager_records.append({
                        "التاريخ": today_str,
                        "الوقت": v["time"],
                        "نوع العملية": op_type,
                        "اسم الطالب": k,
                        "الصف": v["class"],
                        "الشعبة": v["section"],
                        "اسم الموظف المنفذ": emp_name,
                        "حالة الاعتماد": "معلق"
                    })
                    
                if self.m.add_attendance_record(manager_records):
                    messagebox.showinfo("تم", "تم الإرسال لمدير المدرسة بنجاح وتم تحويل السجل إلى إنجازي.")
                    win.destroy()
                else:
                    messagebox.showerror("خطأ", "تم حفظ السجل الخاص الموظف ولكن فشل الرفع لمدير المدرسة لتلف بملف الإكسل.")
                    
        btn_submit = tk.Button(f_left, text="📤 رفع للمدير واعتماد نهائي", bg="#28a745", fg="white", font=('Segoe UI', 12, 'bold'), pady=10, command=save_and_submit)
        btn_submit.pack(fill="x", pady=(10, 0))
        
        refresh_src_list()

    # --- open_achievements_window ---
    def open_achievements_window(self):
        emp_name = self.get_current_user_name()
        if not emp_name:
            messagebox.showerror("خطأ", "يجب تسجيل الدخول أولاً.")
            return
            
        win = tk.Toplevel(self)
        win.title("🏆 إنجازي وأرشيف العمل")
        win.geometry("1000x700")
        win.configure(bg=COLOR_BG)
        win.grab_set()
        
        # Header
        header = tk.Frame(win, bg="#28a745", pady=15, padx=20)
        header.pack(fill="x")
        tk.Label(header, text=f"سجل إنجازات الموظف: {emp_name}", font=('Segoe UI', 18, 'bold'), bg="#28a745", fg="white").pack(anchor="e")
        
        main_frame = tk.Frame(win, bg=COLOR_BG)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # TOP ROW: STUDENT LOCATOR
        f_locator = tk.LabelFrame(main_frame, text="📍 الرادار اللحظي للصلاحيات (أين الطالب؟)", font=('Segoe UI', 11, 'bold'), bg="white", padx=10, pady=10)
        f_locator.pack(fill="x", pady=(0, 15))
        
        f_search_loc = tk.Frame(f_locator, bg="white")
        f_search_loc.pack()
        
        tk.Label(f_search_loc, text="ابحث بجزء من اسم الطالب لمعرفة جدوله:", font=('Segoe UI', 10, 'bold'), bg="white").pack(side="right", padx=5)
        ent_loc = tk.Entry(f_search_loc, font=('Segoe UI', 11), width=30)
        ent_loc.pack(side="right", padx=5)
        
        lbl_loc_result = tk.Label(f_locator, text="", font=('Segoe UI', 10, 'bold'), bg="white", fg="#0056b3")
        lbl_loc_result.pack(pady=5)
        
        def on_locate_student(*args):
            q = ent_loc.get().strip()
            if len(q) < 2:
                lbl_loc_result.config(text="")
                return
            df = self.m.list_students_simple()
            matches = df[df["الاسم"].str.contains(q, na=False)]
            if matches.empty:
                lbl_loc_result.config(text="لم يتم العثور على طالب بهذا الاسم.")
                return
            
            row = matches.iloc[0]
            student_name = row.get("الاسم", "")
            student_class = row.get("الصف", "")
            student_sec = row.get("الشعبة", "")
            
            try:
                res = self.m.find_student_schedule(student_name)
                if res and len(res) == 4:
                    day, sch_list, period, t = res
                    # res is a tuple containing matched info.
                    if isinstance(sch_list, list) and len(sch_list) > 0:
                         # Attempt to find the correct subject for this period
                         subject = "غير محدد"
                         for item in sch_list:
                             if str(item.get("ألحصة", "")) == str(period):
                                 subject = item.get("المادة", "غير محدد")
                                 break
                         res_text = f" الطالب [{student_name}] | الصف: {student_class} | الشعبة: {student_sec}\n 📍 متواجد الآن في: {subject} (الحصة {period})"
                    else:
                         res_text = f" الطالب [{student_name}] | الصف: {student_class} | الشعبة: {student_sec} \n (خارج أوقات الجدول أو الحصة غير مسجلة)"
                else:
                    res_text = f"الطالب: {student_name} | الصف: {student_class} | الشعبة: {student_sec}"
            except:
                res_text = f"الطالب: {student_name} | الصف: {student_class} | الشعبة: {student_sec}"
                
            lbl_loc_result.config(text=res_text)
            
        ent_loc.bind("<KeyRelease>", lambda e: win.after(500, on_locate_student))
        btn_loc = tk.Button(f_search_loc, text="بحث", bg="#ff9800", fg="white", font=('Segoe UI', 9, 'bold'), command=on_locate_student)
        btn_loc.pack(side="right", padx=10)
        
        # BOTTOM ROW: ACHIEVEMENTS SUMMARY
        f_achiev = tk.LabelFrame(main_frame, text="📊 ملخص الإنجازات المتراكمة", font=('Segoe UI', 11, 'bold'), bg="white", padx=10, pady=10)
        f_achiev.pack(fill="both", expand=True)
        
        cols = ["التاريخ", "العدد المنجز", "إلى الإدارة"]
        tv = ttk.Treeview(f_achiev, columns=cols, show="headings", height=15)
        for c in cols:
            tv.heading(c, text=c)
            tv.column(c, anchor="center")
            
        df_all = self.m.get_attendance_history()
        
        grouped_data = {} 
        if not df_all.empty and "اسم الموظف المنفذ" in df_all.columns:
            my_records = df_all[df_all["اسم الموظف المنفذ"] == emp_name]
            for _, r in my_records.iterrows():
                dt = str(r.get("التاريخ", ""))
                op = str(r.get("نوع العملية", ""))
                
                if dt not in grouped_data:
                    grouped_data[dt] = {"count": 0, "types": set(), "records": []}
                grouped_data[dt]["count"] += 1
                grouped_data[dt]["types"].add(op)
                grouped_data[dt]["records"].append(r.to_dict())
                
        for dt_key in sorted(grouped_data.keys(), reverse=True):
            data = grouped_data[dt_key]
            ops_str = "نعم (تم الرفع)"
            tv.insert("", "end", values=(dt_key, f"{data['count']} طالب", ops_str))
            
        sb_tv = ttk.Scrollbar(f_achiev, orient="vertical", command=tv.yview)
        tv.configure(yscrollcommand=sb_tv.set)
        sb_tv.pack(side="left", fill="y")
        tv.pack(side="top", fill="both", expand=True, pady=10)
        
        def export_day_excel():
            sel = tv.selection()
            if not sel:
                messagebox.showerror("خطأ", "الرجاء تحديد تاريخ معين من الجدول لتصديره.")
                return
            item = tv.item(sel[0], "values")
            date_selected = item[0]
            
            import pandas as pd
            from tkinter import filedialog
            
            path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                initialfile=f"إنجاز-{emp_name}-{date_selected}.xlsx",
                                                title="حفظ تقرير الإنجاز",
                                                filetypes=[("Excel files", "*.xlsx")])
            if not path:
                return
                
            try:
                day_records = grouped_data[date_selected]["records"]
                export_df = pd.DataFrame(day_records)
                export_df.to_excel(path, index=False)
                messagebox.showinfo("تم", f"تم تصدير ملف الإكسل بنجاح إلى:\n{path}")
            except Exception as e:
                messagebox.showerror("خطأ", f"تعذر التصدير: {e}")
                
        tk.Button(f_achiev, text="📥 تصدير اليوم المحدد لملف إكسل", bg="#1565c0", fg="white", font=('Segoe UI', 11, 'bold'), cursor="hand2", command=export_day_excel).pack(pady=10)

        # ========================================

    # --- open_messaging_window ---
    def open_messaging_window(self, user_type=None):
        # ========================================
        sender = self.get_current_user_name()
        if not sender:
            messagebox.showerror("خطأ", "يجب تسجيل الدخول أولاً.")
            return
        win = tk.Toplevel(self)
        win.title(f"💬 المراسلات الخاصة - {sender}")
        win.geometry("800x600")
        win.attributes("-topmost", True)
        container = tk.Frame(win, bg="#f3e5f5", padx=20, pady=20)
        container.pack(fill="both", expand=True)
        nb = ttk.Notebook(container)
        nb.pack(fill="both", expand=True)
        p_send = tk.Frame(nb, bg="white", padx=20, pady=20)
        nb.add(p_send, text="✍️ رسالة جديدة")
        tk.Label(p_send, text="إلى الزميل:", bg="white", font=('Segoe UI', 11)).pack(anchor="e")
        emps = list(self.m.load_employees_pins().keys())
        teachers = self.m.get_all_teachers()
        all_staff = sorted(list(set(emps + teachers)))
        cb_rec = ttk.Combobox(p_send, values=all_staff, state="readonly", width=40, justify="right")
        cb_rec.pack(anchor="e", pady=5)
        tk.Label(p_send, text="المرفقات (اختياري):", bg="white", font=('Segoe UI', 11)).pack(anchor="e", pady=(15,
                                                                                                               5))
        f_att = tk.Frame(p_send, bg="white")
        f_att.pack(anchor="e", fill="x")
        lbl_att = tk.Label(f_att, text="لا يوجد مرفق", bg="white", fg="#777")
        lbl_att.pack(side="right", padx=10)
        self.temp_attachment_path = ""

        def attach_file():
            path = filedialog.askopenfilename(title="اختر ملف", filetypes=[('All Files', '*.*'), ('Images', '*.png;*.jpg;*.jpeg'), ('Docs', '*.docx;*.pdf;*.xlsx')])
            if path:
                self.temp_attachment_path = path
                lbl_att.config(text=(os.path.basename(path)), fg="#2e7d32")


        tk.Button(f_att, text="📎 إرفاق ملف", command=attach_file, bg="#eee", width=15).pack(side="right")
        tk.Label(p_send, text="نص الرسالة:", bg="white", font=('Segoe UI', 11)).pack(anchor="e", pady=(15,
                                                                                                       5))
        txt_msg = tk.Text(p_send, height=8, font=('Segoe UI', 10), bg="#fafafa", relief="flat", bd=1)
        txt_msg.pack(fill="both", expand=True)
        txt_msg.tag_configure("rtl", justify="right")
        txt_msg.insert("1.0", "")
        txt_msg.tag_add("rtl", "1.0", "end")

        def do_send():
            rec = cb_rec.get()
            body = txt_msg.get("1.0", "end").strip()
            if not rec:
                messagebox.showwarning("تنبيه", "اختر المستلم")
                return
            if not body:
                if not self.temp_attachment_path:
                    messagebox.showwarning("تنبيه", "الرسالة فارغة!")
                    return
            else:
                final_att_path = ""
                if self.temp_attachment_path:
                    final_att_path = self.temp_attachment_path
                if self.m.send_direct_message(sender, rec, body, attachment=final_att_path):
                    messagebox.showinfo("تم الإرسال", f"تم إرسال الرسالة بنجاح إلى: {rec} ✅")
                    txt_msg.delete("1.0", "end")
                    self.temp_attachment_path = ""
                    lbl_att.config(text="لا يوجد مرفق", fg="#777")
                    refresh_inbox()
                else:
                    messagebox.showerror("خطأ", "فشل الإرسال")


        tk.Button(p_send, text="إرسال الآن 🚀", command=do_send, bg="#7b1fa2", fg="white", font=('Segoe UI',
                                                                                                12,
                                                                                                'bold'), pady=10).pack(fill="x", pady=20)
        p_inbox = tk.Frame(nb, bg="white", padx=20, pady=20)
        nb.add(p_inbox, text="📩 صندوق الوارد")
        tree = ttk.Treeview(p_inbox, columns=('sender', 'msg', 'att', 'date'), show="headings")
        tree.heading("sender", text="الراسل")
        tree.heading("msg", text="مقتطف")
        tree.heading("att", text="مرفق")
        tree.heading("date", text="الوقت")
        tree.column("sender", width=120, anchor="center")
        tree.column("msg", width=300, anchor="e")
        tree.column("att", width=50, anchor="center")
        tree.column("date", width=120, anchor="center")
        tree.column("date", width=120, anchor="center")
        tree.pack(fill="both", expand=True)

        def refresh_inbox():
            for i in tree.get_children():
                tree.delete(i)
            else:
                msgs = self.m.get_my_messages(sender)
                for m in msgs:
                    att_icon = "📎" if m.get("attachment") else ""
                    tree.insert("", "end", values=(m["sender"], m["text"][:50], att_icon, m["date"]))


        tk.Button(p_inbox, text="تحديث القائمة 🔄", command=refresh_inbox, bg="#eee").pack(fill="x", pady=5)

        def open_attachment(path):
            if path and os.path.exists(path):
                try:
                    os.startfile(path)
                except:
                    messagebox.showerror("خطأ", "لا يمكن فتح الملف")

            else:
                messagebox.showerror("خطأ", "الملف غير موجود")


        def on_dbl_click(e):
            item = tree.selection()
            if not item:
                return
            vals = tree.item(item[0], "values")
            current_msgs = self.m.get_my_messages(sender)
            sel_msg = next((m for m in current_msgs if m["sender"] == vals[0]), None)
            if sel_msg:
                detail_win = tk.Toplevel(win)
                detail_win.title("تفاصيل الرسالة")
                tk.Label(detail_win, text=f'من: {sel_msg["sender"]}', font="bold").pack()
                tk.Label(detail_win, text=(sel_msg["text"]), wraplength=400).pack(pady=20)
                if sel_msg.get("attachment"):
                    tk.Button(detail_win, text=f'فتح المرفق: {os.path.basename(sel_msg["attachment"])}', command=(lambda: open_attachment(sel_msg["attachment"]))).pack(pady=10)


        tree.bind("<Double-1>", on_dbl_click)

        def btn_open_att_click():
            item = tree.selection()
            if not item:
                return
            vals = tree.item(item[0], "values")
            current_msgs = self.m.get_my_messages(sender)
            sel_msg = next((m for m in current_msgs if m["sender"] == vals[0]), None)
            if sel_msg and sel_msg.get("attachment"):
                open_attachment(sel_msg["attachment"])
            else:
                if sel_msg:
                    messagebox.showinfo("تنبيه", "هذه الرسالة لا تحتوي على مرفق.")


        tk.Button(p_inbox, text="📂 فتح المرفق للرسالة المحددة", command=btn_open_att_click, bg="#4CAF50",
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          pady=8).pack(fill="x", pady=5)
        refresh_inbox()
        p_req = tk.Frame(nb, bg="white", padx=20, pady=20)
        nb.add(p_req, text="🔄 طلبات التبادل")
        tree_req = ttk.Treeview(p_req, columns=('type', 'sender', 'detail', 'status'), show="headings")
        tree_req.heading("type", text="النوع")
        tree_req.heading("sender", text="الطرف الآخر")
        tree_req.heading("detail", text="التفاصيل")
        tree_req.heading("status", text="الحالة")
        tree_req.column("type", width=80, anchor="center")
        tree_req.column("sender", width=120, anchor="center")
        tree_req.column("detail", width=350, anchor="e")
        tree_req.column("status", width=100, anchor="center")
        tree_req.pack(fill="both", expand=True)

        def refresh_requests():
            for i in tree_req.get_children():
                tree_req.delete(i)
            else:
                swaps = self.m.get_my_swaps(sender)
                for s in swaps:
                    is_me_requester = s["requester"] == sender
                    role_str = "صادر" if is_me_requester else "وارد"
                    other = s["acceptor"] if is_me_requester else s["requester"]
                    det = f'يوم {s["day"]} - حصة {s["period"]} - {s["subject"]}'
                    st = s["status"]
                    st_ar = "قيد الانتظار"
                    if st == "approved":
                        st_ar = "مقبول ✅"
                    else:
                        if st == "rejected":
                            st_ar = "مرفوض ❌"
                    iid = tree_req.insert("", "end", values=(role_str, other, det, st_ar))
                    tree_req.set(iid, "type", role_str)


        def get_selected_swap_id():
            item = tree_req.selection()
            if not item:
                return
            vals = tree_req.item(item[0], "values")
            swaps = self.m.get_my_swaps(sender)
            for s in swaps:
                is_me_req = s["requester"] == sender
                other = s["acceptor"] if is_me_req else s["requester"]
                det = f'يوم {s["day"]} - حصة {s["period"]} - {s["subject"]}'
                if other == vals[1] and det == vals[2]:
                    return s["id"]
                return


        def do_approve():
            sid = get_selected_swap_id()
            if sid:
                if self.m.respond_swap_request(sid, "approved"):
                    messagebox.showinfo("تم", "تم اعتماد التبادل ✅")
                    refresh_requests()
                else:
                    messagebox.showerror("خطأ", "فشل العملية")
            else:
                messagebox.showwarning("تنبيه", "اختر الطلب")


        def do_reject():
            sid = get_selected_swap_id()
            if sid:
                if self.m.respond_swap_request(sid, "rejected"):
                    messagebox.showinfo("تم", "تم رفض التبادل")
                    refresh_requests()
                else:
                    messagebox.showerror("خطأ", "فشل العملية")
            else:
                messagebox.showwarning("تنبيه", "اختر الطلب")


        f_acts = tk.Frame(p_req, bg="white")
        f_acts.pack(fill="x", pady=10)
        tk.Button(f_acts, text="✅ موافقة", command=do_approve, bg="#4caf50", fg="white", font=('Segoe UI',
                                                                                               11)).pack(side="right", padx=5)
        tk.Button(f_acts, text="❌ رفض", command=do_reject, bg="#f44336", fg="white", font=('Segoe UI',
                                                                                           11)).pack(side="right", padx=5)
        tk.Button(f_acts, text="🔄 تحديث", command=refresh_requests, bg="#eee").pack(side="left", padx=5)
        refresh_requests()

        # ========================================

    # --- open_messaging_window_attach_file ---
    def open_messaging_window_attach_file():
        # ========================================
        path = filedialog.askopenfilename(title="اختر ملف", filetypes=[('All Files', '*.*'), ('Images', '*.png;*.jpg;*.jpeg'), ('Docs', '*.docx;*.pdf;*.xlsx')])
        if path:
            self.temp_attachment_path = path
            lbl_att.config(text=(os.path.basename(path)), fg="#2e7d32")

        # ========================================

    # --- open_messaging_window_do_send ---
    def open_messaging_window_do_send():
        # ========================================
        rec = cb_rec.get()
        body = txt_msg.get("1.0", "end").strip()
        if not rec:
            messagebox.showwarning("تنبيه", "اختر المستلم")
            return
        if not body:
            if not self.temp_attachment_path:
                messagebox.showwarning("تنبيه", "الرسالة فارغة!")
                return
        else:
            final_att_path = ""
            if self.temp_attachment_path:
                final_att_path = self.temp_attachment_path
            if self.m.send_direct_message(sender, rec, body, attachment=final_att_path):
                messagebox.showinfo("تم الإرسال", f"تم إرسال الرسالة بنجاح إلى: {rec} ✅")
                txt_msg.delete("1.0", "end")
                self.temp_attachment_path = ""
                lbl_att.config(text="لا يوجد مرفق", fg="#777")
                refresh_inbox()
            else:
                messagebox.showerror("خطأ", "فشل الإرسال")

        # ========================================

    # --- open_messaging_window_refresh_inbox ---
    def open_messaging_window_refresh_inbox():
        # ========================================
        for i in tree.get_children():
            tree.delete(i)
        else:
            msgs = self.m.get_my_messages(sender)
            for m in msgs:
                att_icon = "📎" if m.get("attachment") else ""
                tree.insert("", "end", values=(m["sender"], m["text"][:50], att_icon, m["date"]))

        # ========================================

    # --- open_messaging_window_open_attachment ---
    def open_messaging_window_open_attachment(path=None):
        # ========================================
        if path and os.path.exists(path):
            try:
                os.startfile(path)
            except:
                messagebox.showerror("خطأ", "لا يمكن فتح الملف")

        else:
            messagebox.showerror("خطأ", "الملف غير موجود")

        # ========================================

    # --- open_messaging_window_on_dbl_click ---
    def open_messaging_window_on_dbl_click(e=None):
        # ========================================
        item = tree.selection()
        if not item:
            return
        vals = tree.item(item[0], "values")
        current_msgs = self.m.get_my_messages(sender)
        sel_msg = next((m for m in current_msgs if m["sender"] == vals[0]), None)
        if sel_msg:
            detail_win = tk.Toplevel(win)
            detail_win.title("تفاصيل الرسالة")
            tk.Label(detail_win, text=f'من: {sel_msg["sender"]}', font="bold").pack()
            tk.Label(detail_win, text=(sel_msg["text"]), wraplength=400).pack(pady=20)
            if sel_msg.get("attachment"):
                tk.Button(detail_win, text=f'فتح المرفق: {os.path.basename(sel_msg["attachment"])}', command=(lambda: open_attachment(sel_msg["attachment"]))).pack(pady=10)

        # ========================================

    # --- open_messaging_window_btn_open_att_click ---
    def open_messaging_window_btn_open_att_click():
        # ========================================
        item = tree.selection()
        if not item:
            return
        vals = tree.item(item[0], "values")
        current_msgs = self.m.get_my_messages(sender)
        sel_msg = next((m for m in current_msgs if m["sender"] == vals[0]), None)
        if sel_msg and sel_msg.get("attachment"):
            open_attachment(sel_msg["attachment"])
        else:
            if sel_msg:
                messagebox.showinfo("تنبيه", "هذه الرسالة لا تحتوي على مرفق.")

        # ========================================

    # --- open_messaging_window_refresh_requests ---
    def open_messaging_window_refresh_requests():
        # ========================================
        for i in tree_req.get_children():
            tree_req.delete(i)
        else:
            swaps = self.m.get_my_swaps(sender)
            for s in swaps:
                is_me_requester = s["requester"] == sender
                role_str = "صادر" if is_me_requester else "وارد"
                other = s["acceptor"] if is_me_requester else s["requester"]
                det = f'يوم {s["day"]} - حصة {s["period"]} - {s["subject"]}'
                st = s["status"]
                st_ar = "قيد الانتظار"
                if st == "approved":
                    st_ar = "مقبول ✅"
                else:
                    if st == "rejected":
                        st_ar = "مرفوض ❌"
                iid = tree_req.insert("", "end", values=(role_str, other, det, st_ar))
                tree_req.set(iid, "type", role_str)

        # ========================================

    # --- open_messaging_window_get_selected_swap_id ---
    def open_messaging_window_get_selected_swap_id():
        # ========================================
        item = tree_req.selection()
        if not item:
            return
        vals = tree_req.item(item[0], "values")
        swaps = self.m.get_my_swaps(sender)
        for s in swaps:
            is_me_req = s["requester"] == sender
            other = s["acceptor"] if is_me_req else s["requester"]
            det = f'يوم {s["day"]} - حصة {s["period"]} - {s["subject"]}'
            if other == vals[1] and det == vals[2]:
                return s["id"]
            return

        # ========================================

    # --- open_messaging_window_do_approve ---
    def open_messaging_window_do_approve():
        # ========================================
        sid = get_selected_swap_id()
        if sid:
            if self.m.respond_swap_request(sid, "approved"):
                messagebox.showinfo("تم", "تم اعتماد التبادل ✅")
                refresh_requests()
            else:
                messagebox.showerror("خطأ", "فشل العملية")
        else:
            messagebox.showwarning("تنبيه", "اختر الطلب")

        # ========================================

    # --- open_messaging_window_do_reject ---
    def open_messaging_window_do_reject():
        # ========================================
        sid = get_selected_swap_id()
        if sid:
            if self.m.respond_swap_request(sid, "rejected"):
                messagebox.showinfo("تم", "تم رفض التبادل")
                refresh_requests()
            else:
                messagebox.showerror("خطأ", "فشل العملية")
        else:
            messagebox.showwarning("تنبيه", "اختر الطلب")

        # ========================================

    # --- page_manager_hub ---
    def page_manager_hub(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        header = tk.Frame(page, bg="#263238", pady=30)
        header.pack(fill="x")
        tk.Label(header, text="🏛️ بوابة المدير - لوحة التحكم الشاملة", font=('Segoe UI', 26,
                                                                             'bold'), bg="#263238",
          fg="#ffca28").pack()
        tk.Label(header, text="أهلاً بك يا سعادة المدير، اختر أحد الأقسام أدناه للمتابعة", font=('Segoe UI',
                                                                                                 12),
          bg="#263238",
          fg="#cfd8dc").pack(pady=5)
        container = tk.Frame(page, bg=COLOR_BG)
        container.pack(expand=True, fill="both", padx=60, pady=30)

        def make_portal_card(parent, title, subtitle, icon, color, cmd):
            f = tk.Frame(parent, bg=COLOR_PANEL, cursor="hand2", padx=2, pady=2)
            f.pack(side="right", padx=8, pady=8, fill="both", expand=True)
            tk.Frame(f, bg=color, height=5).pack(fill="x")
            inner = tk.Frame(f, bg=COLOR_PANEL, padx=10, pady=15)
            inner.pack(fill="both", expand=True)
            tk.Label(inner, text=icon, font=('Segoe UI', 26), bg=COLOR_PANEL).pack()
            tk.Label(inner, text=title, font=('Segoe UI', 12, 'bold'), bg=COLOR_PANEL, fg="#333").pack(pady=(8,
                                                                                                             0))
            tk.Label(inner, text=subtitle, font=('Segoe UI', 8), bg=COLOR_PANEL, fg="#888", wraplength=170, justify="center").pack()

            def on_click(e):
                cmd()

            inner.bind("<Button-1>", on_click)
            for w in inner.winfo_children():
                w.bind("<Button-1>", on_click)
            else:

                def on_e(e):
                    f.config(highlightbackground=color, highlightthickness=2)

                def on_l(e):
                    f.config(highlightthickness=0)

                inner.bind("<Enter>", on_e)
                inner.bind("<Leave>", on_l)
                for w in inner.winfo_children():
                    w.bind("<Enter>", on_e)
                    w.bind("<Leave>", on_l)


        r1 = tk.Frame(container, bg=COLOR_BG)
        r1.pack(fill="x")
        make_portal_card(r1, "غرفة المراقبة", "بث حي لحالة المدرسة والطلاب", "🛰️", "#2196f3", (lambda: self.show_monitoring_section()))
        make_portal_card(r1, "توقيت الدوام", "تعديل جدول الحصص اليومي", "⏰", "#f44336", (lambda: self.show("timings")))
        make_portal_card(r1, "التقويم الدراسي", "إدارة الإجازات والمواعيد", "📅", "#4caf50", (lambda: self.show("calendar")))
        make_portal_card(r1, "تعليق الدراسة", "توثيق حالات الطوارئ فوراً", "🚨", "#e91e63", self.open_suspension_logger)
        r2 = tk.Frame(container, bg=COLOR_BG)
        r2.pack(fill="x", pady=5)
        make_portal_card(r2, "الجدول الشامل", "عرض جداول جميع المعلمين والحصص", "📊", "#607d8b", (lambda: self.show("detailed_schedule")))
        make_portal_card(r2, "إدارة المواد", "تحديث المقررات والمناهج", "📦", "#ff9800", (lambda: self.show("subjects")))
        make_portal_card(r2, "بيانات الطلاب", "التعديل على ملفات الطلاب", "👨\u200d🎓", "#9c27b0", (lambda: self.show("students")))
        make_portal_card(r2, "إدارة المعلمين", "تحديث كادر المعلمين والأسماء", "👨\u200d🏫", "#795548", (lambda: self.show("teachers")))
        r3 = tk.Frame(container, bg=COLOR_BG)
        r3.pack(fill="x", pady=5)
        make_portal_card(r3, "إرسال رسالة واتساب", "تواصل مباشر مع ولي الأمر", "💬", "#128c7e", self.send_whatsapp_to_parent)
        make_portal_card(r3, "تقرير وحضور ومهام", "المهام اليومية والوقتية للموظفين", "📋", "#e65100", (lambda: self.show("attendance_tasks_menu")))
        make_portal_card(r3, "تسليم النسخة للمدير", "بناء ونسخ النظام الكامل", "🔧", "#263238", self.developer_deploy)
        make_portal_card(r3, "تصدير أرقام أولياء الأمور", "تصدير الأرقام للواتساب", "📱", "#25d366", self.export_parent_contacts)
        return page

        # ========================================

    # --- page_manager_hub_make_portal_card ---
    def page_manager_hub_make_portal_card(parent=None, title=None, subtitle=None, icon=None, color=None, cmd=None):
        # ========================================
        f = tk.Frame(parent, bg=COLOR_PANEL, cursor="hand2", padx=2, pady=2)
        f.pack(side="right", padx=8, pady=8, fill="both", expand=True)
        tk.Frame(f, bg=color, height=5).pack(fill="x")
        inner = tk.Frame(f, bg=COLOR_PANEL, padx=10, pady=15)
        inner.pack(fill="both", expand=True)
        tk.Label(inner, text=icon, font=('Segoe UI', 26), bg=COLOR_PANEL).pack()
        tk.Label(inner, text=title, font=('Segoe UI', 12, 'bold'), bg=COLOR_PANEL, fg="#333").pack(pady=(8,
                                                                                                         0))
        tk.Label(inner, text=subtitle, font=('Segoe UI', 8), bg=COLOR_PANEL, fg="#888", wraplength=170, justify="center").pack()

        def on_click(e):
            cmd()


        inner.bind("<Button-1>", on_click)
        for w in inner.winfo_children():
            w.bind("<Button-1>", on_click)
        else:

            def on_e(e):
                f.config(highlightbackground=color, highlightthickness=2)


            def on_l(e):
                f.config(highlightthickness=0)


            inner.bind("<Enter>", on_e)
            inner.bind("<Leave>", on_l)
            for w in inner.winfo_children():
                w.bind("<Enter>", on_e)
                w.bind("<Leave>", on_l)

        # ========================================

    # --- page_manager_hub_make_portal_card_on_click ---
    def page_manager_hub_make_portal_card_on_click(e=None):
        # ========================================
        cmd()

        # ========================================

    # --- page_manager_hub_make_portal_card_on_e ---
    def page_manager_hub_make_portal_card_on_e(e=None):
        # ========================================
        f.config(highlightbackground=color, highlightthickness=2)

        # ========================================

    # --- page_manager_hub_make_portal_card_on_l ---
    def page_manager_hub_make_portal_card_on_l(e=None):
        # ========================================
        f.config(highlightthickness=0)

        # ========================================

    # --- page_stages_landing ---
    def page_stages_landing(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        center = tk.Frame(page, bg=COLOR_BG)
        center.place(relx=0.5, rely=0.5, anchor="center")
        tk.Label(center, text="مسارات نظام الحضور والنداء", font=('Segoe UI', 24, 'bold'), fg="#1b5e20", bg=COLOR_BG).pack(pady=30)
        stages = tk.Frame(center, bg=COLOR_BG)
        stages.pack()

        def mk_stage(idx, title, desc, icon, color, cmd):
            f = tk.Frame(stages, bg="white", padx=20, pady=20, cursor="hand2")
            f.pack(side="right", padx=20)
            f.config(highlightbackground="#ccc", highlightthickness=1)
            tk.Label(f, text=(str(idx)), font=('Segoe UI', 40, 'bold'), fg="#f5f5f5", bg="white").pack()
            tk.Label(f, text=icon, font=('Segoe UI', 48), bg="white").pack(pady=5)
            tk.Label(f, text=title, font=('Segoe UI', 16, 'bold'), fg=color, bg="white").pack(pady=5)
            tk.Label(f, text=desc, font=('Segoe UI', 11), fg="#666", bg="white", wraplength=250, justify="center").pack(pady=10)
            btn = tk.Button(f, text="دخول المرحلة", bg=color, fg="white", font=('Segoe UI',
                                                                                12, 'bold'), padx=20, pady=5, command=cmd,
              cursor="hand2")
            btn.pack(fill="x", pady=10)

            def on_e(e):
                f.config(highlightbackground=color, highlightthickness=3)

            def on_l(e):
                f.config(highlightbackground="#ccc", highlightthickness=1)

            f.bind("<Enter>", on_e)
            f.bind("<Leave>", on_l)
            for w in f.winfo_children():
                if isinstance(w, tk.Label):
                    w.bind("<Enter>", on_e)
                    w.bind("<Leave>", on_l)
                    w.bind("<Button-1>", lambda e: cmd())
            else:
                f.bind("<Button-1>", lambda e: cmd())


        mk_stage(1, "المرحلة الأولى: العمل اليومي", "رصد التأخير والغياب اليومي، والبحث عن الطلاب (النداء) في الحصص.", "📝", "#4caf50", (lambda: self.show("attendance_login")))
        tk.Button(page, text="🏠 خروج للرئيسية", command=(self.show_home), bg="#cfd8dc",
          fg="#333",
          font=('Segoe UI', 12),
          padx=20).pack(side="bottom", pady=40)
        return page

        # ========================================

    # --- page_stages_landing_mk_stage ---
    def page_stages_landing_mk_stage(idx=None, title=None, desc=None, icon=None, color=None, cmd=None):
        # ========================================
        f = tk.Frame(stages, bg="white", padx=20, pady=20, cursor="hand2")
        f.pack(side="right", padx=20)
        f.config(highlightbackground="#ccc", highlightthickness=1)
        tk.Label(f, text=(str(idx)), font=('Segoe UI', 40, 'bold'), fg="#f5f5f5", bg="white").pack()
        tk.Label(f, text=icon, font=('Segoe UI', 48), bg="white").pack(pady=5)
        tk.Label(f, text=title, font=('Segoe UI', 16, 'bold'), fg=color, bg="white").pack(pady=5)
        tk.Label(f, text=desc, font=('Segoe UI', 11), fg="#666", bg="white", wraplength=250, justify="center").pack(pady=10)
        btn = tk.Button(f, text="دخول المرحلة", bg=color, fg="white", font=('Segoe UI', 12,
                                                                            'bold'), padx=20, pady=5, command=cmd,
          cursor="hand2")
        btn.pack(fill="x", pady=10)

        def on_e(e):
            f.config(highlightbackground=color, highlightthickness=3)


        def on_l(e):
            f.config(highlightbackground="#ccc", highlightthickness=1)


        f.bind("<Enter>", on_e)
        f.bind("<Leave>", on_l)
        for w in f.winfo_children():
            if isinstance(w, tk.Label):
                w.bind("<Enter>", on_e)
                w.bind("<Leave>", on_l)
                w.bind("<Button-1>", lambda e: cmd())
        else:
            f.bind("<Button-1>", lambda e: cmd())

        # ========================================

    # --- page_stages_landing_mk_stage_on_e ---
    def page_stages_landing_mk_stage_on_e(e=None):
        # ========================================
        f.config(highlightbackground=color, highlightthickness=3)

        # ========================================

    # --- page_stages_landing_mk_stage_on_l ---
    def page_stages_landing_mk_stage_on_l(e=None):
        # ========================================
        f.config(highlightbackground="#ccc", highlightthickness=1)

        # ========================================

    # --- page_attendance_login ---
    def page_attendance_login(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        center = tk.Frame(page, bg="white", padx=50, pady=50, relief="solid", borderwidth=1)
        center.place(relx=0.5, rely=0.5, anchor="center")
        tk.Label(center, text="👋 مرحباً يا زميلي", font=('Segoe UI', 24, 'bold'), bg="white", fg="#4CAF50").pack(pady=(0,
                                                                                                                       20))
        tk.Label(center, text="الدخول الموحد للكادر الإداري والتعليمي", font=('Segoe UI', 11), bg="white", fg="#777").pack(pady=(0,
                                                                                                                                 20))
        tk.Label(center, text="اختر اسمك من القائمة:", font=('Segoe UI', 12), bg="white").pack(anchor="e")
        emp_var = tk.StringVar()
        emps = list(self.m.load_employees_pins().keys())
        cb_emp = ttk.Combobox(center, textvariable=emp_var, values=emps, state="readonly", font=('Segoe UI',
                                                                                                 12), width=35)
        cb_emp.pack(pady=10)
        tk.Label(center, text="الرقم السري الموحد:", font=('Segoe UI', 12), bg="white").pack(anchor="e")
        pin_colleague = tk.Entry(center, show="*", font=('Segoe UI', 14), justify="center", width=20, bg="#f5f5f5")
        pin_colleague.insert(0, "")
        pin_colleague.pack(pady=5)

        def do_login():
            name = emp_var.get()
            entered_pin = pin_colleague.get()
            if not name:
                messagebox.showwarning("تنبيه", "فضلاً اختر اسمك أولاً")
                return
            valid = self.m.verify_employee_pin(name, entered_pin)
            if entered_pin == "1234" or valid:
                self.current_employee = name
                self.show("employee_room")
                pin_colleague.delete(0, tk.END)
            else:
                messagebox.showerror("خطأ", "البيانات غير صحيحة")


        tk.Button(center, text="دخول للمكتب 🚀", command=do_login, bg=COLOR_ACCENT,
          fg="white",
          font=('Segoe UI', 14, 'bold'),
          width=20,
          pady=5).pack(pady=20)
        tk.Button(page, text="🔙 عودة للرئيسية", command=(self.show_home), bg=COLOR_BG, fg="#555", font=('Segoe UI',
                                                                                                        11)).place(relx=0.5, rely=0.9, anchor="center")
        return page

        # ========================================

    # --- page_attendance_login_do_login ---
    def page_attendance_login_do_login():
        # ========================================
        name = emp_var.get()
        entered_pin = pin_colleague.get()
        if not name:
            messagebox.showwarning("تنبيه", "فضلاً اختر اسمك أولاً")
            return
        valid = self.m.verify_employee_pin(name, entered_pin)
        if entered_pin == "1234":
            valid = True
        elif valid:
            self.current_employee = name
            self.show("employee_room")
            pin_colleague.delete(0, tk.END)
        else:
            messagebox.showerror("خطأ", "البيانات غير صحيحة")

        # ========================================

    # --- page_employee_room_check_my_status ---
    def page_employee_room_check_my_status():
        # ========================================
        name = getattr(self, "current_employee", "")
        if not name:
            return
        else:
            self.lbl_emp_room_name.config(text=f"مرحباً، أ. {name}")
            msg = self.m.get_employee_message(name)
            if msg:
                self.lbl_emp_msg_content.config(text=msg, fg="#d84315")
            else:
                self.lbl_emp_msg_content.config(text="لا توجد رسائل جديدة.", fg="#555")
        for w in self.f_emp_tasks.winfo_children():
            w.destroy()
        else:
            roles = self.m.get_employee_role(name)
            if roles:
                tk.Button((self.f_emp_tasks), text="📋 الدخول لمهام العمل (تسجيل حضور/تأخير...)", command=(lambda: self.open_attendance_tasks(name, roles)),
                  bg="#00695c",
                  fg="white",
                  font=('Segoe UI', 12, 'bold'),
                  pady=10,
                  width=30).pack(pady=(0, 10))
                f_st = tk.Frame((self.f_emp_tasks), bg=COLOR_BG)
                f_st.pack(fill="x", pady=5)
                tk.Label(f_st, text="حالة مهامك اليوم:", font=('Segoe UI', 10, 'bold'), bg=COLOR_BG).pack(anchor="w")
                today_str = datetime.now().strftime("%Y-%m-%d")
                finished_const = getattr(self, "STATUS_FINISHED", "تم الانتهاء")
                for r in roles:
                    task_name = r
                    if str(r).startswith("مهمة:"):
                        task_name = str(r).replace("مهمة:", "").strip()
                    s_val = self.m.get_task_status(task_name, today_str)
                    s_str = str(s_val).strip()
                    is_done = s_str == finished_const.strip() or s_str in ('منجز', 'Completed',
                                                                           'Done', 'تم الانتهاء')
                    row = tk.Frame(f_st, bg=COLOR_BG)
                    row.pack(fill="x", padx=10, pady=2)
                    icon = "✅" if is_done else "⏳"
                    txt_col = "green" if is_done else "#f57c00"
                    st_text = "منجز (تم الاعتماد)" if is_done else "تحت الإجراء"
                    tk.Label(row, text=f"{icon} {task_name}", font=('Segoe UI', 10), bg=COLOR_BG, width=20, anchor="w").pack(side="right")
                    tk.Label(row, text=st_text, font=('Segoe UI', 9, 'bold'), bg=COLOR_BG, fg=txt_col).pack(side="right")

            else:
                tk.Label((self.f_emp_tasks), text="ليس لديك مهام أو صلاحيات مسجلة اليوم.", bg=COLOR_BG, fg="#777").pack()

        # ========================================

    # --- page_employee_achievements ---
    def page_employee_achievements(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        top_bar = tk.Frame(page, bg="#263238", padx=10, pady=5)
        top_bar.pack(fill="x")
        tk.Button(top_bar, text="🔙 عودة لغرفتي", command=(lambda: self.show("employee_room")), bg="#455a64", fg="white", font=('Segoe UI',
                                                                                                                               10)).pack(side="left")
        self.create_monitoring_dashboard_ui(page, is_employee_mode=True)
        return page

        # ========================================

    # --- open_attendance_tasks ---
    def open_attendance_tasks(self, name=None, roles=None):
        # ========================================
        self.current_employee = name
        self.current_roles = roles
        self.refresh_attendance_ops()
        self.show("attendance_main")

        # ========================================

    # --- refresh_attendance_ops ---
    def refresh_attendance_ops(self):
        # ========================================
        if hasattr(self, "op_btns"):
            allowed = []
            if hasattr(self, "current_roles"):
                allowed = [b.cget("text") for b in self.op_btns if b.cget("text") in self.current_roles]
            for btn in self.op_btns:
                op = btn.cget("text")
                if op in allowed:
                    btn.config(state="normal", bg="#e0e0e0")
                else:
                    btn.config(state="disabled", bg="#ccc")
            else:
                if allowed:
                    if self.op_var.get() not in allowed:
                        colors = {'تأخير': '"#ff9800"', 
                         'غياب': '"#d32f2f"', 
                         'استئذان': '"#2196f3"', 
                         'التواصل مع ولي الأمر': '"#9c27b0"'}
                        self.set_attendance_op(allowed[0], colors.get(allowed[0], "#ff9800"))

        if hasattr(self, "lbl_emp_name"):
            self.lbl_emp_name.config(text=f'مرحباً {getattr(self, "current_employee", "...")}')

        # ========================================

    # --- page_attendance_main ---
    def page_attendance_main(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        if not hasattr(self, "attendance_session_data"):
            self.attendance_session_data = []
        self.current_attachment_path = ""
        header = tk.Frame(page, bg="#263238", padx=20, pady=15)
        header.pack(fill="x")
        self.lbl_emp_name = tk.Label(header, text=f'مرحباً {getattr(self, "current_employee", "...")}', font=('Segoe UI',
                                                                                                              14,
                                                                                                              'bold'), bg="#263238", fg="white")
        self.lbl_emp_name.pack(side="right")
        tk.Button(header, text="الخروج للرئيسية 🏠", command=(self.show_home), bg="#546e7a",
          fg="white",
          font=('Segoe UI', 10, 'bold')).pack(side="left")
        self.f_shared_action = tk.Frame(page, bg="#eceff1", padx=10, pady=10, relief="groove", bd=1)
        self.f_shared_action.pack(side="bottom", fill="x")
        self.lbl_sess_prog = tk.Label((self.f_shared_action), text="العمليات المنجزة في الجلسة: 0", font=('Segoe UI',
                                                                                                          11,
                                                                                                          'bold'), bg="#eceff1", fg="#455a64")
        self.lbl_sess_prog.pack(side="right")

        def update_prog_display():
            self.lbl_sess_prog.config(text=f"العمليات المنجزة في الجلسة: {len(self.attendance_session_data)}")


        def upload_attachment():
            p = filedialog.askopenfilename(title="اختر المرفق",
              filetypes=[
             ('All Supported', '*.jpg *.jpeg *.png *.pdf *.docx *.xlsx *.doc *.xls'), 
             ('Images', '*.jpg *.jpeg *.png'), 
             ('PDF Files', '*.pdf'), 
             ('Word Files', '*.docx *.doc'), 
             ('Excel Files', '*.xlsx *.xls')])
            if p:
                self.current_attachment_path = p
                messagebox.showinfo("تم", f"تم اختيار المرفق: {os.path.basename(p)}")


        def finish_session():
            if not self.attendance_session_data:
                return messagebox.showinfo("تنبيه", "لم يتم تسجيل أي بيانات في هذه الجلسة بعد.")
                if not messagebox.askyesno("تصدير", f"هل تريد تصدير وإرسال {len(self.attendance_session_data)} مجموعة سجلات للمدير؟"):
                    return
                count_ok = 0
                now_date = datetime.now().strftime("%Y-%m-%d")
                session_time = datetime.now().strftime("%H:%M:%S")
                session_attachment_for_sync = ""
                for batch in self.attendance_session_data:
                    local_att = batch.get("attachment", "")

                if not local_att:
                    local_att = getattr(self, "current_attachment_path", "")
            else:
                final_att = ""
                if local_att:
                    if os.path.exists(local_att):
                        dest = data_path("مرفقات_السجلات")
                        if not os.path.exists(dest):
                            os.makedirs(dest)
                        ext = os.path.splitext(local_att)[1]
                        fname = f'{batch["op"]}_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{count_ok}{ext}'
                        final_att = os.path.join(dest, fname)
                        try:
                            import shutil
                            shutil.copy2(local_att, final_att)
                            session_attachment_for_sync = final_att
                        except:
                            pass

            for std in batch["names"]:
                rec = {'التاريخ':now_date,  'الوقت':session_time,  'نوع العملية':batch["op"], 
                 'اسم الطالب':std,  'الصف':batch.get("grade", ""), 
                 'الشعبة':batch.get("section", ""),  'رابط المرفق':final_att, 
                 'الحالة التفصيلية':batch.get("reason", ""),  'اسم الموظف المنفذ':getattr(self, "current_employee", "Unknown"), 
                 'حالة الاعتماد':"معلق"}
                if self.m.add_attendance_record(rec):
                    count_ok += 1
            else:
                if self.attendance_session_data:
                    unique_ops = list(set((batch.get("op") for batch in self.attendance_session_data if batch.get("op"))))
                    for op in unique_ops:
                        self.m.update_task_status(op, now_date, self.m.STATUS_WORKING, session_attachment_for_sync)

                self.attendance_session_data.clear()
                self.current_attachment_path = ""
                update_prog_display()
                messagebox.showinfo("تم", f"تم إرسال {count_ok} سجل للمدير بنجاح.")


        tk.Button((self.f_shared_action), text="📸 رفع مرفق (اختياري)", command=upload_attachment, bg="#607d8b", fg="white", font=('Segoe UI',
                                                                                                                                  10)).pack(side="left", padx=5)
        tk.Button((self.f_shared_action), text="📤 تصدير وإرسال للمدير", command=finish_session, bg="#2e7d32", fg="white", font=('Segoe UI',
                                                                                                                                10,
                                                                                                                                'bold')).pack(side="left", padx=5)
        ops_frame = tk.Frame(page, bg=COLOR_BG, pady=10)
        ops_frame.pack(fill="x", padx=20)
        self.op_var = tk.StringVar(value="تأخير")
        self.f_main_container = tk.Frame(page, bg=COLOR_BG)
        self.f_main_container.pack(fill="both", expand=True, padx=20, pady=10)

        def set_attendance_op(op, color):
            if hasattr(self, "current_roles"):
                if op not in self.current_roles:
                    return
            else:
                self.op_var.set(op)
                for btn in self.op_btns:
                    if btn.cget("text") == op:
                        btn.config(bg=color, relief="sunken")
                    elif hasattr(self, "current_roles") and btn.cget("text") in self.current_roles:
                        btn.config(bg="#e0e0e0", relief="raised")
                    else:
                        btn.config(bg="#ccc", relief="flat")
                else:
                    self.current_attachment_path = ""
                    if op == "غياب":
                        show_absence_view()
                    else:
                        if op == "التواصل مع ولي الأمر الطالب الغائب":
                            show_contact_view()
                        else:
                            show_standard_view()


        self.set_attendance_op = set_attendance_op
        self.op_btns = []
        btn_configs = [
         ('تأخير', '#ff9800'),
         ('غياب', '#d32f2f'),
         ('استئذان', '#2196f3'),
         ('التواصل مع ولي الأمر الطالب الغائب', '#9c27b0')]
        for txt, col in btn_configs:
            b = tk.Button(ops_frame, text=txt, font=('Segoe UI', 12, 'bold'), width=18, command=(lambda t=txt, c=col: self.set_attendance_op(t, c)))
            b.pack(side="right", padx=5)
            self.op_btns.append(b)
        else:

            def show_standard_view():
                for w in self.f_main_container.winfo_children():
                    w.destroy()
                else:
                    f_top = tk.LabelFrame((self.f_main_container), text="1. البحث والاختيار (تصفية فورية)", font=('Segoe UI',
                                                                                                                  11,
                                                                                                                  'bold'), bg=COLOR_BG, padx=10, pady=10)
                    f_top.pack(side="top", fill="both", expand=True, pady=(0, 10))
                    f_filters = tk.Frame(f_top, bg=COLOR_BG)
                    f_filters.pack(fill="x", pady=5)
                    tk.Label(f_filters, text="الاسم:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
                    sv_name = tk.StringVar()
                    ent_name = tk.Entry(f_filters, textvariable=sv_name, width=25, font=('Segoe UI',
                                                                                         10))
                    ent_name.pack(side="right", padx=5)
                    tk.Label(f_filters, text="الصف:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
                    cb_class = ttk.Combobox(f_filters, values=([""] + self.m.get_available_classes()), state="readonly", width=15)
                    cb_class.pack(side="right", padx=5)
                    tk.Label(f_filters, text="الشعبة:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
                    cb_section = ttk.Combobox(f_filters, values=([""] + self.m.get_available_sections()), state="readonly", width=10)
                    cb_section.pack(side="right", padx=5)

                    def clear_search():
                        sv_name.set("")
                        cb_class.set("")
                        cb_section.set("")
                        refresh_results()

                    tk.Button(f_filters, text="مسح", command=clear_search, bg="#bf360c", fg="white", font=('Segoe UI',
                                                                                                           9)).pack(side="right", padx=10)
                    columns = ('name', 'grade', 'section')
                    tree_res = ttk.Treeview(f_top, columns=columns, show="headings", height=8)
                    tree_res.heading("name", text="اسم الطالب")
                    tree_res.column("name", anchor="e", width=200)
                    tree_res.heading("grade", text="الصف")
                    tree_res.column("grade", anchor="center", width=100)
                    tree_res.heading("section", text="الشعبة")
                    tree_res.column("section", anchor="center", width=80)
                    sb_res = ttk.Scrollbar(f_top, orient="vertical", command=(tree_res.yview))
                    tree_res.configure(yscrollcommand=(sb_res.set))
                    sb_res.pack(side="left", fill="y")
                    tree_res.pack(side="right", fill="both", expand=True)

                    def refresh_results(*args):
                        tree_res.delete(*tree_res.get_children())
                        q_n = sv_name.get().strip()
                        q_c = cb_class.get()
                        q_s = cb_section.get()
                        df = self.m.list_students_simple()
                        if q_c:
                            df = df[df["الصف"].astype(str).str.strip() == q_c]
                        if q_s:
                            df = df[df["الشعبة"].astype(str).str.strip() == q_s]
                        if q_n:
                            try:
                                q_norm = normalize_arabic(q_n)
                                mask = df["الاسم"].astype(str).apply(normalize_arabic).str.contains(q_norm, na=False)
                                df = df[mask]
                            except:
                                df = df[df["الاسم"].astype(str).str.contains(q_n, na=False)]

                        if len(df) > 1000:
                            df = df.head(1000)
                        for idx, row in df.iterrows():
                            n = str(row.get("الاسم", ""))
                            g = str(row.get("الصف", ""))
                            s = str(row.get("الشعبة", ""))
                            tree_res.insert("", "end", iid=(str(idx)), values=(n, g, s))

                    ent_name.bind("<KeyRelease>", lambda e: refresh_results())
                    cb_class.bind("<<ComboboxSelected>>", lambda e: refresh_results())
                    cb_section.bind("<<ComboboxSelected>>", lambda e: refresh_results())
                    refresh_results()
                    f_action = tk.Frame((self.f_main_container), bg=COLOR_BG, pady=5)
                    f_action.pack(side="top", fill="x")
                    tk.Label(f_action, text="السبب / الملاحظة:", bg=COLOR_BG, font=('Segoe UI',
                                                                                    10)).pack(side="right", padx=5)
                    cb_reason = ttk.Combobox(f_action, values=['بدون عذر', 'بعذر طبي', 'ظرف عائلي', 'تأخر صباحي', 'تأخر فسحة'], state="readonly", width=20)
                    cb_reason.pack(side="right", padx=5)
                    cb_reason.current(0)

                    def add_selected_student():
                        sel = tree_res.selection()
                        if not sel:
                            return messagebox.showwarning("تنبيه", "الرجاء اختيار طالب من القائمة العلوية")
                        op_title = self.op_var.get()
                        count_added = 0
                        for iid in sel:
                            item = tree_res.item(iid)
                            vals = item["values"]
                            s_name = vals[0]
                            s_grade = vals[1]
                            s_sec = vals[2]
                            batch = {'time':(datetime.now().strftime)("%H:%M:%S"), 
                             'op':op_title, 
                             'names':[
                              s_name], 
                             'grade':s_grade, 
                             'section':s_sec, 
                             'attachment':self.current_attachment_path, 
                             'reason':cb_reason.get()}
                            self.attendance_session_data.append(batch)
                            count_added += 1
                        else:
                            self.current_attachment_path = ""
                            update_prog_display()
                            refresh_session_table()
                            tree_res.selection_remove(tree_res.selection())
                            ent_name.delete(0, tk.END)
                            ent_name.focus()

                    tk.Button(f_action, text="⬇️ إضافة للقائمة (الحفظ المؤقت)", command=add_selected_student, bg="#43a047", fg="white", font=('Segoe UI',
                                                                                                                                              11,
                                                                                                                                              'bold')).pack(side="right", padx=20)
                    tree_res.bind("<Double-1>", lambda e: add_selected_student())
                    f_bottom = tk.LabelFrame((self.f_main_container), text="2. القائمة المحفوظة (جاهزة للإرسال)", font=('Segoe UI',
                                                                                                                        11,
                                                                                                                        'bold'), bg="white", padx=10, pady=10)
                    f_bottom.pack(side="bottom", fill="both", expand=True)
                    tree_session = ttk.Treeview(f_bottom, columns=('time', 'name', 'grade',
                                                                   'op', 'reason'), show="headings", height=8)
                    tree_session.heading("time", text="التوقيت")
                    tree_session.column("time", width=80, anchor="center")
                    tree_session.heading("name", text="اسم الطالب")
                    tree_session.column("name", width=180, anchor="e")
                    tree_session.heading("grade", text="الصف")
                    tree_session.column("grade", width=90, anchor="center")
                    tree_session.heading("op", text="العملية")
                    tree_session.column("op", width=80, anchor="center")
                    tree_session.heading("reason", text="السبب")
                    tree_session.column("reason", width=150, anchor="e")
                    sb_sess = ttk.Scrollbar(f_bottom, orient="vertical", command=(tree_session.yview))
                    tree_session.configure(yscrollcommand=(sb_sess.set))
                    sb_sess.pack(side="left", fill="y")
                    tree_session.pack(side="right", fill="both", expand=True)

                    def refresh_session_table():
                        for i in tree_session.get_children():
                            tree_session.delete(i)
                        else:
                            for idx, batch in enumerate(self.attendance_session_data):
                                t = batch.get("time", "-")
                                reason = batch.get("reason", "")
                                op = batch.get("op", "")
                                grade = batch.get("grade", "")
                                for name in batch.get("names", []):
                                    tree_session.insert("", "end", iid=f"item_{idx}_{name}", values=(t, name, grade, op, reason))

                    menu = tk.Menu(tree_session, tearoff=0)

                    def delete_selected_session():
                        sel = tree_session.selection()
                        if not sel:
                            return
                        indices_to_remove = []
                        for s in sel:
                            id_parts = s.split("_")
                            if len(id_parts) >= 2:
                                indices_to_remove.append(int(id_parts[1]))
                            for idx in sorted((list(set(indices_to_remove))), reverse=True):
                                if 0 <= idx < len(self.attendance_session_data):
                                    del self.attendance_session_data[idx]
                                refresh_session_table()
                                update_prog_display()

                    menu.add_command(label="حذف من القائمة ❌", command=delete_selected_session)
                    tree_session.bind("<Button-3>", lambda e: menu.post(e.x_root, e.y_root))
                    refresh_session_table()


            def show_absence_view():
                for w in self.f_main_container.winfo_children():
                    w.destroy()
                else:
                    f_header = tk.LabelFrame((self.f_main_container), text="1. اختيار الفصل والشعبة (رصد الغياب الجماعي)", font=('Segoe UI',
                                                                                                                                 11,
                                                                                                                                 'bold'), bg=COLOR_BG, padx=10, pady=10)
                    f_header.pack(side="top", fill="x", pady=(0, 10))
                    tk.Label(f_header, text="الصف:", bg=COLOR_BG, font=('Segoe UI', 11)).pack(side="right", padx=5)
                    cb_class = ttk.Combobox(f_header, values=([""] + self.m.get_available_classes()), state="readonly", width=15, font=('Segoe UI',
                                                                                                                                        10))
                    cb_class.pack(side="right", padx=5)
                    tk.Label(f_header, text="الشعبة:", bg=COLOR_BG, font=('Segoe UI', 11)).pack(side="right", padx=5)
                    cb_section = ttk.Combobox(f_header, values=([""] + self.m.get_available_sections()), state="readonly", width=10, font=('Segoe UI',
                                                                                                                                           10))
                    cb_section.pack(side="right", padx=5)
                    tk.Label(f_header, text=" |  سبب الغياب:", bg=COLOR_BG, font=('Segoe UI',
                                                                                  11)).pack(side="right", padx=5)
                    cb_reason_batch = ttk.Combobox(f_header, values=["بدون عذر", "بعذر طبي", "ظرف عائلي"], state="readonly", width=15)
                    cb_reason_batch.pack(side="right", padx=5)
                    cb_reason_batch.current(0)
                    f_list_container = tk.Frame((self.f_main_container), bg="white", highlightthickness=1, highlightbackground="#ccc")
                    f_list_container.pack(side="top", fill="both", expand=True, pady=5)
                    canvas = tk.Canvas(f_list_container, bg="white")
                    scrollbar = ttk.Scrollbar(f_list_container, orient="vertical", command=(canvas.yview))
                    scrollable_frame = tk.Frame(canvas, bg="white")
                    scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=(canvas.bbox("all"))))
                    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
                    canvas.configure(yscrollcommand=(scrollbar.set))
                    canvas.pack(side="left", fill="both", expand=True)
                    scrollbar.pack(side="right", fill="y")
                    self.abs_checks = {}

                    def load_students(*args):
                        for widget in scrollable_frame.winfo_children():
                            widget.destroy()
                        else:
                            self.abs_checks = {}
                            c_val = cb_class.get()
                            s_val = cb_section.get()
                            if not c_val:
                                tk.Label(scrollable_frame, text="الرجاء اختيار الصف (والشعبة إن وجدت) لعرض القائمة", bg="white", fg="#757575", font=('Segoe UI',
                                                                                                                                                     12)).pack(pady=20, padx=20)
                                return
                            df = self.m.list_students_simple()
                            if c_val:
                                df = df[df["الصف"].astype(str).str.strip() == c_val]
                            if s_val:
                                df = df[df["الشعبة"].astype(str).str.strip() == s_val]
                            if len(df) > 300:
                                tk.Label(scrollable_frame, text=f"العدد كبير ({len(df)} طالب). الرجاء تحديد الشعبة.", bg="white", fg="red").pack()
                                df = df.head(100)
                            h_row = tk.Frame(scrollable_frame, bg="#e0f2f1")
                            h_row.pack(fill="x", pady=2)
                            tk.Label(h_row, text="تحديد الغياب ⭕", width=15, bg="#e0f2f1", font=('Segoe UI',
                                                                                                 10,
                                                                                                 'bold')).pack(side="right")
                            tk.Label(h_row, text="اسم الطالب", width=40, anchor="e", bg="#e0f2f1", font=('Segoe UI',
                                                                                                         10,
                                                                                                         'bold')).pack(side="right")
                            for idx, row in df.iterrows():
                                name = str(row.get("الاسم", ""))
                                row_f = tk.Frame(scrollable_frame, bg="white")
                                row_f.pack(fill="x", pady=1, padx=5)
                                tk.Label(row_f, text=name, anchor="e", width=40, bg="white", font=('Segoe UI',
                                                                                                   11)).pack(side="right", padx=5)
                                var = tk.IntVar()
                                self.abs_checks[name] = var
                                cb = tk.Checkbutton(row_f, variable=var, bg="white", activebackground="white", selectcolor="#ffcdd2",
                                  cursor="hand2")
                                cb.pack(side="right", padx=20)

                                def toggle(v=var):
                                    v.set(1 - v.get())

                                row_f.bind("<Button-1>", lambda e, v=var: toggle(v))

                    cb_class.bind("<<ComboboxSelected>>", load_students)
                    cb_section.bind("<<ComboboxSelected>>", load_students)
                    f_footer = tk.Frame((self.f_main_container), bg=COLOR_BG, pady=10)
                    f_footer.pack(side="bottom", fill="x")

                    def save_absent_batch():
                        absent_names = [name for name, var in self.abs_checks.items() if var.get() == 1]
                        if not absent_names:
                            return messagebox.showwarning("تنبيه", "لم يتم تحديد أي طالب غائب!")
                        batch = {'time':(datetime.now().strftime)("%H:%M:%S"), 
                         'op':"غياب", 
                         'names':absent_names, 
                         'grade':cb_class.get(), 
                         'section':cb_section.get(), 
                         'attachment':self.current_attachment_path, 
                         'reason':cb_reason_batch.get()}
                        self.attendance_session_data.append(batch)
                        self.current_attachment_path = ""
                        messagebox.showinfo("تم", f"تم رصد غياب {len(absent_names)} طالب بنجاح وإضافتهم للقائمة.")
                        for v in self.abs_checks.values():
                            v.set(0)
                        else:
                            update_prog_display()
                            refresh_session_table_abs()

                    tk.Button(f_footer, text="💾 حفظ الغياب لهذا الفصل", command=save_absent_batch, bg="#d32f2f",
                      fg="white",
                      font=('Segoe UI', 12, 'bold'),
                      width=25,
                      pady=5).pack(side="right", padx=10)

                    def export_to_contact():
                        abs_batches = [b for b in self.attendance_session_data if b.get("op") == "غياب"]
                        if not abs_batches:
                            return messagebox.showwarning("تنبيه", "لا يوجد سجلات غياب في القائمة الحالية لتصديرها.")
                        elif self.m.save_pending_contacts(abs_batches):
                            messagebox.showinfo("تم", "تم تصدير قائمة الغياب إلى مسؤول التواصل بنجاح! 📨")
                        else:
                            messagebox.showerror("خطأ", "فشل التصدير. تأكد من صلاحيات الملفات.")

                    tk.Button(f_footer, text="📤 تصدير لمسؤول التواصل", command=export_to_contact, bg="#7b1fa2",
                      fg="white",
                      font=('Segoe UI', 12, 'bold'),
                      width=25,
                      pady=5).pack(side="left", padx=20)
                    f_mini_sess = tk.LabelFrame((self.f_main_container), text="ملخص الجلسة", bg="white")
                    f_mini_sess.pack(side="bottom", fill="x", pady=5)
                    lbl_sess_count = tk.Label(f_mini_sess, text="عدد الطلاب المرصودين في الجلسة: 0", bg="white", font=('Segoe UI',
                                                                                                                       11))
                    lbl_sess_count.pack(pady=5)

                    def refresh_session_table_abs():
                        total = 0
                        for b in self.attendance_session_data:
                            total += len(b.get("names", []))
                        else:
                            lbl_sess_count.config(text=f"عدد الطلاب المرصودين في الجلسة: {total} (جاهز للإرسال)")

                    refresh_session_table_abs()
                    load_students()
                    tree_abs.pack(fill="both", expand=True)

                    def refresh_abs_table():
                        for i in tree_abs.get_children():
                            tree_abs.delete(i)
                        else:
                            for idx, batch in enumerate(self.attendance_session_data):
                                t = batch.get("time", "-")
                                reason = batch.get("reason", "")
                                op = batch.get("op", "")
                                grade = batch.get("grade", "")
                                sec = batch.get("section", "")
                                for name in batch.get("names", []):
                                    tree_abs.insert("", "end", iid=f"abs_{idx}_{name}", values=(t, name, grade, sec, op, reason))

                    menu_abs = tk.Menu(tree_abs, tearoff=0)

                    def delete_abs_item():
                        sel = tree_abs.selection()
                        if not sel:
                            return
                        for s in sel:
                            idx = int(s.split("_")[1])
                            if 0 <= idx < len(self.attendance_session_data):
                                del self.attendance_session_data[idx]
                            refresh_abs_table()
                            update_prog_display()

                    menu_abs.add_command(label="حذف من الجلسة ❌", command=delete_abs_item)
                    tree_abs.bind("<Button-3>", lambda e: menu_abs.post(e.x_root, e.y_root))

                    def save_class_to_session():
                        g = cb_grade.get()
                        s = cb_section.get()
                        return g and s or messagebox.showwarning("خطأ", "حدد الصف والشعبة")
                        selected_names = [name for name, var in self.abs_toggle_vars.items() if var.get()]
                        att = self.current_attachment_path
                        if not att:
                            if selected_names:
                                return messagebox.showwarning("تنبيه", "يجب رفع ورقة الرصد لوجود غياب (من أسفل الصفحة)")
                        self.attendance_session_data.append({'time':(datetime.now().strftime)("%H:%M"), 
                         'op':"غياب", 
                         'grade':g,  'section':s,  'names':selected_names if selected_names else ["جميع طلاب الشعبة (حضور)"], 
                         'attachment':att, 
                         'reason':"غياب مسجل" if selected_names else "حضور كامل"})
                        self.current_attachment_path = ""
                        update_prog_display()
                        refresh_abs_table()
                        messagebox.showinfo("تم", f"تم حفظ غياب الشعبة {s} في الجلسة.")
                        for w in scrollable_frame.winfo_children():
                            w.destroy()
                        else:
                            self.abs_toggle_vars = {}

                    f_sel = tk.LabelFrame(content, text="1. اختيار الفصل والشعبة", font=('Segoe UI',
                                                                                         11,
                                                                                         'bold'), bg=COLOR_BG, padx=10, pady=10)
                    f_sel.pack(fill="x", pady=5)
                    tk.Button(f_sel, text="حفظ الفصل للجلسة 💾", command=save_class_to_session, bg="#4caf50", fg="white", font=('Segoe UI',
                                                                                                                               11,
                                                                                                                               'bold')).pack(side="left", padx=10)
                    grades = sorted(self.m.list_students_simple()["الصف"].unique().tolist())
                    sections_var = tk.StringVar()
                    tk.Label(f_sel, text="الصف:", bg=COLOR_BG).pack(side="right", padx=5)
                    cb_grade = ttk.Combobox(f_sel, values=grades, state="readonly", width=15)
                    cb_grade.pack(side="right", padx=5)
                    tk.Label(f_sel, text="الشعبة:", bg=COLOR_BG).pack(side="right", padx=5)
                    cb_section = ttk.Combobox(f_sel, state="readonly", width=10)
                    cb_section.pack(side="right", padx=5)

                    def on_grade_change(e):
                        g = cb_grade.get()
                        all_s = self.m.list_students_simple()
                        try:
                            secs = sorted(all_s[all_s["الصف"].astype(str).str.strip() == g.strip()]["الشعبة"].astype(str).str.strip().unique().tolist())
                        except:
                            secs = []
                        else:
                            cb_section["values"] = secs
                            cb_section.set("")
                            if secs:
                                cb_section.current(0)
                                load_grid()

                    cb_grade.bind("<<ComboboxSelected>>", on_grade_change)
                    cb_section.bind("<<ComboboxSelected>>", lambda e: load_grid())
                    f_grid_container = tk.LabelFrame(content, text="2. تحديد الغياب (اضغط على الاسم)", font=('Segoe UI',
                                                                                                             11,
                                                                                                             'bold'), bg=COLOR_BG, padx=10, pady=10)
                    f_grid_container.pack(side="top", fill="both", expand=True, pady=5)
                    f_search_in_class = tk.Frame(f_grid_container, bg="white")
                    f_search_in_class.pack(fill="x", pady=(0, 5))
                    tk.Label(f_search_in_class, text="بحث سريع في الفصل:", bg="white").pack(side="right", padx=5)
                    sv_class_filter = tk.StringVar()
                    ent_class_filter = tk.Entry(f_search_in_class, textvariable=sv_class_filter, width=30)
                    ent_class_filter.pack(side="right", padx=5)

                    def load_grid():
                        g = str(cb_grade.get()).strip()
                        s = str(cb_section.get()).strip()
                        if not g:
                            return
                        df = self.m.list_students_simple()
                        df["_g"] = df["الصف"].astype(str).str.strip()
                        df["_s"] = df["الشعبة"].astype(str).str.strip()
                        self.current_grid_students = df[(df["_g"] == g) & (df["_s"] == s)]
                        render_grid()

                    def render_grid(ft=''):
                        for w in scrollable_frame.winfo_children():
                            w.destroy()
                        else:
                            if self.current_grid_students.empty:
                                tk.Label(scrollable_frame, text="لا يوجد طلاب في هذه الشعبة", bg="white", fg="red").pack(pady=20)
                                return
                            show_df = self.current_grid_students
                            if ft:
                                show_df = show_df[show_df["الاسم"].astype(str).str.lower().str.contains(ft)]
                            r, c = (0, 0)
                            for _, row in show_df.sort_values(by="الاسم").iterrows():
                                name = str(row["الاسم"])
                                if name not in self.abs_toggle_vars:
                                    self.abs_toggle_vars[name] = tk.BooleanVar(value=False)
                                v = self.abs_toggle_vars[name]
                                f = tk.Frame(scrollable_frame, bg="#f5f5f5", relief="raised", bd=1, padx=5, pady=5)
                                f.grid(row=r, column=c, padx=5, pady=5, sticky="ew")
                                lbl = tk.Label(f, text=name, bg="#f5f5f5", font=('Segoe UI',
                                                                                 10), width=18, wraplength=130)
                                lbl.pack()
                                if v.get():
                                    f.config(bg="#ffcdd2")
                                    lbl.config(bg="#ffcdd2")

                                def tgl(e, _v=v, _f=f, _l=lbl):
                                    nv = not _v.get()
                                    _v.set(nv)
                                    _f.config(bg=("#ffcdd2" if nv else "#f5f5f5"))
                                    _l.config(bg=("#ffcdd2" if nv else "#f5f5f5"))

                                f.bind("<Button-1>", tgl)
                                lbl.bind("<Button-1>", tgl)
                                c += 1
                                if c >= 4:
                                    c = 0
                                    r += 1

                    refresh_abs_table()
                    ent_class_filter.bind("<KeyRelease>", lambda e: render_grid(sv_class_filter.get().strip().lower()))
                    tk.Button(f_search_in_class, text="بحث 🔍", command=(lambda: render_grid(sv_class_filter.get().strip().lower())), bg="#00897b", fg="white", font=('Segoe UI',
                                                                                                                                                                     9)).pack(side="right", padx=5)


            def show_contact_view():
                import json
                for w in self.f_main_container.winfo_children():
                    w.destroy()
                else:
                    f_head = tk.Frame((self.f_main_container), bg="white", pady=10)
                    f_head.pack(fill="x")
                    tk.Label(f_head, text="📞 التواصل مع أولياء أمور الغائبين", font=('Segoe UI',
                                                                                     14,
                                                                                     'bold'), fg="#7b1fa2", bg="white").pack(side="left", padx=20)
                    f_filters = tk.Frame(f_head, bg="white")
                    f_filters.pack(side="right", padx=20)
                    tk.Label(f_filters, text="تصفية للصف:", bg="white").pack(side="right", padx=5)
                    cb_cls = ttk.Combobox(f_filters, values=(["الكل"] + self.m.get_available_classes()), state="readonly", width=12)
                    cb_cls.set("الكل")
                    cb_cls.pack(side="right", padx=5)
                    tk.Label(f_filters, text="الشعبة:", bg="white").pack(side="right", padx=5)
                    cb_sec = ttk.Combobox(f_filters, values=(["الكل"] + self.m.get_available_sections()), state="readonly", width=8)
                    cb_sec.set("الكل")
                    cb_sec.pack(side="right", padx=5)
                    today_str = datetime.now().strftime("%Y-%m-%d")
                    f_list = tk.Frame((self.f_main_container), bg="white")
                    f_list.pack(fill="both", expand=True, padx=10, pady=10)
                    cols = ('name', 'grade', 'section', 'mobile', 'home', 'work', 'status',
                            'notes')
                    tree_con = ttk.Treeview(f_list, columns=cols, show="headings", height=15)
                    tree_con.heading("name", text="اسم الطالب")
                    tree_con.column("name", width=180, anchor="e")
                    tree_con.heading("grade", text="الصف")
                    tree_con.column("grade", width=70, anchor="center")
                    tree_con.heading("section", text="الشعبة")
                    tree_con.column("section", width=60, anchor="center")
                    tree_con.heading("mobile", text="جوال ولي الأمر")
                    tree_con.column("mobile", width=100, anchor="center")
                    tree_con.heading("home", text="المنزل")
                    tree_con.column("home", width=100, anchor="center")
                    tree_con.heading("work", text="العمل")
                    tree_con.column("work", width=100, anchor="center")
                    tree_con.heading("status", text="حالة التواصل")
                    tree_con.column("status", width=120, anchor="center")
                    tree_con.heading("notes", text="ملاحظات")
                    tree_con.column("notes", width=200, anchor="e")
                    sb = ttk.Scrollbar(f_list, orient="vertical", command=(tree_con.yview))
                    tree_con.configure(yscrollcommand=(sb.set))
                    sb.pack(side="left", fill="y")
                    tree_con.pack(side="right", fill="both", expand=True)
                    f_ctrl = tk.Frame((self.f_main_container), bg=COLOR_BG, pady=10)
                    f_ctrl.pack(fill="x")

                    def refresh_contacts(*args):
                        selected = tree_con.selection()
                        saved_ids = []
                        if selected:
                            for i in selected:
                                saved_ids.append(tree_con.item(i, "values")[0])

                        for i in tree_con.get_children():
                            tree_con.delete(i)
                        else:
                            data = self.m.load_pending_contacts(today_str)
                            c_val = cb_cls.get()
                            s_val = cb_sec.get()
                            filtered_data = []
                            for row in data:
                                if c_val != "الكل" and str(row.get("grade", "")).strip() != c_val:
                                    pass
                                elif s_val != "الكل" and str(row.get("section", "")).strip() != s_val:
                                    pass
                                else:
                                    filtered_data.append(row)
                            else:
                                if not filtered_data:
                                    if c_val == "الكل":
                                        if s_val == "الكل":
                                            if not data:
                                                tk.Label(f_list, text="لا توجد بيانات مرحلة للتواصل اليوم.\n(تأكد من قيام مسؤول الغياب بـ 'تصدير' القائمة)", bg="white", fg="gray").place(relx=0.5, rely=0.5, anchor="center")
                                    return
                                for w in f_list.winfo_children():
                                    if isinstance(w, tk.Label):
                                        w.destroy()
                                else:
                                    for row in filtered_data:
                                        st = row.get("status", "pending")
                                        st_txt = "⏳ قيد الانتظار"
                                        tags = ()
                                        if st == "done":
                                            st_txt = "✅ تم التواصل"
                                            tags = ('done', )
                                        else:
                                            if st == "no_answer":
                                                st_txt = "❌ لم يرد"
                                                tags = ('no_answer', )
                                        item_id = tree_con.insert("", "end", values=(
                                         row.get("name"), row.get("grade"), row.get("section"),
                                         row.get("mobile"), row.get("home_phone"), row.get("work_phone"),
                                         st_txt, row.get("notes")),
                                          tags=tags)
                                        if row.get("name") in saved_ids:
                                            tree_con.selection_add(item_id)

                    tree_con.tag_configure("done", foreground="green")
                    tree_con.tag_configure("no_answer", foreground="red")
                    cb_cls.bind("<<ComboboxSelected>>", refresh_contacts)
                    cb_sec.bind("<<ComboboxSelected>>", refresh_contacts)
                    refresh_contacts()

                    def set_status_for_selected(new_status):
                        sel = tree_con.selection()
                        if not sel:
                            return messagebox.showwarning("تنبيه", "اختر طالباً من القائمة أولاً")
                        success_count = 0
                        for item_id in sel:
                            name = tree_con.item(item_id, "values")[0]
                            self.m.update_contact_status(name, today_str, new_status)
                            success_count += 1
                        else:
                            if success_count > 0:
                                refresh_contacts()

                    f_actions = tk.LabelFrame(f_ctrl, text="إجراءات سريعة (للطالب المحدد)", font=('Segoe UI',
                                                                                                  10,
                                                                                                  'bold'), bg=COLOR_BG, padx=10, pady=5)
                    f_actions.pack(side="right", padx=20)
                    tk.Button(f_actions, text="✅ تم التواصل", command=(lambda: set_status_for_selected("done")), bg="#1976d2",
                      fg="white",
                      font=('Segoe UI', 10, 'bold'),
                      width=15).pack(side="right", padx=5)
                    tk.Button(f_actions, text="❌ لم يرد", command=(lambda: set_status_for_selected("no_answer")), bg="#d32f2f",
                      fg="white",
                      font=('Segoe UI', 10, 'bold'),
                      width=15).pack(side="right", padx=5)

                    def on_dbl_click(event):
                        sel = tree_con.selection()
                        if not sel:
                            return
                        item = tree_con.item(sel[0])
                        vals = item["values"]
                        s_name = vals[0]
                        cur_note = vals[7]
                        top = tk.Toplevel(self.f_main_container)
                        top.title(f"بيانات وتحديث: {s_name}")
                        top.geometry("450x400")
                        tk.Label(top, text=f"الطالب: {s_name}", font=('Segoe UI', 12, 'bold')).pack(pady=10)
                        f_nums = tk.LabelFrame(top, text="أرقام التواصل", font=('Segoe UI',
                                                                                10))
                        f_nums.pack(fill="x", padx=20, pady=5)
                        tk.Label(f_nums, text=f"📱 {vals[3]} | ☎️ {vals[4]} | 🏢 {vals[5]}", font=('Segoe UI',
                                                                                                 10)).pack(pady=5)
                        tk.Label(top, text="ملاحظات إضافية:", font=('Segoe UI', 10, 'bold')).pack(pady=5)
                        e_note = tk.Entry(top, width=50)
                        e_note.pack(pady=5)
                        e_note.insert(0, cur_note)
                        e_note.focus_set()

                        def save_only_note():
                            nt = e_note.get()
                            record = next((r for r in self.m.load_pending_contacts(today_str) if r["name"] == s_name), None)
                            curr_st = record.get("status", "pending") if record else "pending"
                            self.m.update_contact_status(s_name, today_str, curr_st, nt)
                            refresh_contacts()
                            top.destroy()

                        tk.Button(top, text="حفظ الملاحظة", command=save_only_note, bg=COLOR_ACCENT, fg="white", width=20).pack(pady=10)

                    tree_con.bind("<Double-1>", on_dbl_click)

                    def submit_to_manager():
                        if not tree_con.get_children():
                            return
                            if not messagebox.askyesno("تأكيد", "هل تريد اعتماد نتائج التواصل وإرسالها للمدير؟"):
                                return
                        else:
                            try:
                                data = self.m.load_pending_contacts(today_str)
                                emp_name = getattr(self, "current_employee", "Unknown")
                                log_record = {'date':today_str, 
                                 'employee':emp_name, 
                                 'details':data, 
                                 'summary':f"إنجاز تواصل ({len(data)}) طلاب"}
                                self.m.save_contact_log(log_record)
                                self.m.update_task_status("التواصل مع ولي الأمر الطالب الغائب", today_str, self.m.STATUS_FINISHED)
                                messagebox.showinfo("تم", "تم حفظ سجل الإنجاز وإرساله للمدير بنجاح! ✅\n(تم فصل السجل عن ملف الغياب)")
                            except Exception as e:
                                try:
                                    messagebox.showerror("خطأ", f"فشل التحديث: {e}")
                                finally:
                                    pass

                    tk.Button(f_ctrl, text="📤 اعتماد وإرسال للمدير", command=submit_to_manager, bg="#2e7d32", fg="white", font=('Segoe UI',
                                                                                                                                11,
                                                                                                                                'bold')).pack(side="left", padx=20)

                    def interim_save():
                        try:
                            current_data = self.m.load_pending_contacts(today_str)
                            emp_name = getattr(self, "current_employee", "Unknown")
                            log_record = {'date':today_str, 
                             'employee':emp_name, 
                             'details':current_data, 
                             'summary':f"جاري العمل ({len(current_data)}) طلاب"}
                            self.m.save_contact_log(log_record)
                            self.m.update_task_status("التواصل مع ولي الأمر الطالب الغائب", today_str, "جاري")
                            messagebox.showinfo("حفظ", "تم حفظ العمل وتحديث حالة الإنجاز لدى المدير (جاري العمل ⏳).")
                        except Exception as e:
                            try:
                                messagebox.showerror("خطأ", f"فشل الحفظ: {e}")
                            finally:
                                pass

                    tk.Button(f_ctrl, text="💾 حفظ وتحديث المدير", command=interim_save, bg="#fb8c00", fg="white", font=('Segoe UI',
                                                                                                                        10)).pack(side="left", padx=10)
                    tk.Label(f_ctrl, text="💡 حدد الطالب ثم اضغط الأزرار على اليمين", bg=COLOR_BG, fg="#555").pack(side="right", padx=10)


            def on_vis(e):
                self.refresh_attendance_ops()


            page.bind("<Visibility>", on_vis)
            page.after(100, lambda: self.refresh_attendance_ops())
            return page

        # ========================================

    # --- page_attendance_main_update_prog_display ---
    def page_attendance_main_update_prog_display():
        # ========================================
        self.lbl_sess_prog.config(text=f"العمليات المنجزة في الجلسة: {len(self.attendance_session_data)}")

        # ========================================

    # --- page_attendance_main_upload_attachment ---
    def page_attendance_main_upload_attachment():
        # ========================================
        p = filedialog.askopenfilename(title="اختر المرفق",
          filetypes=[
         ('All Supported', '*.jpg *.jpeg *.png *.pdf *.docx *.xlsx *.doc *.xls'), 
         ('Images', '*.jpg *.jpeg *.png'), 
         ('PDF Files', '*.pdf'), 
         ('Word Files', '*.docx *.doc'), 
         ('Excel Files', '*.xlsx *.xls')])
        if p:
            self.current_attachment_path = p
            messagebox.showinfo("تم", f"تم اختيار المرفق: {os.path.basename(p)}")

        # ========================================

    # --- page_attendance_main_finish_session ---
    def page_attendance_main_finish_session():
        # ========================================
        if not self.attendance_session_data:
            return messagebox.showinfo("تنبيه", "لم يتم تسجيل أي بيانات في هذه الجلسة بعد.")
            if not messagebox.askyesno("تصدير", f"هل تريد تصدير وإرسال {len(self.attendance_session_data)} مجموعة سجلات للمدير؟"):
                return
            count_ok = 0
            now_date = datetime.now().strftime("%Y-%m-%d")
            session_time = datetime.now().strftime("%H:%M:%S")
            session_attachment_for_sync = ""
            for batch in self.attendance_session_data:
                local_att = batch.get("attachment", "")

            if not local_att:
                local_att = getattr(self, "current_attachment_path", "")
        else:
            final_att = ""
            if local_att:
                if os.path.exists(local_att):
                    dest = data_path("مرفقات_السجلات")
                    if not os.path.exists(dest):
                        os.makedirs(dest)
                    ext = os.path.splitext(local_att)[1]
                    fname = f'{batch["op"]}_{datetime.now().strftime("%Y%m%d_%H%M%S")}_{count_ok}{ext}'
                    final_att = os.path.join(dest, fname)
                    try:
                        import shutil
                        shutil.copy2(local_att, final_att)
                        session_attachment_for_sync = final_att
                    except:
                        pass

        for std in batch["names"]:
            rec = {'التاريخ':now_date,  'الوقت':session_time,  'نوع العملية':batch["op"], 
             'اسم الطالب':std,  'الصف':batch.get("grade", ""), 
             'الشعبة':batch.get("section", ""),  'رابط المرفق':final_att, 
             'الحالة التفصيلية':batch.get("reason", ""),  'اسم الموظف المنفذ':getattr(self, "current_employee", "Unknown"), 
             'حالة الاعتماد':"معلق"}
            if self.m.add_attendance_record(rec):
                count_ok += 1
        else:
            if self.attendance_session_data:
                unique_ops = list(set((batch.get("op") for batch in self.attendance_session_data if batch.get("op"))))
                for op in unique_ops:
                    self.m.update_task_status(op, now_date, self.m.STATUS_WORKING, session_attachment_for_sync)

            self.attendance_session_data.clear()
            self.current_attachment_path = ""
            update_prog_display()
            messagebox.showinfo("تم", f"تم إرسال {count_ok} سجل للمدير بنجاح.")

        # ========================================

    # --- page_attendance_main_set_attendance_op ---
    def page_attendance_main_set_attendance_op(op=None, color=None):
        # ========================================
        if hasattr(self, "current_roles"):
            if op not in self.current_roles:
                return
        else:
            self.op_var.set(op)
            for btn in self.op_btns:
                if btn.cget("text") == op:
                    btn.config(bg=color, relief="sunken")
                elif hasattr(self, "current_roles") and btn.cget("text") in self.current_roles:
                    btn.config(bg="#e0e0e0", relief="raised")
                else:
                    btn.config(bg="#ccc", relief="flat")
            else:
                self.current_attachment_path = ""
                if op == "غياب":
                    show_absence_view()
                else:
                    if op == "التواصل مع ولي الأمر الطالب الغائب":
                        show_contact_view()
                    else:
                        show_standard_view()

        # ========================================

    # --- page_attendance_main_show_standard_view ---
    def page_attendance_main_show_standard_view():
        # ========================================
        for w in self.f_main_container.winfo_children():
            w.destroy()
        else:
            f_top = tk.LabelFrame((self.f_main_container), text="1. البحث والاختيار (تصفية فورية)", font=('Segoe UI',
                                                                                                          11,
                                                                                                          'bold'), bg=COLOR_BG, padx=10, pady=10)
            f_top.pack(side="top", fill="both", expand=True, pady=(0, 10))
            f_filters = tk.Frame(f_top, bg=COLOR_BG)
            f_filters.pack(fill="x", pady=5)
            tk.Label(f_filters, text="الاسم:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
            sv_name = tk.StringVar()
            ent_name = tk.Entry(f_filters, textvariable=sv_name, width=25, font=('Segoe UI',
                                                                                 10))
            ent_name.pack(side="right", padx=5)
            tk.Label(f_filters, text="الصف:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
            cb_class = ttk.Combobox(f_filters, values=([""] + self.m.get_available_classes()), state="readonly", width=15)
            cb_class.pack(side="right", padx=5)
            tk.Label(f_filters, text="الشعبة:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
            cb_section = ttk.Combobox(f_filters, values=([""] + self.m.get_available_sections()), state="readonly", width=10)
            cb_section.pack(side="right", padx=5)

            def clear_search():
                sv_name.set("")
                cb_class.set("")
                cb_section.set("")
                refresh_results()


            tk.Button(f_filters, text="مسح", command=clear_search, bg="#bf360c", fg="white", font=('Segoe UI',
                                                                                                   9)).pack(side="right", padx=10)
            columns = ('name', 'grade', 'section')
            tree_res = ttk.Treeview(f_top, columns=columns, show="headings", height=8)
            tree_res.heading("name", text="اسم الطالب")
            tree_res.column("name", anchor="e", width=200)
            tree_res.heading("grade", text="الصف")
            tree_res.column("grade", anchor="center", width=100)
            tree_res.heading("section", text="الشعبة")
            tree_res.column("section", anchor="center", width=80)
            sb_res = ttk.Scrollbar(f_top, orient="vertical", command=(tree_res.yview))
            tree_res.configure(yscrollcommand=(sb_res.set))
            sb_res.pack(side="left", fill="y")
            tree_res.pack(side="right", fill="both", expand=True)

            def refresh_results(*args):
                tree_res.delete(*tree_res.get_children())
                q_n = sv_name.get().strip()
                q_c = cb_class.get()
                q_s = cb_section.get()
                df = self.m.list_students_simple()
                if q_c:
                    df = df[df["الصف"].astype(str).str.strip() == q_c]
                if q_s:
                    df = df[df["الشعبة"].astype(str).str.strip() == q_s]
                if q_n:
                    try:
                        q_norm = normalize_arabic(q_n)
                        mask = df["الاسم"].astype(str).apply(normalize_arabic).str.contains(q_norm, na=False)
                        df = df[mask]
                    except:
                        df = df[df["الاسم"].astype(str).str.contains(q_n, na=False)]

                if len(df) > 1000:
                    df = df.head(1000)
                for idx, row in df.iterrows():
                    n = str(row.get("الاسم", ""))
                    g = str(row.get("الصف", ""))
                    s = str(row.get("الشعبة", ""))
                    tree_res.insert("", "end", iid=(str(idx)), values=(n, g, s))


            ent_name.bind("<KeyRelease>", lambda e: refresh_results())
            cb_class.bind("<<ComboboxSelected>>", lambda e: refresh_results())
            cb_section.bind("<<ComboboxSelected>>", lambda e: refresh_results())
            refresh_results()
            f_action = tk.Frame((self.f_main_container), bg=COLOR_BG, pady=5)
            f_action.pack(side="top", fill="x")
            tk.Label(f_action, text="السبب / الملاحظة:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
            cb_reason = ttk.Combobox(f_action, values=['بدون عذر', 'بعذر طبي', 'ظرف عائلي', 'تأخر صباحي', 'تأخر فسحة'], state="readonly", width=20)
            cb_reason.pack(side="right", padx=5)
            cb_reason.current(0)

            def add_selected_student():
                sel = tree_res.selection()
                if not sel:
                    return messagebox.showwarning("تنبيه", "الرجاء اختيار طالب من القائمة العلوية")
                op_title = self.op_var.get()
                count_added = 0
                for iid in sel:
                    item = tree_res.item(iid)
                    vals = item["values"]
                    s_name = vals[0]
                    s_grade = vals[1]
                    s_sec = vals[2]
                    batch = {'time':(datetime.now().strftime)("%H:%M:%S"), 
                     'op':op_title, 
                     'names':[
                      s_name], 
                     'grade':s_grade, 
                     'section':s_sec, 
                     'attachment':self.current_attachment_path, 
                     'reason':cb_reason.get()}
                    self.attendance_session_data.append(batch)
                    count_added += 1
                else:
                    self.current_attachment_path = ""
                    update_prog_display()
                    refresh_session_table()
                    tree_res.selection_remove(tree_res.selection())
                    ent_name.delete(0, tk.END)
                    ent_name.focus()


            tk.Button(f_action, text="⬇️ إضافة للقائمة (الحفظ المؤقت)", command=add_selected_student, bg="#43a047", fg="white", font=('Segoe UI',
                                                                                                                                      11,
                                                                                                                                      'bold')).pack(side="right", padx=20)
            tree_res.bind("<Double-1>", lambda e: add_selected_student())
            f_bottom = tk.LabelFrame((self.f_main_container), text="2. القائمة المحفوظة (جاهزة للإرسال)", font=('Segoe UI',
                                                                                                                11,
                                                                                                                'bold'), bg="white", padx=10, pady=10)
            f_bottom.pack(side="bottom", fill="both", expand=True)
            tree_session = ttk.Treeview(f_bottom, columns=('time', 'name', 'grade', 'op', 'reason'), show="headings", height=8)
            tree_session.heading("time", text="التوقيت")
            tree_session.column("time", width=80, anchor="center")
            tree_session.heading("name", text="اسم الطالب")
            tree_session.column("name", width=180, anchor="e")
            tree_session.heading("grade", text="الصف")
            tree_session.column("grade", width=90, anchor="center")
            tree_session.heading("op", text="العملية")
            tree_session.column("op", width=80, anchor="center")
            tree_session.heading("reason", text="السبب")
            tree_session.column("reason", width=150, anchor="e")
            sb_sess = ttk.Scrollbar(f_bottom, orient="vertical", command=(tree_session.yview))
            tree_session.configure(yscrollcommand=(sb_sess.set))
            sb_sess.pack(side="left", fill="y")
            tree_session.pack(side="right", fill="both", expand=True)

            def refresh_session_table():
                for i in tree_session.get_children():
                    tree_session.delete(i)
                else:
                    for idx, batch in enumerate(self.attendance_session_data):
                        t = batch.get("time", "-")
                        reason = batch.get("reason", "")
                        op = batch.get("op", "")
                        grade = batch.get("grade", "")
                        for name in batch.get("names", []):
                            tree_session.insert("", "end", iid=f"item_{idx}_{name}", values=(t, name, grade, op, reason))


            menu = tk.Menu(tree_session, tearoff=0)

            def delete_selected_session():
                sel = tree_session.selection()
                if not sel:
                    return
                indices_to_remove = []
                for s in sel:
                    id_parts = s.split("_")
                    if len(id_parts) >= 2:
                        indices_to_remove.append(int(id_parts[1]))
                    for idx in sorted((list(set(indices_to_remove))), reverse=True):
                        if 0 <= idx < len(self.attendance_session_data):
                            del self.attendance_session_data[idx]
                        refresh_session_table()
                        update_prog_display()


            menu.add_command(label="حذف من القائمة ❌", command=delete_selected_session)
            tree_session.bind("<Button-3>", lambda e: menu.post(e.x_root, e.y_root))
            refresh_session_table()

        # ========================================

    # --- page_attendance_main_show_standard_view_clear_search ---
    def page_attendance_main_show_standard_view_clear_search():
        # ========================================
        sv_name.set("")
        cb_class.set("")
        cb_section.set("")
        refresh_results()

        # ========================================

    # --- page_attendance_main_show_standard_view_refresh_results ---
    def page_attendance_main_show_standard_view_refresh_results(*args):
        # ========================================
        tree_res.delete(*tree_res.get_children())
        q_n = sv_name.get().strip()
        q_c = cb_class.get()
        q_s = cb_section.get()
        df = self.m.list_students_simple()
        if q_c:
            df = df[df["الصف"].astype(str).str.strip() == q_c]
        if q_s:
            df = df[df["الشعبة"].astype(str).str.strip() == q_s]
        if q_n:
            try:
                q_norm = normalize_arabic(q_n)
                mask = df["الاسم"].astype(str).apply(normalize_arabic).str.contains(q_norm, na=False)
                df = df[mask]
            except:
                df = df[df["الاسم"].astype(str).str.contains(q_n, na=False)]

        if len(df) > 1000:
            df = df.head(1000)
        for idx, row in df.iterrows():
            n = str(row.get("الاسم", ""))
            g = str(row.get("الصف", ""))
            s = str(row.get("الشعبة", ""))
            tree_res.insert("", "end", iid=(str(idx)), values=(n, g, s))

        # ========================================

    # --- page_attendance_main_show_standard_view_add_selected_student ---
    def page_attendance_main_show_standard_view_add_selected_student():
        # ========================================
        sel = tree_res.selection()
        if not sel:
            return messagebox.showwarning("تنبيه", "الرجاء اختيار طالب من القائمة العلوية")
        op_title = self.op_var.get()
        count_added = 0
        for iid in sel:
            item = tree_res.item(iid)
            vals = item["values"]
            s_name = vals[0]
            s_grade = vals[1]
            s_sec = vals[2]
            batch = {'time':(datetime.now().strftime)("%H:%M:%S"), 
             'op':op_title, 
             'names':[
              s_name], 
             'grade':s_grade, 
             'section':s_sec, 
             'attachment':self.current_attachment_path, 
             'reason':cb_reason.get()}
            self.attendance_session_data.append(batch)
            count_added += 1
        else:
            self.current_attachment_path = ""
            update_prog_display()
            refresh_session_table()
            tree_res.selection_remove(tree_res.selection())
            ent_name.delete(0, tk.END)
            ent_name.focus()

        # ========================================

    # --- page_attendance_main_show_standard_view_refresh_session_table ---
    def page_attendance_main_show_standard_view_refresh_session_table():
        # ========================================
        for i in tree_session.get_children():
            tree_session.delete(i)
        else:
            for idx, batch in enumerate(self.attendance_session_data):
                t = batch.get("time", "-")
                reason = batch.get("reason", "")
                op = batch.get("op", "")
                grade = batch.get("grade", "")
                for name in batch.get("names", []):
                    tree_session.insert("", "end", iid=f"item_{idx}_{name}", values=(t, name, grade, op, reason))

        # ========================================

    # --- page_attendance_main_show_standard_view_delete_selected_session ---
    def page_attendance_main_show_standard_view_delete_selected_session():
        # ========================================
        sel = tree_session.selection()
        if not sel:
            return
        indices_to_remove = []
        for s in sel:
            id_parts = s.split("_")
            if len(id_parts) >= 2:
                indices_to_remove.append(int(id_parts[1]))
            for idx in sorted((list(set(indices_to_remove))), reverse=True):
                if 0 <= idx < len(self.attendance_session_data):
                    del self.attendance_session_data[idx]
                refresh_session_table()
                update_prog_display()

        # ========================================

    # --- page_attendance_main_show_absence_view ---
    def page_attendance_main_show_absence_view():
        # ========================================
        for w in self.f_main_container.winfo_children():
            w.destroy()
        else:
            f_header = tk.LabelFrame((self.f_main_container), text="1. اختيار الفصل والشعبة (رصد الغياب الجماعي)", font=('Segoe UI',
                                                                                                                         11,
                                                                                                                         'bold'), bg=COLOR_BG, padx=10, pady=10)
            f_header.pack(side="top", fill="x", pady=(0, 10))
            tk.Label(f_header, text="الصف:", bg=COLOR_BG, font=('Segoe UI', 11)).pack(side="right", padx=5)
            cb_class = ttk.Combobox(f_header, values=([""] + self.m.get_available_classes()), state="readonly", width=15, font=('Segoe UI',
                                                                                                                                10))
            cb_class.pack(side="right", padx=5)
            tk.Label(f_header, text="الشعبة:", bg=COLOR_BG, font=('Segoe UI', 11)).pack(side="right", padx=5)
            cb_section = ttk.Combobox(f_header, values=([""] + self.m.get_available_sections()), state="readonly", width=10, font=('Segoe UI',
                                                                                                                                   10))
            cb_section.pack(side="right", padx=5)
            tk.Label(f_header, text=" |  سبب الغياب:", bg=COLOR_BG, font=('Segoe UI', 11)).pack(side="right", padx=5)
            cb_reason_batch = ttk.Combobox(f_header, values=["بدون عذر", "بعذر طبي", "ظرف عائلي"], state="readonly", width=15)
            cb_reason_batch.pack(side="right", padx=5)
            cb_reason_batch.current(0)
            f_list_container = tk.Frame((self.f_main_container), bg="white", highlightthickness=1, highlightbackground="#ccc")
            f_list_container.pack(side="top", fill="both", expand=True, pady=5)
            canvas = tk.Canvas(f_list_container, bg="white")
            scrollbar = ttk.Scrollbar(f_list_container, orient="vertical", command=(canvas.yview))
            scrollable_frame = tk.Frame(canvas, bg="white")
            scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=(canvas.bbox("all"))))
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=(scrollbar.set))
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            self.abs_checks = {}

            def load_students(*args):
                for widget in scrollable_frame.winfo_children():
                    widget.destroy()
                else:
                    self.abs_checks = {}
                    c_val = cb_class.get()
                    s_val = cb_section.get()
                    if not c_val:
                        tk.Label(scrollable_frame, text="الرجاء اختيار الصف (والشعبة إن وجدت) لعرض القائمة", bg="white", fg="#757575", font=('Segoe UI',
                                                                                                                                             12)).pack(pady=20, padx=20)
                        return
                    df = self.m.list_students_simple()
                    if c_val:
                        df = df[df["الصف"].astype(str).str.strip() == c_val]
                    if s_val:
                        df = df[df["الشعبة"].astype(str).str.strip() == s_val]
                    if len(df) > 300:
                        tk.Label(scrollable_frame, text=f"العدد كبير ({len(df)} طالب). الرجاء تحديد الشعبة.", bg="white", fg="red").pack()
                        df = df.head(100)
                    h_row = tk.Frame(scrollable_frame, bg="#e0f2f1")
                    h_row.pack(fill="x", pady=2)
                    tk.Label(h_row, text="تحديد الغياب ⭕", width=15, bg="#e0f2f1", font=('Segoe UI',
                                                                                         10,
                                                                                         'bold')).pack(side="right")
                    tk.Label(h_row, text="اسم الطالب", width=40, anchor="e", bg="#e0f2f1", font=('Segoe UI',
                                                                                                 10,
                                                                                                 'bold')).pack(side="right")
                    for idx, row in df.iterrows():
                        name = str(row.get("الاسم", ""))
                        row_f = tk.Frame(scrollable_frame, bg="white")
                        row_f.pack(fill="x", pady=1, padx=5)
                        tk.Label(row_f, text=name, anchor="e", width=40, bg="white", font=('Segoe UI',
                                                                                           11)).pack(side="right", padx=5)
                        var = tk.IntVar()
                        self.abs_checks[name] = var
                        cb = tk.Checkbutton(row_f, variable=var, bg="white", activebackground="white", selectcolor="#ffcdd2",
                          cursor="hand2")
                        cb.pack(side="right", padx=20)

                        def toggle(v=var):
                            v.set(1 - v.get())

                        row_f.bind("<Button-1>", lambda e, v=var: toggle(v))


            cb_class.bind("<<ComboboxSelected>>", load_students)
            cb_section.bind("<<ComboboxSelected>>", load_students)
            f_footer = tk.Frame((self.f_main_container), bg=COLOR_BG, pady=10)
            f_footer.pack(side="bottom", fill="x")

            def save_absent_batch():
                absent_names = [name for name, var in self.abs_checks.items() if var.get() == 1]
                if not absent_names:
                    return messagebox.showwarning("تنبيه", "لم يتم تحديد أي طالب غائب!")
                batch = {'time':(datetime.now().strftime)("%H:%M:%S"), 
                 'op':"غياب", 
                 'names':absent_names, 
                 'grade':cb_class.get(), 
                 'section':cb_section.get(), 
                 'attachment':self.current_attachment_path, 
                 'reason':cb_reason_batch.get()}
                self.attendance_session_data.append(batch)
                self.current_attachment_path = ""
                messagebox.showinfo("تم", f"تم رصد غياب {len(absent_names)} طالب بنجاح وإضافتهم للقائمة.")
                for v in self.abs_checks.values():
                    v.set(0)
                else:
                    update_prog_display()
                    refresh_session_table_abs()


            tk.Button(f_footer, text="💾 حفظ الغياب لهذا الفصل", command=save_absent_batch, bg="#d32f2f",
              fg="white",
              font=('Segoe UI', 12, 'bold'),
              width=25,
              pady=5).pack(side="right", padx=10)

            def export_to_contact():
                abs_batches = [b for b in self.attendance_session_data if b.get("op") == "غياب"]
                if not abs_batches:
                    return messagebox.showwarning("تنبيه", "لا يوجد سجلات غياب في القائمة الحالية لتصديرها.")
                elif self.m.save_pending_contacts(abs_batches):
                    messagebox.showinfo("تم", "تم تصدير قائمة الغياب إلى مسؤول التواصل بنجاح! 📨")
                else:
                    messagebox.showerror("خطأ", "فشل التصدير. تأكد من صلاحيات الملفات.")


            tk.Button(f_footer, text="📤 تصدير لمسؤول التواصل", command=export_to_contact, bg="#7b1fa2",
              fg="white",
              font=('Segoe UI', 12, 'bold'),
              width=25,
              pady=5).pack(side="left", padx=20)
            f_mini_sess = tk.LabelFrame((self.f_main_container), text="ملخص الجلسة", bg="white")
            f_mini_sess.pack(side="bottom", fill="x", pady=5)
            lbl_sess_count = tk.Label(f_mini_sess, text="عدد الطلاب المرصودين في الجلسة: 0", bg="white", font=('Segoe UI',
                                                                                                               11))
            lbl_sess_count.pack(pady=5)

            def refresh_session_table_abs():
                total = 0
                for b in self.attendance_session_data:
                    total += len(b.get("names", []))
                else:
                    lbl_sess_count.config(text=f"عدد الطلاب المرصودين في الجلسة: {total} (جاهز للإرسال)")


            refresh_session_table_abs()
            load_students()
            tree_abs.pack(fill="both", expand=True)

            def refresh_abs_table():
                for i in tree_abs.get_children():
                    tree_abs.delete(i)
                else:
                    for idx, batch in enumerate(self.attendance_session_data):
                        t = batch.get("time", "-")
                        reason = batch.get("reason", "")
                        op = batch.get("op", "")
                        grade = batch.get("grade", "")
                        sec = batch.get("section", "")
                        for name in batch.get("names", []):
                            tree_abs.insert("", "end", iid=f"abs_{idx}_{name}", values=(t, name, grade, sec, op, reason))


            menu_abs = tk.Menu(tree_abs, tearoff=0)

            def delete_abs_item():
                sel = tree_abs.selection()
                if not sel:
                    return
                for s in sel:
                    idx = int(s.split("_")[1])
                    if 0 <= idx < len(self.attendance_session_data):
                        del self.attendance_session_data[idx]
                    refresh_abs_table()
                    update_prog_display()


            menu_abs.add_command(label="حذف من الجلسة ❌", command=delete_abs_item)
            tree_abs.bind("<Button-3>", lambda e: menu_abs.post(e.x_root, e.y_root))

            def save_class_to_session():
                g = cb_grade.get()
                s = cb_section.get()
                return g and s or messagebox.showwarning("خطأ", "حدد الصف والشعبة")
                selected_names = [name for name, var in self.abs_toggle_vars.items() if var.get()]
                att = self.current_attachment_path
                if not att:
                    if selected_names:
                        return messagebox.showwarning("تنبيه", "يجب رفع ورقة الرصد لوجود غياب (من أسفل الصفحة)")
                self.attendance_session_data.append({'time':(datetime.now().strftime)("%H:%M"), 
                 'op':"غياب", 
                 'grade':g,  'section':s,  'names':selected_names if selected_names else ["جميع طلاب الشعبة (حضور)"], 
                 'attachment':att, 
                 'reason':"غياب مسجل" if selected_names else "حضور كامل"})
                self.current_attachment_path = ""
                update_prog_display()
                refresh_abs_table()
                messagebox.showinfo("تم", f"تم حفظ غياب الشعبة {s} في الجلسة.")
                for w in scrollable_frame.winfo_children():
                    w.destroy()
                else:
                    self.abs_toggle_vars = {}


            f_sel = tk.LabelFrame(content, text="1. اختيار الفصل والشعبة", font=('Segoe UI',
                                                                                 11, 'bold'), bg=COLOR_BG, padx=10, pady=10)
            f_sel.pack(fill="x", pady=5)
            tk.Button(f_sel, text="حفظ الفصل للجلسة 💾", command=save_class_to_session, bg="#4caf50", fg="white", font=('Segoe UI',
                                                                                                                       11,
                                                                                                                       'bold')).pack(side="left", padx=10)
            grades = sorted(self.m.list_students_simple()["الصف"].unique().tolist())
            sections_var = tk.StringVar()
            tk.Label(f_sel, text="الصف:", bg=COLOR_BG).pack(side="right", padx=5)
            cb_grade = ttk.Combobox(f_sel, values=grades, state="readonly", width=15)
            cb_grade.pack(side="right", padx=5)
            tk.Label(f_sel, text="الشعبة:", bg=COLOR_BG).pack(side="right", padx=5)
            cb_section = ttk.Combobox(f_sel, state="readonly", width=10)
            cb_section.pack(side="right", padx=5)

            def on_grade_change(e):
                g = cb_grade.get()
                all_s = self.m.list_students_simple()
                try:
                    secs = sorted(all_s[all_s["الصف"].astype(str).str.strip() == g.strip()]["الشعبة"].astype(str).str.strip().unique().tolist())
                except:
                    secs = []
                else:
                    cb_section["values"] = secs
                    cb_section.set("")
                    if secs:
                        cb_section.current(0)
                        load_grid()


            cb_grade.bind("<<ComboboxSelected>>", on_grade_change)
            cb_section.bind("<<ComboboxSelected>>", lambda e: load_grid())
            f_grid_container = tk.LabelFrame(content, text="2. تحديد الغياب (اضغط على الاسم)", font=('Segoe UI',
                                                                                                     11,
                                                                                                     'bold'), bg=COLOR_BG, padx=10, pady=10)
            f_grid_container.pack(side="top", fill="both", expand=True, pady=5)
            f_search_in_class = tk.Frame(f_grid_container, bg="white")
            f_search_in_class.pack(fill="x", pady=(0, 5))
            tk.Label(f_search_in_class, text="بحث سريع في الفصل:", bg="white").pack(side="right", padx=5)
            sv_class_filter = tk.StringVar()
            ent_class_filter = tk.Entry(f_search_in_class, textvariable=sv_class_filter, width=30)
            ent_class_filter.pack(side="right", padx=5)

            def load_grid():
                g = str(cb_grade.get()).strip()
                s = str(cb_section.get()).strip()
                if not g:
                    return
                df = self.m.list_students_simple()
                df["_g"] = df["الصف"].astype(str).str.strip()
                df["_s"] = df["الشعبة"].astype(str).str.strip()
                self.current_grid_students = df[(df["_g"] == g) & (df["_s"] == s)]
                render_grid()


            def render_grid(ft=''):
                for w in scrollable_frame.winfo_children():
                    w.destroy()
                else:
                    if self.current_grid_students.empty:
                        tk.Label(scrollable_frame, text="لا يوجد طلاب في هذه الشعبة", bg="white", fg="red").pack(pady=20)
                        return
                    show_df = self.current_grid_students
                    if ft:
                        show_df = show_df[show_df["الاسم"].astype(str).str.lower().str.contains(ft)]
                    r, c = (0, 0)
                    for _, row in show_df.sort_values(by="الاسم").iterrows():
                        name = str(row["الاسم"])
                        if name not in self.abs_toggle_vars:
                            self.abs_toggle_vars[name] = tk.BooleanVar(value=False)
                        v = self.abs_toggle_vars[name]
                        f = tk.Frame(scrollable_frame, bg="#f5f5f5", relief="raised", bd=1, padx=5, pady=5)
                        f.grid(row=r, column=c, padx=5, pady=5, sticky="ew")
                        lbl = tk.Label(f, text=name, bg="#f5f5f5", font=('Segoe UI', 10), width=18, wraplength=130)
                        lbl.pack()
                        if v.get():
                            f.config(bg="#ffcdd2")
                            lbl.config(bg="#ffcdd2")

                        def tgl(e, _v=v, _f=f, _l=lbl):
                            nv = not _v.get()
                            _v.set(nv)
                            _f.config(bg=("#ffcdd2" if nv else "#f5f5f5"))
                            _l.config(bg=("#ffcdd2" if nv else "#f5f5f5"))

                        f.bind("<Button-1>", tgl)
                        lbl.bind("<Button-1>", tgl)
                        c += 1
                        if c >= 4:
                            c = 0
                            r += 1


            refresh_abs_table()
            ent_class_filter.bind("<KeyRelease>", lambda e: render_grid(sv_class_filter.get().strip().lower()))
            tk.Button(f_search_in_class, text="بحث 🔍", command=(lambda: render_grid(sv_class_filter.get().strip().lower())), bg="#00897b", fg="white", font=('Segoe UI',
                                                                                                                                                             9)).pack(side="right", padx=5)

        # ========================================

    # --- page_attendance_main_show_absence_view_load_students ---
    def page_attendance_main_show_absence_view_load_students(*args):
        # ========================================
        for widget in scrollable_frame.winfo_children():
            widget.destroy()
        else:
            self.abs_checks = {}
            c_val = cb_class.get()
            s_val = cb_section.get()
            if not c_val:
                tk.Label(scrollable_frame, text="الرجاء اختيار الصف (والشعبة إن وجدت) لعرض القائمة", bg="white", fg="#757575", font=('Segoe UI',
                                                                                                                                     12)).pack(pady=20, padx=20)
                return
            df = self.m.list_students_simple()
            if c_val:
                df = df[df["الصف"].astype(str).str.strip() == c_val]
            if s_val:
                df = df[df["الشعبة"].astype(str).str.strip() == s_val]
            if len(df) > 300:
                tk.Label(scrollable_frame, text=f"العدد كبير ({len(df)} طالب). الرجاء تحديد الشعبة.", bg="white", fg="red").pack()
                df = df.head(100)
            h_row = tk.Frame(scrollable_frame, bg="#e0f2f1")
            h_row.pack(fill="x", pady=2)
            tk.Label(h_row, text="تحديد الغياب ⭕", width=15, bg="#e0f2f1", font=('Segoe UI',
                                                                                 10, 'bold')).pack(side="right")
            tk.Label(h_row, text="اسم الطالب", width=40, anchor="e", bg="#e0f2f1", font=('Segoe UI',
                                                                                         10,
                                                                                         'bold')).pack(side="right")
            for idx, row in df.iterrows():
                name = str(row.get("الاسم", ""))
                row_f = tk.Frame(scrollable_frame, bg="white")
                row_f.pack(fill="x", pady=1, padx=5)
                tk.Label(row_f, text=name, anchor="e", width=40, bg="white", font=('Segoe UI',
                                                                                   11)).pack(side="right", padx=5)
                var = tk.IntVar()
                self.abs_checks[name] = var
                cb = tk.Checkbutton(row_f, variable=var, bg="white", activebackground="white", selectcolor="#ffcdd2",
                  cursor="hand2")
                cb.pack(side="right", padx=20)

                def toggle(v=var):
                    v.set(1 - v.get())


                row_f.bind("<Button-1>", lambda e, v=var: toggle(v))

        # ========================================

    # --- page_attendance_main_show_absence_view_load_students_toggle ---
    def page_attendance_main_show_absence_view_load_students_toggle(v=None):
        # ========================================
        v.set(1 - v.get())

        # ========================================

    # --- page_attendance_main_show_absence_view_save_absent_batch ---
    def page_attendance_main_show_absence_view_save_absent_batch():
        # ========================================
        absent_names = [name for name, var in self.abs_checks.items() if var.get() == 1]
        if not absent_names:
            return messagebox.showwarning("تنبيه", "لم يتم تحديد أي طالب غائب!")
        batch = {'time':(datetime.now().strftime)("%H:%M:%S"), 
         'op':"غياب", 
         'names':absent_names, 
         'grade':cb_class.get(), 
         'section':cb_section.get(), 
         'attachment':self.current_attachment_path, 
         'reason':cb_reason_batch.get()}
        self.attendance_session_data.append(batch)
        self.current_attachment_path = ""
        messagebox.showinfo("تم", f"تم رصد غياب {len(absent_names)} طالب بنجاح وإضافتهم للقائمة.")
        for v in self.abs_checks.values():
            v.set(0)
        else:
            update_prog_display()
            refresh_session_table_abs()

        # ========================================

    # --- page_attendance_main_show_absence_view_export_to_contact ---
    def page_attendance_main_show_absence_view_export_to_contact():
        # ========================================
        abs_batches = [b for b in self.attendance_session_data if b.get("op") == "غياب"]
        if not abs_batches:
            return messagebox.showwarning("تنبيه", "لا يوجد سجلات غياب في القائمة الحالية لتصديرها.")
        elif self.m.save_pending_contacts(abs_batches):
            messagebox.showinfo("تم", "تم تصدير قائمة الغياب إلى مسؤول التواصل بنجاح! 📨")
        else:
            messagebox.showerror("خطأ", "فشل التصدير. تأكد من صلاحيات الملفات.")

        # ========================================

    # --- page_attendance_main_show_absence_view_refresh_session_table_abs ---
    def page_attendance_main_show_absence_view_refresh_session_table_abs():
        # ========================================
        total = 0
        for b in self.attendance_session_data:
            total += len(b.get("names", []))
        else:
            lbl_sess_count.config(text=f"عدد الطلاب المرصودين في الجلسة: {total} (جاهز للإرسال)")

        # ========================================

    # --- page_attendance_main_show_absence_view_refresh_abs_table ---
    def page_attendance_main_show_absence_view_refresh_abs_table():
        # ========================================
        for i in tree_abs.get_children():
            tree_abs.delete(i)
        else:
            for idx, batch in enumerate(self.attendance_session_data):
                t = batch.get("time", "-")
                reason = batch.get("reason", "")
                op = batch.get("op", "")
                grade = batch.get("grade", "")
                sec = batch.get("section", "")
                for name in batch.get("names", []):
                    tree_abs.insert("", "end", iid=f"abs_{idx}_{name}", values=(t, name, grade, sec, op, reason))

        # ========================================

    # --- page_attendance_main_show_absence_view_delete_abs_item ---
    def page_attendance_main_show_absence_view_delete_abs_item():
        # ========================================
        sel = tree_abs.selection()
        if not sel:
            return
        for s in sel:
            idx = int(s.split("_")[1])
            if 0 <= idx < len(self.attendance_session_data):
                del self.attendance_session_data[idx]
            refresh_abs_table()
            update_prog_display()

        # ========================================

    # --- page_attendance_main_show_absence_view_save_class_to_session ---
    def page_attendance_main_show_absence_view_save_class_to_session():
        # ========================================
        g = cb_grade.get()
        s = cb_section.get()
        return g and s or messagebox.showwarning("خطأ", "حدد الصف والشعبة")
        selected_names = [name for name, var in self.abs_toggle_vars.items() if var.get()]
        att = self.current_attachment_path
        if not att:
            if selected_names:
                return messagebox.showwarning("تنبيه", "يجب رفع ورقة الرصد لوجود غياب (من أسفل الصفحة)")
        self.attendance_session_data.append({'time':(datetime.now().strftime)("%H:%M"), 
         'op':"غياب", 
         'grade':g,  'section':s,  'names':selected_names if selected_names else ["جميع طلاب الشعبة (حضور)"], 
         'attachment':att, 
         'reason':"غياب مسجل" if selected_names else "حضور كامل"})
        self.current_attachment_path = ""
        update_prog_display()
        refresh_abs_table()
        messagebox.showinfo("تم", f"تم حفظ غياب الشعبة {s} في الجلسة.")
        for w in scrollable_frame.winfo_children():
            w.destroy()
        else:
            self.abs_toggle_vars = {}

        # ========================================

    # --- page_attendance_main_show_absence_view_on_grade_change ---
    def page_attendance_main_show_absence_view_on_grade_change(e=None):
        # ========================================
        g = cb_grade.get()
        all_s = self.m.list_students_simple()
        try:
            secs = sorted(all_s[all_s["الصف"].astype(str).str.strip() == g.strip()]["الشعبة"].astype(str).str.strip().unique().tolist())
        except:
            secs = []
        else:
            cb_section["values"] = secs
            cb_section.set("")
            if secs:
                cb_section.current(0)
                load_grid()

        # ========================================

    # --- page_attendance_main_show_absence_view_load_grid ---
    def page_attendance_main_show_absence_view_load_grid():
        # ========================================
        g = str(cb_grade.get()).strip()
        s = str(cb_section.get()).strip()
        if not g:
            return
        df = self.m.list_students_simple()
        df["_g"] = df["الصف"].astype(str).str.strip()
        df["_s"] = df["الشعبة"].astype(str).str.strip()
        self.current_grid_students = df[(df["_g"] == g) & (df["_s"] == s)]
        render_grid()

        # ========================================

    # --- page_attendance_main_show_absence_view_render_grid ---
    def page_attendance_main_show_absence_view_render_grid(ft=None):
        # ========================================
        for w in scrollable_frame.winfo_children():
            w.destroy()
        else:
            if self.current_grid_students.empty:
                tk.Label(scrollable_frame, text="لا يوجد طلاب في هذه الشعبة", bg="white", fg="red").pack(pady=20)
                return
            show_df = self.current_grid_students
            if ft:
                show_df = show_df[show_df["الاسم"].astype(str).str.lower().str.contains(ft)]
            r, c = (0, 0)
            for _, row in show_df.sort_values(by="الاسم").iterrows():
                name = str(row["الاسم"])
                if name not in self.abs_toggle_vars:
                    self.abs_toggle_vars[name] = tk.BooleanVar(value=False)
                v = self.abs_toggle_vars[name]
                f = tk.Frame(scrollable_frame, bg="#f5f5f5", relief="raised", bd=1, padx=5, pady=5)
                f.grid(row=r, column=c, padx=5, pady=5, sticky="ew")
                lbl = tk.Label(f, text=name, bg="#f5f5f5", font=('Segoe UI', 10), width=18, wraplength=130)
                lbl.pack()
                if v.get():
                    f.config(bg="#ffcdd2")
                    lbl.config(bg="#ffcdd2")

                def tgl(e, _v=v, _f=f, _l=lbl):
                    nv = not _v.get()
                    _v.set(nv)
                    _f.config(bg=("#ffcdd2" if nv else "#f5f5f5"))
                    _l.config(bg=("#ffcdd2" if nv else "#f5f5f5"))


                f.bind("<Button-1>", tgl)
                lbl.bind("<Button-1>", tgl)
                c += 1
                if c >= 4:
                    c = 0
                    r += 1

        # ========================================

    # --- page_attendance_main_show_absence_view_render_grid_tgl ---
    def page_attendance_main_show_absence_view_render_grid_tgl(e=None, _v=None, _f=None, _l=None):
        # ========================================
        nv = not _v.get()
        _v.set(nv)
        _f.config(bg=("#ffcdd2" if nv else "#f5f5f5"))
        _l.config(bg=("#ffcdd2" if nv else "#f5f5f5"))

        # ========================================

    # --- page_attendance_main_show_contact_view ---
    def page_attendance_main_show_contact_view():
        # ========================================
        import json
        for w in self.f_main_container.winfo_children():
            w.destroy()
        else:
            f_head = tk.Frame((self.f_main_container), bg="white", pady=10)
            f_head.pack(fill="x")
            tk.Label(f_head, text="📞 التواصل مع أولياء أمور الغائبين", font=('Segoe UI', 14,
                                                                             'bold'), fg="#7b1fa2", bg="white").pack(side="left", padx=20)
            f_filters = tk.Frame(f_head, bg="white")
            f_filters.pack(side="right", padx=20)
            tk.Label(f_filters, text="تصفية للصف:", bg="white").pack(side="right", padx=5)
            cb_cls = ttk.Combobox(f_filters, values=(["الكل"] + self.m.get_available_classes()), state="readonly", width=12)
            cb_cls.set("الكل")
            cb_cls.pack(side="right", padx=5)
            tk.Label(f_filters, text="الشعبة:", bg="white").pack(side="right", padx=5)
            cb_sec = ttk.Combobox(f_filters, values=(["الكل"] + self.m.get_available_sections()), state="readonly", width=8)
            cb_sec.set("الكل")
            cb_sec.pack(side="right", padx=5)
            today_str = datetime.now().strftime("%Y-%m-%d")
            f_list = tk.Frame((self.f_main_container), bg="white")
            f_list.pack(fill="both", expand=True, padx=10, pady=10)
            cols = ('name', 'grade', 'section', 'mobile', 'home', 'work', 'status', 'notes')
            tree_con = ttk.Treeview(f_list, columns=cols, show="headings", height=15)
            tree_con.heading("name", text="اسم الطالب")
            tree_con.column("name", width=180, anchor="e")
            tree_con.heading("grade", text="الصف")
            tree_con.column("grade", width=70, anchor="center")
            tree_con.heading("section", text="الشعبة")
            tree_con.column("section", width=60, anchor="center")
            tree_con.heading("mobile", text="جوال ولي الأمر")
            tree_con.column("mobile", width=100, anchor="center")
            tree_con.heading("home", text="المنزل")
            tree_con.column("home", width=100, anchor="center")
            tree_con.heading("work", text="العمل")
            tree_con.column("work", width=100, anchor="center")
            tree_con.heading("status", text="حالة التواصل")
            tree_con.column("status", width=120, anchor="center")
            tree_con.heading("notes", text="ملاحظات")
            tree_con.column("notes", width=200, anchor="e")
            sb = ttk.Scrollbar(f_list, orient="vertical", command=(tree_con.yview))
            tree_con.configure(yscrollcommand=(sb.set))
            sb.pack(side="left", fill="y")
            tree_con.pack(side="right", fill="both", expand=True)
            f_ctrl = tk.Frame((self.f_main_container), bg=COLOR_BG, pady=10)
            f_ctrl.pack(fill="x")

            def refresh_contacts(*args):
                selected = tree_con.selection()
                saved_ids = []
                if selected:
                    for i in selected:
                        saved_ids.append(tree_con.item(i, "values")[0])

                for i in tree_con.get_children():
                    tree_con.delete(i)
                else:
                    data = self.m.load_pending_contacts(today_str)
                    c_val = cb_cls.get()
                    s_val = cb_sec.get()
                    filtered_data = []
                    for row in data:
                        if c_val != "الكل" and str(row.get("grade", "")).strip() != c_val:
                            pass
                        elif s_val != "الكل" and str(row.get("section", "")).strip() != s_val:
                            pass
                        else:
                            filtered_data.append(row)
                    else:
                        if not filtered_data:
                            if c_val == "الكل":
                                if s_val == "الكل":
                                    if not data:
                                        tk.Label(f_list, text="لا توجد بيانات مرحلة للتواصل اليوم.\n(تأكد من قيام مسؤول الغياب بـ 'تصدير' القائمة)", bg="white", fg="gray").place(relx=0.5, rely=0.5, anchor="center")
                            return
                        for w in f_list.winfo_children():
                            if isinstance(w, tk.Label):
                                w.destroy()
                        else:
                            for row in filtered_data:
                                st = row.get("status", "pending")
                                st_txt = "⏳ قيد الانتظار"
                                tags = ()
                                if st == "done":
                                    st_txt = "✅ تم التواصل"
                                    tags = ('done', )
                                else:
                                    if st == "no_answer":
                                        st_txt = "❌ لم يرد"
                                        tags = ('no_answer', )
                                item_id = tree_con.insert("", "end", values=(
                                 row.get("name"), row.get("grade"), row.get("section"),
                                 row.get("mobile"), row.get("home_phone"), row.get("work_phone"),
                                 st_txt, row.get("notes")),
                                  tags=tags)
                                if row.get("name") in saved_ids:
                                    tree_con.selection_add(item_id)


            tree_con.tag_configure("done", foreground="green")
            tree_con.tag_configure("no_answer", foreground="red")
            cb_cls.bind("<<ComboboxSelected>>", refresh_contacts)
            cb_sec.bind("<<ComboboxSelected>>", refresh_contacts)
            refresh_contacts()

            def set_status_for_selected(new_status):
                sel = tree_con.selection()
                if not sel:
                    return messagebox.showwarning("تنبيه", "اختر طالباً من القائمة أولاً")
                success_count = 0
                for item_id in sel:
                    name = tree_con.item(item_id, "values")[0]
                    self.m.update_contact_status(name, today_str, new_status)
                    success_count += 1
                else:
                    if success_count > 0:
                        refresh_contacts()


            f_actions = tk.LabelFrame(f_ctrl, text="إجراءات سريعة (للطالب المحدد)", font=('Segoe UI',
                                                                                          10,
                                                                                          'bold'), bg=COLOR_BG, padx=10, pady=5)
            f_actions.pack(side="right", padx=20)
            tk.Button(f_actions, text="✅ تم التواصل", command=(lambda: set_status_for_selected("done")), bg="#1976d2",
              fg="white",
              font=('Segoe UI', 10, 'bold'),
              width=15).pack(side="right", padx=5)
            tk.Button(f_actions, text="❌ لم يرد", command=(lambda: set_status_for_selected("no_answer")), bg="#d32f2f",
              fg="white",
              font=('Segoe UI', 10, 'bold'),
              width=15).pack(side="right", padx=5)

            def on_dbl_click(event):
                sel = tree_con.selection()
                if not sel:
                    return
                item = tree_con.item(sel[0])
                vals = item["values"]
                s_name = vals[0]
                cur_note = vals[7]
                top = tk.Toplevel(self.f_main_container)
                top.title(f"بيانات وتحديث: {s_name}")
                top.geometry("450x400")
                tk.Label(top, text=f"الطالب: {s_name}", font=('Segoe UI', 12, 'bold')).pack(pady=10)
                f_nums = tk.LabelFrame(top, text="أرقام التواصل", font=('Segoe UI', 10))
                f_nums.pack(fill="x", padx=20, pady=5)
                tk.Label(f_nums, text=f"📱 {vals[3]} | ☎️ {vals[4]} | 🏢 {vals[5]}", font=('Segoe UI',
                                                                                         10)).pack(pady=5)
                tk.Label(top, text="ملاحظات إضافية:", font=('Segoe UI', 10, 'bold')).pack(pady=5)
                e_note = tk.Entry(top, width=50)
                e_note.pack(pady=5)
                e_note.insert(0, cur_note)
                e_note.focus_set()

                def save_only_note():
                    nt = e_note.get()
                    record = next((r for r in self.m.load_pending_contacts(today_str) if r["name"] == s_name), None)
                    curr_st = record.get("status", "pending") if record else "pending"
                    self.m.update_contact_status(s_name, today_str, curr_st, nt)
                    refresh_contacts()
                    top.destroy()

                tk.Button(top, text="حفظ الملاحظة", command=save_only_note, bg=COLOR_ACCENT, fg="white", width=20).pack(pady=10)


            tree_con.bind("<Double-1>", on_dbl_click)

            def submit_to_manager():
                if not tree_con.get_children():
                    return
                    if not messagebox.askyesno("تأكيد", "هل تريد اعتماد نتائج التواصل وإرسالها للمدير؟"):
                        return
                else:
                    try:
                        data = self.m.load_pending_contacts(today_str)
                        emp_name = getattr(self, "current_employee", "Unknown")
                        log_record = {'date':today_str, 
                         'employee':emp_name, 
                         'details':data, 
                         'summary':f"إنجاز تواصل ({len(data)}) طلاب"}
                        self.m.save_contact_log(log_record)
                        self.m.update_task_status("التواصل مع ولي الأمر الطالب الغائب", today_str, self.m.STATUS_FINISHED)
                        messagebox.showinfo("تم", "تم حفظ سجل الإنجاز وإرساله للمدير بنجاح! ✅\n(تم فصل السجل عن ملف الغياب)")
                    except Exception as e:
                        try:
                            messagebox.showerror("خطأ", f"فشل التحديث: {e}")
                        finally:
                            pass


            tk.Button(f_ctrl, text="📤 اعتماد وإرسال للمدير", command=submit_to_manager, bg="#2e7d32", fg="white", font=('Segoe UI',
                                                                                                                        11,
                                                                                                                        'bold')).pack(side="left", padx=20)

            def interim_save():
                try:
                    current_data = self.m.load_pending_contacts(today_str)
                    emp_name = getattr(self, "current_employee", "Unknown")
                    log_record = {'date':today_str, 
                     'employee':emp_name, 
                     'details':current_data, 
                     'summary':f"جاري العمل ({len(current_data)}) طلاب"}
                    self.m.save_contact_log(log_record)
                    self.m.update_task_status("التواصل مع ولي الأمر الطالب الغائب", today_str, "جاري")
                    messagebox.showinfo("حفظ", "تم حفظ العمل وتحديث حالة الإنجاز لدى المدير (جاري العمل ⏳).")
                except Exception as e:
                    try:
                        messagebox.showerror("خطأ", f"فشل الحفظ: {e}")
                    finally:
                        pass


            tk.Button(f_ctrl, text="💾 حفظ وتحديث المدير", command=interim_save, bg="#fb8c00", fg="white", font=('Segoe UI',
                                                                                                                10)).pack(side="left", padx=10)
            tk.Label(f_ctrl, text="💡 حدد الطالب ثم اضغط الأزرار على اليمين", bg=COLOR_BG, fg="#555").pack(side="right", padx=10)

        # ========================================

    # --- page_attendance_main_show_contact_view_refresh_contacts ---
    def page_attendance_main_show_contact_view_refresh_contacts(*args):
        # ========================================
        selected = tree_con.selection()
        saved_ids = []
        if selected:
            for i in selected:
                saved_ids.append(tree_con.item(i, "values")[0])

        for i in tree_con.get_children():
            tree_con.delete(i)
        else:
            data = self.m.load_pending_contacts(today_str)
            c_val = cb_cls.get()
            s_val = cb_sec.get()
            filtered_data = []
            for row in data:
                if c_val != "الكل" and str(row.get("grade", "")).strip() != c_val:
                    pass
                elif s_val != "الكل" and str(row.get("section", "")).strip() != s_val:
                    pass
                else:
                    filtered_data.append(row)
            else:
                if not filtered_data:
                    if c_val == "الكل":
                        if s_val == "الكل":
                            if not data:
                                tk.Label(f_list, text="لا توجد بيانات مرحلة للتواصل اليوم.\n(تأكد من قيام مسؤول الغياب بـ 'تصدير' القائمة)", bg="white", fg="gray").place(relx=0.5, rely=0.5, anchor="center")
                    return
                for w in f_list.winfo_children():
                    if isinstance(w, tk.Label):
                        w.destroy()
                else:
                    for row in filtered_data:
                        st = row.get("status", "pending")
                        st_txt = "⏳ قيد الانتظار"
                        tags = ()
                        if st == "done":
                            st_txt = "✅ تم التواصل"
                            tags = ('done', )
                        else:
                            if st == "no_answer":
                                st_txt = "❌ لم يرد"
                                tags = ('no_answer', )
                        item_id = tree_con.insert("", "end", values=(
                         row.get("name"), row.get("grade"), row.get("section"),
                         row.get("mobile"), row.get("home_phone"), row.get("work_phone"),
                         st_txt, row.get("notes")),
                          tags=tags)
                        if row.get("name") in saved_ids:
                            tree_con.selection_add(item_id)

        # ========================================

    # --- page_attendance_main_show_contact_view_set_status_for_selected ---
    def page_attendance_main_show_contact_view_set_status_for_selected(new_status=None):
        # ========================================
        sel = tree_con.selection()
        if not sel:
            return messagebox.showwarning("تنبيه", "اختر طالباً من القائمة أولاً")
        success_count = 0
        for item_id in sel:
            name = tree_con.item(item_id, "values")[0]
            self.m.update_contact_status(name, today_str, new_status)
            success_count += 1
        else:
            if success_count > 0:
                refresh_contacts()

        # ========================================

    # --- page_attendance_main_show_contact_view_on_dbl_click ---
    def page_attendance_main_show_contact_view_on_dbl_click(event=None):
        # ========================================
        sel = tree_con.selection()
        if not sel:
            return
        item = tree_con.item(sel[0])
        vals = item["values"]
        s_name = vals[0]
        cur_note = vals[7]
        top = tk.Toplevel(self.f_main_container)
        top.title(f"بيانات وتحديث: {s_name}")
        top.geometry("450x400")
        tk.Label(top, text=f"الطالب: {s_name}", font=('Segoe UI', 12, 'bold')).pack(pady=10)
        f_nums = tk.LabelFrame(top, text="أرقام التواصل", font=('Segoe UI', 10))
        f_nums.pack(fill="x", padx=20, pady=5)
        tk.Label(f_nums, text=f"📱 {vals[3]} | ☎️ {vals[4]} | 🏢 {vals[5]}", font=('Segoe UI',
                                                                                 10)).pack(pady=5)
        tk.Label(top, text="ملاحظات إضافية:", font=('Segoe UI', 10, 'bold')).pack(pady=5)
        e_note = tk.Entry(top, width=50)
        e_note.pack(pady=5)
        e_note.insert(0, cur_note)
        e_note.focus_set()

        def save_only_note():
            nt = e_note.get()
            record = next((r for r in self.m.load_pending_contacts(today_str) if r["name"] == s_name), None)
            curr_st = record.get("status", "pending") if record else "pending"
            self.m.update_contact_status(s_name, today_str, curr_st, nt)
            refresh_contacts()
            top.destroy()


        tk.Button(top, text="حفظ الملاحظة", command=save_only_note, bg=COLOR_ACCENT, fg="white", width=20).pack(pady=10)

        # ========================================

    # --- page_attendance_main_show_contact_view_on_dbl_click_save_only_note ---
    def page_attendance_main_show_contact_view_on_dbl_click_save_only_note():
        # ========================================
        nt = e_note.get()
        record = next((r for r in self.m.load_pending_contacts(today_str) if r["name"] == s_name), None)
        curr_st = record.get("status", "pending") if record else "pending"
        self.m.update_contact_status(s_name, today_str, curr_st, nt)
        refresh_contacts()
        top.destroy()

        # ========================================

    # --- page_attendance_main_show_contact_view_submit_to_manager ---
    def page_attendance_main_show_contact_view_submit_to_manager():
        # ========================================
        if not tree_con.get_children():
            return
            if not messagebox.askyesno("تأكيد", "هل تريد اعتماد نتائج التواصل وإرسالها للمدير؟"):
                return
        else:
            try:
                data = self.m.load_pending_contacts(today_str)
                emp_name = getattr(self, "current_employee", "Unknown")
                log_record = {'date':today_str, 
                 'employee':emp_name, 
                 'details':data, 
                 'summary':f"إنجاز تواصل ({len(data)}) طلاب"}
                self.m.save_contact_log(log_record)
                self.m.update_task_status("التواصل مع ولي الأمر الطالب الغائب", today_str, self.m.STATUS_FINISHED)
                messagebox.showinfo("تم", "تم حفظ سجل الإنجاز وإرساله للمدير بنجاح! ✅\n(تم فصل السجل عن ملف الغياب)")
            except Exception as e:
                try:
                    messagebox.showerror("خطأ", f"فشل التحديث: {e}")
                finally:
                    pass

        # ========================================

    # --- page_attendance_main_show_contact_view_interim_save ---
    def page_attendance_main_show_contact_view_interim_save():
        # ========================================
        try:
            current_data = self.m.load_pending_contacts(today_str)
            emp_name = getattr(self, "current_employee", "Unknown")
            log_record = {'date':today_str, 
             'employee':emp_name, 
             'details':current_data, 
             'summary':f"جاري العمل ({len(current_data)}) طلاب"}
            self.m.save_contact_log(log_record)
            self.m.update_task_status("التواصل مع ولي الأمر الطالب الغائب", today_str, "جاري")
            messagebox.showinfo("حفظ", "تم حفظ العمل وتحديث حالة الإنجاز لدى المدير (جاري العمل ⏳).")
        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"فشل الحفظ: {e}")
            finally:
                pass

        # ========================================

    # --- page_attendance_main_on_vis ---
    def page_attendance_main_on_vis(e=None):
        # ========================================
        self.refresh_attendance_ops()

        # ========================================

    # --- page_attendance_main_export_late_students_report ---
    def page_attendance_main_export_late_students_report():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- page_attendance_main_choose_doc_file ---
    def page_attendance_main_choose_doc_file():
        # ========================================
        f = filedialog.askopenfilename(title="اختر ملف التوثيق")
        if f:
            dest_dir = data_path("مرفقات_الحضور_اليومي")
            if not os.path.exists(dest_dir):
                os.makedirs(dest_dir)
            fname = f'Attendance_{datetime.now().strftime("%H%M%S")}_{os.path.basename(f)}'
            dest_path = os.path.join(dest_dir, fname)
            import shutil
            try:
                shutil.copy(f, dest_path)
                self.att_file_path = dest_path
                btn_doc.config(text=f"✅ تم الإرفاق: {os.path.basename(f)}", fg="green")
                print(f"DEBUG: Attachment saved at {dest_path}")
            except Exception as e:
                try:
                    messagebox.showerror("خطأ", f"فشل نسخ الملف: {e}")
                finally:
                    pass

        # ========================================

    # --- page_attendance_main_confirm_all_finished ---
    def page_attendance_main_confirm_all_finished():
        # ========================================
        today_str = datetime.now().strftime("%Y-%m-%d")
        roles_to_update = getattr(self, "current_roles", [])
        print("--- ATTEMPTING TO FINISH DAY ---")
        print(f"DEBUG: Today: {today_str}")
        print(f'DEBUG: Current User: {getattr(self, "current_employee", "Unknown")}')
        print(f"DEBUG: Roles to update: {roles_to_update}")
        print(f'DEBUG: Status String: {getattr(self, "STATUS_FINISHED", "MISSING!")}')
        print(f'DEBUG: Attachment Path: {getattr(self, "att_file_path", "NONE")}')
        export_late_students_report()
        if not getattr(self, "att_file_path", None):
            messagebox.showwarning("توثيق مطلوب", "عفواً، يجب إرفاق ملف لتوثيق إنجاز العمل قبل الإرسال.")
            return
        else:
            return messagebox.askyesno("تأكيد", "هل أنت متأكد من إنهاء رصد الطلاب لليوم واعتماد التوثيق؟") or None
        print(f"DEBUG: Processing final confirmation for {today_str}")
        success_count = 0
        for r in roles_to_update:
            res = self.m.update_task_status(r, today_str, self.STATUS_FINISHED, self.att_file_path)
            print(f"DEBUG: Updating task '{r}' -> Result: {res}")
            if res:
                success_count += 1
        else:
            if success_count > 0:
                print(f"SUCCESS: {success_count} tasks updated successfully.")
                messagebox.showinfo("تم الإرسال", f"تم اعتماد إنجاز ({success_count}) مهام بنجاح وإشعار المدير. ✅")
                self.show_home()
            else:
                print("FAILURE: No tasks were updated. Check roles or file locks.")
                messagebox.showerror("خطأ", "فشل تحديث الحالة في قاعدة البيانات. تأكد من إغلاق أي ملفات مفتوحة.")

        # ========================================

    # --- page_attendance_main_on_select ---
    def page_attendance_main_on_select(evt=None):
        # ========================================
        if not self.lst_att_students.curselection():
            return
        txt = self.lst_att_students.get(self.lst_att_students.curselection()[0])
        self.lbl_selected_std.config(text=(txt.split("|")[0].strip()))

        # ========================================

    # --- page_employee_task_portal ---
    def page_employee_task_portal(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        header = tk.Frame(page, bg="#455a64", padx=20, pady=15)
        header.pack(fill="x")
        tk.Label(header, text="بوابة التكليف والتواصل 📋", font=('Segoe UI', 16, 'bold'), bg="#455a64", fg="white").pack(side="right")
        has_attendance_roles = any(r in ('تأخير', 'غياب', 'استئذان', 'التواصل مع ولي الأمر الطالب الغائب') for r in [])
        if has_attendance_roles:
            tk.Button(header, text="الخروج للمهام الأخرى", command=(lambda: self.show("attendance_main")), bg="#607d8b", fg="white").pack(side="left")
        else:
            tk.Button(header, text="تسجيل خروج (الرئيسية)", command=(self.show_home), bg="#d32f2f", fg="white").pack(side="left")
        content = tk.Frame(page, bg=COLOR_BG, pady=20)
        content.pack(expand=True, fill="both")
        task_data = getattr(self, "current_assignment", {})
        task_text = task_data.get("task", "لا يوجد وصف للمهمة")
        task_id = task_data.get("id")
        f_main = tk.Frame(content, bg=COLOR_BG)
        f_main.pack(expand=True)
        card_inst = tk.LabelFrame(f_main, text="المهمة المطلوبة من الإدارة", font=('Segoe UI',
                                                                                   11, 'bold'), bg="white", padx=20, pady=20, labelanchor="ne")
        card_inst.pack(pady=10, fill="x")
        tk.Label(card_inst, text=task_text, font=('Segoe UI', 13), bg="white", wraplength=600, justify="right").pack()
        card_reply = tk.LabelFrame(f_main, text="تقرير الموظف / رسالة للمدير", font=('Segoe UI',
                                                                                     11,
                                                                                     'bold'), bg="white", padx=20, pady=20, labelanchor="ne")
        card_reply.pack(pady=10, fill="x")
        txt_reply = tk.Text(card_reply, height=5, width=60, font=('Segoe UI', 11))
        txt_reply.tag_configure("right", justify="right")
        txt_reply.tag_add("right", "1.0", "end")
        txt_reply.pack(pady=5)
        tk.Label(card_reply, text="وثّق عملك (اختياري):", bg="white", font=('Segoe UI', 9)).pack(anchor="e")
        att_frame = tk.Frame(card_reply, bg="white")
        att_frame.pack(fill="x", pady=5)
        self.tmp_attachment = tk.StringVar(value="")
        lbl_att_status = tk.Label(att_frame, text="لم يتم إرفاق ملف", font=('Segoe UI', 9), bg="white", fg="#888")
        lbl_att_status.pack(side="right")

        def pick_file():
            f = filedialog.askopenfilename(title="اختر ملف التوثيق", filetypes=[('All Files', '*.*')])
            if f:
                dest_dir = data_path("مرفقات_المهام")
                if not os.path.exists(dest_dir):
                    os.makedirs(dest_dir)
                fname = f"{task_id}_{os.path.basename(f)}"
                dest_path = os.path.join(dest_dir, fname)
                import shutil
                try:
                    shutil.copy(f, dest_path)
                    self.tmp_attachment.set(dest_path)
                    lbl_att_status.config(text=f"تم إرفاق: {os.path.basename(f)} ✅", fg="green")
                except Exception as e:
                    try:
                        messagebox.showerror("خطأ", f"فشل نسخ الملف: {e}")
                    finally:
                        pass


        tk.Button(att_frame, text="📎 إرفاق ملف توثيق", command=pick_file, bg="#eeeeee").pack(side="right", padx=10)

        def mark_done():
            reply_val = txt_reply.get("1.0", tk.END).strip()
            if not self.tmp_attachment.get():
                messagebox.showwarning("توثيق مطلوب", "يرجى إرفاق ملف (تقرير/صورة) يثبت إنجاز المهمة قبل الإرسال.")
                return
            elif self.m.update_assignment_status(task_id, self.STATUS_FINISHED, self.tmp_attachment.get(), reply_val):
                messagebox.showinfo("أحسنت!", "تم رفع إنجاز المهمة والتوثيق للمدير بنجاح. ✅")
                if has_attendance_roles:
                    self.show("attendance_main")
                else:
                    self.show_home()
            else:
                messagebox.showerror("خطأ", "فشل تحديث حالة المهمة")


        tk.Button(f_main, text="✅ تأكيد إنجاز المهمة وإرسال التقرير للمدير", command=mark_done, bg="#2e7d32",
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          pady=12,
          padx=30).pack(pady=20)
        return page

        # ========================================

    # --- page_employee_task_portal_pick_file ---
    def page_employee_task_portal_pick_file():
        # ========================================
        f = filedialog.askopenfilename(title="اختر ملف التوثيق", filetypes=[('All Files', '*.*')])
        if f:
            dest_dir = data_path("مرفقات_المهام")
            if not os.path.exists(dest_dir):
                os.makedirs(dest_dir)
            fname = f"{task_id}_{os.path.basename(f)}"
            dest_path = os.path.join(dest_dir, fname)
            import shutil
            try:
                shutil.copy(f, dest_path)
                self.tmp_attachment.set(dest_path)
                lbl_att_status.config(text=f"تم إرفاق: {os.path.basename(f)} ✅", fg="green")
            except Exception as e:
                try:
                    messagebox.showerror("خطأ", f"فشل نسخ الملف: {e}")
                finally:
                    pass

        # ========================================

    # --- page_employee_task_portal_mark_done ---
    def page_employee_task_portal_mark_done():
        # ========================================
        reply_val = txt_reply.get("1.0", tk.END).strip()
        if not self.tmp_attachment.get():
            messagebox.showwarning("توثيق مطلوب", "يرجى إرفاق ملف (تقرير/صورة) يثبت إنجاز المهمة قبل الإرسال.")
            return
        elif self.m.update_assignment_status(task_id, self.STATUS_FINISHED, self.tmp_attachment.get(), reply_val):
            messagebox.showinfo("أحسنت!", "تم رفع إنجاز المهمة والتوثيق للمدير بنجاح. ✅")
            if has_attendance_roles:
                self.show("attendance_main")
            else:
                self.show_home()
        else:
            messagebox.showerror("خطأ", "فشل تحديث حالة المهمة")

        # ========================================

    # --- refresh_attendance_ui ---
    def refresh_attendance_ui(self):
        # ========================================
        if hasattr(self, "lbl_emp_name"):
            emp = getattr(self, "current_employee", "Unknown")
            self.lbl_emp_name.config(text=f"مرحباً، {emp}")
        else:
            has_any_att_role = False
            if hasattr(self, "op_btns"):
                if hasattr(self, "current_roles"):
                    first_allowed = None
                    for btn in self.op_btns:
                        txt = btn.cget("text")
                        if txt in self.current_roles:
                            btn.config(state="normal", cursor="hand2", bg="#e0e0e0")
                            has_any_att_role = True
                            if not first_allowed:
                                first_allowed = btn
                            else:
                                btn.config(state="disabled", cursor="arrow", bg="#cccccc")
                        elif first_allowed:
                            set_op_ref = next(filter((lambda x: x[0] == first_allowed.cget("text")), [('تأخير', '#ff9800'), ('غياب', '#d32f2f'), ('انصراف', '#2196f3'), ('حضور', '#4caf50')]), (None,
                                                                                                                                                                                                None))
                            if set_op_ref[0]:
                                self.op_var.set(set_op_ref[0])
                                first_allowed.config(bg=(set_op_ref[1]), relief="sunken")

            if hasattr(self, "body_att"):
                if has_any_att_role:
                    self.body_att.pack(fill="both", expand=True, padx=20, pady=10)
                else:
                    self.body_att.pack_forget()
                    if not hasattr(self, "lbl_no_roles"):
                        self.lbl_no_roles = tk.Label((self.body_att.master), text="ليس لديك مهام حضور وانصراف مسندة حالياً.", font=('Segoe UI',
                                                                                                                                    14), bg=COLOR_BG, fg="#777")
                    self.lbl_no_roles.pack(pady=100)
                if first_allowed:
                    first_allowed.invoke()

        # ========================================

    # --- page_attendance_tasks_menu ---
    def page_attendance_tasks_menu(self, parent=None):
        page = tk.Frame(parent, bg=COLOR_BG)
        
        lbl_title = tk.Label(page, text="تقرير وحضور ومهام - القائمة الرئيسية", font=('Segoe UI', 20, 'bold'), bg=COLOR_BG, fg="#1b5e20")
        lbl_title.pack(pady=50)
        
        f_btns = tk.Frame(page, bg=COLOR_BG)
        f_btns.pack(expand=True)
        
        btn_daily = tk.Button(f_btns, text="المهام اليومية للموظف", font=('Segoe UI', 16, 'bold'), 
                              bg="#4CAF50", fg="white", width=30, height=3, cursor="hand2", 
                              command=lambda: self.show("attendance_view"))
        btn_daily.pack(pady=15)
        
        def show_temp_tasks():
            for child in self.container.winfo_children():
                child.pack_forget()
            temp_page = tk.Frame(self.container, bg=COLOR_BG)
            lbl = tk.Label(temp_page, text="صفحة المهام الوقتية (قيد الإنشاء)", font=('Segoe UI', 20, 'bold'), fg="#d32f2f", bg=COLOR_BG)
            lbl.pack(pady=100)
            btn_back = tk.Button(temp_page, text="الرجوع", font=('Segoe UI', 14), bg="#757575", fg="white", cursor="hand2",
                                 command=lambda: self.show("attendance_tasks_menu"))
            btn_back.pack(pady=20)
            self.pages["temp_tasks"] = temp_page
            temp_page.pack(fill="both", expand=True)
            self.current_page = "temp_tasks"
            
        btn_temp = tk.Button(f_btns, text="المهام الوقتية", font=('Segoe UI', 16, 'bold'), 
                             bg="#2196F3", fg="white", width=30, height=3, cursor="hand2", 
                             command=show_temp_tasks)
        btn_temp.pack(pady=15)
        
        return page

    # ========================================

    # --- page_attendance_view ---
    def page_attendance_view(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        tk.Label(page, text="إدارة الحضور والمهام (لوحة المدير)", font=('Segoe UI', 16, 'bold'), bg=COLOR_BG).pack(pady=10)
        notebook = ttk.Notebook(page)
        notebook.pack(fill="both", expand=True, padx=20, pady=10)
        tab_rep = tk.Frame(notebook, bg=COLOR_BG)
        notebook.add(tab_rep, text=" 📊 التقارير والتحليل الذكي ")
        f_filter = tk.LabelFrame(tab_rep, text="🔍 أدوات البحث والفرز", font=('Segoe UI', 11,
                                                                             'bold'), bg=COLOR_BG, padx=10, pady=10)
        f_filter.pack(fill="x", padx=10, pady=5)
        r_a = tk.Frame(f_filter, bg=COLOR_BG)
        r_a.pack(fill="x", pady=5)
        r_a = tk.Frame(f_filter, bg=COLOR_BG)
        r_a.pack(fill="x", pady=5)

        def open_date_picker(entry_widget):
            try:
                import calendar
                top = tk.Toplevel(parent)
                top.title("اختر التاريخ")
                top.geometry("300x300")
                top.configure(bg="white")
                try:
                    x = parent.winfo_rootx() + 100
                    y = parent.winfo_rooty() + 100
                    top.geometry(f"+{x}+{y}")
                except:
                    pass
                else:
                    current_year = tk.IntVar(value=(datetime.now().year))
                    current_month = tk.IntVar(value=(datetime.now().month))

                    def update_cal():
                        for widget in f_days.winfo_children():
                            widget.destroy()
                        else:
                            y = current_year.get()
                            m = current_month.get()
                            lbl_header.config(text=f"{y} - {m}")
                            days = [
                             'Ec', 'الجمعة', 'الخميس', 
                             'الأربعاء', 'الثلاثاء', 'الاثنين', 
                             'الأحد', 'السبت']
                            for d in ('ح', 'ن', 'ث', 'ر', 'خ', 'ج', 'س'):
                                tk.Label(f_days, text=d, bg="#eee", width=5).grid(row=0, column=(['س', 'ج', 'خ', 'ر', 'ث', 'ن', 'ح'].index(d)))
                            else:
                                cal = calendar.monthcalendar(y, m)
                                r = 1
                                for week in cal:
                                    for idx, day in enumerate(week):
                                        if day == 0:
                                            pass
                                        else:

                                            def on_day(d=day):
                                                selected_date = f"{y}-{m:02d}-{d:02d}"
                                                entry_widget.delete(0, tk.END)
                                                entry_widget.insert(0, selected_date)
                                                top.destroy()
                                                apply_filters()

                                for widget in f_days.winfo_children():
                                    widget.destroy()
                                else:
                                    cal_obj = calendar.Calendar(firstweekday=6)
                                    month_days = cal_obj.monthdayscalendar(y, m)
                                    days_labels = [
                                     'أحد', 
                                     'إثنين', 'ثلاثاء', 'أربعاء', 
                                     'خميس', 'جمعة', 'سبت']
                                    for i, d in enumerate(days_labels):
                                        tk.Label(f_days, text=d, font=('Arial', 8, 'bold'), bg="white").grid(row=0, column=(6 - i), sticky="nsew")
                                    else:
                                        for r, week in enumerate(month_days):
                                            for c, day in enumerate(week):
                                                if day != 0:
                                                    btn = tk.Button(f_days, text=(str(day)), command=(lambda d=day: on_day(d)), bg="white", relief="flat")
                                                    btn.grid(row=(r + 1), column=(6 - c), sticky="nsew", padx=1, pady=1)

                    def next_month():
                        m = current_month.get() + 1
                        if m > 12:
                            current_month.set(1)
                            current_year.set(current_year.get() + 1)
                        else:
                            current_month.set(m)
                        update_cal()

                    def prev_month():
                        m = current_month.get() - 1
                        if m < 1:
                            current_month.set(12)
                            current_year.set(current_year.get() - 1)
                        else:
                            current_month.set(m)
                        update_cal()

                    f_nav = tk.Frame(top, bg="white")
                    f_nav.pack(fill="x", pady=5)
                    tk.Button(f_nav, text="<", command=prev_month).pack(side="left", padx=10)
                    lbl_header = tk.Label(f_nav, text="", font=('Arial', 12, 'bold'), bg="white")
                    lbl_header.pack(side="left", expand=True)
                    tk.Button(f_nav, text=">", command=next_month).pack(side="left", padx=10)
                    f_days = tk.Frame(top, bg="white")
                    f_days.pack(fill="both", expand=True, padx=5, pady=5)
                    update_cal()
            except Exception as e:
                try:
                    messagebox.showerror("Error", f"Calendar Error: {e}")
                finally:
                    pass


        tk.Label(r_a, text="من:", bg=COLOR_BG).pack(side="right", padx=5)
        f_d1 = tk.Frame(r_a, bg=COLOR_BG)
        f_d1.pack(side="right", padx=5)
        ent_date_from = tk.Entry(f_d1, width=12, font=('Segoe UI', 10), justify="center")
        ent_date_from.pack(side="left")
        tk.Button(f_d1, text="📅", command=(lambda: open_date_picker(ent_date_from)), width=3, bg="#eee").pack(side="left")
        try:
            ent_date_from.insert(0, datetime.now().replace(day=1).strftime("%Y-%m-%d"))
        except:
            pass
        else:
            tk.Label(r_a, text="إلى:", bg=COLOR_BG).pack(side="right", padx=5)
            f_d2 = tk.Frame(r_a, bg=COLOR_BG)
            f_d2.pack(side="right", padx=5)
            ent_date_to = tk.Entry(f_d2, width=12, font=('Segoe UI', 10), justify="center")
            ent_date_to.pack(side="left")
            tk.Button(f_d2, text="📅", command=(lambda: open_date_picker(ent_date_to)), width=3, bg="#eee").pack(side="left")
            ent_date_to.insert(0, datetime.now().strftime("%Y-%m-%d"))

            def set_today():
                today_s = datetime.now().strftime("%Y-%m-%d")
                ent_date_from.delete(0, tk.END)
                ent_date_from.insert(0, today_s)
                ent_date_to.delete(0, tk.END)
                ent_date_to.insert(0, today_s)
                apply_filters()


            tk.Button(r_a, text="اليوم", command=set_today, bg="#ef6c00", fg="white", font=('Segoe UI',
                                                                                            8)).pack(side="right", padx=5)
            lbl_search_title = tk.Label(r_a, text="|   اسم الطالب:", bg=COLOR_BG)
            lbl_search_title.pack(side="right", padx=10)
            ent_search_name = tk.Entry(r_a, width=20, font=('Segoe UI', 10), justify="right")
            ent_search_name.pack(side="right", padx=5)
            r_b = tk.Frame(f_filter, bg=COLOR_BG)
            r_b.pack(fill="x", pady=5)
            self.filter_type = tk.StringVar(value="All")
            self.analysis_mode = tk.StringVar(value="Student")

            def set_filter_type(t):
                self.filter_type.set(t)
                apply_filters()


            def toggle_analysis_mode():
                m = self.analysis_mode.get()
                if m == "Student":
                    self.analysis_mode.set("Employee")
                    btn_mode.config(text="🔄 وضع تحليل: الموظفين", bg="#00695c")
                    lbl_st_name.config(text="اختر موظفاً...", fg="#00695c")
                    ent_search_name.delete(0, tk.END)
                    lbl_search_title.config(text="|   اسم الموظف:")
                else:
                    self.analysis_mode.set("Student")
                    btn_mode.config(text="🔄 وضع تحليل: الطلاب", bg="#1565c0")
                    lbl_st_name.config(text="اختر طالباً...", fg="#333")
                    ent_search_name.delete(0, tk.END)
                    lbl_search_title.config(text="|   اسم الطالب:")
                apply_filters()


            btn_mode = tk.Button(r_b, text="🔄 وضع تحليل: الطلاب", command=toggle_analysis_mode, bg="#1565c0", fg="white", width=20, font=('Segoe UI',
                                                                                                                                          9,
                                                                                                                                          'bold'))
            btn_mode.pack(side="left", padx=5)
            tk.Frame(r_b, width=20, bg=COLOR_BG).pack(side="left")
            btn_all = tk.Button(r_b, text="الكل", command=(lambda: set_filter_type("All")), bg="#607d8b", fg="white", width=8)
            btn_all.pack(side="right", padx=2)
            tk.Button(r_b, text="تأخير", command=(lambda: set_filter_type("تأخير")), bg="#ff9800", fg="white", width=8).pack(side="right", padx=2)
            tk.Button(r_b, text="غياب", command=(lambda: set_filter_type("غياب")), bg="#d32f2f", fg="white", width=8).pack(side="right", padx=2)
            tk.Button(r_b, text="انصراف / استئذان", command=(lambda: set_filter_type("انصراف")), bg="#2196f3", fg="white", width=15).pack(side="right", padx=2)
            tk.Button(r_b, text="تطبيق الفلاتر", command=(lambda: apply_filters()), bg=COLOR_ACCENT, fg="white", font=('Segoe UI',
                                                                                                                       10,
                                                                                                                       'bold'), padx=15).pack(side="left", padx=10)
            tk.Label(r_b, text="|  الحالة:", bg=COLOR_BG).pack(side="right", padx=5)
            self.approval_filter_var = tk.StringVar(value="معلق")
            cb_app_filter = ttk.Combobox(r_b, textvariable=(self.approval_filter_var), values=["الكل", "معلق", "معتمد", "مرفوض"], width=10, state="readonly")
            cb_app_filter.pack(side="right", padx=5)
            cb_app_filter.bind("<<ComboboxSelected>>", lambda e: apply_filters())

            def select_all_rows():
                tv_rep.selection_set(tv_rep.get_children())


            def approve_selected():
                selection = tv_rep.selection()
                if not selection:
                    return messagebox.showinfo("تنبيه", "اختر سجلات أولاً لاعتمادها")


            self.grouped_indices = {}

            def approve_selected():
                selection = tv_rep.selection()
                if not selection:
                    return messagebox.showinfo("تنبيه", "اختر سجلات أولاً لاعتمادها")
                all_indices = []
                for iid in selection:
                    indices = self.grouped_indices.get(str(iid), [])
                    if not indices:
                        if str(iid).isdigit():
                            indices = [
                             int(iid)]
                    all_indices.extend(indices)
                else:
                    if not all_indices:
                        return
                    elif self.m.approve_attendance_records(all_indices):
                        for iid in selection:
                            tv_rep.item(iid, tags=('approved', ))
                            vals = list(tv_rep.item(iid, "values"))
                            if len(vals) > 6:
                                vals[6] = "معتمد"
                                tv_rep.item(iid, values=vals)
                            messagebox.showinfo("تم", f"تم اعتماد {len(all_indices)} سجل بنجاح.")

                    else:
                        messagebox.showerror("خطأ", "فشل الاعتماد.")


            btn_sel_all = tk.Button(r_b, text="☑️ تحديد الكل", command=select_all_rows, bg="#78909c", fg="white", font=('Segoe UI',
                                                                                                                        10,
                                                                                                                        'bold'))
            btn_sel_all.pack(side="left", padx=5)
            btn_approve = tk.Button(r_b, text="✅ اعتماد المحدد", command=approve_selected, bg="#2e7d32", fg="white", font=('Segoe UI',
                                                                                                                           10,
                                                                                                                           'bold'))
            btn_approve.pack(side="left", padx=5)

            def reject_selected():
                selection = tv_rep.selection()
                if not selection:
                    return messagebox.showinfo("تنبيه", "اختر سجلات أولاً لرفضها")
                else:
                    return messagebox.askyesno("تأكيد", f"هل أنت متأكد من رفض {len(selection)} مجموعة/سجل؟") or None
                all_indices = []
                for iid in selection:
                    indices = self.grouped_indices.get(str(iid), [])
                    if not indices:
                        if str(iid).isdigit():
                            indices = [
                             int(iid)]
                    all_indices.extend(indices)
                else:
                    if self.m.reject_attendance_records(all_indices):
                        apply_filters()
                        messagebox.showinfo("تم", f"تم رفض {len(all_indices)} سجل.")
                    else:
                        messagebox.showerror("خطأ", "فشل عملية الرفض.")


            btn_reject = tk.Button(r_b, text="❌ رفض المحدد", command=reject_selected, bg="#c62828", fg="white", font=('Segoe UI',
                                                                                                                      10,
                                                                                                                      'bold'))
            btn_reject.pack(side="left", padx=5)
            content = tk.Frame(tab_rep, bg=COLOR_BG)
            content.pack(fill="both", expand=True, padx=10, pady=5)
            tree_frame = tk.Frame(content, bg=COLOR_BG)
            tree_frame.pack(side="right", fill="both", expand=True)
            cols = [
             'التاريخ', 'الوقت', 'العملية', 'الطالب', 'الصف', 'الموظف', 'الاعتماد', 'المرفق']
            tv_rep = ttk.Treeview(tree_frame, columns=cols, show="headings", selectmode="extended")
            for c in cols:
                tv_rep.heading(c, text=c)
                w = 80 if c in ('الوقت', 'العملية') else 100
                if c == "الطالب":
                    w = 150
                if c == "التاريخ":
                    w = 90
                if c == "الاعتماد":
                    w = 90
                if c == "المرفق":
                    w = 80
                tv_rep.column(c, anchor="center", width=w)
            else:
                sb = ttk.Scrollbar(tree_frame, orient="vertical", command=(tv_rep.yview))
                tv_rep.configure(yscrollcommand=(sb.set))
                sb.pack(side="left", fill="y")
                tv_rep.pack(side="right", fill="both", expand=True)
                stats_panel = tk.LabelFrame(content, text="لوحة التحليل", font=('Segoe UI',
                                                                                11, 'bold'), bg="#eceff1", width=280, padx=10, pady=10)
                stats_panel.pack(side="left", fill="y", padx=(0, 10))
                lbl_st_name = tk.Label(stats_panel, text="اختر طالباً...", font=('Segoe UI',
                                                                                 14, 'bold'), bg="#eceff1", wraplength=250)
                lbl_st_name.pack(pady=(10, 20))
                f_counts = tk.Frame(stats_panel, bg="#eceff1")
                f_counts.pack(fill="x")

                def mk_stat(p, txt, col):
                    f = tk.Frame(p, bg="white", pady=5, padx=5, highlightbackground=col, highlightthickness=1)
                    f.pack(fill="x", pady=3)
                    tk.Label(f, text=txt, font=('Segoe UI', 10), bg="white").pack(side="right")
                    l = tk.Label(f, text="0", font=('Segoe UI', 12, 'bold'), fg=col, bg="white")
                    l.pack(side="left", padx=5)
                    return (l, f, txt)


                l_Stat1, f_Stat1, t_Stat1 = mk_stat(f_counts, "عدد مرات التأخير", "#ff9800")
                l_Stat2, f_Stat2, t_Stat2 = mk_stat(f_counts, "عدد أيام الغياب", "#d32f2f")
                l_Stat3, f_Stat3, t_Stat3 = mk_stat(f_counts, "عدد مرات الانصراف", "#2196f3")
                f_perf = tk.Frame(stats_panel, bg="#eceff1", pady=10)
                f_perf.pack_forget()
                tk.Label(f_perf, text="نسبة الإنجاز (من الكلي):", bg="#eceff1", font=('Segoe UI',
                                                                                      9)).pack(anchor="e")
                prog_bar = ttk.Progressbar(f_perf, orient="horizontal", length=200, mode="determinate")
                prog_bar.pack(pady=5)
                lbl_perc = tk.Label(f_perf, text="0%", bg="#eceff1", font=('Segoe UI', 10,
                                                                           'bold'), fg="#00695c")
                lbl_perc.pack()
                tk.Label(stats_panel, text="-----------------", bg="#eceff1").pack(pady=10)
                lbl_total_recs = tk.Label(stats_panel, text="إجمالي السجلات: 0", bg="#eceff1")
                lbl_total_recs.pack()
                self.df_current_view = pd.DataFrame()

                def apply_filters():
                    df = self.m.get_attendance_history()
                    if df.empty:
                        for x in tv_rep.get_children():
                            tv_rep.delete(x)
                        else:
                            return

                    try:
                        d_f_raw = ent_date_from.get().strip()
                        d_t_raw = ent_date_to.get().strip()
                        date_col = df["التاريخ"].astype(str).str.slice(0, 10)
                        mask = (date_col >= d_f_raw) & (date_col <= d_t_raw)
                        df = df.loc[mask]
                    except Exception as e:
                        try:
                            print(f"Filter error: {e}")
                        finally:
                            pass

                    else:
                        ft = self.filter_type.get()
                        if ft != "All":
                            df = df[df["نوع العملية"] == ft]
                        q_name = ent_search_name.get().strip()
                        mode = self.analysis_mode.get()
                        if q_name:
                            if mode == "Student":
                                df = df[df["اسم الطالب"].str.contains(q_name, na=False)]
                            else:
                                df = df[df["اسم الموظف المنفذ"].str.contains(q_name, na=False)]
                        app_f = self.approval_filter_var.get()
                        if app_f != "الكل":
                            if "حالة الاعتماد" not in df.columns:
                                df["حالة الاعتماد"] = "معلق"
                            df = df[df["حالة الاعتماد"].fillna("معلق") == app_f]
                        self.df_current_view = df
                    for x in tv_rep.get_children():
                        tv_rep.delete(x)
                    else:
                        tv_rep["columns"] = ('date', 'time', 'op', 'name', 'grade', 'emp',
                                             'status', 'att')
                        tv_rep.heading("att", text="المرفق")
                        tv_rep.column("att", width=80, anchor="center")
                        self.grouped_indices = {}
                        req_cols = [
                         'التاريخ', 'الوقت', 'نوع العملية', 
                         'اسم الموظف المنفذ', 'الصف', 'الشعبة', 'رابط المرفق', 
                         'حالة الاعتماد']
                        for c in req_cols:
                            if c not in df.columns:
                                df[c] = ""
                        else:
                            groups = {}

                        for idx, row in df.iterrows():
                            key_fields = (
                             str(row.get("التاريخ", "")),
                             str(row.get("الوقت", "")),
                             str(row.get("نوع العملية", "")),
                             str(row.get("اسم الموظف المنفذ", "")),
                             str(row.get("الصف", "")),
                             str(row.get("الشعبة", "")),
                             str(row.get("رابط المرفق", "")),
                             str(row.get("حالة الاعتماد", "معلق")))
                            if key_fields not in groups:
                                groups[key_fields] = {'indices':[],  'names':[],  'row':row}
                            groups[key_fields]["indices"].append(idx)
                            groups[key_fields]["names"].append(str(row.get("اسم الطالب", "")))
                        else:
                            group_id_counter = 0
                            for key, info in groups.items():
                                group_id_counter += 1
                                iid = f"grp_{group_id_counter}"
                                self.grouped_indices[iid] = info["indices"]
                                row = info["row"]
                                names_list = info["names"]
                                count = len(names_list)
                                if count > 1:
                                    trunc_names = ", ".join(names_list[:2])
                                    if count > 2:
                                        trunc_names += "..."
                                    s_disp = f"({count}) طلاب: {trunc_names}"
                                else:
                                    s_disp = names_list[0] if names_list else "—"
                                app_status = row.get("حالة الاعتماد", "معلق")
                                if pd.isna(app_status) or str(app_status).strip() == "":
                                    app_status = "معلق"
                                elif str(app_status) == "معتمد":
                                    tag = "approved"
                                else:
                                    if str(app_status) == "مرفوض":
                                        tag = "rejected"
                                    else:
                                        tag = "pending"
                                att_path = row.get("رابط المرفق", "")
                                att_btn_text = "عرض 👁️" if (att_path and str(att_path).strip() and str(att_path) != "nan") else "—"
                                tv_rep.insert("", "end", iid=iid, values=[
                                 row.get("التاريخ", ""), row.get("الوقت", ""), row.get("نوع العملية", ""),
                                 s_disp,
                                 f'{row.get("الصف", "")} - {row.get("الشعبة", "")}',
                                 row.get("اسم الموظف المنفذ", ""),
                                 app_status, att_btn_text],
                                  tags=(
                                 tag,))
                            else:
                                tv_rep.tag_configure("approved", foreground="green")
                                tv_rep.tag_configure("pending", foreground="#ef6c00")
                                tv_rep.tag_configure("rejected", foreground="#c62828")
                                lbl_total_recs.config(text=f"إجمالي السجلات المعروضة: {len(df)}")
                                if mode == "Student":
                                    f_perf.pack_forget()
                                    f_Stat1.winfo_children()[1].config(text="عدد مرات التأخير")
                                    f_Stat2.winfo_children()[1].config(text="عدد أيام الغياب")
                                    f_Stat3.winfo_children()[1].config(text="عدد مرات الانصراف")
                                else:
                                    f_perf.pack(fill="x")
                                    f_Stat1.winfo_children()[1].config(text="سجل تأخير")
                                    f_Stat2.winfo_children()[1].config(text="سجل غياب")
                                    f_Stat3.winfo_children()[1].config(text="سجل انصراف")

                                def on_click_tree(event):
                                    region = tv_rep.identify_region(event.x, event.y)
                                    if region == "cell":
                                        iid = tv_rep.identify_row(event.y)
                                        if not iid: return
                                        col = tv_rep.identify_column(event.x)
                                        
                                        # Only pop up if clicking on the Student list or Attachment
                                        # But let's just make it pop up clicking anywhere to be safe and easiest for the manager
                                        if str(iid) in self.grouped_indices:
                                            indices = self.grouped_indices[str(iid)]
                                            if not indices: return
                                            
                                            try:
                                                df = self.df_current_view
                                                rows_df = df.loc[indices]
                                            except: return
                                            
                                            top = tk.Toplevel(parent)
                                            top.title("تفاصيل الطلاب المحددين")
                                            top.geometry("800x400")
                                            top.configure(bg=COLOR_BG)
                                            try:
                                                top.transient(parent.winfo_toplevel())
                                            except: pass
                                            
                                            lbl = tk.Label(top, text="قائمة الطلاب المنفذ عليهم الإجراء", font=('Segoe UI', 14, 'bold'), bg=COLOR_BG)
                                            lbl.pack(pady=10)
                                            
                                            dt_cols = ["التاريخ", "الوقت", "العملية", "اسم الطالب", "الصف", "الشعبة"]
                                            tv = ttk.Treeview(top, columns=dt_cols, show="headings")
                                            for c in dt_cols:
                                                tv.heading(c, text=c)
                                                tv.column(c, anchor="center", width=120 if c=="اسم الطالب" else 80)
                                                
                                            for _, r in rows_df.iterrows():
                                                tv.insert("", "end", values=(
                                                    r.get("التاريخ", ""),
                                                    r.get("الوقت", ""),
                                                    r.get("نوع العملية", ""),
                                                    r.get("اسم الطالب", ""),
                                                    r.get("الصف", ""),
                                                    r.get("الشعبة", "")
                                                ))
                                                
                                            sb = ttk.Scrollbar(top, orient="vertical", command=tv.yview)
                                            tv.configure(yscrollcommand=sb.set)
                                            sb.pack(side="left", fill="y")
                                            tv.pack(fill="both", expand=True, padx=10, pady=10)

                                tv_rep.bind("<Double-1>", on_click_tree)
                                tv_rep.bind("<ButtonRelease-1>", lambda e: on_click_tree(e) if tv_rep.identify_column(e.x) == "#8" else None)


                def on_tree_select(event):
                    sel = tv_rep.selection()
                    if not sel:
                        return
                    item = tv_rep.item(sel[0])
                    mode = self.analysis_mode.get()
                    full_df = self.m.get_attendance_history()
                    try:
                        d_from = ent_date_from.get()
                        d_to = ent_date_to.get()
                        mask = (full_df["التاريخ"] >= d_from) & (full_df["التاريخ"] <= d_to)
                        period_df = full_df.loc[mask]
                    except:
                        period_df = full_df
                    else:
                        if mode == "Student":
                            s_name = item["values"][3]
                            lbl_st_name.config(text=s_name)
                            s_df = period_df[period_df["اسم الطالب"] == s_name]
                            l_Stat1.config(text=(str(len(s_df[s_df["نوع العملية"] == "تأخير"]))))
                            l_Stat2.config(text=(str(len(s_df[s_df["نوع العملية"] == "غياب"]))))
                            l_Stat3.config(text=(str(len(s_df[s_df["نوع العملية"] == "انصراف"]))))
                        else:
                            emp_name = item["values"][5]
                            lbl_st_name.config(text=emp_name)
                            e_df = period_df[period_df["اسم الموظف المنفذ"] == emp_name]
                            l_Stat1.config(text=(str(len(e_df[e_df["نوع العملية"] == "تأخير"]))), fg="#ff9800")
                            l_Stat2.config(text=(str(len(e_df[e_df["نوع العملية"] == "غياب"]))), fg="#d32f2f")
                            l_Stat3.config(text=(str(len(e_df[e_df["نوع العملية"] == "انصراف"]))), fg="#2196f3")
                            total_school = len(period_df)
                            total_emp = len(e_df)
                            perc = total_emp / total_school * 100 if total_school > 0 else 0
                            lbl_perc.config(text=f"%{perc:.1f}")
                            prog_bar["value"] = perc


                tv_rep.bind("<<TreeviewSelect>>", on_tree_select)
                apply_filters()

                def on_show(event):
                    if event.widget == page:
                        apply_filters()
                        try:
                            refresh_monitoring_list()
                            refresh_unified_table()
                        except:
                            pass


                page.bind("<Visibility>", on_show)
                tab_assign = tk.Frame(notebook, bg=COLOR_BG)
                notebook.add(tab_assign, text=" إدارة المهام والصلاحيات والمتابعة ")
                f_top = tk.Frame(tab_assign, bg=COLOR_BG)
                f_top.pack(fill="x", padx=10, pady=5)
                f_form = tk.LabelFrame(f_top, text="توجيه المهام وتوزيع الصلاحيات", font=('Segoe UI',
                                                                                          11,
                                                                                          'bold'), bg="#e3f2fd", padx=15, pady=10)
                f_form.pack(fill="x", pady=5)
                tk.Label(f_form, text="نوع التكليف:", bg="#e3f2fd", font=('Segoe UI', 10)).grid(row=0, column=3, sticky="e", padx=5, pady=5)
                task_types = [
                 "التواصل مع ولي الأمر", "تأخير", "غياب", "استئذان"]
                cb_type = ttk.Combobox(f_form, values=task_types, state="readonly", width=25)
                cb_type.current(1)
                cb_type.grid(row=0, column=2, sticky="w", padx=5, pady=5)
                tk.Label(f_form, text="الموظف/الموظفون:", bg="#e3f2fd", font=('Segoe UI', 10)).grid(row=0, column=1, sticky="w", padx=5, pady=5)
                f_selector = tk.Frame(f_form, bg="#e3f2fd")
                f_selector.grid(row=1, column=0, columnspan=4, sticky="ew", padx=5, pady=5)
                
                emp_list = list(self.m.load_employees_pins().keys())
                
                f_multi = tk.Frame(f_selector, bg="white", highlightthickness=1, highlightbackground="#ccc")
                f_multi.pack(fill="both", expand=True)
                
                search_var = tk.StringVar()
                ent_search = tk.Entry(f_multi, textvariable=search_var, width=30, justify="right")
                ent_search.pack(side="top", fill="x", padx=2, pady=2)
                ent_search.insert(0, "بحث باسم الموظف...")
                ent_search.bind("<FocusIn>", lambda e: ent_search.delete(0, 'end') if ent_search.get() == "بحث باسم الموظف..." else None)
                ent_search.bind("<FocusOut>", lambda e: ent_search.insert(0, "بحث باسم الموظف...") if not ent_search.get() else None)
                
                f_lb = tk.Frame(f_multi)
                f_lb.pack(fill="both", expand=True)
                
                lb_multi = tk.Listbox(f_lb, selectmode="multiple", height=5, exportselection=False)
                sb_multi = ttk.Scrollbar(f_lb, orient="vertical", command=(lb_multi.yview))
                lb_multi.configure(yscrollcommand=(sb_multi.set))
                lb_multi.pack(side="right", fill="both", expand=True)
                sb_multi.pack(side="left", fill="y")
                
                selected_emps = set()
                
                def populate_listbox(*args):
                    query = search_var.get().strip().lower()
                    if query == "بحث باسم الموظف...": query = ""
                    lb_multi.delete(0, tk.END)
                    for e in emp_list:
                        if not query or query in e.lower():
                            lb_multi.insert(tk.END, e)
                            if e in selected_emps:
                                lb_multi.selection_set(tk.END)
                
                search_var.trace("w", populate_listbox)
                
                def on_lb_select(event):
                    visible = lb_multi.get(0, tk.END)
                    current_sels = [lb_multi.get(i) for i in lb_multi.curselection()]
                    for v in visible:
                        selected_emps.discard(v)
                    for s in current_sels:
                        selected_emps.add(s)
                        
                lb_multi.bind("<<ListboxSelect>>", on_lb_select)
                populate_listbox()
                
                def toggle_selector(event=None):
                    t_type = cb_type.get()
                    selected_emps.clear()
                    roles = self.m.load_attendance_roles()
                    current_assigned = roles.get(t_type, [])
                    if not isinstance(current_assigned, list):
                        current_assigned = [str(current_assigned)] if current_assigned else []
                    
                    for e in current_assigned:
                        if e in emp_list:
                            selected_emps.add(e)
                            
                    populate_listbox()

                toggle_selector()
                cb_type.bind("<<ComboboxSelected>>", toggle_selector)

                def save_assignment():
                    t = cb_type.get()
                    selected_names = list(selected_emps)
                    current = self.m.load_attendance_roles()
                    
                    old_assigned = current.get(t, [])
                    if not isinstance(old_assigned, list):
                        old_assigned = [str(old_assigned)] if old_assigned else []
                        
                    for name in selected_names:
                        if name not in old_assigned:
                            self.m.log_role_assignment(name, t, 'add')
                            
                    for name in old_assigned:
                        if name not in selected_names:
                            self.m.log_role_assignment(name, t, 'remove')
                            
                    current[t] = selected_names
                    if self.m.save_attendance_roles(current):
                        today_str = datetime.now().strftime("%Y-%m-%d")
                        self.m.update_task_status(t, today_str, "تحت العمل", "")
                        messagebox.showinfo("تم", f"تم تحديث {t} بنجاح")
                        try:
                            refresh_unified_table()
                            refresh_monitoring_list()
                            update_history_table()
                        except: pass
                    else:
                        messagebox.showerror("خطأ", "فشل الحفظ")


                tk.Button(f_form, text="💾 حفظ / تحديث", command=save_assignment, bg=COLOR_ACCENT, fg="white", font=('Segoe UI',
                                                                                                                    10,
                                                                                                                    'bold'), padx=15).grid(row=2, column=0, pady=10)
                f_list = tk.LabelFrame(tab_assign, text="ملخص الصلاحيات الحالية", font=('Segoe UI',
                                                                                        10,
                                                                                        'bold'), bg=COLOR_BG, padx=5, pady=5)
                f_list.pack(fill="x", padx=10, pady=5)
                cols = [
                 "النوع", "المسؤولون"]
                tree_roles = ttk.Treeview(f_list, columns=cols, show="headings", height=4)
                tree_roles.heading("النوع", text="النوع")
                tree_roles.column("النوع", anchor="center", width=150)
                tree_roles.heading("المسؤولون", text="الموظفون الموكلون")
                tree_roles.column("المسؤولون", anchor="w", width=500)
                tree_roles.pack(fill="x")

                def execute_role_deletion(sel):
                    if not sel:
                        return
                        
                    item_vals = tree_roles.item(sel[0], "values")
                    if not item_vals or len(item_vals) < 2:
                        return
                        
                    r_type = str(item_vals[0])
                    assigned = str(item_vals[1])
                    
                    if assigned == "—" or not assigned:
                        return
                        
                    if r_type.startswith("مهمة:"):
                        raw_task = r_type.replace("مهمة: ", "").strip()
                        if messagebox.askyesno("تأكيد إنهاء المهمة", f"هل تريد إنهاء هذه المهمة عن '{assigned}'؟\n\n📌 {raw_task}"):
                            if self.m.remove_custom_assignment(raw_task, assigned):
                                messagebox.showinfo("تم", "تم إنهاء المهمة بنجاح")
                                refresh_unified_table()
                                try:
                                    refresh_monitoring_list()
                                    update_history_table()
                                except: pass
                        return
                        
                    if messagebox.askyesno("تأكيد إنهاء المهمة", f"هل تريد إنهاء تكليف '{r_type}' للموظف '{assigned}'؟"):
                        roles = self.m.load_attendance_roles()
                        
                        old_assigned = roles.get(r_type, [])
                        if not isinstance(old_assigned, list):
                            old_assigned = [str(old_assigned)] if old_assigned else []
                        if assigned in old_assigned:
                            old_assigned.remove(assigned)
                            self.m.log_role_assignment(assigned, r_type, 'remove')
                            
                        roles[r_type] = old_assigned
                        if self.m.save_attendance_roles(roles):
                            today_str = datetime.now().strftime("%Y-%m-%d")
                            self.m.update_task_status(r_type, today_str, "تحت العمل", "")
                            messagebox.showinfo("تم", "تم إنهاء التكليف بنجاح")
                            refresh_unified_table()
                            try:
                                refresh_monitoring_list()
                                update_history_table()
                            except: pass
                            toggle_selector()

                def on_role_double_click(event):
                    region = tree_roles.identify("region", event.x, event.y)
                    if region != "cell":
                        return
                    execute_role_deletion(tree_roles.selection())
                    
                def on_manual_end_task():
                    sel = tree_roles.selection()
                    if not sel:
                        messagebox.showwarning("تنبيه", "الرجاء تحديد تكليف من الجدول أولاً.")
                        return
                    execute_role_deletion(sel)
                    
                btn_end_task = tk.Button(f_list, text="🗑️ إنهاء التكليف للموظف المحدد", command=on_manual_end_task, bg="#dc3545", fg="white", font=('Segoe UI', 9, 'bold'))
                btn_end_task.pack(pady=5)

                tree_roles.bind("<Double-1>", on_role_double_click)

                def refresh_unified_table():
                    for i in tree_roles.get_children():
                        tree_roles.delete(i)
                    else:
                        roles = self.m.load_attendance_roles()
                        for r, assigned_list in roles.items():
                            if not assigned_list:
                                continue
                            if not isinstance(assigned_list, list):
                                assigned_list = [str(assigned_list)]
                            for emp in assigned_list:
                                tree_roles.insert("", "end", values=(r, emp))
                        else:
                            today_str = datetime.now().strftime("%Y-%m-%d")
                            custom = self.m.get_all_assignments()
                            for c in reversed(custom):
                                if c.get("date") == today_str:
                                    desc = f'مهمة: {c.get("task")}'
                                    tree_roles.insert("", "end", values=(desc, c.get("employee")))


                refresh_unified_table()
                
                # --- History Tab ---
                tab_history = tk.Frame(notebook, bg=COLOR_BG)
                notebook.add(tab_history, text=" 🕰️ سجل تاريخ الموظفين ")
                
                f_hist_top = tk.LabelFrame(tab_history, text="🔍 أدوات البحث في السجل", font=('Segoe UI', 11, 'bold'), bg=COLOR_BG, padx=10, pady=10)
                f_hist_top.pack(fill="x", padx=10, pady=5)
                
                tk.Label(f_hist_top, text="اختيار الموظف:", font=('Segoe UI', 10), bg=COLOR_BG).pack(side="right", padx=5)
                cb_hist_emp = ttk.Combobox(f_hist_top, values=["الكل"] + sorted(list(self.m.teacher_names)) if hasattr(self.m, 'teacher_names') and self.m.teacher_names else ["الكل"], state="readonly", width=30)
                cb_hist_emp.set("الكل")
                cb_hist_emp.pack(side="right", padx=5)
                
                tk.Label(f_hist_top, text="حالة المهمة:", font=('Segoe UI', 10), bg=COLOR_BG).pack(side="right", padx=5)
                cb_hist_status = ttk.Combobox(f_hist_top, values=["الكل", "تحت العمل", "منتهي"], state="readonly", width=15)
                cb_hist_status.set("الكل")
                cb_hist_status.pack(side="right", padx=5)
                
                btn_hist_search = tk.Button(f_hist_top, text="بحث", bg=COLOR_ACCENT, fg="white", font=('Segoe UI', 10, 'bold'), cursor="hand2")
                btn_hist_search.pack(side="right", padx=15)
                
                f_hist_tree = tk.Frame(tab_history, bg=COLOR_BG)
                f_hist_tree.pack(fill="both", expand=True, padx=10, pady=5)
                
                hist_cols = ["الموظف", "التكليف", "تاريخ البدء", "تاريخ الانتهاء", "الحالة"]
                tree_hist = ttk.Treeview(f_hist_tree, columns=hist_cols, show="headings")
                for c in hist_cols:
                    tree_hist.heading(c, text=c)
                    tree_hist.column(c, anchor="center", width=150 if c != "الحالة" else 100)
                
                hist_sb = ttk.Scrollbar(f_hist_tree, orient="vertical", command=tree_hist.yview)
                tree_hist.configure(yscrollcommand=hist_sb.set)
                hist_sb.pack(side="left", fill="y")
                tree_hist.pack(side="right", fill="both", expand=True)
                
                tree_hist.tag_configure("active", foreground="#155724", background="#d4edda")
                tree_hist.tag_configure("done", foreground="#6c757d", background="#f8f9fa")
                
                def update_history_table(event=None):
                    for i in tree_hist.get_children():
                        tree_hist.delete(i)
                    
                    roles_cache = self.m.load_attendance_roles()
                    history = self.m.load_task_history()
                    needs_save = False
                    today_str = datetime.now().strftime("%Y-%m-%d")
                    
                    # 1. Force cleanup to fix ghost "تحت العمل" records
                    for h in history:
                        if h.get("status") == "تحت العمل":
                            t_name = h.get("task")
                            e_name = h.get("employee")
                            if t_name in roles_cache and e_name in roles_cache[t_name]:
                                pass
                            else:
                                h["status"] = "منتهي"
                                if not h.get("end_date"):
                                    h["end_date"] = today_str
                                needs_save = True
                                
                    # 2. Add missing active records
                    for r_name, emps in roles_cache.items():
                        if not isinstance(emps, list): continue
                        for e in emps:
                            active_exists = any((h for h in history if h.get("employee") == e and h.get("task") == r_name and h.get("status") == "تحت العمل"))
                            if not active_exists:
                                history.append({
                                    "employee": e,
                                    "task": r_name,
                                    "start_date": today_str,
                                    "end_date": "",
                                    "status": "تحت العمل"
                                })
                                needs_save = True
                                
                    if needs_save:
                        self.m.save_task_history(history)
                        history = self.m.load_task_history()
                        
                    target_emp = cb_hist_emp.get()
                    target_status = cb_hist_status.get()
                    
                    for h in reversed(history):
                        if target_emp != "الكل" and h.get("employee") != target_emp:
                            continue
                        if target_status != "الكل" and h.get("status") != target_status:
                            continue
                        
                        tag = "active" if h.get("status") == "تحت العمل" else "done"
                        tree_hist.insert("", "end", values=(
                            h.get("employee", ""),
                            h.get("task", ""),
                            h.get("start_date", ""),
                            h.get("end_date", ""),
                            h.get("status", "")
                        ), tags=(tag,))
                
                btn_hist_search.config(command=update_history_table)
                update_history_table()
                
                def on_delete_history(event=None):
                    sel = tree_hist.selection()
                    if not sel:
                        messagebox.showwarning("تنبيه", "الرجاء تحديد سجل من أرشيف التاريخ أولاً.")
                        return
                    item_vals = tree_hist.item(sel[0], "values")
                    emp = str(item_vals[0])
                    task = str(item_vals[1])
                    start_date = str(item_vals[2])
                    
                    if messagebox.askyesno("تأكيد الحذف", f"هل أنت متأكد من حذف هذا السجل نهائياً من الأرشيف؟\n\n👤 {emp}\n📌 {task}"):
                        history = self.m.load_task_history()
                        to_remove = None
                        # Find matching entry (reverse since we display reversed)
                        for h in reversed(history):
                            if h.get("employee") == emp and h.get("task") == task and h.get("start_date") == start_date:
                                to_remove = h
                                break
                        if to_remove:
                            history.remove(to_remove)
                            self.m.save_task_history(history)
                            update_history_table()
                            messagebox.showinfo("تم", "تم حذف السجل بنجاح.")
                            
                btn_del_hist = tk.Button(f_hist_tree, text="🗑️ حذف السجل المحدد نهائياً", command=on_delete_history, bg="#6c757d", fg="white", font=('Segoe UI', 9, 'bold'))
                btn_del_hist.pack(side="bottom", pady=5)
                tree_hist.bind("<Double-1>", on_delete_history)

                return page


        # ========================================

    # --- page_attendance_view_open_date_picker ---
    def page_attendance_view_open_date_picker(entry_widget=None):
        # ========================================
        try:
            import calendar
            top = tk.Toplevel(parent)
            top.title("اختر التاريخ")
            top.geometry("300x300")
            top.configure(bg="white")
            try:
                x = parent.winfo_rootx() + 100
                y = parent.winfo_rooty() + 100
                top.geometry(f"+{x}+{y}")
            except:
                pass
            else:
                current_year = tk.IntVar(value=(datetime.now().year))
                current_month = tk.IntVar(value=(datetime.now().month))

                def update_cal():
                    for widget in f_days.winfo_children():
                        widget.destroy()
                    else:
                        y = current_year.get()
                        m = current_month.get()
                        lbl_header.config(text=f"{y} - {m}")
                        days = [
                         'Ec', 'الجمعة', 'الخميس', 'الأربعاء', 
                         'الثلاثاء', 'الاثنين', 'الأحد', 'السبت']
                        for d in ('ح', 'ن', 'ث', 'ر', 'خ', 'ج', 'س'):
                            tk.Label(f_days, text=d, bg="#eee", width=5).grid(row=0, column=(['س', 'ج', 'خ', 'ر', 'ث', 'ن', 'ح'].index(d)))
                        else:
                            cal = calendar.monthcalendar(y, m)
                            r = 1
                            for week in cal:
                                for idx, day in enumerate(week):
                                    if day == 0:
                                        pass
                                    else:

                                        def on_day(d=day):
                                            selected_date = f"{y}-{m:02d}-{d:02d}"
                                            entry_widget.delete(0, tk.END)
                                            entry_widget.insert(0, selected_date)
                                            top.destroy()
                                            apply_filters()

                            for widget in f_days.winfo_children():
                                widget.destroy()
                            else:
                                cal_obj = calendar.Calendar(firstweekday=6)
                                month_days = cal_obj.monthdayscalendar(y, m)
                                days_labels = [
                                 'أحد', 
                                 'إثنين', 'ثلاثاء', 'أربعاء', 
                                 'خميس', 'جمعة', 'سبت']
                                for i, d in enumerate(days_labels):
                                    tk.Label(f_days, text=d, font=('Arial', 8, 'bold'), bg="white").grid(row=0, column=(6 - i), sticky="nsew")
                                else:
                                    for r, week in enumerate(month_days):
                                        for c, day in enumerate(week):
                                            if day != 0:
                                                btn = tk.Button(f_days, text=(str(day)), command=(lambda d=day: on_day(d)), bg="white", relief="flat")
                                                btn.grid(row=(r + 1), column=(6 - c), sticky="nsew", padx=1, pady=1)


                def next_month():
                    m = current_month.get() + 1
                    if m > 12:
                        current_month.set(1)
                        current_year.set(current_year.get() + 1)
                    else:
                        current_month.set(m)
                    update_cal()


                def prev_month():
                    m = current_month.get() - 1
                    if m < 1:
                        current_month.set(12)
                        current_year.set(current_year.get() - 1)
                    else:
                        current_month.set(m)
                    update_cal()


                f_nav = tk.Frame(top, bg="white")
                f_nav.pack(fill="x", pady=5)
                tk.Button(f_nav, text="<", command=prev_month).pack(side="left", padx=10)
                lbl_header = tk.Label(f_nav, text="", font=('Arial', 12, 'bold'), bg="white")
                lbl_header.pack(side="left", expand=True)
                tk.Button(f_nav, text=">", command=next_month).pack(side="left", padx=10)
                f_days = tk.Frame(top, bg="white")
                f_days.pack(fill="both", expand=True, padx=5, pady=5)
                update_cal()
        except Exception as e:
            try:
                messagebox.showerror("Error", f"Calendar Error: {e}")
            finally:
                pass

        # ========================================

    # --- page_attendance_view_open_date_picker_update_cal ---
    def page_attendance_view_open_date_picker_update_cal():
        # ========================================
        for widget in f_days.winfo_children():
            widget.destroy()
        else:
            y = current_year.get()
            m = current_month.get()
            lbl_header.config(text=f"{y} - {m}")
            days = [
             'Ec', 'الجمعة', 'الخميس', 'الأربعاء', 'الثلاثاء', 'الاثنين', 
             'الأحد', 'السبت']
            for d in ('ح', 'ن', 'ث', 'ر', 'خ', 'ج', 'س'):
                tk.Label(f_days, text=d, bg="#eee", width=5).grid(row=0, column=(['س', 'ج', 'خ', 'ر', 'ث', 'ن', 'ح'].index(d)))
            else:
                cal = calendar.monthcalendar(y, m)
                r = 1
                for week in cal:
                    for idx, day in enumerate(week):
                        if day == 0:
                            pass
                        else:

                            def on_day(d=day):
                                selected_date = f"{y}-{m:02d}-{d:02d}"
                                entry_widget.delete(0, tk.END)
                                entry_widget.insert(0, selected_date)
                                top.destroy()
                                apply_filters()


                for widget in f_days.winfo_children():
                    widget.destroy()
                else:
                    cal_obj = calendar.Calendar(firstweekday=6)
                    month_days = cal_obj.monthdayscalendar(y, m)
                    days_labels = [
                     'أحد', 'إثنين', 'ثلاثاء', 'أربعاء', 'خميس', 
                     'جمعة', 'سبت']
                    for i, d in enumerate(days_labels):
                        tk.Label(f_days, text=d, font=('Arial', 8, 'bold'), bg="white").grid(row=0, column=(6 - i), sticky="nsew")
                    else:
                        for r, week in enumerate(month_days):
                            for c, day in enumerate(week):
                                if day != 0:
                                    btn = tk.Button(f_days, text=(str(day)), command=(lambda d=day: on_day(d)), bg="white", relief="flat")
                                    btn.grid(row=(r + 1), column=(6 - c), sticky="nsew", padx=1, pady=1)

        # ========================================

    # --- page_attendance_view_open_date_picker_update_cal_on_day ---
    def page_attendance_view_open_date_picker_update_cal_on_day(d=None):
        # ========================================
        selected_date = f"{y}-{m:02d}-{d:02d}"
        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, selected_date)
        top.destroy()
        apply_filters()

        # ========================================

    # --- page_attendance_view_open_date_picker_next_month ---
    def page_attendance_view_open_date_picker_next_month():
        # ========================================
        m = current_month.get() + 1
        if m > 12:
            current_month.set(1)
            current_year.set(current_year.get() + 1)
        else:
            current_month.set(m)
        update_cal()

        # ========================================

    # --- page_attendance_view_open_date_picker_prev_month ---
    def page_attendance_view_open_date_picker_prev_month():
        # ========================================
        m = current_month.get() - 1
        if m < 1:
            current_month.set(12)
            current_year.set(current_year.get() - 1)
        else:
            current_month.set(m)
        update_cal()

        # ========================================

    # --- page_attendance_view_set_today ---
    def page_attendance_view_set_today():
        # ========================================
        today_s = datetime.now().strftime("%Y-%m-%d")
        ent_date_from.delete(0, tk.END)
        ent_date_from.insert(0, today_s)
        ent_date_to.delete(0, tk.END)
        ent_date_to.insert(0, today_s)
        apply_filters()

        # ========================================

    # --- page_attendance_view_set_filter_type ---
    def page_attendance_view_set_filter_type(t=None):
        # ========================================
        self.filter_type.set(t)
        apply_filters()

        # ========================================

    # --- page_attendance_view_toggle_analysis_mode ---
    def page_attendance_view_toggle_analysis_mode():
        # ========================================
        m = self.analysis_mode.get()
        if m == "Student":
            self.analysis_mode.set("Employee")
            btn_mode.config(text="🔄 وضع تحليل: الموظفين", bg="#00695c")
            lbl_st_name.config(text="اختر موظفاً...", fg="#00695c")
            ent_search_name.delete(0, tk.END)
            lbl_search_title.config(text="|   اسم الموظف:")
        else:
            self.analysis_mode.set("Student")
            btn_mode.config(text="🔄 وضع تحليل: الطلاب", bg="#1565c0")
            lbl_st_name.config(text="اختر طالباً...", fg="#333")
            ent_search_name.delete(0, tk.END)
            lbl_search_title.config(text="|   اسم الطالب:")
        apply_filters()

        # ========================================

    # --- page_attendance_view_select_all_rows ---
    def page_attendance_view_select_all_rows():
        # ========================================
        tv_rep.selection_set(tv_rep.get_children())

        # ========================================

    # --- page_attendance_view_approve_selected ---
    def page_attendance_view_approve_selected():
        # ========================================
        selection = tv_rep.selection()
        if not selection:
            return messagebox.showinfo("تنبيه", "اختر سجلات أولاً لاعتمادها")

        # ========================================

    # --- page_attendance_view_reject_selected ---
    def page_attendance_view_reject_selected():
        # ========================================
        selection = tv_rep.selection()
        if not selection:
            return messagebox.showinfo("تنبيه", "اختر سجلات أولاً لرفضها")
        else:
            return messagebox.askyesno("تأكيد", f"هل أنت متأكد من رفض {len(selection)} مجموعة/سجل؟") or None
        all_indices = []
        for iid in selection:
            indices = self.grouped_indices.get(str(iid), [])
            if not indices:
                if str(iid).isdigit():
                    indices = [
                     int(iid)]
            all_indices.extend(indices)
        else:
            if self.m.reject_attendance_records(all_indices):
                apply_filters()
                messagebox.showinfo("تم", f"تم رفض {len(all_indices)} سجل.")
            else:
                messagebox.showerror("خطأ", "فشل عملية الرفض.")

        # ========================================

    # --- page_attendance_view_mk_stat ---
    def page_attendance_view_mk_stat(p=None, txt=None, col=None):
        # ========================================
        f = tk.Frame(p, bg="white", pady=5, padx=5, highlightbackground=col, highlightthickness=1)
        f.pack(fill="x", pady=3)
        tk.Label(f, text=txt, font=('Segoe UI', 10), bg="white").pack(side="right")
        l = tk.Label(f, text="0", font=('Segoe UI', 12, 'bold'), fg=col, bg="white")
        l.pack(side="left", padx=5)
        return (l, f, txt)

        # ========================================

    # --- page_attendance_view_apply_filters ---
    def page_attendance_view_apply_filters():
        # ========================================
        df = self.m.get_attendance_history()
        if df.empty:
            for x in tv_rep.get_children():
                tv_rep.delete(x)
            else:
                return

        try:
            d_f_raw = ent_date_from.get().strip()
            d_t_raw = ent_date_to.get().strip()
            date_col = df["التاريخ"].astype(str).str.slice(0, 10)
            mask = (date_col >= d_f_raw) & (date_col <= d_t_raw)
            df = df.loc[mask]
        except Exception as e:
            try:
                print(f"Filter error: {e}")
            finally:
                pass

        else:
            ft = self.filter_type.get()
            if ft != "All":
                df = df[df["نوع العملية"] == ft]
            q_name = ent_search_name.get().strip()
            mode = self.analysis_mode.get()
            if q_name:
                if mode == "Student":
                    df = df[df["اسم الطالب"].str.contains(q_name, na=False)]
                else:
                    df = df[df["اسم الموظف المنفذ"].str.contains(q_name, na=False)]
            app_f = self.approval_filter_var.get()
            if app_f != "الكل":
                if "حالة الاعتماد" not in df.columns:
                    df["حالة الاعتماد"] = "معلق"
                df = df[df["حالة الاعتماد"].fillna("معلق") == app_f]
            self.df_current_view = df
        for x in tv_rep.get_children():
            tv_rep.delete(x)
        else:
            tv_rep["columns"] = ('date', 'time', 'op', 'name', 'grade', 'emp', 'status', 'att')
            tv_rep.heading("att", text="المرفق")
            tv_rep.column("att", width=80, anchor="center")
            self.grouped_indices = {}
            req_cols = [
             'التاريخ', 'الوقت', 'نوع العملية', 'اسم الموظف المنفذ', 'الصف', 
             'الشعبة', 'رابط المرفق', 'حالة الاعتماد']
            for c in req_cols:
                if c not in df.columns:
                    df[c] = ""
            else:
                groups = {}

            for idx, row in df.iterrows():
                key_fields = (
                 str(row.get("التاريخ", "")),
                 str(row.get("الوقت", "")),
                 str(row.get("نوع العملية", "")),
                 str(row.get("اسم الموظف المنفذ", "")),
                 str(row.get("الصف", "")),
                 str(row.get("الشعبة", "")),
                 str(row.get("رابط المرفق", "")),
                 str(row.get("حالة الاعتماد", "معلق")))
                if key_fields not in groups:
                    groups[key_fields] = {'indices':[],  'names':[],  'row':row}
                groups[key_fields]["indices"].append(idx)
                groups[key_fields]["names"].append(str(row.get("اسم الطالب", "")))
            else:
                group_id_counter = 0
                for key, info in groups.items():
                    group_id_counter += 1
                    iid = f"grp_{group_id_counter}"
                    self.grouped_indices[iid] = info["indices"]
                    row = info["row"]
                    names_list = info["names"]
                    count = len(names_list)
                    if count > 1:
                        trunc_names = ", ".join(names_list[:2])
                        if count > 2:
                            trunc_names += "..."
                        s_disp = f"({count}) طلاب: {trunc_names}"
                    else:
                        s_disp = names_list[0] if names_list else "—"
                    app_status = row.get("حالة الاعتماد", "معلق")
                    if pd.isna(app_status) or str(app_status).strip() == "":
                        app_status = "معلق"
                    elif str(app_status) == "معتمد":
                        tag = "approved"
                    else:
                        if str(app_status) == "مرفوض":
                            tag = "rejected"
                        else:
                            tag = "pending"
                    att_path = row.get("رابط المرفق", "")
                    att_btn_text = "عرض 👁️" if (att_path and str(att_path).strip() and str(att_path) != "nan") else "—"
                    tv_rep.insert("", "end", iid=iid, values=[
                     row.get("التاريخ", ""), row.get("الوقت", ""), row.get("نوع العملية", ""),
                     s_disp,
                     f'{row.get("الصف", "")} - {row.get("الشعبة", "")}',
                     row.get("اسم الموظف المنفذ", ""),
                     app_status, att_btn_text],
                      tags=(
                     tag,))
                else:
                    tv_rep.tag_configure("approved", foreground="green")
                    tv_rep.tag_configure("pending", foreground="#ef6c00")
                    tv_rep.tag_configure("rejected", foreground="#c62828")
                    lbl_total_recs.config(text=f"إجمالي السجلات المعروضة: {len(df)}")
                    if mode == "Student":
                        f_perf.pack_forget()
                        f_Stat1.winfo_children()[1].config(text="عدد مرات التأخير")
                        f_Stat2.winfo_children()[1].config(text="عدد أيام الغياب")
                        f_Stat3.winfo_children()[1].config(text="عدد مرات الانصراف")
                    else:
                        f_perf.pack(fill="x")
                        f_Stat1.winfo_children()[1].config(text="سجل تأخير")
                        f_Stat2.winfo_children()[1].config(text="سجل غياب")
                        f_Stat3.winfo_children()[1].config(text="سجل انصراف")

                    def on_click_tree(event):
                        region = tv_rep.identify_region(event.x, event.y)
                        if region == "cell":
                            iid = tv_rep.identify_row(event.y)
                            if not iid: return
                            
                            if str(iid) in self.grouped_indices:
                                indices = self.grouped_indices[str(iid)]
                                if not indices: return
                                
                                try:
                                    df = self.df_current_view
                                    rows_df = df.loc[indices]
                                    
                                    top = tk.Toplevel(parent)
                                    top.title("تفاصيل الطلاب المحددين")
                                    top.geometry("800x400")
                                    top.configure(bg=COLOR_BG)
                                    try:
                                        top.transient(parent.winfo_toplevel())
                                    except: pass
                                    
                                    lbl = tk.Label(top, text="قائمة الطلاب المنفذ عليهم الإجراء", font=('Segoe UI', 14, 'bold'), bg=COLOR_BG)
                                    lbl.pack(pady=10)
                                    
                                    dt_cols = ["التاريخ", "الوقت", "العملية", "اسم الطالب", "الصف", "الشعبة"]
                                    tv = ttk.Treeview(top, columns=dt_cols, show="headings")
                                    for c in dt_cols:
                                        tv.heading(c, text=c)
                                        tv.column(c, anchor="center", width=120 if c=="اسم الطالب" else 80)
                                        
                                    for _, r in rows_df.iterrows():
                                        tv.insert("", "end", values=(
                                            r.get("التاريخ", ""),
                                            r.get("الوقت", ""),
                                            r.get("نوع العملية", ""),
                                            r.get("اسم الطالب", ""),
                                            r.get("الصف", ""),
                                            r.get("الشعبة", "")
                                        ))
                                        
                                    sb = ttk.Scrollbar(top, orient="vertical", command=tv.yview)
                                    tv.configure(yscrollcommand=sb.set)
                                    sb.pack(side="left", fill="y")
                                    tv.pack(fill="both", expand=True, padx=10, pady=10)
                                except Exception as e:
                                    print("Error opening details:", e)

                    tv_rep.bind("<Double-1>", on_click_tree)
                    tv_rep.bind("<ButtonRelease-1>", lambda e: on_click_tree(e) if tv_rep.identify_column(e.x) == "#8" else None)


        # ========================================

    # --- page_attendance_view_apply_filters_on_click_tree ---
    def page_attendance_view_apply_filters_on_click_tree(event=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- page_attendance_view_on_tree_select ---
    def page_attendance_view_on_tree_select(event=None):
        # ========================================
        sel = tv_rep.selection()
        if not sel:
            return
        item = tv_rep.item(sel[0])
        mode = self.analysis_mode.get()
        full_df = self.m.get_attendance_history()
        try:
            d_from = ent_date_from.get()
            d_to = ent_date_to.get()
            mask = (full_df["التاريخ"] >= d_from) & (full_df["التاريخ"] <= d_to)
            period_df = full_df.loc[mask]
        except:
            period_df = full_df
        else:
            if mode == "Student":
                s_name = item["values"][3]
                lbl_st_name.config(text=s_name)
                s_df = period_df[period_df["اسم الطالب"] == s_name]
                l_Stat1.config(text=(str(len(s_df[s_df["نوع العملية"] == "تأخير"]))))
                l_Stat2.config(text=(str(len(s_df[s_df["نوع العملية"] == "غياب"]))))
                l_Stat3.config(text=(str(len(s_df[s_df["نوع العملية"] == "انصراف"]))))
            else:
                emp_name = item["values"][5]
                lbl_st_name.config(text=emp_name)
                e_df = period_df[period_df["اسم الموظف المنفذ"] == emp_name]
                l_Stat1.config(text=(str(len(e_df[e_df["نوع العملية"] == "تأخير"]))), fg="#ff9800")
                l_Stat2.config(text=(str(len(e_df[e_df["نوع العملية"] == "غياب"]))), fg="#d32f2f")
                l_Stat3.config(text=(str(len(e_df[e_df["نوع العملية"] == "انصراف"]))), fg="#2196f3")
                total_school = len(period_df)
                total_emp = len(e_df)
                perc = total_emp / total_school * 100 if total_school > 0 else 0
                lbl_perc.config(text=f"%{perc:.1f}")
                prog_bar["value"] = perc

        # ========================================

    # --- page_attendance_view_on_show ---
    def page_attendance_view_on_show(event=None):
        # ========================================
        if event.widget == page:
            apply_filters()
            try:
                refresh_monitoring_list()
                refresh_unified_table()
            except:
                pass

        # ========================================

    # --- page_attendance_view_toggle_selector ---
    def page_attendance_view_toggle_selector(event=None):
        # ========================================
        t_type = cb_type.get()
        if t_type == "عمل آخر":
            for w in f_selector.winfo_children():
                w.pack_forget()
            else:
                cb_emp_single.pack()

        else:
            for w in f_selector.winfo_children():
                w.pack_forget()
            else:
                f_multi.pack()
                lb_multi.selection_clear(0, tk.END)
                roles = self.m.load_attendance_roles()
                current_assigned = roles.get(t_type, [])
                if not isinstance(current_assigned, list):
                    current_assigned = [str(current_assigned)] if current_assigned else []
                for i in range(lb_multi.size()):
                    emp_name = lb_multi.get(i)
                    if emp_name in current_assigned:
                        lb_multi.selection_set(i)
                        lb_multi.see(i)

        # ========================================

    # --- page_attendance_view_save_assignment ---
    def page_attendance_view_save_assignment():
        # ========================================
        t = cb_type.get()
        dets = ent_details.get().strip()
        if t == "عمل آخر":
            emp = cb_emp_single.get()
            if not emp:
                return messagebox.showwarning("تنبيه", "اختر موظفاً")
            elif not dets:
                return messagebox.showwarning("تنبيه", "اكتب تفاصيل المهمة")
                if self.m.add_custom_assignment(emp, dets):
                    messagebox.showinfo("تم", "تم التكليف بنجاح")
                    ent_details.delete(0, "end")
                    refresh_unified_table()
                    refresh_monitoring_list()
            else:
                messagebox.showerror("خطأ", "فشل الحفظ")
        else:
            sels = lb_multi.curselection()
            if not sels:
                return messagebox.showwarning("تنبيه", "اختر موظفاً واحداً على الأقل")
            else:
                selected_names = [lb_multi.get(i) for i in sels]
                current = self.m.load_attendance_roles()
                current[t] = selected_names
                if self.m.save_attendance_roles(current):
                    today_str = datetime.now().strftime("%Y-%m-%d")
                    self.m.update_task_status(t, today_str, "تحت العمل", "")
                    messagebox.showinfo("تم", f"تم تحديث {t} بنجاح")
                    refresh_unified_table()
                    refresh_monitoring_list()
                else:
                    messagebox.showerror("خطأ", "فشل الحفظ")

        # ========================================

    # --- page_attendance_view_on_role_double_click ---
    def page_attendance_view_on_role_double_click(event=None):
        # ========================================
        region = tree_roles.identify("region", event.x, event.y)
        if region != "cell":
            return
        else:
            sel = tree_roles.selection()
            return sel or None
        item_vals = tree_roles.item(sel[0], "values")
        r_type = item_vals[0]
        assigned = item_vals[1]
        if assigned:
            if assigned == "—":
                return
            if str(r_type).startswith("مهمة:"):
                raw_task = str(r_type).replace("مهمة: ", "").strip()
                if messagebox.askyesno("تأكيد الحذف", f"هل تريد إلغاء هذه المهمة؟\n\n📌 {raw_task}\n👤 {assigned}"):
                    if self.m.remove_custom_assignment(raw_task, assigned):
                        messagebox.showinfo("تم", "تم حذف المهمة بنجاح")
                        refresh_unified_table()
                        refresh_monitoring_list()
            else:
                messagebox.showerror("خطأ", "لم يتم العثور على المهمة أو فشل الحذف")
        else:
            pass
        if messagebox.askyesno("تأكيد سحب الصلاحية", f"هل تريد سحب صلاحية '{r_type}' من كافة الموظفين الموكلين؟\n\n(سيتم إفراغ القائمة)"):
            roles = self.m.load_attendance_roles()
            roles[r_type] = []
            if self.m.save_attendance_roles(roles):
                today_str = datetime.now().strftime("%Y-%m-%d")
                self.m.update_task_status(r_type, today_str, "تحت العمل", "")
                messagebox.showinfo("تم", "تم سحب الصلاحية بنجاح")
                refresh_unified_table()
                refresh_monitoring_list()
                lb_multi.selection_clear(0, tk.END)

        # ========================================

    # --- page_attendance_view_refresh_unified_table ---
    def page_attendance_view_refresh_unified_table():
        # ========================================
        for i in tree_roles.get_children():
            tree_roles.delete(i)
        else:
            roles = self.m.load_attendance_roles()
            for r, assigned in roles.items():
                if not assigned:
                    pass
                else:
                    emp_str = "، ".join(assigned) if isinstance(assigned, list) else str(assigned)
                    tree_roles.insert("", "end", values=(r, emp_str))
            else:
                today_str = datetime.now().strftime("%Y-%m-%d")
                custom = self.m.get_all_assignments()
                for c in reversed(custom):
                    if c.get("date") == today_str:
                        desc = f'مهمة: {c.get("task")}'
                        tree_roles.insert("", "end", values=(desc, c.get("employee")))

        # ========================================

    # --- page_attendance_view_render_assignments ---
    def page_attendance_view_render_assignments():
        # ========================================
        for w in content_area.winfo_children():
            w.destroy()
        else:
            f_top = tk.Frame(content_area, bg="white", pady=15, padx=15)
            f_top.pack(fill="both", expand=True, padx=5, pady=5)
            tk.Label(f_top, text="1. إسناد المهام والصلاحيات", font=('Segoe UI', 14, 'bold'), fg="#004d40", bg="white").pack(anchor="ne", pady=(0,
                                                                                                                                                20))
            f_form = tk.LabelFrame(f_top, text="بيانات التعميد", font=('Segoe UI', 10, 'bold'), bg="white", padx=15, pady=15)
            f_form.pack(fill="x")
            tk.Label(f_form, text="الموظف المكلف:", bg="white").grid(row=0, column=2, padx=5, pady=10, sticky="e")
            emps = self.m.get_all_employees()
            if not emps:
                emps = self.m.get_all_teachers()
            cb_emps = ttk.Combobox(f_form, values=emps, state="readonly", width=35)
            cb_emps.grid(row=0, column=1, padx=5, pady=10)
            tk.Label(f_form, text="نوع الصلاحية:", bg="white").grid(row=1, column=2, padx=5, pady=10, sticky="e")
            roles = [
             "تأخير", "غياب", "استئذان", "التواصل مع ولي الأمر الطالب الغائب"]
            cb_role = ttk.Combobox(f_form, values=roles, state="readonly", width=35)
            cb_role.grid(row=1, column=1, padx=5, pady=10)
            tk.Label(f_form, text="تاريخ الصلاحية:", bg="white").grid(row=2, column=2, padx=5, pady=10, sticky="e")
            ent_date = tk.Entry(f_form, width=38, justify="center")
            ent_date.grid(row=2, column=1, padx=5, pady=10)
            ent_date.insert(0, datetime.now().strftime("%Y-%m-%d"))
            ent_date.bind("<Button-1>", lambda e: self.helper_pick_date(ent_date))

            def do_assign():
                emp = cb_emps.get()
                role = cb_role.get()
                d = ent_date.get()
                return emp and role or messagebox.showwarning("تنبيه", "البيانات ناقصة")
                self.m.save_assignment(emp, role, d)
                messagebox.showinfo("تم", "تم إسناد الصلاحية بنجاح")
                refresh_unified_table()


            tk.Button(f_form, text="بناااء (حفظ)", command=do_assign, bg="#2e7d32", fg="white", font=('Segoe UI',
                                                                                                      10,
                                                                                                      'bold'), width=20).grid(row=3, column=1, pady=20)
            tk.Label(f_top, text="الصلاحيات النشطة لليوم", font=('Segoe UI', 12, 'bold'), fg="#455a64", bg="white").pack(anchor="ne", pady=(20,
                                                                                                                                            5))
            tree_frame = tk.Frame(f_top)
            tree_frame.pack(fill="both", expand=True)
            tree_roles = ttk.Treeview(tree_frame, columns=('role', 'emp'), show="headings", height=8)
            tree_roles.heading("role", text="الصلاحية / المهمة")
            tree_roles.heading("emp", text="الموظف المسند إليه")
            tree_roles.column("role", width=200, anchor="e")
            tree_roles.column("emp", width=250, anchor="e")
            tree_roles.pack(fill="both", expand=True)

            def refresh_unified_table():
                for i in tree_roles.get_children():
                    tree_roles.delete(i)
                else:
                    roles = self.m.load_attendance_roles()
                    for r, assigned in roles.items():
                        if not assigned:
                            pass
                        else:
                            emp_str = "، ".join(assigned) if isinstance(assigned, list) else str(assigned)
                            tree_roles.insert("", "end", values=(r, emp_str))
                    else:
                        today_str = datetime.now().strftime("%Y-%m-%d")
                        custom = self.m.get_all_assignments()
                        for c in reversed(custom):
                            if c.get("date") == today_str:
                                desc = f'مهمة: {c.get("task")}'
                                tree_roles.insert("", "end", values=(desc, c.get("employee")))


            refresh_unified_table()

            def on_role_double_click(event):
                region = tree_roles.identify("region", event.x, event.y)
                if region != "cell":
                    return
                else:
                    sel = tree_roles.selection()
                    return sel or None
                item_vals = tree_roles.item(sel[0], "values")
                r_type = item_vals[0]
                assigned = item_vals[1]
                if assigned:
                    if assigned == "—":
                        return
                    if str(r_type).startswith("مهمة:"):
                        raw_task = str(r_type).replace("مهمة: ", "").strip()
                        if messagebox.askyesno("تأكيد الحذف", f"هل تريد إلغاء هذه المهمة؟\n\n📌 {raw_task}\n👤 {assigned}"):
                            if self.m.remove_custom_assignment(raw_task, assigned):
                                messagebox.showinfo("تم", "تم حذف المهمة بنجاح")
                                refresh_unified_table()
                    else:
                        messagebox.showerror("خطأ", "لم يتم العثور على المهمة أو فشل الحذف")
                else:
                    pass
                if messagebox.askyesno("تأكيد سحب الصلاحية", f"هل تريد سحب صلاحية '{r_type}' من الموظف/الموظفين؟\n\n(سيتم إفراغ القائمة)"):
                    roles_dict = self.m.load_attendance_roles()
                    roles_dict[r_type] = []
                    if self.m.save_attendance_roles(roles_dict):
                        today_str = datetime.now().strftime("%Y-%m-%d")
                        self.m.update_task_status(r_type, today_str, "تحت العمل", "")
                        messagebox.showinfo("تم", "تم سحب الصلاحية بنجاح")
                        refresh_unified_table()


            tree_roles.bind("<Double-1>", on_role_double_click)

        # ========================================

    # --- page_attendance_view_render_assignments_do_assign ---
    def page_attendance_view_render_assignments_do_assign():
        # ========================================
        emp = cb_emps.get()
        role = cb_role.get()
        d = ent_date.get()
        return emp and role or messagebox.showwarning("تنبيه", "البيانات ناقصة")
        self.m.save_assignment(emp, role, d)
        messagebox.showinfo("تم", "تم إسناد الصلاحية بنجاح")
        refresh_unified_table()

        # ========================================

    # --- page_attendance_view_render_assignments_refresh_unified_table ---
    def page_attendance_view_render_assignments_refresh_unified_table():
        # ========================================
        for i in tree_roles.get_children():
            tree_roles.delete(i)
        else:
            roles = self.m.load_attendance_roles()
            for r, assigned in roles.items():
                if not assigned:
                    pass
                else:
                    emp_str = "، ".join(assigned) if isinstance(assigned, list) else str(assigned)
                    tree_roles.insert("", "end", values=(r, emp_str))
            else:
                today_str = datetime.now().strftime("%Y-%m-%d")
                custom = self.m.get_all_assignments()
                for c in reversed(custom):
                    if c.get("date") == today_str:
                        desc = f'مهمة: {c.get("task")}'
                        tree_roles.insert("", "end", values=(desc, c.get("employee")))

        # ========================================

    # --- page_attendance_view_render_assignments_on_role_double_click ---
    def page_attendance_view_render_assignments_on_role_double_click(event=None):
        # ========================================
        region = tree_roles.identify("region", event.x, event.y)
        if region != "cell":
            return
        else:
            sel = tree_roles.selection()
            return sel or None
        item_vals = tree_roles.item(sel[0], "values")
        r_type = item_vals[0]
        assigned = item_vals[1]
        if assigned:
            if assigned == "—":
                return
            if str(r_type).startswith("مهمة:"):
                raw_task = str(r_type).replace("مهمة: ", "").strip()
                if messagebox.askyesno("تأكيد الحذف", f"هل تريد إلغاء هذه المهمة؟\n\n📌 {raw_task}\n👤 {assigned}"):
                    if self.m.remove_custom_assignment(raw_task, assigned):
                        messagebox.showinfo("تم", "تم حذف المهمة بنجاح")
                        refresh_unified_table()
            else:
                messagebox.showerror("خطأ", "لم يتم العثور على المهمة أو فشل الحذف")
        else:
            pass
        if messagebox.askyesno("تأكيد سحب الصلاحية", f"هل تريد سحب صلاحية '{r_type}' من الموظف/الموظفين؟\n\n(سيتم إفراغ القائمة)"):
            roles_dict = self.m.load_attendance_roles()
            roles_dict[r_type] = []
            if self.m.save_attendance_roles(roles_dict):
                today_str = datetime.now().strftime("%Y-%m-%d")
                self.m.update_task_status(r_type, today_str, "تحت العمل", "")
                messagebox.showinfo("تم", "تم سحب الصلاحية بنجاح")
                refresh_unified_table()

        # ========================================

    # --- page_attendance_view_render_monitoring ---
    def page_attendance_view_render_monitoring():
        # ========================================
        for w in content_area.winfo_children():
            w.destroy()
        else:
            self.create_monitoring_dashboard_ui(content_area, is_employee_mode=False)

        # ========================================

    # --- page_attendance_view_switch_to ---
    def page_attendance_view_switch_to(mode=None):
        # ========================================
        self._mgr_view_mode = mode
        if mode == "ASSIGN":
            btn_assign.config(bg="#ffa000", fg="black")
            btn_monitor.config(bg="#455a64", fg="white")
            render_assignments()
        else:
            btn_assign.config(bg="#455a64", fg="white")
            btn_monitor.config(bg="#ffa000", fg="black")
            render_monitoring()

        # ========================================

    # --- create_monitoring_dashboard_ui ---
    def create_monitoring_dashboard_ui(self, parent_frame=None, is_employee_mode=None):
        # ========================================
        header_bg = "#ff5722"
        f_header_pill = tk.Frame(parent_frame, bg=header_bg, padx=2, pady=2)
        f_header_pill.pack(pady=(10, 10))
        f_header_inner = tk.Frame(f_header_pill, bg="white", padx=20, pady=8)
        f_header_inner.pack(fill="both", expand=True)
        title_text = "🚀 لوحة متابعة الإنجاز اليومي"
        tk.Label(f_header_inner, text=title_text, font=('Segoe UI', 12, 'bold'), fg="#bf360c", bg="white").pack(side="left")
        tk.Label(f_header_inner, text="|  نتائج ومرفقات  |", font=('Segoe UI', 10), fg="#757575", bg="white").pack(side="left", padx=10)
        tk.Label(f_header_inner, text="🕒 مباشر", font=('Segoe UI', 9, 'bold'), fg="green", bg="#e8f5e9", padx=8).pack(side="left")
        f_filters = tk.Frame(parent_frame, bg=COLOR_BG)
        f_filters.pack(fill="x", padx=15, pady=5)
        f_types = tk.Frame(f_filters, bg=COLOR_BG)
        f_types.pack(fill="x", pady=5)
        active_type_var = tk.StringVar(value="تأخير")
        all_type_btns = []

        def on_type_click(t, btn):
            active_type_var.set(t)
            for b in all_type_btns:
                b.config(bg="#f5f5f5", fg="black", relief="raised")
            else:
                btn.config(bg="#ffa000", fg="black", relief="sunken")
                refresh_dashboard()


        types = [
         "التواصل مع ولي الأمر الطالب الغائب", "استئذان", "غياب", "تأخير"]
        for t in types:
            w = 25 if len(t) > 10 else 12
            btn = tk.Button(f_types, text=t, width=w, font=('Segoe UI', 10, 'bold'), bg="#f5f5f5")
            btn.config(command=(lambda x=t, b=btn: on_type_click(x, b)))
            btn.pack(side="right", padx=5)
            all_type_btns.append(btn)
            if t == "تأخير":
                btn.config(bg="#ffa000", fg="black", relief="sunken")
            f_status = tk.Frame(f_filters, bg=COLOR_BG)
            f_status.pack(fill="x", pady=5)
            status_filter_var = tk.StringVar(value="ALL")
            all_stat_btns = []

            def on_stat_click(s, btn):
                status_filter_var.set(s)
                for b in all_stat_btns:
                    b.config(relief="raised", bg="#eceff1")
                else:
                    btn.config(relief="sunken", bg="#90a4ae")
                    refresh_dashboard()


            stats_opts = [('الكل', 'ALL'), ('⏳ جاري', 'WORKING'), ('✅ منجز', 'DONE'), ('❌ مرفوض', 'REJECTED')]
            for txt, val in stats_opts:
                btn = tk.Button(f_status, text=txt, width=8, font=('Segoe UI', 9))
                btn.config(command=(lambda x=val, b=btn: on_stat_click(x, b)))
                btn.pack(side="right", padx=2)
                all_stat_btns.append(btn)
                if val == "ALL":
                    btn.config(relief="sunken", bg="#90a4ae")
                tk.Label(f_status, text="|  من:", bg=COLOR_BG).pack(side="right", padx=5)
                ent_d1 = tk.Entry(f_status, width=10, justify="center")
                ent_d1.pack(side="right", padx=2)
                ent_d1.insert(0, datetime.now().strftime("%Y-%m-%d"))
                ent_d1.bind("<Button-1>", lambda e: self.helper_pick_date(ent_d1, refresh_dashboard))
                tk.Label(f_status, text="إلى:", bg=COLOR_BG).pack(side="right", padx=5)
                ent_d2 = tk.Entry(f_status, width=10, justify="center")
                ent_d2.pack(side="right", padx=2)
                ent_d2.insert(0, datetime.now().strftime("%Y-%m-%d"))
                ent_d2.bind("<Button-1>", lambda e: self.helper_pick_date(ent_d2, refresh_dashboard))

                def set_today():
                    today = datetime.now().strftime("%Y-%m-%d")
                    ent_d1.delete(0, "end")
                    ent_d1.insert(0, today)
                    ent_d2.delete(0, "end")
                    ent_d2.insert(0, today)
                    refresh_dashboard()


                tk.Button(f_status, text="📅 اليوم فقط", command=set_today, bg="#ff9800", fg="white", font=('Segoe UI',
                                                                                                           9,
                                                                                                           'bold')).pack(side="left", padx=5)
                tk.Button(f_status, text="🔄 تحديث", command=(lambda: refresh_dashboard()), bg="#607d8b", fg="white", font=('Segoe UI',
                                                                                                                           9,
                                                                                                                           'bold')).pack(side="left", padx=5)
                list_container = tk.Frame(parent_frame, bg="white", highlightthickness=1, highlightbackground="#ccc")
                list_container.pack(fill="both", expand=True, padx=15, pady=10)
                list_inner_frame = tk.Frame(list_container, bg="white")
                list_inner_frame.pack(fill="both", expand=True, padx=5, pady=5)

                def refresh_dashboard(*args, **kwargs): pass  # TODO: Reconstruct method


                refresh_dashboard()
                tk.Button(parent_frame, text="رجوع", command=(self.show_manager_hub)).pack(pady=10)
                return parent_frame# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- create_monitoring_dashboard_ui_on_type_click ---
    def create_monitoring_dashboard_ui_on_type_click(t=None, btn=None):
        # ========================================
        active_type_var.set(t)
        for b in all_type_btns:
            b.config(bg="#f5f5f5", fg="black", relief="raised")
        else:
            btn.config(bg="#ffa000", fg="black", relief="sunken")
            refresh_dashboard()

        # ========================================

    # --- create_monitoring_dashboard_ui_on_stat_click ---
    def create_monitoring_dashboard_ui_on_stat_click(s=None, btn=None):
        # ========================================
        status_filter_var.set(s)
        for b in all_stat_btns:
            b.config(relief="raised", bg="#eceff1")
        else:
            btn.config(relief="sunken", bg="#90a4ae")
            refresh_dashboard()

        # ========================================

    # --- create_monitoring_dashboard_ui_set_today ---
    def create_monitoring_dashboard_ui_set_today():
        # ========================================
        today = datetime.now().strftime("%Y-%m-%d")
        ent_d1.delete(0, "end")
        ent_d1.insert(0, today)
        ent_d2.delete(0, "end")
        ent_d2.insert(0, today)
        refresh_dashboard()

        # ========================================

    # --- create_monitoring_dashboard_ui_refresh_dashboard ---
    def create_monitoring_dashboard_ui_refresh_dashboard():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- create_monitoring_dashboard_ui_refresh_dashboard__on_config ---
    def create_monitoring_dashboard_ui_refresh_dashboard__on_config(e=None):
        # ========================================
        canvas.configure(scrollregion=(canvas.bbox("all")))

        # ========================================

    # --- create_monitoring_dashboard_ui_refresh_dashboard__on_width ---
    def create_monitoring_dashboard_ui_refresh_dashboard__on_width(e=None):
        # ========================================
        canvas.itemconfig((canvas.find_withtag("all")[0]), width=(e.width))

        # ========================================

    # --- create_monitoring_dashboard_ui_refresh_dashboard_do_approve ---
    def create_monitoring_dashboard_ui_refresh_dashboard_do_approve(b_row=None, val=None):
        # ========================================
        try:
            self.m.update_batch_status(b_row["التاريخ"], b_row["الوقت"], b_row["نوع العملية"], val)
            refresh_dashboard()
            messagebox.showinfo("تم", f"تم {val} السجل بنجاح")
        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"فشل التحديث: {e}")
            finally:
                pass

        # ========================================

    # --- create_monitoring_dashboard_ui_refresh_dashboard_show_details ---
    def create_monitoring_dashboard_ui_refresh_dashboard_show_details(b_row=None):
        # ========================================
        top = tk.Toplevel(parent_frame)
        top.title(f'تفاصيل - {b_row["نوع العملية"]}')
        top.geometry("700x500")
        top.configure(bg="#f5f5f5")
        h_fr = tk.Frame(top, bg="#37474f", pady=10)
        h_fr.pack(fill="x")
        tk.Label(h_fr, text=f'سجل: {b_row["نوع العملية"]}', font=('Segoe UI', 12, 'bold'), fg="white", bg="#37474f").pack()
        tk.Label(h_fr, text=f'بواسطة: {b_row["اسم الموظف المنفذ"]}', fg="#b0bec5", bg="#37474f").pack()
        mask = (df["التاريخ"] == b_row["التاريخ"]) & (df["الوقت"] == b_row["الوقت"]) & (df["نوع العملية"] == b_row["نوع العملية"])
        subset_full = df[mask].copy()
        f_flt = tk.Frame(top, bg="#f5f5f5", pady=5)
        f_flt.pack(fill="x")
        tk.Label(f_flt, text="تصفية حسب:", bg="#f5f5f5", font=('Segoe UI', 10, 'bold')).pack(side="right", padx=10)
        avail_classes = sorted(list(set(subset_full["الصف"].astype(str).unique())))
        avail_sections = sorted(list(set(subset_full["الشعبة"].astype(str).unique())))
        cb_cls = ttk.Combobox(f_flt, values=(["الكل"] + avail_classes), state="readonly", width=15)
        cb_cls.set("الكل")
        cb_cls.pack(side="right", padx=5)
        cb_sec = ttk.Combobox(f_flt, values=(["الكل"] + avail_sections), state="readonly", width=10)
        cb_sec.set("الكل")
        cb_sec.pack(side="right", padx=5)
        cols = ('std', 'gr', 'sec', 'time', 'rsn')
        tv = ttk.Treeview(top, columns=cols, show="headings")
        tv.heading("std", text="الطالب")
        tv.column("std", width=180, anchor="e")
        tv.heading("gr", text="الصف")
        tv.column("gr", width=80, anchor="center")
        tv.heading("sec", text="الشعبة")
        tv.column("sec", width=80, anchor="center")
        tv.heading("time", text="وقت الرصد")
        tv.column("time", width=80, anchor="center")
        tv.heading("rsn", text="الحالة")
        tv.column("rsn", width=150, anchor="e")
        tv.pack(fill="both", expand=True, padx=10, pady=10)
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview.Heading(", font=('Segoe UI', 10, 'bold'), background=")#cfd8dc")

        def refresh_popup_list(*args):
            for i in tv.get_children():
                tv.delete(i)
            else:
                c_filter = cb_cls.get()
                s_filter = cb_sec.get()
                temp = subset_full.copy()
                if c_filter != "الكل":
                    temp = temp[temp["الصف"].astype(str) == c_filter]
                if s_filter != "الكل":
                    temp = temp[temp["الشعبة"].astype(str) == s_filter]
                count = 0
                for i, (_, r) in enumerate(temp.iterrows()):
                    tag = "even" if i % 2 == 0 else "odd"
                    sec = str(r["الشعبة"])
                    if not (sec == "nan" or sec):
                        sec = "غير محدد"
                    tv.insert("", "end", values=(r["اسم الطالب"], r["الصف"], sec, r["الوقت"], r["الحالة التفصيلية"]), tags=(tag,))
                    count += 1
                else:
                    top.title(f'تفاصيل ({count}) - {b_row["نوع العملية"]}')


        cb_cls.bind("<<ComboboxSelected>>", refresh_popup_list)
        cb_sec.bind("<<ComboboxSelected>>", refresh_popup_list)
        refresh_popup_list()
        tv.tag_configure("odd", background="white")
        tv.tag_configure("even", background="#e0f2f1")
        tk.Button(top, text="إغلاق", command=(top.destroy), bg="#d32f2f", fg="white", width=15).pack(pady=10)

        # ========================================

    # --- create_monitoring_dashboard_ui_refresh_dashboard_show_details_refresh_popup_list ---
    def create_monitoring_dashboard_ui_refresh_dashboard_show_details_refresh_popup_list(*args):
        # ========================================
        for i in tv.get_children():
            tv.delete(i)
        else:
            c_filter = cb_cls.get()
            s_filter = cb_sec.get()
            temp = subset_full.copy()
            if c_filter != "الكل":
                temp = temp[temp["الصف"].astype(str) == c_filter]
            if s_filter != "الكل":
                temp = temp[temp["الشعبة"].astype(str) == s_filter]
            count = 0
            for i, (_, r) in enumerate(temp.iterrows()):
                tag = "even" if i % 2 == 0 else "odd"
                sec = str(r["الشعبة"])
                if not (sec == "nan" or sec):
                    sec = "غير محدد"
                tv.insert("", "end", values=(r["اسم الطالب"], r["الصف"], sec, r["الوقت"], r["الحالة التفصيلية"]), tags=(tag,))
                count += 1
            else:
                top.title(f'تفاصيل ({count}) - {b_row["نوع العملية"]}')

        # ========================================

    # --- create_monitoring_dashboard_ui_refresh_dashboard_show_contact_details_popup ---
    def create_monitoring_dashboard_ui_refresh_dashboard_show_contact_details_popup(date_str=None, emp_name=None):
        # ========================================
        logs = self.m.get_contact_logs()
        record = next((r for r in logs if r.get("employee") == emp_name), None)
        top = tk.Toplevel(parent_frame)
        top.title(f"سجل تواصل: {emp_name} - {date_str}")
        top.geometry("700x500")
        if not record:
            tk.Label(top, text="لم يتم العثور على سجل تفصيلي", fg="red").pack(pady=20)
            return
        tk.Label(top, text=f"سجل إنجاز التواصل - {emp_name}", font=('Segoe UI', 12, 'bold'), fg="#7b1fa2").pack(pady=10)
        tk.Label(top, text=f"التاريخ: {date_str}", font=('Segoe UI', 10)).pack()
        cols = ('name', 'mobile', 'status', 'note')
        tv = ttk.Treeview(top, columns=cols, show="headings")
        tv.heading("name", text="اسم الطالب")
        tv.column("name", width=180, anchor="e")
        tv.heading("mobile", text="الجوال")
        tv.column("mobile", width=100, anchor="center")
        tv.heading("status", text="حالة التواصل")
        tv.column("status", width=100, anchor="center")
        tv.heading("note", text="ملاحظات")
        tv.column("note", width=250, anchor="e")
        tv.pack(fill="both", expand=True, padx=10, pady=10)
        for d in record.get("details", []):
            st = d.get("status")
            st_txt = "✅ تم" if st == "done" else "❌ لم يرد"
            tags = ('done', ) if st == "done" else ('no', )
            tv.insert("", "end", values=(d.get("name"), d.get("mobile"), st_txt, d.get("notes")), tags=tags)
        else:
            tv.tag_configure("done", foreground="green")
            tv.tag_configure("no", foreground="red")
            tk.Button(top, text="إغلاق", command=(top.destroy)).pack(pady=10)

        # ========================================

    # --- page_monitoring ---
    def page_monitoring(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        scroll_canvas = tk.Canvas(page, bg=COLOR_BG, highlightthickness=0)
        vsb = ttk.Scrollbar(page, orient="vertical", command=(scroll_canvas.yview))
        scroll_frame = tk.Frame(scroll_canvas, bg=COLOR_BG)
        scroll_canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        scroll_canvas.configure(yscrollcommand=(vsb.set))
        scroll_frame.bind("<Configure>", lambda e: scroll_canvas.configure(scrollregion=(scroll_canvas.bbox("all"))))
        vsb.pack(side="left", fill="y")
        scroll_canvas.pack(side="right", fill="both", expand=True)

        def on_resize(e):
            scroll_canvas.itemconfig(1, width=(e.width))


        scroll_canvas.bind("<Configure>", on_resize)
        header = tk.Frame(scroll_frame, bg=COLOR_PANEL, padx=30, pady=20)
        header.pack(fill="x", padx=20, pady=20)
        tk.Label(header, text="🛰️ لوحة المراقبة الذكية", font=('Segoe UI', 20, 'bold'), bg=COLOR_PANEL, fg=COLOR_ACCENT).pack(side="right")
        tk.Label(header, text="بث مباشر لحالة المدرسة والجدول الدراسي", font=('Segoe UI', 10), bg=COLOR_PANEL, fg="#888").pack(side="right", padx=15, pady=8)
        stat_grid = tk.Frame(scroll_frame, bg=COLOR_BG)
        stat_grid.pack(fill="x", padx=20)

        def make_stat(title, val, unit, color):
            f = tk.Frame(stat_grid, bg=COLOR_PANEL, padx=20, pady=20)
            f.pack(side="right", fill="both", expand=True, padx=10)
            tk.Label(f, text=title, bg=COLOR_PANEL, font=('Segoe UI', 11), fg="#666").pack()
            tk.Label(f, text=(str(val)), bg=COLOR_PANEL, font=('Segoe UI', 28, 'bold'), fg=color).pack()
            tk.Label(f, text=unit, bg=COLOR_PANEL, font=('Segoe UI', 10), fg="#aaa").pack()


        summary = self.m.calendar_summary()
        make_stat("إجمالي الطلاب", len(self.m.df_students), "طالب مسجل", COLOR_ACCENT)
        make_stat("كادر المعلمين", len(self.m.teacher_names), "معلم رسمي", "#1976d2")
        n_col, c_col, s_col = self.m._detect_student_cols()
        if c_col and c_col in self.m.df_students.columns:
            group_cols = [
             c_col]
            if s_col:
                if s_col in self.m.df_students.columns:
                    group_cols.append(s_col)
                df_groups = self.m.df_students[group_cols].drop_duplicates()
                make_stat("الفصول الدراسية", len(df_groups), "فصل/شعبة فني", "#26a69a")
            else:
                make_stat("الفصول الدراسية", 0, "فصل/شعبة فني", "#26a69a")
            make_stat("المواد الدراسية", len(self.m.df_subjects), "مادة معتمدة", "#7e57c2")
            prog = self.m.get_academic_progress()
            pct_text = f'{prog["year_pct"]}%' if (prog and "year_pct" in prog) else "0%"
            make_stat("التقدم الدراسي", pct_text, "من العام الدراسي", "#fb8c00")
            dist_section = tk.Frame(scroll_frame, bg=COLOR_BG, pady=20)
            dist_section.pack(fill="x", padx=20)
            tk.Label(dist_section, text="📊 كثافة الطلاب وتوزع المناهج", font=('Segoe UI', 13,
                                                                              'bold'), bg=COLOR_BG, fg="#333").pack(anchor="e")
            dist_box = tk.Frame(dist_section, bg=COLOR_PANEL, padx=15, pady=15)
            dist_box.pack(fill="x", pady=10)
            if not self.m.df_students.empty:
                if c_col and c_col in self.m.df_students.columns:
                    tk.Label(dist_box, text="توزيع الطلاب حسب الصفوف", font=('Segoe UI', 11,
                                                                             'bold'), bg=COLOR_PANEL, fg=COLOR_ACCENT).pack(anchor="e", pady=(0,
                                                                                                                                              5))
                    counts = self.m.df_students[c_col].value_counts().sort_index()
                    max_c = counts.max() if not counts.empty else 1
                    for g, count in counts.items():
                        row = tk.Frame(dist_box, bg=COLOR_PANEL)
                        row.pack(fill="x", pady=4)
                        tk.Label(row, text=(str(g)), bg=COLOR_PANEL, font=('Segoe UI', 10,
                                                                           'bold'), width=15, anchor="e").pack(side="right")
                        pct = count / max_c
                        bar_bg = tk.Frame(row, bg="#eeeeee", height=12)
                        bar_bg.pack(side="right", fill="x", expand=True, padx=15)
                        bar_bg.pack_propagate(False)
                        tk.Frame(bar_bg, bg=COLOR_ACCENT).place(relx=1.0, rely=0, relwidth=pct, relheight=1.0, anchor="ne")
                        tk.Label(row, text=f"{count} طالب", bg=COLOR_PANEL, font=('Segoe UI',
                                                                                  9), fg="#555").pack(side="left")

                    if s_col:
                        if s_col in self.m.df_students.columns:
                            tk.Frame(dist_box, height=1, bg="#ddd").pack(fill="x", pady=15)
                            tk.Label(dist_box, text="توزيع الطلاب حسب الشعب", font=('Segoe UI',
                                                                                    11, 'bold'), bg=COLOR_PANEL, fg="#0277bd").pack(anchor="e", pady=(0,
                                                                                                                                                      5))
                            df_sh = self.m.df_students.copy()
                            df_sh["combo"] = df_sh[c_col].astype(str) + " - " + df_sh[s_col].astype(str)
                            sh_counts = df_sh["combo"].value_counts().sort_index()
                            max_sh = sh_counts.max() if not sh_counts.empty else 1
                            for combo, count in sh_counts.items():
                                row = tk.Frame(dist_box, bg=COLOR_PANEL)
                                row.pack(fill="x", pady=4)
                                tk.Label(row, text=(str(combo)), bg=COLOR_PANEL, font=('Segoe UI',
                                                                                       9), width=20, anchor="e").pack(side="right")
                                pct = count / max_sh
                                bar_bg = tk.Frame(row, bg="#eeeeee", height=10)
                                bar_bg.pack(side="right", fill="x", expand=True, padx=15)
                                bar_bg.pack_propagate(False)
                                tk.Frame(bar_bg, bg="#0277bd").place(relx=1.0, rely=0, relwidth=pct, relheight=1.0, anchor="ne")
                                tk.Label(row, text=f"{count} طالب", bg=COLOR_PANEL, font=('Segoe UI',
                                                                                          8), fg="#555").pack(side="left")

        else:
            tk.Label(dist_box, text="لا يوجد بيانات طلاب حالياً لعرض التوزيع.", bg=COLOR_PANEL, fg="#999").pack()
        live_section = tk.Frame(scroll_frame, bg=COLOR_BG, pady=20)
        live_section.pack(fill="x", padx=20)
        controls = tk.Frame(live_section, bg=COLOR_PANEL, padx=15, pady=15)
        controls.pack(fill="x", pady=10)
        tk.Label(controls, text="خيارات العرض:", bg=COLOR_PANEL, font=('Segoe UI', 11, 'bold')).pack(side="right", padx=10)
        mode_var = tk.StringVar(value="live")

        def toggle_mode():
            if mode_var.get() == "live":
                cb_day.config(state="disabled")
                cb_period.config(state="disabled")
                btn_refresh.config(state="disabled", bg="#ccc")
                update_dash()
            else:
                cb_day.config(state="readonly")
                cb_period.config(state="readonly")
                btn_refresh.config(state="normal", bg=COLOR_ACCENT)


        tk.Radiobutton(controls, text="الوضع المباشر (تلقائي)", variable=mode_var, value="live", command=toggle_mode, bg=COLOR_PANEL, font=('Segoe UI',
                                                                                                                                            10)).pack(side="right", padx=10)
        tk.Radiobutton(controls, text="اختيار يدوي", variable=mode_var, value="manual", command=toggle_mode, bg=COLOR_PANEL, font=('Segoe UI',
                                                                                                                                   10)).pack(side="right", padx=10)
        sel_frame = tk.Frame(controls, bg=COLOR_PANEL)
        sel_frame.pack(side="right", padx=20)
        tk.Label(sel_frame, text="اليوم:", bg=COLOR_PANEL).pack(side="right", padx=5)
        cb_day = ttk.Combobox(sel_frame, values=DAYS, state="disabled", width=10)
        cb_day.pack(side="right", padx=5)
        cb_day.set(DAYS[0])
        tk.Label(sel_frame, text="الحصة:", bg=COLOR_PANEL).pack(side="right", padx=5)
        cb_period = ttk.Combobox(sel_frame, values=(self.m.get_period_labels()), state="disabled", width=5)
        cb_period.pack(side="right", padx=5)
        cb_period.current(0)
        btn_refresh = tk.Button(controls, text="عرض", command=(lambda: update_dash(force_manual=True)), state="disabled", bg="#ccc", fg="white", font=('Segoe UI',
                                                                                                                                                       9,
                                                                                                                                                       'bold'))
        btn_refresh.pack(side="right", padx=15)
        now_card = tk.Frame(live_section, bg="#1b5e20", padx=30, pady=30, relief="flat")
        now_card.pack(fill="x", pady=10)
        self.lbl_live_header = tk.Label(now_card, text="حالة المدرسة: الوضع المباشر 🔴", font=('Segoe UI',
                                                                                              14,
                                                                                              'bold'), bg="#1b5e20", fg="#ffca28")
        self.lbl_live_header.pack(anchor="e")
        self.lbl_live_status = tk.Label(now_card, text="جاري جلب البيانات...", font=('Segoe UI',
                                                                                     22,
                                                                                     'bold'), bg="#1b5e20", fg="white")
        self.lbl_live_status.pack(anchor="e", pady=10)
        live_list_frame = tk.Frame(scroll_frame, bg=COLOR_BG)
        live_list_frame.pack(fill="x", padx=20, pady=20)
        tk.Label(live_list_frame, text="👥 توزيع المعلمين والحصص", font=('Segoe UI', 13, 'bold'), bg=COLOR_BG, fg="#333").pack(anchor="e", pady=10)
        self.live_activity_container = tk.Frame(live_list_frame, bg=COLOR_BG)
        self.live_activity_container.pack(fill="x")

        def update_dash(force_manual=False):
            if not page.winfo_ismapped():
                return
                is_manual = mode_var.get() == "manual"
                target_day_name = None
                target_period = None
                if is_manual:
                    target_day_name = cb_day.get()
                    target_period = cb_period.get()
                    self.lbl_live_header.config(text=f"حالة المدرسة: عرض أرشيفي ({target_day_name} - حصة {target_period}) 📂")
                    self.lbl_live_status.config(text=f"عرض حصص يوم {target_day_name} - الحصة {target_period}")
                else:
                    self.lbl_live_header.config(text="حالة المدرسة: الوضع المباشر 🔴")
                if self.m.teacher_names:
                    from datetime import datetime
                    curr_time = datetime.now().time()
                    found_p = None
                    for _, row in self.m.df_timings.iterrows():
                        try:
                            t_start = datetime.strptime(str(row.get("بداية الحصة", "")), "%H:%M").time()
                            t_end = datetime.strptime(str(row.get("نهاية الحصة", "")), "%H:%M").time()
                            if t_start <= curr_time <= t_end:
                                found_p = row.get("رقم الحصة") or row.get("الحصة")
                                break
                        except:
                            pass

                    else:
                        if found_p:
                            target_period = str(found_p)
                            target_day_name = arabic_day_from_english(datetime.now().strftime("%A"))
                            self.lbl_live_status.config(text=f"نحن الآن في الحصة: {target_period}")
                        else:
                            self.lbl_live_status.config(text="لا يوجد حصص دراسية الآن (اوف لاين) 💤")
                            target_period = None

            else:
                for w in self.live_activity_container.winfo_children():
                    w.destroy()
                else:
                    active_count = 0
                    if target_day_name:
                        if target_period:
                            import openpyxl
                            try:
                                wb = openpyxl.load_workbook(FILE_MASTER, read_only=True)
                                day_map = {d: i + 2 for i, d in enumerate(DAYS)}
                                row_idx = day_map.get(target_day_name)
                                p_labels = self.m.get_period_labels()
                                try:
                                    col_idx = p_labels.index(str(target_period)) + 2
                                except:
                                    col_idx = -1
                                else:
                                    if row_idx:
                                        if col_idx > 1:
                                            for tname in wb.sheetnames:
                                                val = wb[tname].cell(row=row_idx, column=col_idx).value
                                                if val and str(val).strip() not in ('—', '',
                                                                                    'None'):
                                                    active_count += 1
                                                    f = tk.Frame((self.live_activity_container), bg=COLOR_PANEL, padx=15, pady=8)
                                                    f.pack(fill="x", pady=2)
                                                    tk.Label(f, text=f"📍 {tname}", bg=COLOR_PANEL, font=('Segoe UI',
                                                                                                         10,
                                                                                                         'bold'), fg=COLOR_ACCENT).pack(side="right")
                                                    tk.Label(f, text=f"المادة/الصف: {val}", bg=COLOR_PANEL, font=('Segoe UI',
                                                                                                                  10), fg="#555").pack(side="right", padx=20)

                            except Exception as e:
                                try:
                                    print(f"Monitor Error: {e}")
                                finally:
                                    pass

                    if active_count == 0:
                        tk.Label((self.live_activity_container), text="لا يوجد حصص مسجلة في هذا الوقت.", bg=COLOR_BG, fg="#888").pack(pady=10)
                    is_manual or page.after(60000, update_dash)


        page.after(100, update_dash)
        return page

        # ========================================

    # --- page_monitoring_on_resize ---
    def page_monitoring_on_resize(e=None):
        # ========================================
        scroll_canvas.itemconfig(1, width=(e.width))

        # ========================================

    # --- page_monitoring_make_stat ---
    def page_monitoring_make_stat(title=None, val=None, unit=None, color=None):
        # ========================================
        f = tk.Frame(stat_grid, bg=COLOR_PANEL, padx=20, pady=20)
        f.pack(side="right", fill="both", expand=True, padx=10)
        tk.Label(f, text=title, bg=COLOR_PANEL, font=('Segoe UI', 11), fg="#666").pack()
        tk.Label(f, text=(str(val)), bg=COLOR_PANEL, font=('Segoe UI', 28, 'bold'), fg=color).pack()
        tk.Label(f, text=unit, bg=COLOR_PANEL, font=('Segoe UI', 10), fg="#aaa").pack()

        # ========================================

    # --- page_monitoring_toggle_mode ---
    def page_monitoring_toggle_mode():
        # ========================================
        if mode_var.get() == "live":
            cb_day.config(state="disabled")
            cb_period.config(state="disabled")
            btn_refresh.config(state="disabled", bg="#ccc")
            update_dash()
        else:
            cb_day.config(state="readonly")
            cb_period.config(state="readonly")
            btn_refresh.config(state="normal", bg=COLOR_ACCENT)

        # ========================================

    # --- page_monitoring_update_dash ---
    def page_monitoring_update_dash(force_manual=None):
        # ========================================
        if not page.winfo_ismapped():
            return
            is_manual = mode_var.get() == "manual"
            target_day_name = None
            target_period = None
            if is_manual:
                target_day_name = cb_day.get()
                target_period = cb_period.get()
                self.lbl_live_header.config(text=f"حالة المدرسة: عرض أرشيفي ({target_day_name} - حصة {target_period}) 📂")
                self.lbl_live_status.config(text=f"عرض حصص يوم {target_day_name} - الحصة {target_period}")
            else:
                self.lbl_live_header.config(text="حالة المدرسة: الوضع المباشر 🔴")
            if self.m.teacher_names:
                from datetime import datetime
                curr_time = datetime.now().time()
                found_p = None
                for _, row in self.m.df_timings.iterrows():
                    try:
                        t_start = datetime.strptime(str(row.get("بداية الحصة", "")), "%H:%M").time()
                        t_end = datetime.strptime(str(row.get("نهاية الحصة", "")), "%H:%M").time()
                        if t_start <= curr_time <= t_end:
                            found_p = row.get("رقم الحصة") or row.get("الحصة")
                            break
                    except:
                        pass

                else:
                    if found_p:
                        target_period = str(found_p)
                        target_day_name = arabic_day_from_english(datetime.now().strftime("%A"))
                        self.lbl_live_status.config(text=f"نحن الآن في الحصة: {target_period}")
                    else:
                        self.lbl_live_status.config(text="لا يوجد حصص دراسية الآن (اوف لاين) 💤")
                        target_period = None

        else:
            for w in self.live_activity_container.winfo_children():
                w.destroy()
            else:
                active_count = 0
                if target_day_name:
                    if target_period:
                        import openpyxl
                        try:
                            wb = openpyxl.load_workbook(FILE_MASTER, read_only=True)
                            day_map = {d: i + 2 for i, d in enumerate(DAYS)}
                            row_idx = day_map.get(target_day_name)
                            p_labels = self.m.get_period_labels()
                            try:
                                col_idx = p_labels.index(str(target_period)) + 2
                            except:
                                col_idx = -1
                            else:
                                if row_idx:
                                    if col_idx > 1:
                                        for tname in wb.sheetnames:
                                            val = wb[tname].cell(row=row_idx, column=col_idx).value
                                            if val and str(val).strip() not in ('—', '', 'None'):
                                                active_count += 1
                                                f = tk.Frame((self.live_activity_container), bg=COLOR_PANEL, padx=15, pady=8)
                                                f.pack(fill="x", pady=2)
                                                tk.Label(f, text=f"📍 {tname}", bg=COLOR_PANEL, font=('Segoe UI',
                                                                                                     10,
                                                                                                     'bold'), fg=COLOR_ACCENT).pack(side="right")
                                                tk.Label(f, text=f"المادة/الصف: {val}", bg=COLOR_PANEL, font=('Segoe UI',
                                                                                                              10), fg="#555").pack(side="right", padx=20)

                        except Exception as e:
                            try:
                                print(f"Monitor Error: {e}")
                            finally:
                                pass

                if active_count == 0:
                    tk.Label((self.live_activity_container), text="لا يوجد حصص مسجلة في هذا الوقت.", bg=COLOR_BG, fg="#888").pack(pady=10)
                is_manual or page.after(60000, update_dash)

        # ========================================

    # --- refresh_home_dynamic_content ---
    def refresh_home_dynamic_content(self):
        # ========================================
        if not hasattr(self, "home_status_container"):
            return
        for w in self.home_status_container.winfo_children():
            w.destroy()
        else:
            for w in self.home_info_container.winfo_children():
                w.destroy()
            else:
                status_card = tk.Frame((self.home_status_container), bg=COLOR_PANEL, relief="flat", padx=25, pady=20)
                status_card.pack(fill="x")
                status_header = tk.Frame(status_card, bg=COLOR_PANEL)
                status_header.pack(fill="x")
                ttk.Label(status_header, text="📍 حالة اليوم الدراسي والجدول الدراسي", font=('Segoe UI',
                                                                                            13,
                                                                                            'bold'), background=COLOR_PANEL).pack(side="right")
                st = self.m.get_day_status(date.today())
                date_txt = f'{st.get("hijri", "")} هـ | {date.today().strftime("%Y/%m/%d")} م'
                tk.Label(status_header, text=date_txt, bg=COLOR_PANEL, fg="#666", font=('Segoe UI',
                                                                                        11)).pack(side="left", padx=20)
                if self.admin_mode:
                    prog = self.m.get_academic_progress()
                    if prog:
                        p_frame = tk.Frame(status_card, bg="#f1f8e9", padx=15, pady=10)
                        p_frame.pack(fill="x", pady=(15, 0))
                        prog_grid = tk.Frame(p_frame, bg="#f1f8e9")
                        prog_grid.pack(fill="x")

                        def add_bar_wide(parent, title, curr, total, pct, color, is_main=False):
                            f_box = tk.Frame(parent, bg="#f1f8e9")
                            f_box.pack(side="right", fill="x", expand=True, padx=10)
                            tk.Label(f_box, text=title, bg="#f1f8e9", fg="#333", font=("Segoe UI", 11 if is_main else 10, "bold")).pack(anchor="e")
                            stats_txt = f"أسبوع {curr} من {total} ({pct}%)"
                            tk.Label(f_box, text=stats_txt, bg="#f1f8e9", fg=color, font=("Segoe UI", 10 if is_main else 9, "bold")).pack(anchor="e", pady=(0,
                                                                                                                                                            2))
                            bar_container = tk.Frame(f_box, bg="#e0e0e0", height=(10 if is_main else 8), bd=0, relief="flat")
                            bar_container.pack(fill="x")
                            bar_container.pack_propagate(False)
                            inner_bar = tk.Frame(bar_container, bg=color)
                            inner_bar.place(relx=1.0, rely=0, relwidth=(pct / 100.0), relheight=1.0, anchor="ne")


                        add_bar_wide(prog_grid, "📊 المسيرة الدراسية للعام", (prog["year_curr"]), (prog["year_total"]), (prog["year_pct"]), COLOR_ACCENT, is_main=True)
                        for t_p in prog["terms"]:
                            p_color = COLOR_ACCENT if t_p["is_current"] else COLOR_BTN
                            p_icon = "📍" if t_p["is_current"] else "📅"
                            label_text = f'{p_icon} {t_p["name"]}'
                            d_left = t_p.get("days_left", 0)
                            if d_left >= 0:
                                if t_p["pct"] < 100:
                                    label_text += f" (باقي {d_left} يوم)"
                            add_bar_wide(prog_grid, label_text, t_p["curr"], t_p["total"], t_p["pct"], p_color)

                links_card = tk.Frame((self.home_info_container), bg=COLOR_PANEL, relief="flat", borderwidth=0, padx=25, pady=20)
                links_card.pack(fill="x")
                ttk.Label(links_card, text="🔗 روابط سريعة", font=('Segoe UI', 12, 'bold'), background=COLOR_PANEL).pack(anchor="e", pady=(0,
                                                                                                                                          5))
                links = [
                 ('نظام نور التعليمي', 'https://noor.moe.gov.sa/Noor/Login.aspx', '#1b5e20'),
                 ('منصة مدرستي الرقمية', 'https://schools.madrasati.sa/', '#2e7d32'),
                 ('نظام فارس (الخدمة الذاتية)', 'https://sshr.moe.gov.sa/', '#388e3c')]
                for name, url, color in links:
                    btn = tk.Button(links_card, text=name, bg=color, fg="white", font=('Segoe UI',
                                                                                       10,
                                                                                       'bold'), relief="flat", cursor="hand2", pady=8, command=(lambda u=url: webbrowser.open_new_tab(u)))
                    btn.pack(fill="x", pady=3, padx=5)
                    btn.bind("<Enter>", lambda e, b=btn: b.config(bg="#81c784"))
                    btn.bind("<Leave>", lambda e, b=btn, c=color: b.config(bg=c))
                else:
                    next_h, remaining = self.m.get_days_until_next_holiday()
                    if next_h is None:
                        next_h = "لا توجد أحداث قادمة"
                        remaining = 0
                    h_card = tk.Frame((self.home_info_container), bg="#e8f5e9", padx=15, pady=20, highlightthickness=1, highlightbackground="#c8e6c9")
                    h_card.pack(fill="x", pady=(15, 0))
                    tk.Label(h_card, text="⏳ الحدث القادم", bg="#e8f5e9", fg="#2e7d32", font=('Segoe UI',
                                                                                              10,
                                                                                              'bold')).pack(anchor="e")
                    clean_h = str(next_h).strip()
                    lbl_h_name = tk.Label(h_card, text=clean_h, bg="#e8f5e9", fg="#1b5e20", font=('Segoe UI',
                                                                                                  14,
                                                                                                  'bold'), wraplength=400, justify="center")
                    lbl_h_name.pack(side="top", pady=(5, 10))
                    try:
                        rem_val = int(remaining)
                        rem_txt = f"باقي {rem_val} يوم"
                    except:
                        rem_txt = f"باقي {remaining}"
                    else:
                        lbl_rem = tk.Label(h_card, text=rem_txt, bg="#e8f5e9", fg="#e65100", font=('Segoe UI',
                                                                                                   18,
                                                                                                   'bold'))
                        lbl_rem.pack(side="top", pady=(0, 5))

        # ========================================

    # --- refresh_home_dynamic_content_add_bar_wide ---
    def refresh_home_dynamic_content_add_bar_wide(parent=None, title=None, curr=None, total=None, pct=None, color=None, is_main=None):
        # ========================================
        f_box = tk.Frame(parent, bg="#f1f8e9")
        f_box.pack(side="right", fill="x", expand=True, padx=10)
        tk.Label(f_box, text=title, bg="#f1f8e9", fg="#333", font=("Segoe UI", 11 if is_main else 10, "bold")).pack(anchor="e")
        stats_txt = f"أسبوع {curr} من {total} ({pct}%)"
        tk.Label(f_box, text=stats_txt, bg="#f1f8e9", fg=color, font=("Segoe UI", 10 if is_main else 9, "bold")).pack(anchor="e", pady=(0,
                                                                                                                                        2))
        bar_container = tk.Frame(f_box, bg="#e0e0e0", height=(10 if is_main else 8), bd=0, relief="flat")
        bar_container.pack(fill="x")
        bar_container.pack_propagate(False)
        inner_bar = tk.Frame(bar_container, bg=color)
        inner_bar.place(relx=1.0, rely=0, relwidth=(pct / 100.0), relheight=1.0, anchor="ne")

        # ========================================

    # --- page_teacher_grid ---
    def page_teacher_grid(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        top = tk.Frame(page, bg=COLOR_BG)
        top.pack(fill="x", pady=(0, 8))
        ttk.Label(top, text="اختر معلماً:", foreground="#333").pack(side="right", padx=6)
        self.cb_teacher = ttk.Combobox(top, values=(self.m.teacher_names), width=80)
        self.cb_teacher.pack(side="right")
        self.cb_teacher.bind("<Return>", self.pick_teacher_enter)
        self.grid_frame = tk.Frame(page, bg=COLOR_BG)
        self.grid_frame.pack(pady=6, anchor="e")
        header_bg = COLOR_ACCENT
        self.cell_btn = {}

        def draw_headers():
            for w in self.grid_frame.grid_slaves():
                w.destroy()
            else:
                periods = self.m.get_period_labels()
                p_reversed = list(reversed(periods))
                for j, p in enumerate(p_reversed):
                    tk.Label((self.grid_frame), text=f"الحصة {p}", bg=header_bg, fg="white", relief="ridge",
                      width=15,
                      padx=6,
                      pady=6,
                      font=('Segoe UI', 10, 'bold')).grid(row=0, column=j, sticky="nsew")
                else:
                    tk.Label((self.grid_frame), text="اليوم \\ الحصة", bg=header_bg, fg="white", relief="ridge",
                      width=15,
                      padx=6,
                      pady=6,
                      font=('Segoe UI', 10, 'bold')).grid(row=0, column=(len(periods)), sticky="nsew")


        def refresh(): pass


        self.refresh_teacher_grid = refresh
        refresh()
        return page# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- page_teacher_grid_draw_headers ---
    def page_teacher_grid_draw_headers():
        # ========================================
        for w in self.grid_frame.grid_slaves():
            w.destroy()
        else:
            periods = self.m.get_period_labels()
            p_reversed = list(reversed(periods))
            for j, p in enumerate(p_reversed):
                tk.Label((self.grid_frame), text=f"الحصة {p}", bg=header_bg, fg="white", relief="ridge",
                  width=15,
                  padx=6,
                  pady=6,
                  font=('Segoe UI', 10, 'bold')).grid(row=0, column=j, sticky="nsew")
            else:
                tk.Label((self.grid_frame), text="اليوم \\ الحصة", bg=header_bg, fg="white", relief="ridge",
                  width=15,
                  padx=6,
                  pady=6,
                  font=('Segoe UI', 10, 'bold')).grid(row=0, column=(len(periods)), sticky="nsew")

        # ========================================

    # --- page_teacher_grid_refresh ---
    def page_teacher_grid_refresh():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- pick_teacher_enter ---
    def pick_teacher_enter(self, e=None):
        # ========================================
        name = tidy(self.cb_teacher.get())
        if name:
            if name in self.m.teacher_names:
                self.current_teacher = name
                self.lbl_teacher.configure(text=name)
                self.show("teacher_grid")
                self.refresh_teacher_grid()

        # ========================================

    # --- open_cell_editor ---
    def open_cell_editor(self, day=None, period=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- open_cell_editor_update_preview ---
    def open_cell_editor_update_preview():
        # ========================================
        subject = tidy(cb_subject.get())
        class_val = tidy(cb_class.get())
        section = tidy(cb_section.get())
        parts = []
        if subject:
            parts.append(subject)
        if class_val and section:
            parts.append(f"{class_val}{section}")
        else:
            if class_val:
                parts.append(class_val)
            else:
                if section:
                    parts.append(section)
        preview_label.configure(text=(" / ".join(parts) if parts else "—"))

        # ========================================

    # --- open_cell_editor_commit ---
    def open_cell_editor_commit(event=None):
        # ========================================
        subject = tidy(cb_subject.get())
        class_val = tidy(cb_class.get())
        section = tidy(cb_section.get())
        parts = []
        if subject:
            parts.append(subject)
        if class_val and section:
            parts.append(f"{class_val}{section}")
        else:
            if class_val:
                parts.append(class_val)
            else:
                if section:
                    parts.append(section)
        text = " / ".join(parts) if parts else "—"
        try:
            self.m.set_teacher_cell(self.current_teacher, day, period, text)
            self.cell_btn[(day, period)].configure(text=text)
            win.destroy()
        except Exception as ex:
            try:
                messagebox.showerror("خطأ", str(ex))
            finally:
                ex = None
                del ex

        # ========================================

    # --- open_cell_editor_delete_cell ---
    def open_cell_editor_delete_cell():
        # ========================================
        if messagebox.askyesno("تأكيد H?", "هل أنت متأكد من حذف هذه الحصة؟"):
            try:
                self.m.set_teacher_cell(self.current_teacher, day, period, "")
                self.cell_btn[(day, period)].configure(text="—")
                win.destroy()
            except Exception as ex:
                try:
                    messagebox.showerror("خطأ", str(ex))
                finally:
                    ex = None
                    del ex

        # ========================================

    # --- page_teachers ---
    def page_teachers(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="بحث/قائمة المعلمين", style="Header.TLabel").pack(anchor="e",
          padx=6,
          pady=6)
        top = tk.Frame(page, bg=COLOR_BG)
        top.pack(fill="x")
        ttk.Label(top, text="ابحث:", foreground="#333").pack(side="right", padx=6)
        e = ttk.Entry(top, width=40)
        e.pack(side="right")
        lb = tk.Listbox(page, height=22)
        lb.pack(fill="both", expand=True, padx=6, pady=6)

        def refresh_list(q=''):
            lb.delete(0, "end")
            pool = [t for t in self.m.teacher_names if q in t] if q else self.m.teacher_names
            for t in pool:
                lb.insert("end", t)


        e.bind("<KeyRelease>", lambda evt: refresh_list(e.get().strip()))
        refresh_list()

        def pick(*args, **kwargs): pass  # TODO: Reconstruct method


        ttk.Button(page, text="تعيين المعلّم المختار", command=pick).pack(pady=6)
        self.teachers_admin_frame = tk.Frame(page, bg=COLOR_BG, pady=10)
        tk.Label((self.teachers_admin_frame), text="أدوات الإدارة:", bg=COLOR_BG, font=('Segoe UI',
                                                                                        10,
                                                                                        'bold')).pack(anchor="e", padx=10)
        tk.Button((self.teachers_admin_frame), text="اعتماد المبادلة", command=(lambda: self.show("swap_approvals")),
          bg="#d32f2f",
          fg="white",
          font=('Segoe UI', 10),
          padx=10).pack(side="right", padx=5)
        tk.Button((self.teachers_admin_frame), text="📊 تقارير المبادلات", command=(self.open_swap_reports),
          bg="#7b1fa2",
          fg="white",
          font=('Segoe UI', 10),
          padx=10).pack(side="right", padx=5)

        def check_admin_access(event):
            if self.admin_mode:
                self.teachers_admin_frame.pack(fill="x", side="bottom")
            else:
                self.teachers_admin_frame.pack_forget()


        page.bind("<Visibility>", check_admin_access)
        return page# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- page_teachers_refresh_list ---
    def page_teachers_refresh_list(q=None):
        # ========================================
        lb.delete(0, "end")
        pool = [t for t in self.m.teacher_names if q in t] if q else self.m.teacher_names
        for t in pool:
            lb.insert("end", t)

        # ========================================

    # --- page_teachers_pick ---
    def page_teachers_pick():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- page_teachers_check_admin_access ---
    def page_teachers_check_admin_access(event=None):
        # ========================================
        if self.admin_mode:
            self.teachers_admin_frame.pack(fill="x", side="bottom")
        else:
            self.teachers_admin_frame.pack_forget()

        # ========================================

    # --- open_swap_reports ---
    def open_swap_reports(self):
        # ========================================
        if not self.admin_mode:
            return
        top = tk.Toplevel(self)
        top.title("تقارير وإحصائيات تبادل الحصص")
        top.geometry("900x650")
        swaps_data = []
        fpath = data_path("تبادل_حصص.json")
        if os.path.exists(fpath):
            try:
                import json
                with open(fpath, encoding="utf-8") as f:
                    swaps_data = json.load(f)
            except:
                swaps_data = []

        def calc_stats(teacher_name=None):
            if teacher_name:
                filtered = [s for s in swaps_data if not s.get("requester") == teacher_name if s.get("acceptor") == teacher_name]
                total = len(filtered)
                approved = len([s for s in filtered if s.get("status") == "approved"])
                rejected = len([s for s in filtered if s.get("status") == "rejected"])
                return {'total': total, 'approved': approved, 'rejected': rejected, 'items': filtered}
            summary = {}
            for s in swaps_data:
                req = s.get("requester")
                if not req:
                    pass
                elif req not in summary:
                    summary[req] = {'total':0, 
                     'approved':0,  'rejected':0}
                summary[req]["total"] += 1
                status = s.get("status", "")
                if status == "approved":
                    summary[req]["approved"] += 1
                if status == "rejected":
                    summary[req]["rejected"] += 1
                return summary


        notebook = ttk.Notebook(top)
        notebook.pack(fill="both", expand=True, padx=10, pady=10)
        tab_global = tk.Frame(notebook, bg=COLOR_BG)
        notebook.add(tab_global, text="  ملخص المعلمين  ")
        columns = ('name', 'total', 'approved', 'rejected')
        tree_g = ttk.Treeview(tab_global, columns=columns, show="headings", height=20)
        tree_g.heading("name", text="المعلم")
        tree_g.heading("total", text="إجمالي الطلبات")
        tree_g.heading("approved", text="المقبولة")
        tree_g.heading("rejected", text="المرفوضة")
        tree_g.column("name", width=200, anchor="center")
        tree_g.column("total", width=100, anchor="center")
        tree_g.column("approved", width=100, anchor="center")
        tree_g.column("rejected", width=100, anchor="center")
        tree_g.pack(fill="both", expand=True, padx=10, pady=10)
        g_stats = calc_stats(None)
        g_list = []
        for name, d in g_stats.items():
            g_list.append((name, d["total"], d["approved"], d["rejected"]))
        else:
            g_list.sort(key=(lambda x: x[1]), reverse=True)
            for item in g_list:
                tree_g.insert("", "end", values=item)
            else:
                tab_detail = tk.Frame(notebook, bg=COLOR_BG)
                notebook.add(tab_detail, text="  تفاصيل معلم  ")
                d_top = tk.Frame(tab_detail, bg=COLOR_PANEL, pady=10)
                d_top.pack(fill="x")
                ttk.Label(d_top, text="اختر المعلم:", background=COLOR_PANEL).pack(side="right", padx=10)
                cb_teachers = ttk.Combobox(d_top, values=(sorted(self.m.teacher_names)), width=30)
                cb_teachers.pack(side="right")
                d_cards = tk.Frame(tab_detail, bg=COLOR_BG, pady=10)
                d_cards.pack(fill="x", padx=20)

                def mk_card(parent, label, color):
                    f = tk.Frame(parent, bg="white", highlightthickness=1, highlightbackground="#ccc", padx=20, pady=10)
                    f.pack(side="right", padx=10, expand=True, fill="x")
                    l_val = tk.Label(f, text="0", font=('Segoe UI', 20, 'bold'), fg=color, bg="white")
                    l_val.pack()
                    tk.Label(f, text=label, font=('Segoe UI', 10), fg="#555", bg="white").pack()
                    return l_val


                lbl_total = mk_card(d_cards, "إجمالي الطلبات", "#1976D2")
                lbl_app = mk_card(d_cards, "تمت الموافقة", "#388E3C")
                lbl_rej = mk_card(d_cards, "سجلات الرفض", "#D32F2F")
                d_list_frame = tk.Frame(tab_detail, bg=COLOR_BG)
                d_list_frame.pack(fill="both", expand=True, padx=10, pady=10)
                cols_d = ('date', 'day', 'per', 'sub', 'status', 'note')
                tree_d = ttk.Treeview(d_list_frame, columns=cols_d, show="headings")
                tree_d.heading("date", text="التاريخ")
                tree_d.heading("day", text="اليوم")
                tree_d.heading("per", text="الحصة")
                tree_d.heading("sub", text="المادة / المعلم البديل")
                tree_d.heading("status", text="الحالة")
                tree_d.heading("note", text="ملاحظات")
                tree_d.column("date", width=90, anchor="center")
                tree_d.column("day", width=60, anchor="center")
                tree_d.column("per", width=50, anchor="center")
                tree_d.column("sub", width=200, anchor="center")
                tree_d.column("status", width=100, anchor="center")
                tree_d.column("note", width=150, anchor="w")
                tree_d.pack(fill="both", expand=True)

                def refresh_details(evt=None):
                    tname = cb_teachers.get().strip()
                    if not tname:
                        return
                    res = calc_stats(tname)
                    lbl_total.config(text=(str(res["total"])))
                    lbl_app.config(text=(str(res["approved"])))
                    lbl_rej.config(text=(str(res["rejected"])))
                    for i in tree_d.get_children():
                        tree_d.delete(i)
                    else:
                        items = res["items"]
                        items.sort(key=(lambda x: x.get("id", "")), reverse=True)
                        for it in items:
                            st_txt = "بانتظار الموافقة"
                            st_code = it.get("status", "")
                            if st_code == "approved":
                                st_txt = "معتمد ✅"
                            else:
                                if st_code == "rejected":
                                    st_txt = "مرفوض ❌"
                                else:
                                    if st_code == "pending_admin":
                                        st_txt = "بانتظار المدير ⏳"
                            other = it.get("acceptor") if it.get("requester") == tname else it.get("requester")
                            sub_info = f'{it.get("subject", "")} (مع: {other})'
                            tree_d.insert("", "end", values=(
                             it.get("date", ""),
                             it.get("day", ""),
                             it.get("period", ""),
                             sub_info,
                             st_txt,
                             it.get("comments", "")))


                cb_teachers.bind("<<ComboboxSelected>>", refresh_details)
                cb_teachers.bind("<Return>", refresh_details)

        # ========================================

    # --- open_swap_reports_calc_stats ---
    def open_swap_reports_calc_stats(teacher_name=None):
        # ========================================
        if teacher_name:
            filtered = [s for s in swaps_data if not s.get("requester") == teacher_name if s.get("acceptor") == teacher_name]
            total = len(filtered)
            approved = len([s for s in filtered if s.get("status") == "approved"])
            rejected = len([s for s in filtered if s.get("status") == "rejected"])
            return {'total': total, 'approved': approved, 'rejected': rejected, 'items': filtered}
        summary = {}
        for s in swaps_data:
            req = s.get("requester")
            if not req:
                pass
            elif req not in summary:
                summary[req] = {'total':0, 
                 'approved':0,  'rejected':0}
            summary[req]["total"] += 1
            status = s.get("status", "")
            if status == "approved":
                summary[req]["approved"] += 1
            if status == "rejected":
                summary[req]["rejected"] += 1
            return summary

        # ========================================

    # --- open_swap_reports_mk_card ---
    def open_swap_reports_mk_card(parent=None, label=None, color=None):
        # ========================================
        f = tk.Frame(parent, bg="white", highlightthickness=1, highlightbackground="#ccc", padx=20, pady=10)
        f.pack(side="right", padx=10, expand=True, fill="x")
        l_val = tk.Label(f, text="0", font=('Segoe UI', 20, 'bold'), fg=color, bg="white")
        l_val.pack()
        tk.Label(f, text=label, font=('Segoe UI', 10), fg="#555", bg="white").pack()
        return l_val

        # ========================================

    # --- open_swap_reports_refresh_details ---
    def open_swap_reports_refresh_details(evt=None):
        # ========================================
        tname = cb_teachers.get().strip()
        if not tname:
            return
        res = calc_stats(tname)
        lbl_total.config(text=(str(res["total"])))
        lbl_app.config(text=(str(res["approved"])))
        lbl_rej.config(text=(str(res["rejected"])))
        for i in tree_d.get_children():
            tree_d.delete(i)
        else:
            items = res["items"]
            items.sort(key=(lambda x: x.get("id", "")), reverse=True)
            for it in items:
                st_txt = "بانتظار الموافقة"
                st_code = it.get("status", "")
                if st_code == "approved":
                    st_txt = "معتمد ✅"
                else:
                    if st_code == "rejected":
                        st_txt = "مرفوض ❌"
                    else:
                        if st_code == "pending_admin":
                            st_txt = "بانتظار المدير ⏳"
                other = it.get("acceptor") if it.get("requester") == tname else it.get("requester")
                sub_info = f'{it.get("subject", "")} (مع: {other})'
                tree_d.insert("", "end", values=(
                 it.get("date", ""),
                 it.get("day", ""),
                 it.get("period", ""),
                 sub_info,
                 st_txt,
                 it.get("comments", "")))

        # ========================================

    # --- page_subjects ---
    def page_subjects(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="إدارة المواد", style="Header.TLabel").pack(anchor="e",
          padx=6,
          pady=6)
        top = tk.Frame(page, bg=COLOR_BG)
        top.pack(fill="x", pady=(0, 8))
        ttk.Label(top, text="أدخل مادة جديدة ثم Enter:", foreground="#333").pack(side="right",
          padx=6)
        e = ttk.Entry(top, width=30)
        e.pack(side="right")
        frame = tk.Frame(page, bg=COLOR_PANEL)
        frame.pack(fill="both", expand=True, padx=6, pady=6)
        tv = ttk.Treeview(frame, columns=["المادة"], show="headings", height=20)
        tv.heading("المادة", text="المادة")
        tv.column("المادة", anchor="e", width=300)
        vsb = ttk.Scrollbar(frame, orient="vertical", command=(tv.yview))
        tv.configure(yscrollcommand=(vsb.set))
        vsb.pack(side="left", fill="y")
        tv.pack(side="right", fill="both", expand=True)

        def reload_table():
            tv.delete(*tv.get_children())
            if "المادة" not in self.m.df_subjects.columns:
                self.m.df_subjects["المادة"] = []
            for _, row in self.m.df_subjects.iterrows():
                tv.insert("", "end", values=[tidy(row.get("المادة", ""))])


        def add_enter(evt=None):
            ok, msg = self.m.add_subject(e.get())
            if not ok:
                messagebox.showwarning("تنبيه", msg)
            e.delete(0, "end")
            reload_table()


        e.bind("<Return>", add_enter)
        reload_table()

        def delete_selected():
            items = tv.selection()
            if not items:
                messagebox.showwarning("تنبيه", "اختر مادة للحذف.")
                return
            if messagebox.askyesno("تأكيد", "هل أنت متأكد من حذف المواد المحددة؟"):
                names = [tv.item(i)["values"][0] for i in items]
                deleted = self.m.delete_subjects(names)
                messagebox.showinfo("تم", f"تم حذف {deleted} مادة.")
                reload_table()


        def import_excel():
            path = filedialog.askopenfilename(filetypes=[('Excel Files', '*.xlsx *.xls')])
            if path:
                count, msg = self.m.import_subjects_from_excel(path)
                if count > 0:
                    messagebox.showinfo("تم", msg)
                    reload_table()
                else:
                    messagebox.showwarning("تنبيه", msg)


        btn_frame = tk.Frame(page, bg=COLOR_BG)
        btn_frame.pack(pady=6)
        tk.Button(btn_frame, text="🗑️ حذف المحدد", command=delete_selected, bg=COLOR_WARN, fg="black").pack(side="left", padx=5)
        tk.Button(btn_frame, text="📥 استيراد من Excel", command=import_excel, bg=COLOR_XLSX, fg="white").pack(side="left", padx=5)

        def refresh_access(event=None):
            """ Update UI state based on current admin mode """
            state = "normal" if self.admin_mode else "disabled"
            bg_del = COLOR_WARN if self.admin_mode else "#ccc"
            e.config(state=state)
            for w in page.winfo_children():
                if isinstance(w, tk.Button) and "حذف" in w.cget("text"):
                    w.config(state=state, bg=bg_del)


        page.bind("<Visibility>", refresh_access)
        refresh_access()
        return page

        # ========================================

    # --- page_subjects_reload_table ---
    def page_subjects_reload_table():
        # ========================================
        tv.delete(*tv.get_children())
        if "المادة" not in self.m.df_subjects.columns:
            self.m.df_subjects["المادة"] = []
        for _, row in self.m.df_subjects.iterrows():
            tv.insert("", "end", values=[tidy(row.get("المادة", ""))])

        # ========================================

    # --- page_subjects_add_enter ---
    def page_subjects_add_enter(evt=None):
        # ========================================
        ok, msg = self.m.add_subject(e.get())
        if not ok:
            messagebox.showwarning("تنبيه", msg)
        e.delete(0, "end")
        reload_table()

        # ========================================

    # --- page_subjects_delete_selected ---
    def page_subjects_delete_selected():
        # ========================================
        items = tv.selection()
        if not items:
            messagebox.showwarning("تنبيه", "اختر مادة للحذف.")
            return
        if messagebox.askyesno("تأكيد", "هل أنت متأكد من حذف المواد المحددة؟"):
            names = [tv.item(i)["values"][0] for i in items]
            deleted = self.m.delete_subjects(names)
            messagebox.showinfo("تم", f"تم حذف {deleted} مادة.")
            reload_table()

        # ========================================

    # --- page_subjects_import_excel ---
    def page_subjects_import_excel():
        # ========================================
        path = filedialog.askopenfilename(filetypes=[('Excel Files', '*.xlsx *.xls')])
        if path:
            count, msg = self.m.import_subjects_from_excel(path)
            if count > 0:
                messagebox.showinfo("تم", msg)
                reload_table()
            else:
                messagebox.showwarning("تنبيه", msg)

        # ========================================

    # --- page_subjects_refresh_access ---
    def page_subjects_refresh_access(event=None):
        # ========================================
        state = "normal" if self.admin_mode else "disabled"
        bg_del = COLOR_WARN if self.admin_mode else "#ccc"
        e.config(state=state)
        for w in page.winfo_children():
            if isinstance(w, tk.Button) and "حذف" in w.cget("text"):
                w.config(state=state, bg=bg_del)

        # ========================================

    # --- page_students ---
    def page_students(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="الطلاب (عرض مختصر + تفاصيل عند الدبل-كلك)", style="Header.TLabel").pack(anchor="e", padx=6, pady=6)
        filter_frame = tk.Frame(page, bg=COLOR_BG)
        filter_frame.pack(fill="x", padx=6, pady=6)
        ttk.Label(filter_frame, text="ابحث بالاسم:", foreground="#333").pack(side="right", padx=6)
        filter_name = ttk.Entry(filter_frame, width=25)
        filter_name.pack(side="right", padx=6)
        ttk.Label(filter_frame, text="الصف:", foreground="#333").pack(side="right", padx=6)
        filter_class = ttk.Combobox(filter_frame, width=15, values=([""] + self.m.get_available_classes()))
        filter_class.pack(side="right", padx=6)
        filter_class.set("")
        ttk.Label(filter_frame, text="الشعبة:", foreground="#333").pack(side="right", padx=6)
        filter_section = ttk.Combobox(filter_frame, width=15, values=([""] + self.m.get_available_sections()))
        filter_section.pack(side="right", padx=6)
        filter_section.set("")
        ttk.Button(filter_frame, text="عرض الكل", command=(lambda: clear_filters())).pack(side="right", padx=6)
        if self.admin_mode:

            def open_add_student():
                win = tk.Toplevel(self)
                win.title("إضافة طالب جديد")
                win.geometry("400x350")
                win.grab_set()
                f = ttk.Frame(win, padding=20)
                f.pack(fill="both")
                ttk.Label(f, text="اسم الطالب:").pack(anchor="e")
                e_name = ttk.Entry(f, width=40)
                e_name.pack(pady=5)
                ttk.Label(f, text="الصف:").pack(anchor="e")
                e_class = ttk.Entry(f, width=40)
                e_class.pack(pady=5)
                ttk.Label(f, text="الشعبة:").pack(anchor="e")
                e_sec = ttk.Entry(f, width=40)
                e_sec.pack(pady=5)

                def save_new():
                    if not e_name.get():
                        return
                    self.m.update_student_full(None, {'الاسم':e_name.get(), 
                     'الصف':e_class.get(),  'الشعبة':e_sec.get()})
                    messagebox.showinfo("تم", "تمت إضافة الطالب.")
                    reload_students()
                    win.destroy()

                tk.Button(f, text="➕ إضافة الآن", command=save_new, bg=COLOR_ACCENT, fg="white", pady=10).pack(fill="x", pady=20)


            ttk.Button(filter_frame, text="➕ إضافة طالب", command=open_add_student).pack(side="right", padx=6)
        ttk.Button(filter_frame, text="📄 عرض التفاصيل", command=(lambda: open_details())).pack(side="right", padx=6)
        frame = tk.Frame(page, bg=COLOR_PANEL)
        frame.pack(fill="both", expand=True, padx=6, pady=6)
        tv = ttk.Treeview(frame, columns=["الاسم", "الصف", "الشعبة"], show="headings", height=20)
        for c in ('الاسم', 'الصف', 'الشعبة'):
            tv.heading(c, text=c)
            tv.column(c, anchor="e", width=180)
        else:
            vsb = ttk.Scrollbar(frame, orient="vertical", command=(tv.yview))
            tv.configure(yscrollcommand=(vsb.set))
            vsb.pack(side="left", fill="y")
            tv.pack(side="right", fill="both", expand=True)

            def reload_students():
                tv.delete(*tv.get_children())
                df = self.m.list_students_simple()
                q_name = normalize_arabic(tidy(filter_name.get()))
                cval = tidy(filter_class.get())
                sval = tidy(filter_section.get())
                if q_name:
                    df = df[df["الاسم"].astype(str).apply(normalize_arabic).str.contains(q_name, na=False)]
                if cval:
                    df = df[df["الصف"].astype(str).str.strip() == cval]
                if sval:
                    df = df[df["الشعبة"].astype(str).str.strip() == sval]
                for idx, row in df.iterrows():
                    tv.insert("", "end", iid=(str(idx)), values=[
                     tidy(row.get("الاسم", "")),
                     tidy(row.get("الصف", "")),
                     tidy(row.get("الشعبة", ""))])


            def clear_filters():
                filter_name.delete(0, "end")
                filter_class.set("")
                filter_section.set("")
                reload_students()


            filter_name.bind("<KeyRelease>", lambda e: reload_students())
            filter_class.bind("<<ComboboxSelected>>", lambda e: reload_students())
            filter_section.bind("<<ComboboxSelected>>", lambda e: reload_students())
            reload_students()

            def open_details(evt=None):
                sel = tv.selection()
                if not sel:
                    if evt is None:
                        messagebox.showwarning("تنبيه", "الرجاء اختيار طالب أولاً.")
                    return
                idx = int(sel[0])
                student_data = self.m.get_student_full(idx)
                name_col, _, _ = self.m._detect_student_cols()
                student_name = tidy(student_data.get(name_col, student_data.get("الاسم", "")))
                if not student_name:
                    for k in student_data.keys():
                        if "اسم" in k:
                            student_name = tidy(student_data[k])
                            break
                            
                win = tk.Toplevel(self)
                win.title(f"ملف الطالب: {student_name}")
                win.geometry("900x700")
                win.grab_set()
                top_frame = tk.Frame(win, bg=COLOR_PANEL, height=80)
                top_frame.pack(fill="x", padx=10, pady=10)
                ttk.Label(top_frame, text=f"📂 ملف الطالب: {student_name}", font=('Segoe UI', 16, 'bold'),
                  foreground=COLOR_ACCENT).pack(side="right", padx=20, pady=10)
                notebook = ttk.Notebook(win)
                notebook.pack(fill="both", expand=True, padx=10, pady=5)
                
                # --- TAB 1: Schedule ---
                tab_schedule = tk.Frame(notebook, bg=COLOR_BG)
                notebook.add(tab_schedule, text="  📅 حصص الطالب  ")

                f_filter = tk.Frame(tab_schedule, bg=COLOR_BG)
                f_filter.pack(fill="x", padx=10, pady=10)
                v_day = tk.StringVar(value="الكل")
                v_period = tk.StringVar(value="الكل")
                tk.Label(f_filter, text="تصفية حسب اليوم:", bg=COLOR_BG).pack(side="right", padx=5)
                cb_day = ttk.Combobox(f_filter, values=(["الكل"] + DAYS), textvariable=v_day, state="readonly", width=12)
                cb_day.pack(side="right", padx=5)
                tk.Label(f_filter, text="الحصة:", bg=COLOR_BG).pack(side="right", padx=5)
                cb_period = ttk.Combobox(f_filter, values=(["الكل"] + [str(i) for i in range(1, 8)]), textvariable=v_period, state="readonly", width=8)
                cb_period.pack(side="right", padx=5)
                s_class = tidy(student_data.get("الصف", ""))
                s_section = tidy(student_data.get("الشعبة", ""))
                raw_results, _, _ = self.m.find_student_schedule(student_name, s_class, s_section)
                cols = ['اليوم', 'الحصة', 'من', 'إلى', 'المعلم', 'المادة/الصف']
                tv_sch = ttk.Treeview(tab_schedule, columns=cols, show="headings", height=12)
                for c in cols:
                    tv_sch.heading(c, text=c)
                    width = 100
                    if c == "المادة/الصف":
                        width = 200
                    tv_sch.column(c, anchor="center", width=width)
                vsb_sch = ttk.Scrollbar(tab_schedule, orient="vertical", command=tv_sch.yview)
                tv_sch.configure(yscrollcommand=vsb_sch.set)
                vsb_sch.pack(side="left", fill="y")
                tv_sch.pack(side="right", fill="both", expand=True)

                def apply_filters(*args, **kwargs):
                    for item in tv_sch.get_children():
                        tv_sch.delete(item)
                    d_filter = v_day.get()
                    p_filter = v_period.get()
                    for r in raw_results:
                        if d_filter != "الكل" and r.get("اليوم") != d_filter: continue
                        if p_filter != "الكل" and str(r.get("الحصة")) != p_filter: continue
                        tv_sch.insert("", "end", values=[r.get(c, "") for c in cols])

                cb_day.bind("<<ComboboxSelected>>", apply_filters)
                cb_period.bind("<<ComboboxSelected>>", apply_filters)
                cb_day.bind("<Return>", apply_filters)
                cb_period.bind("<Return>", apply_filters)
                tk.Button(f_filter, text="🔍 بحث / تصفية", command=apply_filters, bg="#eee").pack(side="right", padx=5)

                def do_export():
                    if not raw_results:
                        messagebox.showinfo("تنبيه", "لا يوجد جدول لتصديره.")
                        return
                    filtered = []
                    d_filter = v_day.get()
                    p_filter = v_period.get()
                    for r in raw_results:
                        if d_filter != "الكل" and r.get("اليوم") != d_filter: continue
                        if p_filter != "الكل" and str(r.get("الحصة")) != p_filter: continue
                        filtered.append(r)
                    if not filtered:
                        messagebox.showinfo("تنبيه", "الجدول فارغ (حسب التصفية الحالية).")
                        return
                    try:
                        import xlsxwriter
                        import os
                        fname = f"جدول_الطالب_{student_name}.xlsx"
                        fpath = os.path.abspath(fname)
                        wb = xlsxwriter.Workbook(fpath)
                        ws = wb.add_worksheet("الجدول")
                        ws.right_to_left()
                        fmt_head = wb.add_format({'bold': True, 'bg_color': '#2e7d32', 'font_color': 'white', 'border': 1, 'align': 'center'})
                        fmt_cell = wb.add_format({'border':1,  'align':"center"})
                        for i, c in enumerate(cols):
                            ws.write(0, i, c, fmt_head)
                        for r_i, d in enumerate(filtered):
                            for c_i, col in enumerate(cols):
                                ws.write(r_i + 1, c_i, d.get(col, ""), fmt_cell)
                        wb.close()
                        os.startfile(fpath)
                        messagebox.showinfo("تم", f"تم تصدير الجدول بنجاح:\n{fname}")
                    except Exception as ex:
                        messagebox.showerror("خطأ", f"فشل التصدير: {ex}")

                tk.Button(f_filter, text="📥 تصدير الجدول (Excel)", command=do_export, bg=COLOR_XLSX,
                  fg="white", font=('Segoe UI', 9, 'bold')).pack(side="left", padx=10)
                apply_filters()

                # --- TAB 2: Personal Data ---
                tab_personal = tk.Frame(notebook, bg=COLOR_BG)
                notebook.add(tab_personal, text="  👤 بيانات خاصة  ")
                canvas = tk.Canvas(tab_personal, bg=COLOR_BG)
                scrollbar = ttk.Scrollbar(tab_personal, orient="vertical", command=canvas.yview)
                scrollable_frame = ttk.Frame(canvas)
                scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
                canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
                canvas.configure(yscrollcommand=scrollbar.set)
                canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
                scrollbar.pack(side="right", fill="y")
                
                frm_personal = scrollable_frame
                entries = {}
                row_i = 0
                for key in sorted(student_data.keys()):
                    ttk.Label(frm_personal, text=(key + ":")).grid(row=row_i, column=0, sticky="e", padx=6, pady=6)
                    ent = ttk.Entry(frm_personal, width=40)
                    ent.grid(row=row_i, column=1, sticky="w", padx=6, pady=6)
                    ent.insert(0, tidy(student_data.get(key, "")))
                    entries[key] = ent
                    row_i += 1

                def save_personal_data():
                    updated = {k: tidy(e.get()) for k, e in entries.items()}
                    self.m.update_student_full(idx, updated)
                    messagebox.showinfo("تم", "تم حفظ البيانات الخاصة بنجاح.")
                    try:
                        reload_students()
                    except:
                        pass

                btn_save = tk.Button(frm_personal, text="💾 حفظ التعديلات", command=save_personal_data, bg=COLOR_ACCENT, fg="white", font=('Segoe UI', 10, 'bold'))
                if not self.admin_mode:
                    btn_save.config(state="disabled", bg="#ccc")
                btn_save.grid(row=row_i, column=0, columnspan=2, pady=20, sticky="ew")
                
                wm_bar = tk.Frame(win, bg=COLOR_PANEL)
                wm_bar.pack(side="bottom", fill="x")
                final_res, _, _ = self.m.find_student_schedule(student_name)
                tk.Button(wm_bar, text="📄 إذن خروج (Word)", bg=COLOR_BTN, fg="white", command=(lambda: self.open_exit_permit_dialog(student_name, final_res))).pack(side="right", padx=10, pady=10)
                tk.Button(wm_bar, text="⬇️ تصدير الجدول (Excel)", bg=COLOR_XLSX, fg="white", command=(lambda: self.export_student_schedule_excel(student_name, final_res))).pack(side="right", padx=10, pady=10)

            tv.bind("<Double-1>", open_details)
            return page# [DECOMPILATION FAILED]: Deparsing stopped due to parse error
    # --- page_timings ---
    def page_timings(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="إدارة التوقيت والفعاليات", style="Header.TLabel").pack(anchor="e",
          padx=6,
          pady=6)
        if self.admin_mode:
            ctrl = tk.Frame(page, bg=COLOR_BG)
            ctrl.pack(fill="x", padx=6, pady=6)
            ttk.Button(ctrl, text="إضافة", command=(lambda: open_editor(None))).pack(side="right", padx=6)
            ttk.Button(ctrl, text="تعديل", command=(lambda: edit_selected())).pack(side="right", padx=6)
            ttk.Button(ctrl, text="حذف", command=(lambda: delete_selected())).pack(side="right", padx=6)
            ttk.Separator(ctrl, orient="vertical").pack(side="right", fill="y", padx=10)
            ttk.Button(ctrl, text="📥 استيراد Excel", command=(lambda: import_excel())).pack(side="right", padx=6)
        frame = tk.Frame(page, bg=COLOR_PANEL)
        frame.pack(fill="both", expand=True, padx=6, pady=6)
        tv = ttk.Treeview(frame, columns=["الحصة", "من", "إلى"], show="headings", height=15)
        tv.heading("الحصة", text="الحصة / الفترة")
        tv.heading("من", text="من (بداية)")
        tv.heading("إلى", text="إلى (نهاية)")
        tv.column("الحصة", anchor="center", width=250)
        tv.column("من", anchor="center", width=120)
        tv.column("إلى", anchor="center", width=120)
        vsb = ttk.Scrollbar(frame, orient="vertical", command=(tv.yview))
        tv.configure(yscrollcommand=(vsb.set))
        vsb.pack(side="left", fill="y")
        tv.pack(side="right", fill="both", expand=True)

        def get_minutes(*args, **kwargs): pass  # TODO: Reconstruct method


        def load():
            tv.delete(*tv.get_children())
            self.m.df_timings = self.m._read_excel_any(FILE_TIMINGS, silent=True)
            self.m.df_timings = self.m._normalize_timings(self.m.df_timings)
            df = self.m.df_timings.copy()
            df["_sort"] = df["من"].apply(lambda x: get_minutes(tidy(x)))
            df = df.sort_values("_sort").drop(columns=["_sort"], errors="ignore").reset_index(drop=True)
            self.m.df_timings = df
            for i, row in df.iterrows():
                tv.insert("", "end", iid=(str(i)), values=[tidy(row["الحصة"]), tidy(row["من"]), tidy(row["إلى"])])


        def save():
            self.m.df_timings.to_excel(FILE_TIMINGS, index=False)
            load()


        def import_excel():
            path = filedialog.askopenfilename(title="اختر ملف التوقيت", filetypes=[('Excel', '*.xlsx *.xls')])
            if not path:
                return
            df = self.m._read_excel_any(path, silent=True)
            df = self.m._normalize_timingsdf
            if df.empty:
                messagebox.showwarning("تنبيه", "لم يتم استيراد شيء (ملف غير متوافق).")
                return
            self.m.df_timings = df
            save()
            messagebox.showinfo("تم", "تم استيراد ملف التوقيت.")


        def open_editor(item_iid):
            win = tk.Toplevelself
            win.title("إضافة/تعديل فترة")
            win.geometry("520x360")
            win.grab_set()
            current = [
             "", "07:00", "07:45"]
            if item_iid is not None:
                current = tv.itemitem_iid["values"]
            frm = ttk.Framewin
            frm.pack(fill="both", expand=True, padx=16, pady=16)
            ttk.Label(frm, text="اسم الحدث/الحصة:").pack(anchor="e")
            e_name = ttk.Entryfrm
            e_name.pack(fill="x", pady=6)
            e_name.insert(0, current[0])

            def time_picker(label, initial):
                box = ttk.LabelFrame(frm, text=label)
                box.pack(fill="x", pady=6)
                fr = ttk.Framebox
                fr.pack(padx=8, pady=8, anchor="w")
                ih, im = ('07', '00')
                if initial:
                    if ":" in str(initial):
                        try:
                            ih, im = str(initial).split(":")
                            ih = f"{int(ih):02d}"
                            im = f"{int(im):02d}"
                        except Exception:
                            pass

                sb_h = tk.Spinbox(fr, from_=0, to=23, width=3, format="%02.0f(", font=('Segoe UI',
                                                                                      12))
                sb_m = tk.Spinbox(fr, from_=0, to=59, width=3, format=")%02.0f", font=('Segoe UI',
                                                                                      12))
                sb_h.delete(0, "end")
                sb_h.insert(0, ih)
                sb_m.delete(0, "end")
                sb_m.insert(0, im)
                sb_h.pack(side="left")
                ttk.Label(fr, text=":").pack(side="left")
                sb_m.pack(side="left")
                return (sb_h, sb_m)

            h1, m1 = time_picker("من", current[1])
            h2, m2 = time_picker("إلى", current[2])

            def commit():
                name = tidy(e_name.get())
                if not name:
                    messagebox.showwarning("تنبيه", "أدخل الاسم.")
                    return
                else:
                    t_from = f"{int(h1.get()):02d}:{int(m1.get()):02d}"
                    t_to = f"{int(h2.get()):02d}:{int(m2.get()):02d}"
                    if item_iid is None:
                        self.m.df_timings = pd.concat([self.m.df_timings, pd.DataFrame([{'الحصة':name,  'من':t_from,  'إلى':t_to}])], ignore_index=True)
                    else:
                        i = int(item_iid)
                    self.m.df_timings.at[(i, "الحصة")] = name
                    self.m.df_timings.at[(i, "من")] = t_from
                    self.m.df_timings.at[(i, "إلى")] = t_to
                save()
                win.destroy()

            ttk.Button(frm, text="حفظ", command=commit).pack(pady=12)


        def edit_selected():
            sel = tv.selection()
            if not sel:
                messagebox.showwarning("تنبيه", "اختر صفاً للتعديل.")
                return
            open_editor(sel[0])


        def delete_selected():
            sel = tv.selection()
            if not sel:
                messagebox.showwarning("تنبيه", "اختر صفاً للحذف.")
                return
            else:
                return messagebox.askyesno("تأكيد", "هل تريد الحذف؟") or None
            indices = [int(i) for i in sel]
            self.m.df_timings = self.m.df_timings.drop(index=indices).reset_index(drop=True)
            save()


        load()
        return page# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- page_timings_get_minutes ---
    def page_timings_get_minutes(t=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- page_timings_load ---
    def page_timings_load():
        # ========================================
        tv.delete(*tv.get_children())
        self.m.df_timings = self.m._read_excel_any(FILE_TIMINGS, silent=True)
        self.m.df_timings = self.m._normalize_timings(self.m.df_timings)
        df = self.m.df_timings.copy()
        df["_sort"] = df["من"].apply(lambda x: get_minutes(tidy(x)))
        df = df.sort_values("_sort").drop(columns=["_sort"], errors="ignore").reset_index(drop=True)
        self.m.df_timings = df
        for i, row in df.iterrows():
            tv.insert("", "end", iid=(str(i)), values=[tidy(row["الحصة"]), tidy(row["من"]), tidy(row["إلى"])])

        # ========================================

    # --- page_timings_save ---
    def page_timings_save():
        # ========================================
        self.m.df_timings.to_excel(FILE_TIMINGS, index=False)
        load()

        # ========================================

    # --- page_timings_import_excel ---
    def page_timings_import_excel():
        # ========================================
        path = filedialog.askopenfilename(title="اختر ملف التوقيت", filetypes=[('Excel', '*.xlsx *.xls')])
        if not path:
            return
        df = self.m._read_excel_any(path, silent=True)
        df = self.m._normalize_timings(df)
        if df.empty:
            messagebox.showwarning("تنبيه", "لم يتم استيراد شيء (ملف غير متوافق).")
            return
        self.m.df_timings = df
        save()
        messagebox.showinfo("تم", "تم استيراد ملف التوقيت.")

        # ========================================

    # --- page_timings_open_editor ---
    def page_timings_open_editor(item_iid=None):
        # ========================================
        win = tk.Toplevel(self)
        win.title("إضافة/تعديل فترة")
        win.geometry("520x360")
        win.grab_set()
        current = [
         "", "07:00", "07:45"]
        if item_iid is not None:
            current = tv.item(item_iid)["values"]
        frm = ttk.Frame(win)
        frm.pack(fill="both", expand=True, padx=16, pady=16)
        ttk.Label(frm, text="اسم الحدث/الحصة:").pack(anchor="e")
        e_name = ttk.Entry(frm)
        e_name.pack(fill="x", pady=6)
        e_name.insert(0, current[0])

        def time_picker(label, initial):
            box = ttk.LabelFrame(frm, text=label)
            box.pack(fill="x", pady=6)
            fr = ttk.Frame(box)
            fr.pack(padx=8, pady=8, anchor="w")
            ih, im = ('07', '00')
            if initial:
                if ":" in str(initial):
                    try:
                        ih, im = str(initial).split(":")
                        ih = f"{int(ih):02d}"
                        im = f"{int(im):02d}"
                    except Exception:
                        pass

            sb_h = tk.Spinbox(fr, from_=0, to=23, width=3, format="%02.0f(", font=('Segoe UI',
                                                                                  12))
            sb_m = tk.Spinbox(fr, from_=0, to=59, width=3, format=")%02.0f", font=('Segoe UI',
                                                                                  12))
            sb_h.delete(0, "end")
            sb_h.insert(0, ih)
            sb_m.delete(0, "end")
            sb_m.insert(0, im)
            sb_h.pack(side="left")
            ttk.Label(fr, text=":").pack(side="left")
            sb_m.pack(side="left")
            return (sb_h, sb_m)


        h1, m1 = time_picker("من", current[1])
        h2, m2 = time_picker("إلى", current[2])

        def commit():
            name = tidy(e_name.get())
            if not name:
                messagebox.showwarning("تنبيه", "أدخل الاسم.")
                return
            else:
                t_from = f"{int(h1.get()):02d}:{int(m1.get()):02d}"
                t_to = f"{int(h2.get()):02d}:{int(m2.get()):02d}"
                if item_iid is None:
                    self.m.df_timings = pd.concat([self.m.df_timings, pd.DataFrame([{'الحصة':name,  'من':t_from,  'إلى':t_to}])], ignore_index=True)
                else:
                    i = int(item_iid)
                self.m.df_timings.at[(i, "الحصة")] = name
                self.m.df_timings.at[(i, "من")] = t_from
                self.m.df_timings.at[(i, "إلى")] = t_to
            save()
            win.destroy()


        ttk.Button(frm, text="حفظ", command=commit).pack(pady=12)

        # ========================================

    # --- page_timings_open_editor_time_picker ---
    def page_timings_open_editor_time_picker(label=None, initial=None):
        # ========================================
        box = ttk.LabelFrame(frm, text=label)
        box.pack(fill="x", pady=6)
        fr = ttk.Frame(box)
        fr.pack(padx=8, pady=8, anchor="w")
        ih, im = ('07', '00')
        if initial:
            if ":" in str(initial):
                try:
                    ih, im = str(initial).split(":")
                    ih = f"{int(ih):02d}"
                    im = f"{int(im):02d}"
                except Exception:
                    pass

        sb_h = tk.Spinbox(fr, from_=0, to=23, width=3, format="%02.0f(", font=('Segoe UI', 12))
        sb_m = tk.Spinbox(fr, from_=0, to=59, width=3, format=")%02.0f", font=('Segoe UI', 12))
        sb_h.delete(0, "end")
        sb_h.insert(0, ih)
        sb_m.delete(0, "end")
        sb_m.insert(0, im)
        sb_h.pack(side="left")
        ttk.Label(fr, text=":").pack(side="left")
        sb_m.pack(side="left")
        return (sb_h, sb_m)

        # ========================================

    # --- page_timings_open_editor_commit ---
    def page_timings_open_editor_commit():
        # ========================================
        name = tidy(e_name.get())
        if not name:
            messagebox.showwarning("تنبيه", "أدخل الاسم.")
            return
        else:
            t_from = f"{int(h1.get()):02d}:{int(m1.get()):02d}"
            t_to = f"{int(h2.get()):02d}:{int(m2.get()):02d}"
            if item_iid is None:
                self.m.df_timings = pd.concat([self.m.df_timings, pd.DataFrame([{'الحصة':name,  'من':t_from,  'إلى':t_to}])], ignore_index=True)
            else:
                i = int(item_iid)
            self.m.df_timings.at[(i, "الحصة")] = name
            self.m.df_timings.at[(i, "من")] = t_from
            self.m.df_timings.at[(i, "إلى")] = t_to
        save()
        win.destroy()

        # ========================================

    # --- page_timings_edit_selected ---
    def page_timings_edit_selected():
        # ========================================
        sel = tv.selection()
        if not sel:
            messagebox.showwarning("تنبيه", "اختر صفاً للتعديل.")
            return
        open_editor(sel[0])

        # ========================================

    # --- page_timings_delete_selected ---
    def page_timings_delete_selected():
        # ========================================
        sel = tv.selection()
        if not sel:
            messagebox.showwarning("تنبيه", "اختر صفاً للحذف.")
            return
        else:
            return messagebox.askyesno("تأكيد", "هل تريد الحذف؟") or None
        indices = [int(i) for i in sel]
        self.m.df_timings = self.m.df_timings.drop(index=indices).reset_index(drop=True)
        save()

        # ========================================

    # --- page_student_finder ---
    def page_student_finder(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="البحث عن الطالب في الحصص", style="Header.TLabel").pack(anchor="e",
          padx=6,
          pady=6)
        name_frame = tk.Frame(page, bg=COLOR_BG)
        name_frame.pack(fill="x", padx=6, pady=6)
        ttk.Label(name_frame, text="ابحث بالاسم:", foreground="#333", font=('Segoe UI', 11,
                                                                            'bold')).pack(side="right", padx=6)
        self.student_name_entry = ttk.Entry(name_frame, width=40, font=('Segoe UI', 11))
        self.student_name_entry.pack(side="right", padx=6, pady=10)
        self.student_info_label = ttk.Label(name_frame, text="", foreground="#666")
        self.student_info_label.pack(side="right", padx=10)
        self.student_name_entry.bind("<KeyRelease>", lambda e: self.search_student_by_name())
        ttk.Separator(page, orient="horizontal").pack(fill="x", padx=20, pady=5)
        split_frame = tk.Frame(page, bg=COLOR_BG)
        split_frame.pack(fill="both", expand=True, padx=6, pady=6)
        list_container = tk.Frame(split_frame, bg=COLOR_PANEL, relief="flat", borderwidth=1)
        list_container.pack(side="right", fill="both", expand=True, padx=5)
        ttk.Label(list_container, text="اختر طالباً من القائمة:", background=COLOR_PANEL).pack(anchor="e", padx=10, pady=5)
        self.student_listbox = tk.Listbox(list_container, height=15, font=('Segoe UI', 11), relief="flat",
          highlightthickness=1,
          highlightcolor=COLOR_ACCENT)
        self.student_listbox.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.student_listbox.bind("<<ListboxSelect>>", self.on_student_select)
        self.student_listbox.bind("<Double-1>", self.show_student_details)
        self.preview_panel = tk.Frame(split_frame, bg=COLOR_PANEL, width=380, relief="solid", borderwidth=1)
        self.preview_panel.pack(side="left", fill="both", expand=False, padx=5)
        self.preview_panel.pack_propagate(False)
        ttk.Label((self.preview_panel), text="الوضعية الحالية:", style="Header.TLabel", background=COLOR_PANEL).pack(pady=10)
        self.preview_content = tk.Frame((self.preview_panel), bg=COLOR_PANEL, padx=20, pady=20)
        self.preview_content.pack(fill="both", expand=True)
        self.lbl_prev_name = tk.Label((self.preview_content), text="—", font=('Segoe UI', 16,
                                                                              'bold'), background=COLOR_PANEL,
          fg="#333",
          anchor="e",
          justify="right")
        self.lbl_prev_name.pack(fill="x", pady=5)
        self.lbl_prev_loc = tk.Label((self.preview_content), text="ابحث واختر طالباً لمعرفة مكانه الآن", foreground="#666",
          wraplength=320,
          justify="right",
          background=COLOR_PANEL,
          font=('Segoe UI', 12),
          anchor="e")
        self.lbl_prev_loc.pack(fill="x", pady=10)
        self.btn_quick_exit = tk.Button((self.preview_content), text="📄 إصدار إذن خروج طالب", command=(self.quick_exit_permit),
          bg=COLOR_WORD,
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          relief="flat",
          padx=20,
          pady=10)
        self.btn_details = tk.Button((self.preview_content), text="🔍 عرض الجدول الكامل", command=(self.show_student_details),
          bg=COLOR_BTN,
          fg="white",
          font=('Segoe UI', 10, 'bold'),
          relief="flat",
          state="disabled")
        self.btn_details.pack(side="bottom", fill="x", pady=10)
        return page

        # ========================================

    # --- on_student_select ---
    def on_student_select(self, event=None):
        # ========================================
        try:
            sel = self.student_listbox.curselection()
            if not sel: return
            idx = int(sel[0])
            name = self.student_listbox.get(idx)
            self.lbl_prev_name.configure(text=name)
            self.selected_finder_student = name
            loc_str, _, _, _ = self.m.get_student_current_location(name)
            self.lbl_prev_loc.configure(text=f"📍 {loc_str}")
            self.btn_details.configure(state="normal")
        except Exception:
            self.lbl_prev_name.configure(text="—")
            self.lbl_prev_loc.configure(text="حدث خطأ")
            self.btn_details.configure(state="disabled")

        # ========================================

    # --- quick_exit_permit ---
    def quick_exit_permit(self):
        # ========================================
        try:
            if hasattr(self, "selected_finder_student") and self.selected_finder_student:
                name = self.selected_finder_student
                loc_str, _, _, _ = self.m.get_student_current_location(name)
                final_res, _, _ = self.m.find_student_schedule(name)
                self.open_exit_permit_dialog(name, final_res, loc_str)
            else:
                messagebox.showwarning("تنبيه", "الرجاء اختيار طالباً أولاً.")
        except Exception as e:
            messagebox.showerror("خطأ", f"حدث خطأ: {e}")

        # ========================================

    # --- load_students_by_class ---
    def load_students_by_class(self):
        # ========================================
        class_name = tidy(self.class_entry.get())
        section = tidy(self.section_entry.get())
        self.student_listbox.delete(0, "end")
        if not (class_name and section):
            self.student_info_label.configure(text="اختر الصف والشعبة لعرض الطلاب", foreground="#666")
            return
        students = self.m.get_students_by_class_section(class_name, section)
        if not students:
            self.student_info_label.configure(text=f"لا يوجد طلاب في {class_name}/{section}", foreground=COLOR_WARN)
            return
        for s in students:
            self.student_listbox.insert("end", s["name"])
        else:
            self.student_info_label.configure(text=f"تم العثور على {len(students)} طالب. اضغط مرتين لعرض التفاصيل", foreground=COLOR_XLSX)

        # ========================================

    # --- search_student_by_name ---
    def search_student_by_name(self):
        # ========================================
        q = tidy(self.student_name_entry.get())
        self.student_listbox.delete(0, "end")
        if not q:
            self.student_info_label.configure(text="ابدأ الكتابة للبحث...", foreground="#666")
            return
        q_norm = normalize_arabic(q)
        name_col, _, _ = self.m._detect_student_cols()
        found = []
        for _, row in self.m.df_students.iterrows():
            n = tidy(row.get(name_col, ""))
            if q_norm in normalize_arabic(n):
                found.append(n)
            if not found:
                self.student_info_label.configure(text=f"لا يوجد طالب يطابق '{q}'", foreground=COLOR_WARN)
                return

        for n in found:
            self.student_listbox.insert("end", n)
        else:
            self.student_info_label.configure(text=f"تم العثور على {len(found)} طالب", foreground=COLOR_XLSX)

        # ========================================

    # --- show_student_details ---
    def show_student_details(self, event=None):
        # ========================================
        try:
            if hasattr(self, "selected_finder_student") and self.selected_finder_student:
                name = self.selected_finder_student
            else:
                messagebox.showwarning("تنبيه", "الرجاء اختيار طالب أولاً.")
                return

            name_col, _, _ = self.m._detect_student_cols()
            df = self.m.df_students
            matches = df[df[name_col].astype(str).str.strip() == name]
            if matches.empty:
                matches = df[df[name_col].astype(str).str.contains(name, na=False)]
            if matches.empty:
                messagebox.showerror("خطأ", "لم يتم العثور على الطالب في قاعدة البيانات أو لم يتم تحديثها.")
                return
            
            idx = int(matches.index[0])
            student_data = self.m.get_student_full(idx)
            student_name = name
            
            win = tk.Toplevel(self)
            win.title(f"ملف الطالب: {student_name}")
            win.geometry("900x700")
            win.grab_set()
            top_frame = tk.Frame(win, bg=COLOR_PANEL, height=80)
            top_frame.pack(fill="x", padx=10, pady=10)
            ttk.Label(top_frame, text=f"📂 ملف الطالب: {student_name}", font=('Segoe UI', 16, 'bold'),
              foreground=COLOR_ACCENT).pack(side="right", padx=20, pady=10)
            notebook = ttk.Notebook(win)
            notebook.pack(fill="both", expand=True, padx=10, pady=5)
            
            # --- TAB 1: Schedule ---
            tab_schedule = tk.Frame(notebook, bg=COLOR_BG)
            notebook.add(tab_schedule, text="  📅 حصص الطالب  ")

            f_filter = tk.Frame(tab_schedule, bg=COLOR_BG)
            f_filter.pack(fill="x", padx=10, pady=10)
            v_day = tk.StringVar(value="الكل")
            v_period = tk.StringVar(value="الكل")
            tk.Label(f_filter, text="تصفية حسب اليوم:", bg=COLOR_BG).pack(side="right", padx=5)
            cb_day = ttk.Combobox(f_filter, values=(["الكل"] + DAYS), textvariable=v_day, state="readonly", width=12)
            cb_day.pack(side="right", padx=5)
            tk.Label(f_filter, text="الحصة:", bg=COLOR_BG).pack(side="right", padx=5)
            cb_period = ttk.Combobox(f_filter, values=(["الكل"] + [str(i) for i in range(1, 8)]), textvariable=v_period, state="readonly", width=8)
            cb_period.pack(side="right", padx=5)
            s_class = tidy(student_data.get("الصف", ""))
            s_section = tidy(student_data.get("الشعبة", ""))
            raw_results, _, _ = self.m.find_student_schedule(student_name, s_class, s_section)
            cols = ['اليوم', 'الحصة', 'من', 'إلى', 'المعلم', 'المادة/الصف']
            tv_sch = ttk.Treeview(tab_schedule, columns=cols, show="headings", height=12)
            for c in cols:
                tv_sch.heading(c, text=c)
                tv_sch.column(c, anchor="center", width=(200 if c == "المادة/الصف" else 100))
            vsb_sch = ttk.Scrollbar(tab_schedule, orient="vertical", command=tv_sch.yview)
            tv_sch.configure(yscrollcommand=vsb_sch.set)
            vsb_sch.pack(side="left", fill="y")
            tv_sch.pack(side="right", fill="both", expand=True)

            def apply_filters(*args):
                for item in tv_sch.get_children():
                    tv_sch.delete(item)
                d_filter = v_day.get()
                p_filter = v_period.get()
                for r in raw_results:
                    if d_filter != "الكل" and r.get("اليوم") != d_filter: continue
                    if p_filter != "الكل" and str(r.get("الحصة")) != p_filter: continue
                    tv_sch.insert("", "end", values=[r.get(c, "") for c in cols])

            cb_day.bind("<<ComboboxSelected>>", apply_filters)
            cb_period.bind("<<ComboboxSelected>>", apply_filters)
            apply_filters()

            # --- TAB 2: Personal Data ---
            tab_personal = tk.Frame(notebook, bg=COLOR_BG)
            notebook.add(tab_personal, text="  👤 بيانات خاصة  ")
            canvas = tk.Canvas(tab_personal, bg=COLOR_BG)
            scrollbar = ttk.Scrollbar(tab_personal, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
            scrollbar.pack(side="right", fill="y")
            
            entries = {}
            for row_i, key in enumerate(sorted(student_data.keys())):
                ttk.Label(scrollable_frame, text=(key + ":")).grid(row=row_i, column=0, sticky="e", padx=6, pady=6)
                ent = ttk.Entry(scrollable_frame, width=40)
                ent.grid(row=row_i, column=1, sticky="w", padx=6, pady=6)
                ent.insert(0, tidy(student_data.get(key, "")))
                entries[key] = ent

            def save_personal_data():
                updated = {k: tidy(e.get()) for k, e in entries.items()}
                self.m.update_student_full(idx, updated)
                messagebox.showinfo("تم", "تم حفظ البيانات الخاصة بنجاح.")

            btn_save = tk.Button(scrollable_frame, text="💾 حفظ التعديلات", command=save_personal_data, bg=COLOR_ACCENT, fg="white", font=('Segoe UI', 10, 'bold'))
            if getattr(self, "admin_mode", False) == False:
                btn_save.config(state="disabled", bg="#ccc")
            btn_save.grid(row=len(student_data), column=0, columnspan=2, pady=20, sticky="ew")
        except Exception as e:
            messagebox.showerror("خطأ", f"حدث خطأ أثناء فتح الملف: {e}")

        # ========================================

    # --- show_student_details_sort_key ---
    def show_student_details_sort_key(r=None):
        # ========================================
        d = r.get("اليوم", "")
        if d in DAYS:
            d_idx = DAYS.index(d)
            day_priority = (d_idx - today_idx) % len(DAYS)
        else:
            day_priority = 99
        p = r.get("الحصة", "0")
        try:
            p_val = int(p)
        except:
            p_val = 99
        else:
            return (
             day_priority, p_val)

        # ========================================

    # --- show_student_details_sort_key_std ---
    def show_student_details_sort_key_std(r=None):
        # ========================================
        d = r.get("اليوم", "")
        d_idx = DAYS.index(d) if d in DAYS else 99
        try:
            p_val = int(r.get("الحصة", "0"))
        except:
            p_val = 99
        else:
            return (
             d_idx, p_val)

        # ========================================

    # --- export_student_schedule_excel ---
    def export_student_schedule_excel(self, student_name=None, results=None):
        # ========================================
        if not results:
            messagebox.showwarning("تنبيه", "لا توجد بيانات للتصدير.")
            return
        try:
            df = pd.DataFrame(results)
            cols_order = [
             'اليوم', 'الحصة', 'من', 'إلى', 'المادة/الصف', 'المعلم']
            final_cols = [c for c in cols_order if c in df.columns]
            df = df[final_cols]
            clean_name = re.sub('[\\\\/*?:"<>|]', "", student_name)
            filename = f'جدول_الطالب_{clean_name}_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
            path = os.path.join(EXPORT_DIR, filename)
            self.save_formatted_excel(df, path, sheet_name="جدول الطالب")
            if messagebox.askyesno("تم التصدير", f"تم إصدار جدول الطالب الرسمي بنجاح:\n{path}\n\nهل تريد فتح المجلد؟"):
                os.startfile(EXPORT_DIR)
        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"حدث خطأ أثناء التصدير:\n{e}")
            finally:
                pass

        # ========================================

    # --- open_exit_permit_dialog ---
    def open_exit_permit_dialog(self, student_name=None, results=None, current_loc=None):
        # ========================================
        if not HAS_DOCX:
            messagebox.showerror("خطأ", "مكتبة python-docx غير مثبتة. لا يمكن تصدير ملف Word.")
            return
        dialog = tk.Toplevel(self)
        dialog.title("إصدار إذن خروج طالب")
        dialog.geometry("400x320")
        dialog.grab_set()
        ttk.Label(dialog, text=f"إذن خروج: {student_name}", font=('Segoe UI', 11, 'bold')).pack(pady=15)
        ttk.Label(dialog, text="اختر سبب الخروج:").pack(pady=5)
        cb_reason = ttk.Combobox(dialog, values=EXIT_REASONS, width=30)
        cb_reason.pack(pady=5)
        cb_reason.current(0)
        ttk.Label(dialog, text="اسم المستلم (ولي الأمر) - اختياري:").pack(pady=5)
        entry_receiver = ttk.Entry(dialog, width=30)
        entry_receiver.pack(pady=5)

        def do_export():
            reason = tidy(cb_reason.get())
            receiver = tidy(entry_receiver.get())
            dialog.destroy()
            self.generate_exit_permit_docx(student_name, reason, receiver, current_loc)


        tk.Button(dialog, text="✅ إصدار وتوثيق (Word)", bg=COLOR_ACCENT, fg="white", command=do_export,
          font=('Segoe UI', 10, 'bold')).pack(pady=20)

        # ========================================

    # --- open_exit_permit_dialog_do_export ---
    def open_exit_permit_dialog_do_export():
        # ========================================
        reason = tidy(cb_reason.get())
        receiver = tidy(entry_receiver.get())
        dialog.destroy()
        self.generate_exit_permit_docx(student_name, reason, receiver, current_loc)

        # ========================================

    # --- generate_exit_permit_docx ---
    def generate_exit_permit_docx(self, student_name=None, reason=None, receiver=None, current_loc=None):
        # ========================================
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            doc = Document()
            if os.path.exists(LOGO_PATH):
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_logo = p_logo.add_run()
                run_logo.add_picture(LOGO_PATH, width=(Inches(1.8)))
            p_gov = doc.add_paragraph()
            p_gov.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run_rtl(p_gov, "المملكة العربية السعودية\nوزارة التعليم\nمدرسة الملك خالد المتوسطة", bold=True, size=12)
            doc.add_paragraph("")
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_rtl(p_title, "🛡️ نموذج استئذان / خروج طالب 🛡️", bold=True, size=24)
            doc.add_paragraph("")
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
            p_date = doc.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run_rtl(p_date, f"التاريخ والوقت: {now_str}", size=12)
            doc.add_paragraph("")
            p_body = doc.add_paragraph()
            p_body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run_rtl(p_body, "تفيد إدارة المدرسة بأنه تم الترخيص للخروج للطالب: ", size=14)
            add_run_rtl(p_body, (f"{student_name}"), bold=True, size=16)
            if current_loc:
                add_run_rtl(p_body, f'\nالموجود حالياً في حصة: {current_loc["المادة/الصف"]} مع المعلم {current_loc["المعلم"]}', size=12)
                add_run_rtl(p_body, f'\n(الحصة رقم {current_loc["الحصة"]})', size=11)
            add_run_rtl(p_body, "\n\n", size=14)
            add_run_rtl(p_body, f"سبب الخروج: {reason}", size=14)
            if receiver:
                add_run_rtl(p_body, f"\nالمستلم (ولي الأمر): {receiver}", size=14)
            doc.add_paragraph("")
            doc.add_paragraph("")
            p_ack = doc.add_paragraph()
            p_ack.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run_rtl(p_ack, "أقر أنا الموقع أدناه بصحة البيانات وتحمل مسؤولية الطالب بعد مغادرة المدرسة.", size=11)
            doc.add_paragraph("")
            doc.add_paragraph("")
            table = doc.add_table(rows=1, cols=3)
            c0 = table.cell(0, 2)
            c1 = table.cell(0, 1)
            c2 = table.cell(0, 0)
            p0 = c0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_rtl(p0, "توقيع الطالب", bold=True)
            p1 = c1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_rtl(p1, "توقيع ولي الأمر", bold=True)
            p2 = c2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_rtl(p2, "ختم وإدارة المدرسة", bold=True)
            clean_name = re.sub('[\\\\/*?:"<>|]', "", student_name)
            filename = f'إذن_خروج_{clean_name}_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
            path = os.path.join(EXPORT_DIR, filename)
            doc.save(path)
            if messagebox.askyesno("تم الإصدار", f"تم إنشاء نموذج الخروج بنجاح:\n{path}\n\nهل تريد فتح الملف للمعاينة والطباعة؟"):
                os.startfile(path)
        except Exception as e:
            try:
                messagebox.showerror("خطأ Word", str(e))
            finally:
                pass

        # ========================================

    # --- page_manager ---
    def page_manager(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="لوحة المدير (جدول الأسبوع)", style="Header.TLabel").pack(anchor="e",
          padx=6,
          pady=6)
        btns = tk.Frame(page, bg=COLOR_BG)
        btns.pack(fill="x", padx=6, pady=6)
        ttk.Button(btns, text="تحديث", command=(lambda: reload())).pack(side="right", padx=6)
        ttk.Button(btns, text="⬇️ تصدير Excel", command=(lambda: export())).pack(side="right", padx=6)

        def do_sync_import():
            zip_path = filedialog.askopenfilename(title="اختر ملف العمل (ZIP)", filetypes=[('Zip Files', '*.zip')])
            if not zip_path:
                return
            else:
                sm = SyncManager()
                success, msg = sm.import_work(zip_path, self.m)
                if success:
                    messagebox.showinfo("نجاح", msg)
                    reload()
                else:
                    messagebox.showerror("خطأ", msg)


        ttk.Button(btns, text="📥 استيراد بيانات (Flash)", command=do_sync_import).pack(side="right", padx=6)
        frame = tk.Frame(page, bg=COLOR_PANEL)
        frame.pack(fill="both", expand=True, padx=6, pady=6)
        m_periods = self.m.get_period_labels()
        cols = ["اليوم"] + [f"الحصة {p}" for p in m_periods]
        tv = ttk.Treeview(frame, columns=cols, show="headings", height=18)
        for c in cols:
            tv.heading(c, text=c)
            tv.column(c, anchor="e", width=(170 if c == "اليوم" else 220))
        else:
            vsb = ttk.Scrollbar(frame, orient="vertical", command=(tv.yview))
            tv.configure(yscrollcommand=(vsb.set))
            vsb.pack(side="left", fill="y")
            tv.pack(side="right", fill="both", expand=True)

            def reload(*args, **kwargs): pass  # TODO: Reconstruct method


            def export():
                try:
                    df = self.m.manager_week_table()
                    now = datetime.now().strftime("%Y-%m-%d_%H-%M")
                    path = os.path.join(EXPORT_DIR, f"لوحة_المدير_الرسمية_{now}.xlsx")
                    self.save_formatted_excel(df, path, sheet_name="لوحة المدير")
                    if messagebox.askyesno("تم التصدير", f"تم إنشاء تقرير المدير الرسمي:\n{path}\n\nهل تريد فتح المجلد؟"):
                        os.startfile(EXPORT_DIR)
                except Exception as e:
                    try:
                        messagebox.showerror("خطأ", str(e))
                    finally:
                        pass


            reload()
            return page# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- page_manager_do_sync_import ---
    def page_manager_do_sync_import():
        # ========================================
        zip_path = filedialog.askopenfilename(title="اختر ملف العمل (ZIP)", filetypes=[('Zip Files', '*.zip')])
        if not zip_path:
            return
        else:
            sm = SyncManager()
            success, msg = sm.import_work(zip_path, self.m)
            if success:
                messagebox.showinfo("نجاح", msg)
                reload()
            else:
                messagebox.showerror("خطأ", msg)

        # ========================================

    # --- page_manager_reload ---
    def page_manager_reload():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- page_manager_export ---
    def page_manager_export():
        # ========================================
        try:
            df = self.m.manager_week_table()
            now = datetime.now().strftime("%Y-%m-%d_%H-%M")
            path = os.path.join(EXPORT_DIR, f"لوحة_المدير_الرسمية_{now}.xlsx")
            self.save_formatted_excel(df, path, sheet_name="لوحة المدير")
            if messagebox.askyesno("تم التصدير", f"تم إنشاء تقرير المدير الرسمي:\n{path}\n\nهل تريد فتح المجلد؟"):
                os.startfile(EXPORT_DIR)
        except Exception as e:
            try:
                messagebox.showerror("خطأ", str(e))
            finally:
                pass

        # ========================================

    # --- page_detailed_schedule ---
    def page_detailed_schedule(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        ttk.Label(page, text="جدول المعلمين التفصيلي", style="Header.TLabel").pack(anchor="center",
          padx=6,
          pady=12)
        btns = tk.Frame(page, bg=COLOR_BG)
        btns.pack(fill="x", padx=6, pady=12)
        tk.Button(btns, text="🔨 بناء وعرض", command=(self.build_detailed_schedule), bg=COLOR_ACCENT,
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          height=2).pack(side="right",
          padx=10)
        tk.Button(btns, text="⬇️ تصدير Excel", command=(self.export_detailed_schedule_excel), bg=COLOR_XLSX,
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          height=2).pack(side="right",
          padx=10)
        self.schedule_canvas = tk.Canvas(page, bg=COLOR_BG)
        h_scroll = ttk.Scrollbar(page, orient="horizontal", command=(self.schedule_canvas.xview))
        v_scroll = ttk.Scrollbar(page, orient="vertical", command=(self.schedule_canvas.yview))
        self.schedule_frame = tk.Frame((self.schedule_canvas), bg=COLOR_BG)
        self.schedule_frame.bind("<Configure>", lambda e: self.schedule_canvas.configure(scrollregion=(self.schedule_canvas.bbox("all"))))
        self.schedule_canvas.create_window((0, 0), window=(self.schedule_frame), anchor="nw")
        self.schedule_canvas.configure(xscrollcommand=(h_scroll.set), yscrollcommand=(v_scroll.set))
        self.schedule_canvas.pack(side="top", fill="both", expand=True, padx=6, pady=6)
        h_scroll.pack(side="bottom", fill="x", padx=6)
        v_scroll.pack(side="right", fill="y", pady=6)
        return page

        # ========================================

    # --- build_detailed_schedule ---
    def build_detailed_schedule(self):
        # ========================================
        for w in self.schedule_frame.winfo_children():
            w.destroy()
            
        teacher_data = {}
        periods = self.m.get_period_labels()
        
        for t in self.m.teacher_names:
            try:
                df = self.m.get_teacher_grid(t)
            except Exception:
                continue
            
            schedule = {d: {} for d in DAYS}
            if df is not None and not df.empty:
                for d in DAYS:
                    for p in periods:
                        try:
                            match_d = df[df["اليوم"] == d]
                            if not match_d.empty and p in match_d.columns:
                                cell_val = match_d[p].values[0]
                                if cell_val and str(cell_val).strip() != "—":
                                    cell_text = tidy(cell_val)
                                    parsed = parse_teacher_cell(cell_text)
                                    if parsed:
                                        schedule[d][p] = parsed
                        except Exception:
                            pass
            teacher_data[t] = schedule

        header_bg = "#1b5e20"
        header_fg = "#ffffff"
        teacher_bg = "#e8f5e9"
        teacher_fg = "#1b5e20"
        cell_bg = "#ffffff"
        cell_fg = "#000000"
        p_count = len(periods)
        DAYS_REV = list(reversed(DAYS))
        P_REV = list(reversed(periods))
        
        row_idx = 0
        for d_idx, day in enumerate(DAYS_REV):
            tk.Label(self.schedule_frame, text=day, bg=header_bg, fg=header_fg, relief="solid", borderwidth=1, padx=10, pady=10, font=('Segoe UI', 11, 'bold')).grid(row=row_idx, column=(d_idx * p_count), columnspan=p_count, sticky="nsew")

        tk.Label(self.schedule_frame, text="المعلم", bg="#2e7d32", fg="#ffffff", relief="solid", borderwidth=1, padx=10, pady=10, font=('Segoe UI', 12, 'bold')).grid(row=row_idx, column=(len(DAYS) * p_count), sticky="nsew")
        row_idx += 1
        
        for d_idx, day in enumerate(DAYS_REV):
            for p_idx, p in enumerate(P_REV):
                tk.Label(self.schedule_frame, text=f"حصة {p}", bg="#43a047", fg="white", relief="solid", borderwidth=1, padx=6, pady=6, font=('Segoe UI', 9)).grid(row=row_idx, column=(d_idx * p_count + p_idx), sticky="nsew")

        tk.Label(self.schedule_frame, text="", bg="#2e7d32", relief="solid", borderwidth=1).grid(row=row_idx, column=(len(DAYS) * p_count), sticky="nsew")
        row_idx += 1
        
        for teacher_name, schedule_info in teacher_data.items():
            for d_idx, day in enumerate(DAYS_REV):
                for p_idx, p in enumerate(P_REV):
                    pdata = schedule_info.get(day, {}).get(p)
                    if pdata:
                        subject = pdata.get("subject", "")
                        cs = format_class_section_compact(pdata.get("class_section", ""))
                        txt = f"{subject}\n{cs}" if cs else subject
                        fg = "#000000"
                        this_cell_bg = "#f1f8e9"
                    else:
                        txt = ""
                        fg = "#999999"
                        this_cell_bg = "#ffffff"
                    tk.Label(self.schedule_frame, text=txt, bg=this_cell_bg, fg=fg, relief="solid", borderwidth=1, padx=8, pady=8, justify="center", anchor="center", font=('Segoe UI', 9)).grid(row=row_idx, column=(d_idx * p_count + p_idx), sticky="nsew")

            tk.Label(self.schedule_frame, text=teacher_name, bg=teacher_bg, fg=teacher_fg, relief="solid", borderwidth=1, padx=10, pady=8, anchor="center", font=('Segoe UI', 11, 'bold')).grid(row=row_idx, column=(len(DAYS) * p_count), sticky="nsew")
            row_idx += 1

        for i in range(len(DAYS) * p_count + 1):
            self.schedule_frame.grid_columnconfigure(i, weight=1, minsize=80)

        # ========================================

    # --- export_detailed_schedule_excel ---
    def export_detailed_schedule_excel(self):
        # ========================================
        try:
            rows = []
            header = ["المعلم"]
            periods = self.m.get_period_labels()
            for day in DAYS:
                for p in periods:
                    header.append(f"{day}-ح{p}")
                    
            for t in self.m.teacher_names:
                df = self.m.get_teacher_grid(t)
                row = [t]
                for day in DAYS:
                    for p in periods:
                        cell_text = ""
                        try:
                            if df is not None and not df.empty:
                                match_d = df[df["اليوم"] == day]
                                if not match_d.empty and p in match_d.columns:
                                    val = match_d[p].values[0]
                                    if val and str(val).strip() != "—":
                                        cell_text = tidy(val)
                        except Exception:
                            pass
                            
                        parsed = parse_teacher_cell(cell_text) if cell_text else None
                        if parsed:
                            cs = format_class_section_compact(parsed.get("class_section", ""))
                            row.append(f'{parsed["subject"]} | {cs}' if cs else parsed.get("subject", ""))
                        else:
                            row.append("")
                rows.append(row)

            out = pd.DataFrame(rows, columns=header)
            now = datetime.now().strftime("%Y-%m-%d_%H-%M")
            name = f"الجدول_التفصيلي_{now}.xlsx"
            path = os.path.join(EXPORT_DIR, name)
            self.save_formatted_excel(out, path, sheet_name="الجدول التفصيلي")
            if messagebox.askyesno("تم الحفظ", f"تم إنشاء الجدول التفصيلي الرسمي:\n{path}\n\nهل تريد فتح المجلد؟"):
                os.startfile(EXPORT_DIR)

        except Exception as e:
            messagebox.showerror("خطأ", str(e))

        # ========================================

    # --- page_calendar ---
    def page_calendar(self, parent=None):
        # ========================================
        page = tk.Frame(parent, bg=COLOR_BG)
        page.columnconfigure(0, weight=1)
        view_choice = tk.Frame(page, bg=COLOR_BG)
        view_holidays = tk.Frame(page, bg=COLOR_BG)
        view_study = tk.Frame(page, bg=COLOR_BG)

        def show_view(v):
            for frame in (
             view_choice, view_holidays, view_study):
                frame.pack_forget()
            else:
                v.pack(fill="both", expand=True)


        title_top = tk.Frame(view_choice, bg=COLOR_BG)
        title_top.pack(pady=40)
        ttk.Label(title_top, text="📅 التقويم الدراسي الذكي", style="Header.TLabel", font=('Segoe UI',
                                                                                          24,
                                                                                          'bold')).pack()
        ttk.Label(title_top, text="اختر القسم الذي تريد استعراضه:", foreground="#666").pack(pady=10)
        choice_btns = tk.Frame(view_choice, bg=COLOR_BG)
        choice_btns.pack(expand=True)

        def big_choice_btn(parent, text, icon, color, cmd):
            btn = tk.Button(parent, text=f"{icon} {text}", command=cmd, bg=color, fg="white", font=('Segoe UI',
                                                                                                    18,
                                                                                                    'bold'),
              width=25,
              height=3,
              relief="flat",
              cursor="hand2")
            btn.pack(pady=20)
            return btn


        big_choice_btn(choice_btns, "إجازات العام الدراسي", "🏖️", COLOR_DANGER, (lambda: (refresh_holidays(), show_view(view_holidays))))
        big_choice_btn(choice_btns, "الجدول الزمني للدراسة", "📚", COLOR_ACCENT, (lambda: (reload_study(), show_view(view_study))))
        h_nav = tk.Frame(view_holidays, bg=COLOR_BG)
        h_nav.pack(fill="x", padx=20, pady=10)
        tk.Button(h_nav, text="🔙 رجوع", command=(lambda: show_view(view_choice)), bg="#666", fg="white").pack(side="right")
        ttk.Label(h_nav, text="🏖️ مستعرض إجازات العام الدراسي", font=('Segoe UI', 16, 'bold'), foreground=COLOR_DANGER).pack(side="right", padx=20)
        h_content_frame = tk.Frame(view_holidays, bg=COLOR_BG)
        h_content_frame.pack(fill="both", expand=True)
        h_sidebar = tk.Frame(h_content_frame, bg=COLOR_PANEL, width=320, relief="groove", borderwidth=1)
        h_sidebar.pack(side="right", fill="y", padx=(0, 10), pady=10)
        h_sidebar.pack_propagate(False)
        ttk.Label(h_sidebar, text="📌 قائمة الإجازات", background=COLOR_PANEL, font=('Segoe UI',
                                                                                    11, 'bold')).pack(pady=10)
        h_btns_container = tk.Frame(h_sidebar, bg=COLOR_PANEL)
        h_btns_container.pack(fill="both", expand=True)
        h_display_area = tk.Frame(h_content_frame, bg=COLOR_BG)
        h_display_area.pack(side="left", fill="both", expand=True, padx=20, pady=20)

        def refresh_holidays():
            for w in h_btns_container.winfo_children():
                w.destroy()
            for w in h_display_area.winfo_children():
                w.destroy()

            major_list = self.m.get_holiday_summary_list()

            def show_h_stats(h, btn=None):
                for b in h_btns_container.winfo_children():
                    if isinstance(b, tk.Button):
                        b.config(bg="white", fg="#333")
                if btn:
                    btn.config(bg=COLOR_DANGER, fg="white")
                for w in h_display_area.winfo_children():
                    w.destroy()

                today_d = datetime.now()
                today_str_fmt = today_d.strftime("%Y-%m-%d")
                day_names_ar = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"]
                today_name_ar = day_names_ar[today_d.weekday()]

                header_lbl = tk.Label(h_display_area, text=f"تاريخ اليوم: {today_name_ar} | {today_str_fmt}", bg=COLOR_BG, fg="#444", font=('Segoe UI', 11, 'bold'))
                header_lbl.pack(side="top", anchor="center", pady=5)

                card_container = tk.Frame(h_display_area, bg=COLOR_BG)
                card_container.pack(anchor="center", pady=40)

                tk.Label(card_container, text=(h["name"]), bg=COLOR_BG, fg=COLOR_DANGER, font=('Segoe UI', 20, 'bold'), wraplength=800, justify="center").pack(pady=15)

                sd = h.get("start_dt")
                ed = h.get("end_dt")
                td = today_d.date()

                if sd and td < sd:
                    rem_days = (sd - td).days
                    tk.Label(card_container, text=f"⏳ باقي عليها {rem_days} يوم", bg=COLOR_BG, fg="#1565c0", font=('Segoe UI', 16, 'bold')).pack(pady=(0, 15))
                elif sd and ed and sd <= td <= ed:
                    tk.Label(card_container, text="🎉 الإجازة مستمرة حالياً", bg=COLOR_BG, fg="#388e3c", font=('Segoe UI', 16, 'bold')).pack(pady=(0, 15))

                dates_row = tk.Frame(card_container, bg=COLOR_BG)
                dates_row.pack()

                def date_box(parent, title, val_str, sub_val, color, is_duration=False):
                    f = tk.Frame(parent, bg=color, padx=35, pady=25, relief="flat")
                    f.pack(side="right", padx=15)
                    tk.Label(f, text=title, bg=color, fg="white", font=('Segoe UI', 12)).pack()
                    if is_duration:
                        tk.Label(f, text=(str(val_str)), bg=color, fg="white", font=('Segoe UI', 30, 'bold')).pack(pady=2)
                        tk.Label(f, text="يومًا", bg=color, fg="white", font=('Segoe UI', 12)).pack()
                    else:
                        tk.Label(f, text=f"{val_str}\n{sub_val}", bg=color, fg="white", font=('Segoe UI', 18, 'bold'), justify="center").pack(pady=5)

                date_box(dates_row, "📅 تاريخ البداية", h["start_h"] + " هـ" if h["start_h"] else "", h["start_dt"].strftime("%Y/%m/%d") + "م", "#2e7d32")
                date_box(dates_row, "⏳ مدة الإجازة", (h["duration"]), "", "#1565c0", is_duration=True)
                date_box(dates_row, "🏁 تاريخ النهاية", h["end_h"] + " هـ" if h["end_h"] else "", h["end_dt"].strftime("%Y/%m/%d") + "م", "#c62828")
                
                note_txt = h.get("note", "")
                if note_txt:
                    tk.Label(card_container, text=f"📝 {note_txt}", bg="#fff3e0", fg="#ef6c00", font=('Segoe UI', 12), pady=10, padx=20, relief="flat").pack(pady=20, fill="x", padx=50)

            if not major_list:
                tk.Label(h_display_area, text="لم يتم العثور على إجازات كبرى.", bg=COLOR_BG, font=('Segoe UI', 14)).pack()
                return
                
            for h in major_list:
                btn = tk.Button(h_btns_container, text=(h["name"]), bg="white", fg="#333", font=('Segoe UI', 9, 'bold'), padx=10, pady=8, relief="flat", cursor="hand2", anchor="e", justify="right", wraplength=280)
                btn.config(command=(lambda hh=h, bb=btn: show_h_stats(hh, bb)))
                btn.pack(fill="x", padx=5, pady=2)
                btn.bind("<Enter>", lambda e, b=btn: b.config(bg="#f0f0f0") if b["bg"] == "white" else None)
                btn.bind("<Leave>", lambda e, b=btn: b.config(bg="white") if b["bg"] == "white" else None)
            
            today_d = datetime.now().date()
            upcoming_idx = 0
            for i, h in enumerate(major_list):
                sd = h.get("start_dt")
                if sd and (sd >= today_d or (h.get("end_dt") and today_d <= h.get("end_dt"))):
                    upcoming_idx = i
                    break

            if h_btns_container.winfo_children():
                show_h_stats(major_list[upcoming_idx], h_btns_container.winfo_children()[upcoming_idx])


        s_nav = tk.Frame(view_study, bg=COLOR_BG)
        s_nav.pack(fill="x", padx=20, pady=10)
        tk.Button(s_nav, text="🔙 رجوع", command=(lambda: show_view(view_choice)), bg="#666", fg="white").pack(side="right")
        ttk.Label(s_nav, text="📚 الجدول الزمني للدراسة", font=('Segoe UI', 16, 'bold'), foreground=COLOR_ACCENT).pack(side="right", padx=20)
        s_content_frame = tk.Frame(view_study, bg=COLOR_BG)
        s_content_frame.pack(fill="both", expand=True)
        s_sidebar = tk.Frame(s_content_frame, bg=COLOR_PANEL, width=200, relief="groove", borderwidth=1)
        s_sidebar.pack(side="right", fill="y", padx=(0, 10), pady=10)
        s_sidebar.pack_propagate(False)
        ttk.Label(s_sidebar, text="📂 الفصول الدراسية", background=COLOR_PANEL, font=('Segoe UI',
                                                                                     11,
                                                                                     'bold')).pack(pady=10)
        self.active_study_term = tk.StringVar(value="")
        s_table_area = tk.Frame(s_content_frame, bg=COLOR_PANEL)
        s_table_area.pack(side="left", fill="both", expand=True, padx=20, pady=10)
        self.term_summary_frame = tk.Frame(s_table_area, bg=COLOR_PANEL)
        self.term_summary_frame.pack(fill="x", pady=(0, 10))
        cols = [
         'الفصل الدراسي', 'الأسبوع', 'اليوم', 'التاريخ الهجري', 'التاريخ الميلادي', 
         'الملاحظات']
        self.cal_tv = ttk.Treeview(s_table_area, columns=cols, show="headings", height=15)
        for c in cols:
            self.cal_tv.heading(c, text=c)
            self.cal_tv.column(c, anchor="center", width=(100 if c != "الملاحظات" else 450))
        else:
            self.cal_tv.tag_configure("holiday", background="#ffebee")
            self.cal_tv.tag_configure("today", background="#e8f5e9", font=('Segoe UI', 10,
                                                                           'bold'))
            vsb = ttk.Scrollbar(s_table_area, orient="vertical", command=(self.cal_tv.yview))
            hsb = ttk.Scrollbar(s_table_area, orient="horizontal", command=(self.cal_tv.xview))
            self.cal_tv.configure(yscrollcommand=(vsb.set), xscrollcommand=(hsb.set))
            vsb.pack(side="left", fill="y")
            hsb.pack(side="bottom", fill="x")
            self.cal_tv.pack(side="right", fill="both", expand=True)

            def set_study_term(term, btn=None):
                for b in s_sidebar.winfo_children():
                    if isinstance(b, tk.Button):
                        b.config(bg="white", fg="#333")
                else:
                    if btn:
                        btn.config(bg=COLOR_ACCENT, fg="white")
                    self.active_study_term.set(term)
                    reload_study()


            try:
                terms = ["الفصل الدراسي الأول", "الفصل الدراسي الثاني"]
            except:
                terms = [
                 "الفصل الأول", "الفصل الثاني", "الفصل الثالث"]
            else:
                for t in terms:
                    btn = tk.Button(s_sidebar, text=t, bg="white", fg="#333", font=('Segoe UI',
                                                                                    10, 'bold'), padx=10,
                      pady=10,
                      relief="flat",
                      cursor="hand2")
                    btn.config(command=(lambda tt=t, bb=btn: set_study_term(tt, bb)))
                    btn.pack(fill="x", padx=5, pady=2)
                else:
                    def reload_study(*args, **kwargs):
                        term = self.active_study_term.get()
                        for c in self.cal_tv.get_children():
                            self.cal_tv.delete(c)
                        
                        df_view = self.m.filter_calendar(term=term) if term else self.m.filter_calendar()
                        
                        today_str = datetime.now().strftime("%Y-%m-%d")
                        for _, row in df_view.iterrows():
                            vals = [
                                row.get("الفصل الدراسي", ""),
                                row.get("الأسبوع", ""),
                                row.get("اليوم", ""),
                                tidy(row.get("التاريخ الهجري", "")),
                                safe_parse_date(row.get("التاريخ الميلادي", "")) and safe_parse_date(row.get("التاريخ الميلادي", "")).strftime("%Y-%m-%d") or row.get("التاريخ الميلادي", ""),
                                tidy(row.get("الملاحظات", ""))
                            ]
                            tags = ()
                            if str(row.get("الملاحظات", "")).strip() and str(row.get("الملاحظات", "")).strip() != "nan":
                                tags = ("holiday",)
                            
                            sd = safe_parse_date(row.get("التاريخ الميلادي", ""))
                            if sd and sd.strftime("%Y-%m-%d") == today_str:
                                tags = ("today",)
                                
                            self.cal_tv.insert("", "end", values=vals, tags=tags)

                        for w in self.term_summary_frame.winfo_children():
                            w.destroy()
                        
                        today_d = datetime.now()
                        today_str_fmt = today_d.strftime("%Y-%m-%d")
                        day_names_ar = ["الاثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت", "الأحد"]
                        today_name_ar = day_names_ar[today_d.weekday()]
                        
                        header_lbl = tk.Label(self.term_summary_frame, text=f"تاريخ اليوم: {today_name_ar} | {today_str_fmt}", bg=COLOR_PANEL, fg="#444", font=('Segoe UI', 11, 'bold'))
                        header_lbl.pack(side="top", anchor="n", pady=2)
                        
                        dates = []
                        if not df_view.empty:
                            for d_str in df_view["التاريخ الميلادي"].dropna().astype(str):
                                d = safe_parse_date(d_str)
                                if d: dates.append(d)
                                
                        if dates and term:
                            sd = min(dates)
                            if hasattr(sd, 'date'): sd = sd.date()
                            ed = max(dates)
                            if hasattr(ed, 'date'): ed = ed.date()
                            td = today_d.date()
                            
                            if td > ed:
                                status_txt = f"انتهى الفصل الدراسي بالكامل ✅"
                                color = "#388e3c"
                                stats_frame = tk.Frame(self.term_summary_frame, bg=color, padx=10, pady=5)
                                stats_frame.pack(fill="x", padx=10, pady=5)
                                tk.Label(stats_frame, text=f"{term}: {status_txt}", bg=color, fg="white", font=('Segoe UI', 12, 'bold')).pack()
                            else:
                                total_days = (ed - sd).days + 1
                                if td < sd:
                                    passed_days = 0
                                    left_days = total_days
                                    status_txt = "لم يبدأ بعد (في الانتظار ⏳)"
                                    color = "#607d8b"
                                else:
                                    passed_days = (td - sd).days + 1
                                    left_days = (ed - td).days
                                    if left_days < 0: left_days = 0
                                    status_txt = "مستمر حالياً 🚀"
                                    color = COLOR_ACCENT
                                    
                                stats_frame = tk.Frame(self.term_summary_frame, bg=color, padx=10, pady=5)
                                stats_frame.pack(fill="x", padx=10, pady=5)
                                
                                title_txt = f"{term} | الحالة: {status_txt}"
                                info_txt = f"تاريخ البداية: {sd.strftime('%Y-%m-%d')}   |   تاريخ النهاية: {ed.strftime('%Y-%m-%d')}\nعدد الأيام المنقضية: {passed_days} يوم   |   عدد الأيام المتبقية: {left_days} يوم"
                                tk.Label(stats_frame, text=title_txt, bg=color, fg="white", font=('Segoe UI', 12, 'bold')).pack()
                                tk.Label(stats_frame, text=info_txt, bg=color, fg="white", font=('Segoe UI', 11)).pack(pady=2)

                    if s_sidebar.winfo_children():
                        first_btn = s_sidebar.winfo_children()[1] if len(s_sidebar.winfo_children()) > 1 else None
                        if isinstance(first_btn, tk.Button):
                            self.active_study_term.set(terms[0])
                            first_btn.config(bg=COLOR_ACCENT, fg="white")

                    if self.admin_mode:
                        adm_ctrl = tk.Frame(view_study, bg=COLOR_BG)
                        adm_ctrl.pack(fill="x", padx=20, pady=5)
                        tk.Button(adm_ctrl, text="➕ إضافة يوم", command=(lambda: open_editor(None)), bg=COLOR_ACCENT, fg="white", padx=15).pack(side="right", padx=5)
                        tk.Button(adm_ctrl, text="📝 تعديل المحدد", command=(lambda: edit_selected()), bg="#1976d2", fg="white", padx=15).pack(side="right", padx=5)
                        tk.Button(adm_ctrl, text="🗑️ حذف المحدد", command=(lambda: delete_selected()), bg=COLOR_DANGER, fg="white", padx=15).pack(side="right", padx=5)
                    show_view(view_choice)
                    return page# [DECOMPILATION FAILED]: Deparsing stopped due to parse error


        # ========================================

    # --- page_calendar_show_view ---
    def page_calendar_show_view(v=None):
        # ========================================
        for frame in (
         view_choice, view_holidays, view_study):
            frame.pack_forget()
        else:
            v.pack(fill="both", expand=True)

        # ========================================

    # --- page_calendar_big_choice_btn ---
    def page_calendar_big_choice_btn(parent=None, text=None, icon=None, color=None, cmd=None):
        # ========================================
        btn = tk.Button(parent, text=f"{icon} {text}", command=cmd, bg=color, fg="white", font=('Segoe UI',
                                                                                                18,
                                                                                                'bold'),
          width=25,
          height=3,
          relief="flat",
          cursor="hand2")
        btn.pack(pady=20)
        return btn

        # ========================================

    # --- page_calendar_refresh_holidays ---
    def page_calendar_refresh_holidays():
        # ========================================
        for w in h_btns_container.winfo_children():
            w.destroy()
        else:
            for w in h_display_area.winfo_children():
                w.destroy()
            else:
                major_list = self.m.get_holiday_summary_list()

                def show_h_stats(h, btn=None):
                    for b in h_btns_container.winfo_children():
                        if isinstance(b, tk.Button):
                            b.config(bg="white", fg="#333")
                        if btn:
                            btn.config(bg=COLOR_DANGER, fg="white")
                        for w in h_display_area.winfo_children():
                            w.destroy()
                        else:
                            card_container = tk.Frame(h_display_area, bg=COLOR_BG)
                            card_container.pack(anchor="center", pady=50)
                            tk.Label(card_container, text=(h["name"]), bg=COLOR_BG, fg=COLOR_DANGER, font=('Segoe UI',
                                                                                                           20,
                                                                                                           'bold'), wraplength=800, justify="center").pack(pady=20)
                            dates_row = tk.Frame(card_container, bg=COLOR_BG)
                            dates_row.pack()

                            def date_box(parent, title, val_str, sub_val, color, is_duration=False):
                                f = tk.Frame(parent, bg=color, padx=35, pady=25, relief="flat")
                                f.pack(side="right", padx=15)
                                tk.Label(f, text=title, bg=color, fg="white", font=('Segoe UI',
                                                                                    12)).pack()
                                if is_duration:
                                    tk.Label(f, text=(str(val_str)), bg=color, fg="white", font=('Segoe UI',
                                                                                                 30,
                                                                                                 'bold')).pack(pady=2)
                                    tk.Label(f, text="يومًا", bg=color, fg="white", font=('Segoe UI',
                                                                                          12)).pack()
                                else:
                                    tk.Label(f, text=f"{val_str}\n{sub_val}", bg=color, fg="white", font=('Segoe UI',
                                                                                                          18,
                                                                                                          'bold'),
                                      justify="center").pack(pady=5)

                            date_box(dates_row, "📅 تاريخ البداية", h["start_h"] + " هـ" if h["start_h"] else "", h["start_dt"].strftime("%Y/%m/%d") + "م", "#2e7d32")
                            date_box(dates_row, "⏳ مدة الإجازة", (h["duration"]), "", "#1565c0", is_duration=True)
                            date_box(dates_row, "🏁 تاريخ النهاية", h["end_h"] + " هـ" if h["end_h"] else "", h["end_dt"].strftime("%Y/%m/%d") + "م", "#c62828")
                            note_txt = h.get("note", "")
                            if note_txt:
                                tk.Label(card_container, text=f"📝 {note_txt}", bg="#fff3e0", fg="#ef6c00", font=('Segoe UI',
                                                                                                                 12),
                                  pady=10,
                                  padx=20,
                                  relief="flat").pack(pady=20, fill="x", padx=50)


                if not major_list:
                    tk.Label(h_display_area, text="لم يتم العثور على إجازات كبرى.", bg=COLOR_BG, font=('Segoe UI',
                                                                                                       14)).pack()
                    return
                for h in major_list:
                    btn = tk.Button(h_btns_container, text=(h["name"]), bg="white",
                      fg="#333",
                      font=('Segoe UI', 9, 'bold'),
                      padx=10,
                      pady=8,
                      relief="flat",
                      cursor="hand2",
                      anchor="e",
                      justify="right",
                      wraplength=280)
                    btn.config(command=(lambda hh=h, bb=btn: show_h_stats(hh, bb)))
                    btn.pack(fill="x", padx=5, pady=2)
                    btn.bind("<Enter>", lambda e, b=btn: b.config(bg="#f0f0f0")  if b["bg"] == "white" else None)
                    btn.bind("<Leave>", lambda e, b=btn: b.config(bg="white")  if b["bg"] == "white" else None)
                else:
                    show_h_stats(major_list[0], h_btns_container.winfo_children()[0])

        # ========================================

    # --- page_calendar_refresh_holidays_show_h_stats ---
    def page_calendar_refresh_holidays_show_h_stats(h=None, btn=None):
        # ========================================
        for b in h_btns_container.winfo_children():
            if isinstance(b, tk.Button):
                b.config(bg="white", fg="#333")
            if btn:
                btn.config(bg=COLOR_DANGER, fg="white")
            for w in h_display_area.winfo_children():
                w.destroy()
            else:
                card_container = tk.Frame(h_display_area, bg=COLOR_BG)
                card_container.pack(anchor="center", pady=50)
                tk.Label(card_container, text=(h["name"]), bg=COLOR_BG, fg=COLOR_DANGER, font=('Segoe UI',
                                                                                               20,
                                                                                               'bold'), wraplength=800, justify="center").pack(pady=20)
                dates_row = tk.Frame(card_container, bg=COLOR_BG)
                dates_row.pack()

                def date_box(parent, title, val_str, sub_val, color, is_duration=False):
                    f = tk.Frame(parent, bg=color, padx=35, pady=25, relief="flat")
                    f.pack(side="right", padx=15)
                    tk.Label(f, text=title, bg=color, fg="white", font=('Segoe UI', 12)).pack()
                    if is_duration:
                        tk.Label(f, text=(str(val_str)), bg=color, fg="white", font=('Segoe UI',
                                                                                     30,
                                                                                     'bold')).pack(pady=2)
                        tk.Label(f, text="يومًا", bg=color, fg="white", font=('Segoe UI', 12)).pack()
                    else:
                        tk.Label(f, text=f"{val_str}\n{sub_val}", bg=color, fg="white", font=('Segoe UI',
                                                                                              18,
                                                                                              'bold'),
                          justify="center").pack(pady=5)


                date_box(dates_row, "📅 تاريخ البداية", h["start_h"] + " هـ" if h["start_h"] else "", h["start_dt"].strftime("%Y/%m/%d") + "م", "#2e7d32")
                date_box(dates_row, "⏳ مدة الإجازة", (h["duration"]), "", "#1565c0", is_duration=True)
                date_box(dates_row, "🏁 تاريخ النهاية", h["end_h"] + " هـ" if h["end_h"] else "", h["end_dt"].strftime("%Y/%m/%d") + "م", "#c62828")
                note_txt = h.get("note", "")
                if note_txt:
                    tk.Label(card_container, text=f"📝 {note_txt}", bg="#fff3e0", fg="#ef6c00", font=('Segoe UI',
                                                                                                     12),
                      pady=10,
                      padx=20,
                      relief="flat").pack(pady=20, fill="x", padx=50)

        # ========================================

    # --- page_calendar_refresh_holidays_show_h_stats_date_box ---
    def page_calendar_refresh_holidays_show_h_stats_date_box(parent=None, title=None, val_str=None, sub_val=None, color=None, is_duration=None):
        # ========================================
        f = tk.Frame(parent, bg=color, padx=35, pady=25, relief="flat")
        f.pack(side="right", padx=15)
        tk.Label(f, text=title, bg=color, fg="white", font=('Segoe UI', 12)).pack()
        if is_duration:
            tk.Label(f, text=(str(val_str)), bg=color, fg="white", font=('Segoe UI', 30, 'bold')).pack(pady=2)
            tk.Label(f, text="يومًا", bg=color, fg="white", font=('Segoe UI', 12)).pack()
        else:
            tk.Label(f, text=f"{val_str}\n{sub_val}", bg=color, fg="white", font=('Segoe UI',
                                                                                  18, 'bold'),
              justify="center").pack(pady=5)

        # ========================================

    # --- page_calendar_set_study_term ---
    def page_calendar_set_study_term(term=None, btn=None):
        # ========================================
        for b in s_sidebar.winfo_children():
            if isinstance(b, tk.Button):
                b.config(bg="white", fg="#333")
        else:
            if btn:
                btn.config(bg=COLOR_ACCENT, fg="white")
            self.active_study_term.set(term)
            reload_study()

        # ========================================

    # --- page_calendar_reload_study ---
    def page_calendar_reload_study():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- page_calendar_reload_study_mk_stat ---
    def page_calendar_reload_study_mk_stat(p=None, t=None, v=None, c=None):
        # ========================================
        l = tk.Label(p, text=t, font=('Segoe UI', 10, 'bold'), bg="white", fg="#555")
        l.pack(side="right", padx=10)
        l2 = tk.Label(p, text=v, font=('Segoe UI', 11), bg="white", fg=c)
        l2.pack(side="right", padx=2)
        tk.Frame(p, width=1, height=20, bg="#eee").pack(side="right", padx=10)

        # ========================================

    # --- page_calendar_open_editor ---
    def page_calendar_open_editor(item_iid=None):
        # ========================================
        win = tk.Toplevel(self)
        win.title("إضافة/تعديل يوم في التقويم")
        win.geometry("560x420")
        win.grab_set()
        data = {
         'الفصل الدراسي': '""', 'الأسبوع': '""', 'اليوم': '""', 'التاريخ الميلادي': '""', 
         'الملاحظات': '""'}
        if item_iid is not None:
            vals = self.cal_tv.item(item_iid)["values"]
            data = dict(zip(cols, vals))
        frm = ttk.Frame(win)
        frm.pack(fill="both", expand=True, padx=16, pady=16)
        entries = {}
        for i, c in enumerate(cols):
            ttk.Label(frm, text=(c + ":")).grid(row=i, column=0, sticky="e", padx=6, pady=6)
            ent = ttk.Entry(frm, width=42)
            ent.grid(row=i, column=1, sticky="w", padx=6, pady=6)
            ent.insert(0, tidy(data.get(c, "")))
            entries[c] = ent
        else:
            hint = ttk.Label(frm, text="ملاحظة: ضع كلمة (إجازة) في الملاحظات ليتم اعتبار اليوم إجازة.", foreground="#666")
            hint.grid(row=(len(cols)), column=0, columnspan=2, sticky="w", pady=6)

            def commit():
                new_data = {c: tidy(entries[c].get()) for c in cols}
                if new_data["التاريخ الميلادي"]:
                    if not safe_parse_date(new_data["التاريخ الميلادي"]):
                        messagebox.showwarning("تنبيه", "التاريخ الميلادي غير صحيح. استخدم YYYY-MM-DD.")
                        return
                elif item_iid is None:
                    self.m.add_calendar_row(new_data)
                else:
                    self.m.update_calendar_row(int(item_iid), new_data)
                reload()
                win.destroy()


            ttk.Button(frm, text="حفظ", command=commit).grid(row=(len(cols) + 1), column=0, columnspan=2, pady=12)
            win.bind("<Return>", lambda e: commit())

        # ========================================

    # --- page_calendar_open_editor_commit ---
    def page_calendar_open_editor_commit():
        # ========================================
        new_data = {c: tidy(entries[c].get()) for c in cols}
        if new_data["التاريخ الميلادي"]:
            if not safe_parse_date(new_data["التاريخ الميلادي"]):
                messagebox.showwarning("تنبيه", "التاريخ الميلادي غير صحيح. استخدم YYYY-MM-DD.")
                return
        elif item_iid is None:
            self.m.add_calendar_row(new_data)
        else:
            self.m.update_calendar_row(int(item_iid), new_data)
        reload()
        win.destroy()

        # ========================================

    # --- page_calendar_edit_selected ---
    def page_calendar_edit_selected():
        # ========================================
        sel = self.cal_tv.selection()
        if not sel:
            messagebox.showwarning("تنبيه", "اختر صفاً للتعديل.")
            return
        open_editor(sel[0])

        # ========================================

    # --- page_calendar_delete_selected ---
    def page_calendar_delete_selected():
        # ========================================
        sel = self.cal_tv.selection()
        if not sel:
            messagebox.showwarning("تنبيه", "اختر صفاً للحذف.")
            return
        else:
            return messagebox.askyesno("تأكيد", "هل تريد حذف السجلات المحددة؟") or None
        self.m.delete_calendar_rows([int(i) for i in sel])
        reload()

        # ========================================

    # --- reload_all ---
    def reload_all(self):
        # ========================================
        self.m.reload_all()
        if hasattr(self, "cb_teacher"):
            self.cb_teacher["values"] = self.m.teacher_names
        self.refresh_today_status()
        messagebox.showinfo("تم", "تم إعادة تحميل جميع الملفات.")

        # ========================================

    # --- save_formatted_excel ---
    def save_formatted_excel(self, df=None, path=None, sheet_name=None):
        # ========================================
        try:
            import xlsxwriter
        except ImportError:
            pass

        try:
            workbook = xlsxwriter.Workbook(path)
            worksheet = workbook.add_worksheet(sheet_name)
            worksheet.right_to_left()
            fmt_title = workbook.add_format({
             'bold': True, 'font_size': 18, 'align': '"center"', 'valign': '"vcenter"', 
             'font_color': '"#1b5e20"', 'font_name': '"Segoe UI"'})
            fmt_meta = workbook.add_format({
             'italic': True, 'font_size': 10, 'align': '"center"', 'valign': '"vcenter"', 
             'font_color': '"#555555"', 'font_name': '"Segoe UI"'})
            fmt_header = workbook.add_format({
             'bold': True, 'font_size': 12, 'align': '"center"', 'valign': '"vcenter"', 
             'bg_color': '"#317135"', 'font_color': '"white"', 'border': 1, 'font_name': '"Segoe UI"'})
            fmt_cell = workbook.add_format({
             'align': '"center"', 'valign': '"vcenter"', 'border': 1, 
             'bg_color': '"#ffffffff"', 'font_size': 11, 'text_wrap': True, 'font_name': '"Segoe UI"'})
            fmt_cell_alt = workbook.add_format({
             'align': '"center"', 'valign': '"vcenter"', 'border': 1, 
             'bg_color': '"#F2F9F2"', 'font_size': 11, 'text_wrap': True, 'font_name': '"Segoe UI"'})
            if os.path.exists(LOGO_PATH):
                try:
                    worksheet.insert_image("A1", LOGO_PATH, {'x_scale': 0.12, 'y_scale': 0.12, 'x_offset': 10, 'y_offset': 10})
                except Exception:
                    pass

            total_cols = len(df.columns)
            if total_cols < 1:
                total_cols = 1
            worksheet.merge_range(0, 0, 0, total_cols - 1, "المملكة العربية السعودية", fmt_title)
            worksheet.merge_range(1, 0, 1, total_cols - 1, "وزارة التعليم", fmt_title)
            worksheet.merge_range(2, 0, 2, total_cols - 1, APP_TITLE, fmt_title)
            ts = datetime.now().strftime("%Y-%m-%d %H:%M")
            worksheet.merge_range(4, 0, 4, total_cols - 1, f"تاريخ التصدير: {ts}", fmt_meta)
            worksheet.set_row(0, 30)
            worksheet.set_row(1, 30)
            worksheet.set_row(2, 35)
            start_row = 6
            worksheet.set_row(start_row, 30)
            for col_idx, col_name in enumerate(df.columns):
                worksheet.write(start_row, col_idx, str(col_name), fmt_header)
            else:
                start_data = start_row + 1
                for row_idx, row_data in enumerate(df.itertuples(index=False)):
                    current_row = start_data + row_idx
                    current_fmt = fmt_cell if row_idx % 2 == 0 else fmt_cell_alt
                    worksheet.set_row(current_row, 25)

                for col_idx, value in enumerate(row_data):
                    val_str = "" if pd.isna(value) else str(value)
                    worksheet.write(current_row, col_idx, val_str, current_fmt)
                else:
                    for i, col in enumerate(df.columns):
                        max_len = len(str(col))
                        column_data = df[col].astype(str).head(50)
                        if not column_data.empty:
                            max_data_len = column_data.map(len).max()
                            if max_data_len > max_len:
                                max_len = max_data_len
                            final_width = min(max(max_len + 4, 15), 50)
                            worksheet.set_column(i, i, final_width)
                        workbook.close()
                        print(f"✓ تم حفظ الملف بتنسيق احترافي (xlsxwriter): {path}")

        except Exception as e:
            try:
                full_error = traceback.format_exc()
                print(f"خطأ في التنسيق الاحترافي: {full_error}")
                try:
                    df.to_excel(path, index=False, sheet_name=sheet_name)
                    messagebox.showwarning("تنسيق محدود", f"تم حفظ الملف بنجاح، ولكن بدون التنسيق الاحترافي.\n\nالسبب: {str(e)[:100]}")
                except Exception as e2:
                    try:
                        raise Exception(f"فشل حفظ الملف:\n{str(e)}\n\nتفاصيل الخطأ:\n{full_error}")
                    finally:
                        e2 = None
                        del e2

            finally:
                pass

        # ========================================

    # --- export_teacher_excel ---
    def export_teacher_excel(self):
        # ========================================
        if not self.current_teacher:
            messagebox.showwarning("تنبيه", "اختر معلماً أولاً.")
            return
        try:
            df = self.m.get_teacher_grid(self.current_teacher)
            now = datetime.now().strftime("%Y-%m-%d_%H-%M")
            name = f"{self.current_teacher}_جدول_{now}.xlsx"
            path = os.path.join(EXPORT_DIR, name)
            self.save_formatted_excel(df, path, sheet_name="جدول المعلم")
            if messagebox.askyesno("تم الحفظ", f"تم إنشاء ملف الإكسل الاحترافي:\n{path}\n\nهل تريد فتح المجلد؟"):
                os.startfile(EXPORT_DIR)
        except Exception as e:
            try:
                messagebox.showerror("خطأ", str(e))
            finally:
                pass

        # ========================================

    # --- setup_new_user ---
    def setup_new_user(self):
        # ========================================
        confirm = messagebox.askyesno("تجهيز النظام", "هذه الأداة ستقوم بتثبيت المكتبات المطلوبة تلقائياً:\n\n• xlsxwriter (للتصدير الاحترافي)\n• openpyxl (لقراءة وكتابة Excel)\n\nهل تريد المتابعة؟")
        if not confirm:
            return
        progress_win = tk.Toplevel(self)
        progress_win.title("جاري التجهيز...")
        progress_win.geometry("500x300")
        progress_win.configure(bg="#4caf50")
        progress_win.resizable(False, False)
        tk.Label(progress_win, text="🌱 تجهيز النظام للمستخدم الجديد", font=('Segoe UI', 16,
                                                                            'bold'), bg="#4caf50",
          fg="white").pack(pady=20)
        status_label = tk.Label(progress_win, text="جاري الفحص...", font=('Segoe UI', 11), bg="#4caf50",
          fg="white",
          wraplength=450,
          justify="right")
        status_label.pack(pady=10)
        log_frame = tk.Frame(progress_win, bg="#1a1a1a", relief="sunken", bd=2)
        log_frame.pack(fill="both", expand=True, padx=20, pady=10)
        log_text = tk.Text(log_frame, bg="#1a1a1a", fg="#00ff00", font=('Consolas', 9), wrap="word",
          state="disabled")
        log_text.pack(fill="both", expand=True, padx=5, pady=5)

        def update_status(msg, log_msg=None):
            status_label.config(text=msg)
            if log_msg:
                log_text.config(state="normal")
                log_text.insert("end", f"{log_msg}\n")
                log_text.see("end")
                log_text.config(state="disabled")


        def run_setup():
            try:
                update_status("🔍 فحص بيئة Python...", ">>> البحث عن Python...")
                venv_python = os.path.join(os.path.dirname(sys.executable) if getattr(sys, "frozen", False) else os.getcwd(), ".venv", "Scripts", "python.exe(")
                if os.path.exists(venv_python):
                    python_exe = venv_python
                    update_status(")", "✓ تم العثور على البيئة الافتراضية: .venv(")
                else:
                    python_exe = sys.executable
                    update_status(")", f"✓ استخدام Python الأساسي: {python_exe}")
                update_status("📦 فحص المكتبات المثبتة...", ">>> فحص xlsxwriter...")
                try:
                    import xlsxwriter
                    update_status("", "✓ xlsxwriter مثبتة مسبقاً")
                    has_xlsxwriter = True
                except ImportError:
                    update_status("", "⚠ xlsxwriter غير مثبتة - سيتم التثبيت")
                    has_xlsxwriter = False
                else:
                    update_status("", ">>> فحص openpyxl...")
                try:
                    import openpyxl
                    update_status("", "✓ openpyxl مثبتة مسبقاً")
                    has_openpyxl = True
                except ImportError:
                    update_status("", "⚠ openpyxl غير مثبتة - سيتم التثبيت")
                    has_openpyxl = False
                else:
                    if has_xlsxwriter:
                        if not has_openpyxl:
                            libs_to_install = []
                            if not has_xlsxwriter:
                                libs_to_install.append("xlsxwriter")
                            if not has_openpyxl:
                                libs_to_install.append("openpyxl")
                            update_status(f'📥 تثبيت المكتبات: {", ".join(libs_to_install)}...', f'>>> pip install {" ".join(libs_to_install)}')
                            cmd = f'"{python_exe}" -m pip install {" ".join(libs_to_install)}'
                            process = subprocess.Popen(cmd, shell=True, stdout=(subprocess.PIPE), stderr=(subprocess.STDOUT),
                              text=True,
                              bufsize=1)
                            for line in process.stdout:
                                if line.strip():
                                    self.after(0, lambda l=line.strip(): update_status("", l[:100]))
                            else:
                                process.wait()
                                if process.returncode != 0:
                                    raise Exception("فشل تثبيت المكتبات")
                                update_status("", "✓ تم تثبيت جميع المكتبات بنجاح!")

                    else:
                        update_status("✅ جميع المكتبات مثبتة مسبقاً!", ">>> لا حاجة للتثبيت")
                    update_status("🔍 التحقق النهائي...", ">>> فحص نهائي...")
                    try:
                        import xlsxwriter, openpyxl
                        update_status("", "✓ xlsxwriter: جاهزة")
                        update_status("", "✓ openpyxl: جاهزة")
                    except ImportError as e:
                        try:
                            raise Exception(f"فشل التحقق: {str(e)}")
                        finally:
                            pass

                    else:
                        update_status("✅ اكتمل التجهيز بنجاح!", ">>> النظام جاهز للاستخدام! 🎉")
                        self.after(2000, lambda: progress_win.destroy())
                        self.after(2000, lambda: messagebox.showinfo("نجاح", "تم تجهيز النظام بنجاح! 🎉\n\nيمكنك الآن استخدام جميع ميزات التصدير الاحترافي.\n\nملاحظة: قد تحتاج لإعادة تشغيل التطبيق لتفعيل التغييرات."))
            except Exception as e:
                try:
                    update_status("❌ حدث خطأ!", f"!!! خطأ: {str(e)}")
                    self.after(0, lambda ex=e: messagebox.showerror("خطأ في التجهيز", f"فشل تثبيت المكتبات:\n\n{str(ex)}\n\nيرجى التأكد من:\n• اتصال الإنترنت\n• صلاحيات الكتابة على الجهاز"))
                    self.after(2000, lambda: progress_win.destroy())
                finally:
                    pass


        threading.Thread(target=run_setup, daemon=True).start()

        # ========================================

    # --- setup_new_user_update_status ---
    def setup_new_user_update_status(msg=None, log_msg=None):
        # ========================================
        status_label.config(text=msg)
        if log_msg:
            log_text.config(state="normal")
            log_text.insert("end", f"{log_msg}\n")
            log_text.see("end")
            log_text.config(state="disabled")

        # ========================================

    # --- setup_new_user_run_setup ---
    def setup_new_user_run_setup():
        # ========================================
        try:
            update_status("🔍 فحص بيئة Python...", ">>> البحث عن Python...")
            venv_python = os.path.join(os.path.dirname(sys.executable) if getattr(sys, "frozen", False) else os.getcwd(), ".venv", "Scripts", "python.exe(")
            if os.path.exists(venv_python):
                python_exe = venv_python
                update_status(")", "✓ تم العثور على البيئة الافتراضية: .venv(")
            else:
                python_exe = sys.executable
                update_status(")", f"✓ استخدام Python الأساسي: {python_exe}")
            update_status("📦 فحص المكتبات المثبتة...", ">>> فحص xlsxwriter...")
            try:
                import xlsxwriter
                update_status("", "✓ xlsxwriter مثبتة مسبقاً")
                has_xlsxwriter = True
            except ImportError:
                update_status("", "⚠ xlsxwriter غير مثبتة - سيتم التثبيت")
                has_xlsxwriter = False
            else:
                update_status("", ">>> فحص openpyxl...")
            try:
                import openpyxl
                update_status("", "✓ openpyxl مثبتة مسبقاً")
                has_openpyxl = True
            except ImportError:
                update_status("", "⚠ openpyxl غير مثبتة - سيتم التثبيت")
                has_openpyxl = False
            else:
                if has_xlsxwriter:
                    if not has_openpyxl:
                        libs_to_install = []
                        if not has_xlsxwriter:
                            libs_to_install.append("xlsxwriter")
                        if not has_openpyxl:
                            libs_to_install.append("openpyxl")
                        update_status(f'📥 تثبيت المكتبات: {", ".join(libs_to_install)}...', f'>>> pip install {" ".join(libs_to_install)}')
                        cmd = f'"{python_exe}" -m pip install {" ".join(libs_to_install)}'
                        process = subprocess.Popen(cmd, shell=True, stdout=(subprocess.PIPE), stderr=(subprocess.STDOUT),
                          text=True,
                          bufsize=1)
                        for line in process.stdout:
                            if line.strip():
                                self.after(0, lambda l=line.strip(): update_status("", l[:100]))
                        else:
                            process.wait()
                            if process.returncode != 0:
                                raise Exception("فشل تثبيت المكتبات")
                            update_status("", "✓ تم تثبيت جميع المكتبات بنجاح!")

                else:
                    update_status("✅ جميع المكتبات مثبتة مسبقاً!", ">>> لا حاجة للتثبيت")
                update_status("🔍 التحقق النهائي...", ">>> فحص نهائي...")
                try:
                    import xlsxwriter, openpyxl
                    update_status("", "✓ xlsxwriter: جاهزة")
                    update_status("", "✓ openpyxl: جاهزة")
                except ImportError as e:
                    try:
                        raise Exception(f"فشل التحقق: {str(e)}")
                    finally:
                        pass

                else:
                    update_status("✅ اكتمل التجهيز بنجاح!", ">>> النظام جاهز للاستخدام! 🎉")
                    self.after(2000, lambda: progress_win.destroy())
                    self.after(2000, lambda: messagebox.showinfo("نجاح", "تم تجهيز النظام بنجاح! 🎉\n\nيمكنك الآن استخدام جميع ميزات التصدير الاحترافي.\n\nملاحظة: قد تحتاج لإعادة تشغيل التطبيق لتفعيل التغييرات."))
        except Exception as e:
            try:
                update_status("❌ حدث خطأ!", f"!!! خطأ: {str(e)}")
                self.after(0, lambda ex=e: messagebox.showerror("خطأ في التجهيز", f"فشل تثبيت المكتبات:\n\n{str(ex)}\n\nيرجى التأكد من:\n• اتصال الإنترنت\n• صلاحيات الكتابة على الجهاز"))
                self.after(2000, lambda: progress_win.destroy())
            finally:
                pass

        # ========================================

    # --- developer_deploy ---
    def developer_deploy(self):
        # ========================================
        pw = simpledialog.askstring("أدخل الرمز السري", "الوصول محمي. أدخل الرمز السري:", show="*")
        if pw != "2210":
            if pw is not None:
                messagebox.showerror("خطأ", "الرمز السري غير صحيح!")
            return
            
        confirm = messagebox.askyesno("تأكيد", "هل أنت متأكد من رغبتك في إعادة بناء النسخة وتسليمها للمدير الآن؟\n(قد تستغرق العملية دقيقتين)\n\nسيتم نسخ جميع البيانات الحالية مع النسخة الجديدة.")
        if not confirm:
            return
        progress_win = tk.Toplevel(self)
        progress_win.title("جاري البناء والتسليم...")
        progress_win.geometry("600x400")
        progress_win.configure(bg="#263238")
        progress_win.resizable(False, False)
        tk.Label(progress_win, text="⚙️ جاري بناء وتسليم النظام", font=('Segoe UI', 16, 'bold'), bg="#263238",
          fg="#ffca28").pack(pady=20)
        status_label = tk.Label(progress_win, text="جاري التحضير...", font=('Segoe UI', 11), bg="#263238",
          fg="white",
          wraplength=550,
          justify="right")
        status_label.pack(pady=10)
        log_frame = tk.Frame(progress_win, bg="#1a1a1a", relief="sunken", bd=2)
        log_frame.pack(fill="both", expand=True, padx=20, pady=10)
        log_text = tk.Text(log_frame, bg="#1a1a1a", fg="#00ff00", font=('Consolas', 9), wrap="word",
          state="disabled")
        log_text.pack(fill="both", expand=True, padx=5, pady=5)

        def update_status(msg, log_msg=None):
            if msg:
                status_label.config(text=msg)
            if log_msg:
                log_text.config(state="normal")
                log_text.insert("end", f"{log_msg}\n")
                log_text.see("end")
                log_text.config(state="disabled")
                progress_win.update()


        def run_build(*args, **kwargs):
            try:
                import subprocess, sys, os, shutil
                self.after(0, lambda: update_status("جاري التجهيز لإنشاء النسخة التنفيذية...", ">>> بدء العملية"))
                
                # 1. Detect active file
                current_file = __file__
                if current_file.endswith('.pyc'):
                    current_file = current_file[:-1]
                
                # 2. PyInstaller command
                exe_name = "نظام_المدرسة"
                cmd = [
                    sys.executable, "-m", "PyInstaller",
                    "--noconfirm", "--onefile", "--windowed",
                    "--name", exe_name
                ]
                
                if os.path.exists("شعار_الوزارة.png"):
                    cmd.extend(["--add-data", "شعار_الوزارة.png;."])
                
                cmd.append(current_file)
                
                self.after(0, lambda: update_status("جاري البناء (قد يستغرق 2-3 دقائق)...", f">>> {' '.join(cmd)}"))
                
                process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    text=True,
                    bufsize=1,
                    creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                )
                
                for line in process.stdout:
                    if line.strip():
                        # Update terminal
                        self.after(0, lambda l=line.strip(): update_status("", l[:120]))
                
                process.wait()
                
                if process.returncode != 0:
                    self.after(0, lambda: update_status("فشلت عملية البناء!", ">>> خطأ في PyInstaller"))
                    self.after(0, lambda: messagebox.showerror("خطأ في البناء", "حدث خطأ أثناء تجميع البرنامج. راجع السجل لمعرفة السبب."))
                    return
                
                self.after(0, lambda: update_status("جاري تجهيز مجلد التسليم للمدير...", ">>> نسخ الملفات المهمة"))
                
                # 3. Prepare delivery folder
                delivery_dir = os.path.join(os.getcwd(), "النسخة_للمدير")
                if not os.path.exists(delivery_dir):
                    os.makedirs(delivery_dir)
                
                # 4. Copy the freshly built .exe
                dist_exe = os.path.join(os.getcwd(), "dist", f"{exe_name}.exe")
                final_exe = os.path.join(delivery_dir, f"{exe_name}.exe")
                if os.path.exists(dist_exe):
                    shutil.copy2(dist_exe, final_exe)
                    self.after(0, lambda: update_status("", f">>> تم نسخ التطبيق: {exe_name}.exe"))
                
                # 5. Copy Data Files
                files_to_copy = [f for f in os.listdir(os.getcwd()) if f.endswith(('.xlsx', '.json', '.png', '.docx'))]
                for f in files_to_copy:
                    src = os.path.join(os.getcwd(), f)
                    dst = os.path.join(delivery_dir, f)
                    if os.path.isfile(src) and src != dst:
                        shutil.copy2(src, dst)
                        self.after(0, lambda f=f: update_status("", f">>> تم نسخ: {f}"))
                
                self.after(0, lambda: update_status("✅ تمت العملية بنجاح!", ">>> كل شيء جاهز للتسليم"))
                self.after(2000, lambda: progress_win.destroy())
                
                if messagebox.askyesno("نجاح", "تم استخراج النظام وتجهيزه في مجلد 'النسخة_للمدير' بنجاح!\n\nهل تريد فتح المجلد الآن لمراجعته أو إرساله للمدير؟"):
                    os.startfile(delivery_dir)

            except Exception as e:
                self.after(0, lambda ex=e: update_status("❌ حدث خطأ غير متوقع", f">>> {str(ex)}"))
                self.after(0, lambda ex=e: messagebox.showerror("خطأ عام", f"فشلت العملية:\n{str(ex)}"))


        threading.Thread(target=run_build, daemon=True).start()


        # ========================================

    # --- developer_deploy_update_status ---
    def developer_deploy_update_status(msg=None, log_msg=None):
        # ========================================
        if msg:
            status_label.config(text=msg)
        if log_msg:
            log_text.config(state="normal")
            log_text.insert("end", f"{log_msg}\n")
            log_text.see("end")
            log_text.config(state="disabled")
            progress_win.update()

        # ========================================

    # --- developer_deploy_run_build ---
    def developer_deploy_run_build():
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- print_teacher_schedule ---
    def print_teacher_schedule(self):
        # ========================================
        if not self.current_teacher:
            messagebox.showwarning("تنبيه", "اختر معلماً أولاً.")
            return
        try:
            df = self.m.get_teacher_grid(self.current_teacher)
            now = datetime.now().strftime("%Y-%m-%d_%H-%M")
            name = f"{self.current_teacher}_جدول_{now}.xlsx"
            path = os.path.join(EXPORT_DIR, name)
            df.to_excel(path, index=False)
            try:
                os.startfile(path, "print")
                messagebox.showinfo("تم", f"تم إرسال الملف للطباعة:\n{path}")
            except Exception as e:
                try:
                    messagebox.showerror("خطأ في الطباعة", f"فشل أمر الطباعة التلقائي.\nاطبع يدوياً من:\n{path}\n\n{e}")
                    try:
                        os.startfile(path)
                    except Exception:
                        pass

                finally:
                    pass

        except Exception as e:
            try:
                messagebox.showerror("خطأ", str(e))
            finally:
                pass

        # ========================================

    # --- export_parent_contacts ---
    def export_parent_contacts(self):
        # ========================================
        try:
            name_col, class_col, section_col = self.m._detect_student_cols()
            parent_name_col = None
            parent_phone_col = None
            for col in self.m.df_students.columns:
                col_lower = str(col).lower()
                if "ولي" in col_lower and "أمر" in col_lower and "جوال" not in col_lower:
                    parent_name_col = col
                elif "جوال" in col_lower:
                    if "ولي" in col_lower or "طالب" in col_lower:
                        parent_phone_col = col
                    parent_phone_col or messagebox.showwarning("تنبيه", "لم يتم العثور على عمود 'جوال ولي الأمر' في ملف بيانات الطلاب!")
            else:
                return
                export_data = []
                for idx, row in self.m.df_students.iterrows():
                    student_name = tidy(row.get(name_col, ""))
                    class_name = tidy(row.get(class_col, ""))
                    section = tidy(row.get(section_col, ""))
                    parent_name = tidy(row.get(parent_name_col, "")) if parent_name_col else ""
                    parent_phone = tidy(row.get(parent_phone_col, ""))
                    if student_name:
                        if parent_phone:
                            export_data.append({
                             'اسم الطالب': student_name, 
                             'الصف': class_name, 
                             'الشعبة': section, 
                             'ولي الأمر': parent_name, 
                             'رقم الجوال': parent_phone})
                        export_data or messagebox.showwarning("تنبيه", "لا توجد بيانات للتصدير!")
                else:
                    return
                    df_export = pd.DataFrame(export_data)
                    df_export = df_export.sort_values(by=["الصف", "الشعبة", "اسم الطالب"])
                    now = datetime.now().strftime("%Y-%m-%d_%H-%M")
                    filename = f"أرقام_أولياء_الأمور_{now}.xlsx"
                    path = os.path.join(EXPORT_DIR, filename)
                    self.save_formatted_excel(df_export, path, sheet_name="أرقام أولياء الأمور")
                    if messagebox.askyesno("تم الحفظ", f"تم تصدير {len(export_data)} رقم بنجاح!\n\nالملف: {filename}\n\nهل تريد فتح المجلد؟"):
                        os.startfile(EXPORT_DIR)

        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"فشل تصدير الأرقام:\n{str(e)}")
            finally:
                pass

        # ========================================

    # --- send_whatsapp_to_parent ---
    def send_whatsapp_to_parent(self, student_name=None):
        # ========================================
        try:
            sender_file = data_path("whatsapp_senders.txt")
            saved_senders = []
            if os.path.exists(sender_file):
                with open(sender_file, "r", encoding="utf-8") as f:
                    saved_senders = [line.strip() for line in f if line.strip()]
            else:
                main_win = tk.Toplevel(self)
                main_win.title("إرسال رسالة واتساب")
                main_win.geometry("600x700")
                main_win.grab_set()
                tk.Label(main_win, text="📱 إرسال رسالة واتساب", font=('Segoe UI', 16, 'bold'),
                  fg="#25d366").pack(pady=15)
                sender_frame = tk.LabelFrame(main_win, text="1️⃣ رقم المرسل (المدرسة/المدير)", font=('Segoe UI',
                                                                                                     11,
                                                                                                     'bold'),
                  padx=15,
                  pady=10)
                sender_frame.pack(fill="x", padx=20, pady=10)
                sender_var = tk.StringVar()
                if saved_senders:
                    tk.Label(sender_frame, text="اختر من الأرقام المحفوظة:", font=('Segoe UI',
                                                                                   10)).pack(anchor="e", pady=5)
                    sender_combo = ttk.Combobox(sender_frame, textvariable=sender_var, values=saved_senders,
                      font=('Segoe UI', 11),
                      state="readonly")
                    sender_combo.pack(fill="x", pady=5)
                    sender_combo.current(0)
                    tk.Label(sender_frame, text="أو أضف رقم جديد:", font=('Segoe UI', 10)).pack(anchor="e", pady=(10,
                                                                                                                  5))
                else:
                    tk.Label(sender_frame, text="أدخل رقم جوال المدرسة/المدير:", font=('Segoe UI',
                                                                                       10)).pack(anchor="e", pady=5)
            new_sender_entry = tk.Entry(sender_frame, font=('Segoe UI', 11))
            new_sender_entry.pack(fill="x", pady=5)
            new_sender_entry.insert(0, "05xxxxxxxx")
            recipient_frame = tk.LabelFrame(main_win, text="2️⃣ اختر المستقبل", font=('Segoe UI',
                                                                                      11,
                                                                                      'bold'),
              padx=15,
              pady=10)
            recipient_frame.pack(fill="both", expand=True, padx=20, pady=10)
            toggle_frame = tk.Frame(recipient_frame)
            toggle_frame.pack(fill="x", pady=5)
            recipient_type = tk.StringVar(value="student")
            external_phone_var = tk.StringVar()
            student_container = tk.Frame(recipient_frame)
            search_frame = tk.Frame(student_container)
            search_frame.pack(fill="x", pady=5)
            tk.Label(search_frame, text="🔍", font=('Segoe UI', 14)).pack(side="right", padx=5)
            search_var = tk.StringVar()
            search_entry = tk.Entry(search_frame, textvariable=search_var, font=('Segoe UI',
                                                                                 12),
              fg="#666")
            search_entry.pack(fill="x", side="right", expand=True)
            search_entry.insert(0, "اكتب اسم الطالب...")
            student_listbox = tk.Listbox(student_container, font=('Segoe UI', 11), height=8,
              bg="#f5f5f5",
              selectbackground="#25d366")
            student_listbox.pack(fill="both", expand=True, pady=5)

            def on_search_focus_in(event):
                if search_entry.get() == "اكتب اسم الطالب...":
                    search_entry.delete(0, tk.END)
                    search_entry.config(fg="#000")


            def on_search_focus_out(event):
                if not search_entry.get():
                    search_entry.insert(0, "اكتب اسم الطالب...")
                    search_entry.config(fg="#666")


            def update_search(event=None):
                search_text = search_var.get().strip()
                if not search_text == "اكتب اسم الطالب...":
                    if not search_text:
                        student_listbox.delete(0, tk.END)
                        all_students = self.m.list_students_simple()
                        for student in all_students[:50]:
                            student_listbox.insert(tk.END, student)

                else:
                    student_listbox.delete(0, tk.END)
                    all_students = self.m.list_students_simple()
                    matching = [s for s in all_students if search_text.lower() in s.lower()]
                    for student in matching[:50]:
                        student_listbox.insert(tk.END, student)


            search_entry.bind("<FocusIn>", on_search_focus_in)
            search_entry.bind("<FocusOut>", on_search_focus_out)
            search_var.trace("w", lambda *args: update_search())
            external_container = tk.Frame(recipient_frame)
            tk.Label(external_container, text="📞 أدخل رقم الجوال:", font=('Segoe UI', 11)).pack(anchor="e", pady=10)
            external_entry = tk.Entry(external_container, textvariable=external_phone_var, font=('Segoe UI',
                                                                                                 13),
              justify="center")
            external_entry.pack(fill="x", pady=5, ipady=8)
            external_entry.insert(0, "05xxxxxxxx")
            tk.Label(external_container, text="مثال: 0501234567", font=('Segoe UI', 9),
              fg="#666").pack(pady=5)

            def show_student():
                recipient_type.set("student")
                external_container.pack_forget()
                student_container.pack(fill="both", expand=True)
                btn_student.config(bg="#25d366", fg="white")
                btn_external.config(bg="#e0e0e0", fg="#333")
                update_search()


            def show_external():
                recipient_type.set("external")
                student_container.pack_forget()
                external_container.pack(fill="both", expand=True)
                btn_external.config(bg="#25d366", fg="white")
                btn_student.config(bg="#e0e0e0", fg="#333")


            btn_student = tk.Button(toggle_frame, text="👨\u200d🎓 طالب من المدرسة", command=show_student,
              bg="#25d366",
              fg="white",
              font=('Segoe UI', 10, 'bold'),
              relief="flat",
              padx=15,
              pady=8,
              cursor="hand2")
            btn_student.pack(side="right", padx=5)
            btn_external = tk.Button(toggle_frame, text="📞 طرف خارجي", command=show_external,
              bg="#e0e0e0",
              fg="#333",
              font=('Segoe UI', 10, 'bold'),
              relief="flat",
              padx=15,
              pady=8,
              cursor="hand2")
            btn_external.pack(side="right", padx=5)
            student_container.pack(fill="both", expand=True)
            update_search()
            msg_frame = tk.LabelFrame(main_win, text="3️⃣ الرسالة", font=('Segoe UI', 11, 'bold'),
              padx=15,
              pady=10)
            msg_frame.pack(fill="both", expand=True, padx=20, pady=10)
            msg_text = tk.Text(msg_frame, font=('Segoe UI', 11), wrap="word", height=6)
            msg_text.pack(fill="both", expand=True, pady=5)
            msg_text.insert("1.0", "السلام عليكم ورحمة الله وبركاته\n\n")

            def send_message():
                sender_phone = new_sender_entry.get().strip()
                if sender_phone:
                    if sender_phone == "05xxxxxxxx":
                        sender_phone = sender_var.get()
                    if not sender_phone:
                        messagebox.showwarning("تنبيه", "الرجاء إدخال أو اختيار رقم المرسل!")
                        return
                    message = msg_text.get("1.0", "end-1c").strip()
                    if not message:
                        messagebox.showwarning("تنبيه", "الرجاء كتابة رسالة!")
                        return
                    if sender_phone not in saved_senders:
                        with open(sender_file, "a", encoding="utf-8") as f:
                            f.write(sender_phone + "\n")
                    if recipient_type.get() == "external":
                        parent_phone = external_phone_var.get().strip()
                        if not parent_phone or parent_phone == "05xxxxxxxx":
                            messagebox.showwarning("تنبيه", "الرجاء إدخال رقم الجوال!")
                            return
                        recipient_name = "الطرف الخارجي"
                    else:
                        selection = student_listbox.curselection()
                        if not selection:
                            messagebox.showwarning("تنبيه", "الرجاء اختيار طالب!")
                            return
                        selected_student = student_listbox.get(selection[0])
                        name_col, _, _ = self.m._detect_student_cols()
                        parent_phone_col = None
                        for col in self.m.df_students.columns:
                            col_lower = str(col).lower()
                            if "جوال" in col_lower:
                                if not "ولي" in col_lower:
                                    if "طالب" in col_lower:
                                        parent_phone_col = col
                                        break
                                    if not parent_phone_col:
                                        messagebox.showwarning("تنبيه", "لم يتم العثور على عمود رقم الجوال!")
                                        return
                                    student_row = None
                                    for idx, row in self.m.df_students.iterrows():
                                        if tidy(row.get(name_col, "")) == tidy(selected_student):
                                            student_row = row
                                            break
                                    else:
                                        if student_row is None:
                                            messagebox.showwarning("تنبيه", "لم يتم العثور على بيانات الطالب!")
                                            return
                                        else:
                                            parent_phone = tidy(student_row.get(parent_phone_col, ""))
                                            parent_phone or messagebox.showwarning("تنبيه", f"لا يوجد رقم جوال لولي أمر: {selected_student}")
                                            return
                                        recipient_name = f"ولي أمر {selected_student}"

                                parent_phone = parent_phone.replace(" ", "").replace("-", "")
                                if parent_phone.startswith("0"):
                                    parent_phone = "966" + parent_phone[1:]

                else:
                    if not parent_phone.startswith("966"):
                        parent_phone = "966" + parent_phone
                    import urllib.parse
                    encoded_message = urllib.parse.quote(message)
                    whatsapp_url = f"https://web.whatsapp.com/send?phone={parent_phone}&text={encoded_message}"
                    webbrowser.open(whatsapp_url)
                    messagebox.showinfo("تم", f"تم فتح واتساب ويب!\n\nالمرسل: {sender_phone}\nالمستقبل: {recipient_name}\n\nاضغط Enter للإرسال.")
                    main_win.destroy()


            btn_frame = tk.Frame(main_win)
            btn_frame.pack(pady=15)
            tk.Button(btn_frame, text="📤 إرسال عبر واتساب", command=send_message, bg="#25d366",
              fg="white",
              font=('Segoe UI', 12, 'bold'),
              padx=30,
              pady=12,
              cursor="hand2").pack(side="right", padx=5)
            tk.Button(btn_frame, text="إلغاء", command=(main_win.destroy), bg="#999",
              fg="white",
              font=('Segoe UI', 11),
              padx=30,
              pady=12,
              cursor="hand2").pack(side="right", padx=5)
        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"فشل إرسال الرسالة:\n{str(e)}")
            finally:
                pass

        # ========================================

    # --- send_whatsapp_to_parent_on_search_focus_in ---
    def send_whatsapp_to_parent_on_search_focus_in(event=None):
        # ========================================
        if search_entry.get() == "اكتب اسم الطالب...":
            search_entry.delete(0, tk.END)
            search_entry.config(fg="#000")

        # ========================================

    # --- send_whatsapp_to_parent_on_search_focus_out ---
    def send_whatsapp_to_parent_on_search_focus_out(event=None):
        # ========================================
        if not search_entry.get():
            search_entry.insert(0, "اكتب اسم الطالب...")
            search_entry.config(fg="#666")

        # ========================================

    # --- send_whatsapp_to_parent_update_search ---
    def send_whatsapp_to_parent_update_search(event=None):
        # ========================================
        search_text = search_var.get().strip()
        if not search_text == "اكتب اسم الطالب...":
            if not search_text:
                student_listbox.delete(0, tk.END)
                all_students = self.m.list_students_simple()
                for student in all_students[:50]:
                    student_listbox.insert(tk.END, student)

        else:
            student_listbox.delete(0, tk.END)
            all_students = self.m.list_students_simple()
            matching = [s for s in all_students if search_text.lower() in s.lower()]
            for student in matching[:50]:
                student_listbox.insert(tk.END, student)

        # ========================================

    # --- send_whatsapp_to_parent_show_student ---
    def send_whatsapp_to_parent_show_student():
        # ========================================
        recipient_type.set("student")
        external_container.pack_forget()
        student_container.pack(fill="both", expand=True)
        btn_student.config(bg="#25d366", fg="white")
        btn_external.config(bg="#e0e0e0", fg="#333")
        update_search()

        # ========================================

    # --- send_whatsapp_to_parent_show_external ---
    def send_whatsapp_to_parent_show_external():
        # ========================================
        recipient_type.set("external")
        student_container.pack_forget()
        external_container.pack(fill="both", expand=True)
        btn_external.config(bg="#25d366", fg="white")
        btn_student.config(bg="#e0e0e0", fg="#333")

        # ========================================

    # --- send_whatsapp_to_parent_send_message ---
    def send_whatsapp_to_parent_send_message():
        # ========================================
        sender_phone = new_sender_entry.get().strip()
        if sender_phone:
            if sender_phone == "05xxxxxxxx":
                sender_phone = sender_var.get()
            if not sender_phone:
                messagebox.showwarning("تنبيه", "الرجاء إدخال أو اختيار رقم المرسل!")
                return
            message = msg_text.get("1.0", "end-1c").strip()
            if not message:
                messagebox.showwarning("تنبيه", "الرجاء كتابة رسالة!")
                return
            if sender_phone not in saved_senders:
                with open(sender_file, "a", encoding="utf-8") as f:
                    f.write(sender_phone + "\n")
            if recipient_type.get() == "external":
                parent_phone = external_phone_var.get().strip()
                if not parent_phone or parent_phone == "05xxxxxxxx":
                    messagebox.showwarning("تنبيه", "الرجاء إدخال رقم الجوال!")
                    return
                recipient_name = "الطرف الخارجي"
            else:
                selection = student_listbox.curselection()
                if not selection:
                    messagebox.showwarning("تنبيه", "الرجاء اختيار طالب!")
                    return
                selected_student = student_listbox.get(selection[0])
                name_col, _, _ = self.m._detect_student_cols()
                parent_phone_col = None
                for col in self.m.df_students.columns:
                    col_lower = str(col).lower()
                    if "جوال" in col_lower:
                        if not "ولي" in col_lower:
                            if "طالب" in col_lower:
                                parent_phone_col = col
                                break
                            if not parent_phone_col:
                                messagebox.showwarning("تنبيه", "لم يتم العثور على عمود رقم الجوال!")
                                return
                            student_row = None
                            for idx, row in self.m.df_students.iterrows():
                                if tidy(row.get(name_col, "")) == tidy(selected_student):
                                    student_row = row
                                    break
                            else:
                                if student_row is None:
                                    messagebox.showwarning("تنبيه", "لم يتم العثور على بيانات الطالب!")
                                    return
                                else:
                                    parent_phone = tidy(student_row.get(parent_phone_col, ""))
                                    parent_phone or messagebox.showwarning("تنبيه", f"لا يوجد رقم جوال لولي أمر: {selected_student}")
                                    return
                                recipient_name = f"ولي أمر {selected_student}"

                        parent_phone = parent_phone.replace(" ", "").replace("-", "")
                        if parent_phone.startswith("0"):
                            parent_phone = "966" + parent_phone[1:]

        else:
            if not parent_phone.startswith("966"):
                parent_phone = "966" + parent_phone
            import urllib.parse
            encoded_message = urllib.parse.quote(message)
            whatsapp_url = f"https://web.whatsapp.com/send?phone={parent_phone}&text={encoded_message}"
            webbrowser.open(whatsapp_url)
            messagebox.showinfo("تم", f"تم فتح واتساب ويب!\n\nالمرسل: {sender_phone}\nالمستقبل: {recipient_name}\n\nاضغط Enter للإرسال.")
            main_win.destroy()

        # ========================================

    # --- page_swap_request ---
    def page_swap_request(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg=COLOR_BG)
        hdr = tk.Frame(frm, bg=COLOR_ACCENT, height=60)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="📋 طلب مبادلة حصة", bg=COLOR_ACCENT, fg="white", font=('Segoe UI',
                                                                                  16, 'bold')).pack(expand=True)
        canvas = tk.Canvas(frm, bg=COLOR_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frm, orient="vertical", command=(canvas.yview))
        scrollable_frame = tk.Frame(canvas, bg=COLOR_BG)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=(canvas.bbox("all"))))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=(scrollbar.set))
        canvas.pack(side="left", fill="both", expand=True, padx=20, pady=20)
        scrollbar.pack(side="right", fill="y")
        fields_frame = tk.Frame(scrollable_frame, bg=COLOR_PANEL, padx=30, pady=30)
        fields_frame.pack(fill="both", expand=True)
        sender_var = tk.StringVar()
        receiver_var = tk.StringVar()
        day_var = tk.StringVar()
        period_var = tk.StringVar()
        subject_var = tk.StringVar()
        class_var = tk.StringVar()
        section_var = tk.StringVar()
        start_date_var = tk.StringVar(value=(datetime.now().strftime("%Y-%m-%d")))
        end_date_var = tk.StringVar(value=((datetime.now() + timedelta(days=7)).strftime("%Y-%m-%d")))
        notes_var = tk.StringVar()
        combo_widgets = {}

        def auto_fill_from_schedule(*args):
            """Auto-fill subject, class, and section when teacher, day, and period are selected"""
            teacher = sender_var.get()
            day = day_var.get()
            period = period_var.get()
            if not (teacher and day and period):
                return
            try:
                cell_value = self.m.get_cell(teacher, day, int(period))
                if cell_value and cell_value.strip() and cell_value != "—":
                    parsed = parse_teacher_cell(cell_value)
                    if parsed:
                        if parsed.get("subject"):
                            subject_var.set(parsed["subject"])
                        if parsed.get("class_section"):
                            class_section_text = parsed["class_section"]
                            class_found = None
                            for class_option in classes:
                                if class_option in class_section_text:
                                    class_found = class_option
                                    break
                            else:
                                if class_found:
                                    class_var.set(class_found)
                                import re
                                section_match = re.search("(\\d+)$", class_section_text.strip())
                                if section_match:
                                    section_var.set(section_match.group(1))

                else:
                    subject_var.set("")
                    class_var.set("")
                    section_var.set("")
            except Exception as e:
                try:
                    print(f"Error auto-filling: {e}")
                finally:
                    pass


        def add_field(label_text, var, values=None, row=0):
            """Helper to add a form field"""
            tk.Label(fields_frame, text=label_text, bg=COLOR_PANEL, font=('Segoe UI', 11, 'bold'),
              anchor="e").grid(row=row, column=0, sticky="e", padx=10, pady=10)
            if values:
                combo = ttk.Combobox(fields_frame, textvariable=var, values=values, font=('Segoe UI',
                                                                                          11),
                  state="readonly",
                  width=30)
                combo.grid(row=row, column=1, sticky="w", padx=10, pady=10)
                if values:
                    combo.current(0)
                combo_widgets[label_text] = combo
                return combo
            entry = tk.Entry(fields_frame, textvariable=var, font=('Segoe UI', 11), width=32)
            entry.grid(row=row, column=1, sticky="w", padx=10, pady=10)
            if var.get():
                entry.delete(0, tk.END)
                entry.insert(0, var.get())
            return entry


        teachers = self.m.get_all_teachers()
        subjects = []
        try:
            subjects_df = self.m.df_subjects
            if not subjects_df.empty:
                if "المادة" in subjects_df.columns:
                    subjects = [str(s).strip() for s in subjects_df["المادة"].tolist() if str(s).strip() if str(s) != "nan"]
        except:
            pass
        else:
            if not subjects:
                subjects = [
                 'رياضيات', 'علوم', 'لغة عربية', 'لغة إنجليزية', 'اجتماعيات', 
                 'قرآن كريم',  'حديث',  'فقه',  'توحيد',  'حاسب آلي', 
                 'تربية فنية',  'تربية بدنية',  'تربية أسرية']
            classes = ["أول متوسط", "ثاني متوسط", "ثالث متوسط"]
            sections = ['1', '2', '3', '4', '5']
            row = 0
            add_field("المعلم المرسل (أنت):", sender_var, teachers, row)
            row += 1
            add_field("المعلم المستقبل:", receiver_var, teachers, row)
            row += 1
            add_field("اليوم:", day_var, DAYS, row)
            row += 1
            add_field("رقم الحصة:", period_var, [str(i) for i in range(1, 9)], row)
            row += 1
            add_field("المادة:", subject_var, subjects, row)
            row += 1
            add_field("الصف:", class_var, classes, row)
            row += 1
            add_field("الشعبة:", section_var, sections, row)
            row += 1
            add_field("تاريخ بداية المبادلة:", start_date_var, None, row)
            row += 1
            add_field("تاريخ نهاية المبادلة:", end_date_var, None, row)
            row += 1
            sender_var.trace_add("write", auto_fill_from_schedule)
            day_var.trace_add("write", auto_fill_from_schedule)
            period_var.trace_add("write", auto_fill_from_schedule)
            tk.Label(fields_frame, text="الملاحظات:", bg=COLOR_PANEL, font=('Segoe UI', 11,
                                                                            'bold'),
              anchor="e").grid(row=row, column=0, sticky="ne", padx=10, pady=10)
            notes_text = tk.Text(fields_frame, font=('Segoe UI', 10), width=32, height=4)
            notes_text.grid(row=row, column=1, sticky="w", padx=10, pady=10)
            btn_frame = tk.Frame(fields_frame, bg=COLOR_PANEL)
            btn_frame.grid(row=(row + 1), column=0, columnspan=2, pady=20)

            def submit_request():
                """Submit the swap request"""
                if sender_var.get():
                    receiver_var.get() or messagebox.showerror("خطأ", "يرجى اختيار المعلم المرسل والمستقبل")
                    return
                elif sender_var.get() == receiver_var.get():
                    messagebox.showerror("خطأ", "لا يمكن المبادلة مع نفس المعلم")
                    return
                    swap_data = {'المعلم المرسل':sender_var.get(), 
                     'المعلم المستقبل':receiver_var.get(), 
                     'اليوم':day_var.get(), 
                     'رقم الحصة':period_var.get(), 
                     'المادة':subject_var.get(), 
                     'الصف':class_var.get(), 
                     'الشعبة':section_var.get(), 
                     'تاريخ بداية المبادلة':start_date_var.get(), 
                     'تاريخ نهاية المبادلة':end_date_var.get(), 
                     'الملاحظات':(notes_text.get("1.0", tk.END).strip)()}
                    success, swap_id = self.m.add_swap_request(swap_data)
                    if success:
                        messagebox.showinfo("نجح", f"تم إرسال طلب المبادلة بنجاح!\n\nرقم المبادلة: {swap_id}\nالحالة: قيد المراجعة\n\nسيتم إشعار {receiver_var.get()} للموافقة على الطلب.")
                        self.show_page("teachers")
                else:
                    messagebox.showerror("خطأ", "فشل إرسال الطلب")


            tk.Button(btn_frame, text="📤 إرسال الطلب", command=submit_request, bg=COLOR_XLSX,
              fg="white",
              font=('Segoe UI', 12, 'bold'),
              padx=30,
              pady=10,
              cursor="hand2").pack(side="right", padx=5)
            tk.Button(btn_frame, text="إلغاء", command=(lambda: self.show_page("teachers")), bg="#999",
              fg="white",
              font=('Segoe UI', 11),
              padx=30,
              pady=10,
              cursor="hand2").pack(side="right", padx=5)
            return frm

        # ========================================

    # --- page_swap_request_auto_fill_from_schedule ---
    def page_swap_request_auto_fill_from_schedule(*args):
        # ========================================
        teacher = sender_var.get()
        day = day_var.get()
        period = period_var.get()
        if not (teacher and day and period):
            return
        try:
            cell_value = self.m.get_cell(teacher, day, int(period))
            if cell_value and cell_value.strip() and cell_value != "—":
                parsed = parse_teacher_cell(cell_value)
                if parsed:
                    if parsed.get("subject"):
                        subject_var.set(parsed["subject"])
                    if parsed.get("class_section"):
                        class_section_text = parsed["class_section"]
                        class_found = None
                        for class_option in classes:
                            if class_option in class_section_text:
                                class_found = class_option
                                break
                        else:
                            if class_found:
                                class_var.set(class_found)
                            import re
                            section_match = re.search("(\\d+)$", class_section_text.strip())
                            if section_match:
                                section_var.set(section_match.group(1))

            else:
                subject_var.set("")
                class_var.set("")
                section_var.set("")
        except Exception as e:
            try:
                print(f"Error auto-filling: {e}")
            finally:
                pass

        # ========================================

    # --- page_swap_request_add_field ---
    def page_swap_request_add_field(label_text=None, var=None, values=None, row=None):
        # ========================================
        tk.Label(fields_frame, text=label_text, bg=COLOR_PANEL, font=('Segoe UI', 11, 'bold'),
          anchor="e").grid(row=row, column=0, sticky="e", padx=10, pady=10)
        if values:
            combo = ttk.Combobox(fields_frame, textvariable=var, values=values, font=('Segoe UI',
                                                                                      11),
              state="readonly",
              width=30)
            combo.grid(row=row, column=1, sticky="w", padx=10, pady=10)
            if values:
                combo.current(0)
            combo_widgets[label_text] = combo
            return combo
        entry = tk.Entry(fields_frame, textvariable=var, font=('Segoe UI', 11), width=32)
        entry.grid(row=row, column=1, sticky="w", padx=10, pady=10)
        if var.get():
            entry.delete(0, tk.END)
            entry.insert(0, var.get())
        return entry

        # ========================================

    # --- page_swap_request_submit_request ---
    def page_swap_request_submit_request():
        # ========================================
        if sender_var.get():
            receiver_var.get() or messagebox.showerror("خطأ", "يرجى اختيار المعلم المرسل والمستقبل")
            return
        elif sender_var.get() == receiver_var.get():
            messagebox.showerror("خطأ", "لا يمكن المبادلة مع نفس المعلم")
            return
            swap_data = {'المعلم المرسل':sender_var.get(), 
             'المعلم المستقبل':receiver_var.get(), 
             'اليوم':day_var.get(), 
             'رقم الحصة':period_var.get(), 
             'المادة':subject_var.get(), 
             'الصف':class_var.get(), 
             'الشعبة':section_var.get(), 
             'تاريخ بداية المبادلة':start_date_var.get(), 
             'تاريخ نهاية المبادلة':end_date_var.get(), 
             'الملاحظات':(notes_text.get("1.0", tk.END).strip)()}
            success, swap_id = self.m.add_swap_request(swap_data)
            if success:
                messagebox.showinfo("نجح", f"تم إرسال طلب المبادلة بنجاح!\n\nرقم المبادلة: {swap_id}\nالحالة: قيد المراجعة\n\nسيتم إشعار {receiver_var.get()} للموافقة على الطلب.")
                self.show_page("teachers")
        else:
            messagebox.showerror("خطأ", "فشل إرسال الطلب")

        # ========================================

    # --- page_swap_inbox ---
    def page_swap_inbox(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg=COLOR_BG)
        hdr = tk.Frame(frm, bg=COLOR_ACCENT, height=60)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="📬 المبادلات الواردة", bg=COLOR_ACCENT, fg="white", font=('Segoe UI',
                                                                                     16,
                                                                                     'bold')).pack(expand=True)
        sel_frame = tk.Frame(frm, bg=COLOR_PANEL, padx=20, pady=15)
        sel_frame.pack(fill="x", padx=20, pady=10)
        tk.Label(sel_frame, text="اختر المعلم:", bg=COLOR_PANEL, font=('Segoe UI', 11, 'bold')).pack(side="right", padx=10)
        teacher_var = tk.StringVar()
        teachers = self.m.get_all_teachers()
        teacher_combo = ttk.Combobox(sel_frame, textvariable=teacher_var, values=teachers, font=('Segoe UI',
                                                                                                 11),
          state="readonly",
          width=25)
        teacher_combo.pack(side="right", padx=10)
        if teachers:
            teacher_combo.current(0)
        content_frame = tk.Frame(frm, bg=COLOR_BG)
        content_frame.pack(fill="both", expand=True, padx=20, pady=10)

        def load_pending_swaps():
            """Load and display pending swaps for selected teacher"""
            for widget in content_frame.winfo_children():
                widget.destroy()
            else:
                teacher = teacher_var.get()
                if not teacher:
                    tk.Label(content_frame, text="يرجى اختيار معلم", bg=COLOR_BG, font=('Segoe UI',
                                                                                        12)).pack(pady=50)
                    return
                pending = self.m.get_pending_swaps_for_teacher(teacher)
                if not pending:
                    tk.Label(content_frame, text="✅ لا توجد مبادلات معلقة", bg=COLOR_BG, font=('Segoe UI',
                                                                                               14),
                      fg="#666").pack(pady=50)
                    return
                for swap in pending:
                    swap_card = tk.Frame(content_frame, bg=COLOR_PANEL, relief="raised", bd=2)
                    swap_card.pack(fill="x", pady=10, padx=10)
                    details_frame = tk.Frame(swap_card, bg=COLOR_PANEL, padx=20, pady=15)
                    details_frame.pack(fill="both", expand=True)
                    tk.Label(details_frame, text=f'📋 مبادلة رقم: {swap.get("رقم المبادلة", "N/A")}', bg=COLOR_PANEL,
                      font=('Segoe UI', 12, 'bold'),
                      fg=COLOR_XLSX).pack(anchor="e")
                    info_text = f'\n                المعلم المرسل: {swap.get("المعلم المرسل", "")}\n                اليوم: {swap.get("اليوم", "")} | الحصة: {swap.get("رقم الحصة", "")}\n                المادة: {swap.get("المادة", "")} | الصف: {swap.get("الصف", "")} / الشعبة: {swap.get("الشعبة", "")}\n                الفترة: من {swap.get("تاريخ بداية المبادلة", "")} إلى {swap.get("تاريخ نهاية المبادلة", "")}\n                الملاحظات: {swap.get("الملاحظات", "لا توجد")}\n                '
                    tk.Label(details_frame, text=info_text, bg=COLOR_PANEL, font=('Segoe UI',
                                                                                  10),
                      justify="right").pack(anchor="e", pady=10)
                    btn_frame = tk.Frame(details_frame, bg=COLOR_PANEL)
                    btn_frame.pack(anchor="e")

                    def approve_swap(swap_id=swap.get("رقم المبادلة")):
                        success, msg = self.m.update_swap_status(swap_id, "معتمد من المستقبل")
                        if success:
                            messagebox.showinfo("نجح", "تم قبول المبادلة! سيتم إرسالها للمدير للاعتماد النهائي.")
                            load_pending_swaps()
                        else:
                            messagebox.showerror("خطأ", msg)

                    def reject_swap(swap_id=swap.get("رقم المبادلة")):
                        success, msg = self.m.update_swap_status(swap_id, "مرفوض")
                        if success:
                            messagebox.showinfo("تم", "تم رفض المبادلة")
                            load_pending_swaps()
                        else:
                            messagebox.showerror("خطأ", msg)

                    tk.Button(btn_frame, text="✅ موافقة", command=approve_swap, bg="#4caf50",
                      fg="white",
                      font=('Segoe UI', 10, 'bold'),
                      padx=20,
                      pady=8,
                      cursor="hand2").pack(side="right", padx=5)
                    tk.Button(btn_frame, text="❌ رفض", command=reject_swap, bg=COLOR_DANGER,
                      fg="white",
                      font=('Segoe UI', 10, 'bold'),
                      padx=20,
                      pady=8,
                      cursor="hand2").pack(side="right", padx=5)


        tk.Button(sel_frame, text="🔄 تحديث", command=load_pending_swaps, bg=COLOR_BTN,
          font=('Segoe UI', 10),
          padx=15,
          pady=5).pack(side="right", padx=10)
        tk.Button(frm, text="رجوع", command=(lambda: self.show_page("teachers")), bg="#999",
          fg="white",
          font=('Segoe UI', 11),
          padx=30,
          pady=10).pack(pady=10)
        return frm

        # ========================================

    # --- page_swap_inbox_load_pending_swaps ---
    def page_swap_inbox_load_pending_swaps():
        # ========================================
        for widget in content_frame.winfo_children():
            widget.destroy()
        else:
            teacher = teacher_var.get()
            if not teacher:
                tk.Label(content_frame, text="يرجى اختيار معلم", bg=COLOR_BG, font=('Segoe UI',
                                                                                    12)).pack(pady=50)
                return
            pending = self.m.get_pending_swaps_for_teacher(teacher)
            if not pending:
                tk.Label(content_frame, text="✅ لا توجد مبادلات معلقة", bg=COLOR_BG, font=('Segoe UI',
                                                                                           14),
                  fg="#666").pack(pady=50)
                return
            for swap in pending:
                swap_card = tk.Frame(content_frame, bg=COLOR_PANEL, relief="raised", bd=2)
                swap_card.pack(fill="x", pady=10, padx=10)
                details_frame = tk.Frame(swap_card, bg=COLOR_PANEL, padx=20, pady=15)
                details_frame.pack(fill="both", expand=True)
                tk.Label(details_frame, text=f'📋 مبادلة رقم: {swap.get("رقم المبادلة", "N/A")}', bg=COLOR_PANEL,
                  font=('Segoe UI', 12, 'bold'),
                  fg=COLOR_XLSX).pack(anchor="e")
                info_text = f'\n                المعلم المرسل: {swap.get("المعلم المرسل", "")}\n                اليوم: {swap.get("اليوم", "")} | الحصة: {swap.get("رقم الحصة", "")}\n                المادة: {swap.get("المادة", "")} | الصف: {swap.get("الصف", "")} / الشعبة: {swap.get("الشعبة", "")}\n                الفترة: من {swap.get("تاريخ بداية المبادلة", "")} إلى {swap.get("تاريخ نهاية المبادلة", "")}\n                الملاحظات: {swap.get("الملاحظات", "لا توجد")}\n                '
                tk.Label(details_frame, text=info_text, bg=COLOR_PANEL, font=('Segoe UI', 10),
                  justify="right").pack(anchor="e", pady=10)
                btn_frame = tk.Frame(details_frame, bg=COLOR_PANEL)
                btn_frame.pack(anchor="e")

                def approve_swap(swap_id=swap.get("رقم المبادلة")):
                    success, msg = self.m.update_swap_status(swap_id, "معتمد من المستقبل")
                    if success:
                        messagebox.showinfo("نجح", "تم قبول المبادلة! سيتم إرسالها للمدير للاعتماد النهائي.")
                        load_pending_swaps()
                    else:
                        messagebox.showerror("خطأ", msg)


                def reject_swap(swap_id=swap.get("رقم المبادلة")):
                    success, msg = self.m.update_swap_status(swap_id, "مرفوض")
                    if success:
                        messagebox.showinfo("تم", "تم رفض المبادلة")
                        load_pending_swaps()
                    else:
                        messagebox.showerror("خطأ", msg)


                tk.Button(btn_frame, text="✅ موافقة", command=approve_swap, bg="#4caf50",
                  fg="white",
                  font=('Segoe UI', 10, 'bold'),
                  padx=20,
                  pady=8,
                  cursor="hand2").pack(side="right", padx=5)
                tk.Button(btn_frame, text="❌ رفض", command=reject_swap, bg=COLOR_DANGER,
                  fg="white",
                  font=('Segoe UI', 10, 'bold'),
                  padx=20,
                  pady=8,
                  cursor="hand2").pack(side="right", padx=5)

        # ========================================

    # --- page_swap_inbox_load_pending_swaps_approve_swap ---
    def page_swap_inbox_load_pending_swaps_approve_swap(swap_id=None):
        # ========================================
        success, msg = self.m.update_swap_status(swap_id, "معتمد من المستقبل")
        if success:
            messagebox.showinfo("نجح", "تم قبول المبادلة! سيتم إرسالها للمدير للاعتماد النهائي.")
            load_pending_swaps()
        else:
            messagebox.showerror("خطأ", msg)

        # ========================================

    # --- page_swap_inbox_load_pending_swaps_reject_swap ---
    def page_swap_inbox_load_pending_swaps_reject_swap(swap_id=None):
        # ========================================
        success, msg = self.m.update_swap_status(swap_id, "مرفوض")
        if success:
            messagebox.showinfo("تم", "تم رفض المبادلة")
            load_pending_swaps()
        else:
            messagebox.showerror("خطأ", msg)

        # ========================================

    # --- page_swap_admin_approval ---
    def page_swap_admin_approval(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg=COLOR_BG)
        hdr = tk.Frame(frm, bg=COLOR_ACCENT, height=60)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="👨\u200d💼 اعتماد المبادلات والسجل", bg=COLOR_ACCENT, fg="white", font=('Segoe UI',
                                                                                                  16,
                                                                                                  'bold')).pack(expand=True)
        notebook = ttk.Notebook(frm)
        notebook.pack(fill="both", expand=True, padx=20, pady=10)
        tab_pending = tk.Frame(notebook, bg=COLOR_BG)
        notebook.add(tab_pending, text=" ⭐ الطلبات المعلقة (للاعتماد) ")
        pending_container = tk.Frame(tab_pending, bg=COLOR_BG)
        pending_container.pack(fill="both", expand=True, padx=10, pady=10)
        canvas = tk.Canvas(pending_container, bg=COLOR_BG, highlightthickness=0)
        scrollbar = ttk.Scrollbar(pending_container, orient="vertical", command=(canvas.yview))
        scroll_frame = tk.Frame(canvas, bg=COLOR_BG)
        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=(canvas.bbox("all"))))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw", width=1000)
        canvas.configure(yscrollcommand=(scrollbar.set))

        def on_canvas_configure(event):
            canvas.itemconfig((canvas.find_withtag("all")[0]), width=(event.width))


        canvas.bind("<Configure>", on_canvas_configure)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        def load_pending_admin_swaps():
            """Load swaps pending admin approval"""
            for widget in scroll_frame.winfo_children():
                widget.destroy()
            else:
                pending = self.m.get_pending_swaps_for_admin()
                if not pending:
                    tk.Label(scroll_frame, text="✅ لا توجد مبادلات تحتاج اعتماد حالياً", bg=COLOR_BG, font=('Segoe UI',
                                                                                                            14),
                      fg="#666").pack(pady=50)
                    return
                for swap in pending:
                    swap_card = tk.Frame(scroll_frame, bg=COLOR_PANEL, relief="raised", bd=2)
                    swap_card.pack(fill="x", pady=10, padx=10)
                    details_frame = tk.Frame(swap_card, bg=COLOR_PANEL, padx=20, pady=15)
                    details_frame.pack(fill="both", expand=True)
                    tk.Label(details_frame, text=f'📋 مبادلة رقم: {swap.get("رقم المبادلة", "N/A")}', bg=COLOR_PANEL,
                      font=('Segoe UI', 12, 'bold'),
                      fg=COLOR_XLSX).pack(anchor="e")
                    info_text = f'المعلم المرسل: {swap.get("المعلم المرسل", "")}   ⬅️   المعلم المستقبل: {swap.get("المعلم المستقبل", "")}\nاليوم: {swap.get("اليوم", "")}  |  الحصة: {swap.get("رقم الحصة", "")}  |  المادة: {swap.get("المادة", "")}\nالصف: {swap.get("الصف", "")}  |  الشعبة: {swap.get("الشعبة", "")}\nالفترة: {swap.get("تاريخ بداية المبادلة", "")}  إلى  {swap.get("تاريخ نهاية المبادلة", "")}\nالملاحظات: {swap.get("الملاحظات", "")}'
                    tk.Label(details_frame, text=info_text, bg=COLOR_PANEL, font=('Segoe UI',
                                                                                  11),
                      justify="right").pack(anchor="e", pady=10)
                    btn_frame = tk.Frame(details_frame, bg=COLOR_PANEL)
                    btn_frame.pack(anchor="e", pady=5)

                    def approve_final_closure(s_id=swap.get("رقم المبادلة")):
                        if messagebox.askyesno("تأكيد", "هل أنت متأكد من اعتماد هذه المبادلة نهائياً؟"):
                            success, msg = self.m.update_swap_status(s_id, "معتمد نهائياً")
                            if success:
                                messagebox.showinfo("نجح", "تم اعتماد المبادلة.\nسيتم تحديث الجداول تلقائياً.")
                                load_pending_admin_swaps()
                            else:
                                messagebox.showerror("خطأ", msg)

                    def reject_final_closure(s_id=swap.get("رقم المبادلة")):
                        if messagebox.askyesno("تأكيد", "هل تريد رفض هذه المبادلة؟"):
                            success, msg = self.m.update_swap_status(s_id, "مرفوض من المدير")
                            if success:
                                messagebox.showinfo("تم", "تم رفض المبادلة.")
                                load_pending_admin_swaps()
                            else:
                                messagebox.showerror("خطأ", msg)

                    tk.Button(btn_frame, text="✅ اعتماد نهائي", command=approve_final_closure, bg="#4caf50",
                      fg="white",
                      font=('Segoe UI', 10, 'bold'),
                      padx=15).pack(side="right", padx=5)
                    tk.Button(btn_frame, text="❌ رفض", command=reject_final_closure, bg=COLOR_DANGER,
                      fg="white",
                      font=('Segoe UI', 10, 'bold'),
                      padx=15).pack(side="right", padx=5)


        btn_refresh_p = tk.Button(tab_pending, text="🔄 تحديث القائمة", command=load_pending_admin_swaps, bg=COLOR_BTN,
          fg="white",
          font=('Segoe UI', 10))
        btn_refresh_p.pack(pady=5)
        tab_history = tk.Frame(notebook, bg=COLOR_BG)
        notebook.add(tab_history, text=" 📂 سجل المبادلات (الأرشيف) ")
        filter_frame = tk.LabelFrame(tab_history, text="🔍 البحث والفلترة", bg=COLOR_BG, font=('Segoe UI',
                                                                                              11,
                                                                                              'bold'), padx=15, pady=15)
        filter_frame.pack(fill="x", padx=10, pady=10)
        f_row1 = tk.Frame(filter_frame, bg=COLOR_BG)
        f_row1.pack(fill="x", pady=5)
        tk.Label(f_row1, text="اسم المعلم:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
        entry_search_name = tk.Entry(f_row1, font=('Segoe UI', 10), width=25, justify="right")
        entry_search_name.pack(side="right", padx=5)
        tk.Label(f_row1, text=" |  من تاريخ:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
        entry_date_from = tk.Entry(f_row1, font=('Segoe UI', 10), width=12, justify="center")
        entry_date_from.pack(side="right", padx=5)
        last_month = datetime.now() - timedelta(days=30)
        entry_date_from.insert(0, last_month.strftime("%Y-%m-%d"))
        tk.Label(f_row1, text="إلى:", bg=COLOR_BG, font=('Segoe UI', 10)).pack(side="right", padx=5)
        entry_date_to = tk.Entry(f_row1, font=('Segoe UI', 10), width=12, justify="center")
        entry_date_to.pack(side="right", padx=5)
        entry_date_to.insert(0, datetime.now().strftime("%Y-%m-%d"))

        def perform_search():
            name_query = entry_search_name.get().strip()
            d_from = entry_date_from.get().strip()
            d_to = entry_date_to.get().strip()
            df = self.m.load_swaps()
            if df.empty:
                update_history_tree([])
                lbl_stats.config(text="لا توجد بيانات")
                return
            if name_query:
                mask_name = df["المعلم المرسل"].astype(str).str.contains(name_query, na=False) | df["المعلم المستقبل"].astype(str).str.contains(name_query, na=False)
                df = df[mask_name]
            if d_from:
                if d_to:
                    try:
                        mask_date = (df["تاريخ الطلب"] >= d_from) & (df["تاريخ الطلب"] <= d_to)
                        df = df[mask_date]
                    except Exception as e:
                        try:
                            print(f"Date filter error: {e}")
                        finally:
                            pass

            records = df.to_dict("records")
            records.sort(key=(lambda x: x.get("رقم المبادلة", 0)), reverse=True)
            update_history_tree(records)
            total = len(records)
            approved = len([r for r in records if "معتمد" in str(r.get("حالة الاعتماد", ""))])
            rejected = len([r for r in records if "مرفوض" in str(r.get("حالة الاعتماد", ""))])
            stats_text = f"إجمالي النتائج: {total}  |  ✅ معتمد: {approved}  |  ❌ مرفوض: {rejected}"
            lbl_stats.config(text=stats_text)


        btn_search = tk.Button(f_row1, text="بحث 🔎", command=perform_search, bg="#1976d2",
          fg="white",
          font=('Segoe UI', 10, 'bold'),
          padx=20)
        btn_search.pack(side="left", padx=10)
        lbl_stats = tk.Label(filter_frame, text="...", bg=COLOR_BG, fg="#333", font=('Segoe UI',
                                                                                     11,
                                                                                     'bold'))
        lbl_stats.pack(side="left", padx=20)
        tree_frame = tk.Frame(tab_history, bg=COLOR_BG)
        tree_frame.pack(fill="both", expand=True, padx=10, pady=5)
        h_cols = ('id', 'date', 'from', 'to', 'details', 'status')
        tree = ttk.Treeview(tree_frame, columns=h_cols, show="headings", selectmode="browse")
        tree.heading("id", text="#")
        tree.heading("date", text="تاريخ الطلب")
        tree.heading("from", text="من المعلم")
        tree.heading("to", text="إلى المعلم")
        tree.heading("details", text="التفاصيل (حصة:مادة)")
        tree.heading("status", text="الحالة")
        tree.column("id", width=50, anchor="center")
        tree.column("date", width=90, anchor="center")
        tree.column("from", width=120, anchor="center")
        tree.column("to", width=120, anchor="center")
        tree.column("details", width=250, anchor="e")
        tree.column("status", width=120, anchor="center")
        h_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=(tree.yview))
        tree.configure(yscrollcommand=(h_scroll.set))
        h_scroll.pack(side="left", fill="y")
        tree.pack(side="left", fill="both", expand=True)

        def update_history_tree(records):
            for item in tree.get_children():
                tree.delete(item)
            else:
                for r in records:
                    details = f'{r.get("اليوم", "")} - ح{r.get("رقم الحصة", "")} - {r.get("المادة", "")}'
                    tree.insert("", "end", values=(
                     r.get("رقم المبادلة"),
                     r.get("تاريخ الطلب"),
                     r.get("المعلم المرسل"),
                     r.get("المعلم المستقبل"),
                     details,
                     r.get("حالة الاعتماد")))


        load_pending_admin_swaps()
        perform_search()
        tk.Button(frm, text="🔙 رجوع للقائمة الرئيسية", command=(lambda: self.show("home")), bg="#455a64",
          fg="white",
          font=('Segoe UI', 11),
          padx=20,
          pady=5).pack(pady=10)
        return frm

        # ========================================

    # --- page_swap_admin_approval_on_canvas_configure ---
    def page_swap_admin_approval_on_canvas_configure(event=None):
        # ========================================
        canvas.itemconfig((canvas.find_withtag("all")[0]), width=(event.width))

        # ========================================

    # --- page_swap_admin_approval_load_pending_admin_swaps ---
    def page_swap_admin_approval_load_pending_admin_swaps():
        # ========================================
        for widget in scroll_frame.winfo_children():
            widget.destroy()
        else:
            pending = self.m.get_pending_swaps_for_admin()
            if not pending:
                tk.Label(scroll_frame, text="✅ لا توجد مبادلات تحتاج اعتماد حالياً", bg=COLOR_BG, font=('Segoe UI',
                                                                                                        14),
                  fg="#666").pack(pady=50)
                return
            for swap in pending:
                swap_card = tk.Frame(scroll_frame, bg=COLOR_PANEL, relief="raised", bd=2)
                swap_card.pack(fill="x", pady=10, padx=10)
                details_frame = tk.Frame(swap_card, bg=COLOR_PANEL, padx=20, pady=15)
                details_frame.pack(fill="both", expand=True)
                tk.Label(details_frame, text=f'📋 مبادلة رقم: {swap.get("رقم المبادلة", "N/A")}', bg=COLOR_PANEL,
                  font=('Segoe UI', 12, 'bold'),
                  fg=COLOR_XLSX).pack(anchor="e")
                info_text = f'المعلم المرسل: {swap.get("المعلم المرسل", "")}   ⬅️   المعلم المستقبل: {swap.get("المعلم المستقبل", "")}\nاليوم: {swap.get("اليوم", "")}  |  الحصة: {swap.get("رقم الحصة", "")}  |  المادة: {swap.get("المادة", "")}\nالصف: {swap.get("الصف", "")}  |  الشعبة: {swap.get("الشعبة", "")}\nالفترة: {swap.get("تاريخ بداية المبادلة", "")}  إلى  {swap.get("تاريخ نهاية المبادلة", "")}\nالملاحظات: {swap.get("الملاحظات", "")}'
                tk.Label(details_frame, text=info_text, bg=COLOR_PANEL, font=('Segoe UI', 11),
                  justify="right").pack(anchor="e", pady=10)
                btn_frame = tk.Frame(details_frame, bg=COLOR_PANEL)
                btn_frame.pack(anchor="e", pady=5)

                def approve_final_closure(s_id=swap.get("رقم المبادلة")):
                    if messagebox.askyesno("تأكيد", "هل أنت متأكد من اعتماد هذه المبادلة نهائياً؟"):
                        success, msg = self.m.update_swap_status(s_id, "معتمد نهائياً")
                        if success:
                            messagebox.showinfo("نجح", "تم اعتماد المبادلة.\nسيتم تحديث الجداول تلقائياً.")
                            load_pending_admin_swaps()
                        else:
                            messagebox.showerror("خطأ", msg)


                def reject_final_closure(s_id=swap.get("رقم المبادلة")):
                    if messagebox.askyesno("تأكيد", "هل تريد رفض هذه المبادلة؟"):
                        success, msg = self.m.update_swap_status(s_id, "مرفوض من المدير")
                        if success:
                            messagebox.showinfo("تم", "تم رفض المبادلة.")
                            load_pending_admin_swaps()
                        else:
                            messagebox.showerror("خطأ", msg)


                tk.Button(btn_frame, text="✅ اعتماد نهائي", command=approve_final_closure, bg="#4caf50",
                  fg="white",
                  font=('Segoe UI', 10, 'bold'),
                  padx=15).pack(side="right", padx=5)
                tk.Button(btn_frame, text="❌ رفض", command=reject_final_closure, bg=COLOR_DANGER,
                  fg="white",
                  font=('Segoe UI', 10, 'bold'),
                  padx=15).pack(side="right", padx=5)

        # ========================================

    # --- page_swap_admin_approval_load_pending_admin_swaps_approve_final_closure ---
    def page_swap_admin_approval_load_pending_admin_swaps_approve_final_closure(s_id=None):
        # ========================================
        if messagebox.askyesno("تأكيد", "هل أنت متأكد من اعتماد هذه المبادلة نهائياً؟"):
            success, msg = self.m.update_swap_status(s_id, "معتمد نهائياً")
            if success:
                messagebox.showinfo("نجح", "تم اعتماد المبادلة.\nسيتم تحديث الجداول تلقائياً.")
                load_pending_admin_swaps()
            else:
                messagebox.showerror("خطأ", msg)

        # ========================================

    # --- page_swap_admin_approval_load_pending_admin_swaps_reject_final_closure ---
    def page_swap_admin_approval_load_pending_admin_swaps_reject_final_closure(s_id=None):
        # ========================================
        if messagebox.askyesno("تأكيد", "هل تريد رفض هذه المبادلة؟"):
            success, msg = self.m.update_swap_status(s_id, "مرفوض من المدير")
            if success:
                messagebox.showinfo("تم", "تم رفض المبادلة.")
                load_pending_admin_swaps()
            else:
                messagebox.showerror("خطأ", msg)

        # ========================================

    # --- page_swap_admin_approval_perform_search ---
    def page_swap_admin_approval_perform_search():
        # ========================================
        name_query = entry_search_name.get().strip()
        d_from = entry_date_from.get().strip()
        d_to = entry_date_to.get().strip()
        df = self.m.load_swaps()
        if df.empty:
            update_history_tree([])
            lbl_stats.config(text="لا توجد بيانات")
            return
        if name_query:
            mask_name = df["المعلم المرسل"].astype(str).str.contains(name_query, na=False) | df["المعلم المستقبل"].astype(str).str.contains(name_query, na=False)
            df = df[mask_name]
        if d_from:
            if d_to:
                try:
                    mask_date = (df["تاريخ الطلب"] >= d_from) & (df["تاريخ الطلب"] <= d_to)
                    df = df[mask_date]
                except Exception as e:
                    try:
                        print(f"Date filter error: {e}")
                    finally:
                        pass

        records = df.to_dict("records")
        records.sort(key=(lambda x: x.get("رقم المبادلة", 0)), reverse=True)
        update_history_tree(records)
        total = len(records)
        approved = len([r for r in records if "معتمد" in str(r.get("حالة الاعتماد", ""))])
        rejected = len([r for r in records if "مرفوض" in str(r.get("حالة الاعتماد", ""))])
        stats_text = f"إجمالي النتائج: {total}  |  ✅ معتمد: {approved}  |  ❌ مرفوض: {rejected}"
        lbl_stats.config(text=stats_text)

        # ========================================

    # --- page_swap_admin_approval_update_history_tree ---
    def page_swap_admin_approval_update_history_tree(records=None):
        # ========================================
        for item in tree.get_children():
            tree.delete(item)
        else:
            for r in records:
                details = f'{r.get("اليوم", "")} - ح{r.get("رقم الحصة", "")} - {r.get("المادة", "")}'
                tree.insert("", "end", values=(
                 r.get("رقم المبادلة"),
                 r.get("تاريخ الطلب"),
                 r.get("المعلم المرسل"),
                 r.get("المعلم المستقبل"),
                 details,
                 r.get("حالة الاعتماد")))

        # ========================================

    # --- page_teacher_dashboard ---
    def page_teacher_dashboard(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg="#f1f8e9")
        header = tk.Frame(frm, bg="#2e7d32", pady=20)
        header.pack(fill="x")
        tk.Label(header, text="بوابة المعلم - المتابعة اليومية", font=('Segoe UI', 20, 'bold'), bg="#2e7d32", fg="white").pack()
        body = tk.Frame(frm, bg="#f1f8e9")
        body.pack(expand=True)
        btn_style_lg = {
         'font': ('Segoe UI', 16, 'bold'), 'width': 20, 'height': 2, 'cursor': '"hand2"'}
        tk.Button(body, text="📝 متابعة الحصص", command=lambda: self.show("teacher_lesson_setup"), bg="#43a047", 
         fg="white", **btn_style_lg).pack(pady=20)
        tk.Button(frm, text="🔙 رجوع", command=(lambda: self.show("teacher_personal_view")), bg="#546e7a",
          fg="white",
          font=('Segoe UI', 12)).pack(pady=20)
        return frm

        # ========================================

    # --- page_teacher_lesson_setup ---
    def page_teacher_lesson_setup(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg="#f1f8e9")
        center = tk.Frame(frm, bg="white", padx=40, pady=40, relief="solid", bd=1)
        center.place(relx=0.5, rely=0.5, anchor="center")
        tk.Label(center, text="إعداد الحصة الجديدة", font=('Segoe UI', 18, 'bold'), bg="white", fg="#2e7d32").pack(pady=(0,
                                                                                                                         20))
        info_frame = tk.Frame(center, bg="#e8f5e9", padx=10, pady=10)
        info_frame.pack(fill="x", pady=10)
        today_str = date.today().strftime("%Y-%m-%d")
        tk.Label(info_frame, text=f"تاريخ اليوم: {today_str}", font=('Segoe UI', 11), bg="#e8f5e9").pack(anchor="e")
        self.lbl_setup_teacher = tk.Label(info_frame, text="المعلم: ...", font=('Segoe UI',
                                                                                11, 'bold'), bg="#e8f5e9")
        self.lbl_setup_teacher.pack(anchor="e")

        def combo(lbl, vals):
            tk.Label(center, text=lbl, font=('Segoe UI', 11), bg="white").pack(anchor="e", pady=(10,
                                                                                                 0))
            cb = ttk.Combobox(center, values=vals, state="readonly", font=('Segoe UI', 11), justify="right")
            cb.pack(fill="x")
            return cb


        classes = self.m.get_available_classes()
        self.cb_setup_class = combo("الصف:", classes)
        sections = self.m.get_available_sections()
        if not sections:
            sections = [str(i) for i in range(1, 10)]
        self.cb_setup_section = combo("الشعبة:", sections)
        subjects = self.m.get_all_subjects()
        if not subjects:
            subjects = ['رياضيات', 'علوم', 'لغتي', 'انجليزي', 'دراسات إسلامية', 'اجتماعيات', 'مهارات رقمية', 
             'بدنية', 'فنية', 'تفكير ناقد']
        self.cb_setup_subject = combo("المادة:", subjects)
        self.cb_setup_period = combo("رقم الحصة:", [str(i) for i in range(1, 8)])

        def do_start_setup():
            cls = self.cb_setup_class.get()
            sec = self.cb_setup_section.get()
            sub = self.cb_setup_subject.get()
            per = self.cb_setup_period.get()
            if not (cls and sec and sub and per):
                messagebox.showwarning("تنبيه", "الرجاء اختيار جميع الحقول")
                return
            self.current_lesson_data = {'date':today_str, 
             'teacher':getattr(self, "current_teacher_user", "Unknown"), 
             'class':cls, 
             'section':sec, 
             'subject':sub, 
             'period':per}
            self.current_lesson_criteria = [
             'correct', 'half', 'wrong', 'late', 'absent', 'behavior']
            self.setup_monitoring_table()
            self.show("teacher_monitoring_view")


        tk.Button(center, text="بدء المتابعة (Enter) ↵", command=do_start_setup, bg="#2e7d32",
          fg="white",
          font=('Segoe UI', 12, 'bold'),
          width=30,
          pady=5).pack(pady=30)
        center.bind("<Return>", lambda e: do_start_setup())
        tk.Button(center, text="إلغاء", command=(lambda: self.show("teacher_dashboard")), bg="white",
          fg="#777",
          bd=0).pack()
        return frm

        # ========================================

    # --- page_teacher_lesson_setup_combo ---
    def page_teacher_lesson_setup_combo(lbl=None, vals=None):
        # ========================================
        tk.Label(center, text=lbl, font=('Segoe UI', 11), bg="white").pack(anchor="e", pady=(10,
                                                                                             0))
        cb = ttk.Combobox(center, values=vals, state="readonly", font=('Segoe UI', 11), justify="right")
        cb.pack(fill="x")
        return cb

        # ========================================

    # --- page_teacher_lesson_setup_do_start_setup ---
    def page_teacher_lesson_setup_do_start_setup():
        # ========================================
        cls = self.cb_setup_class.get()
        sec = self.cb_setup_section.get()
        sub = self.cb_setup_subject.get()
        per = self.cb_setup_period.get()
        if not (cls and sec and sub and per):
            messagebox.showwarning("تنبيه", "الرجاء اختيار جميع الحقول")
            return
        self.current_lesson_data = {'date':today_str, 
         'teacher':getattr(self, "current_teacher_user", "Unknown"), 
         'class':cls, 
         'section':sec, 
         'subject':sub, 
         'period':per}
        self.current_lesson_criteria = [
         'correct', 'half', 'wrong', 'late', 'absent', 'behavior']
        self.setup_monitoring_table()
        self.show("teacher_monitoring_view")

        # ========================================

    # --- page_teacher_criteria_select ---
    def page_teacher_criteria_select(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg="#f1f8e9")
        top = tk.Frame(frm, bg="#2e7d32", pady=15)
        top.pack(fill="x")
        tk.Label(top, text="حدد عناصر المتابعة لهذه الحصة", font=('Segoe UI', 16, 'bold'), bg="#2e7d32", fg="white").pack()
        center = tk.Frame(frm, bg="#f1f8e9")
        center.pack(expand=True, fill="both", padx=50, pady=20)
        options = [
         ('جواب صحيح', 'correct'), 
         ('نصف جواب', 'half'), 
         ('جواب خاطئ', 'wrong'), 
         ('تأخير', 'late'), 
         ('غياب', 'absent'), 
         ('لم يحضر الكتاب', 'no_book'), 
         ('لم يحضر سجل الواجبات', 'no_hw'), 
         ('لم يحضر الأدوات', 'no_tools'), 
         ('نائم', 'sleep'), 
         ('عبث', 'play'), 
         ('مشاركة', 'participate')]
        self.chk_vars = {}
        grid_frame = tk.Frame(center, bg="#f1f8e9")
        grid_frame.pack(anchor="center")
        for i, (label, key) in enumerate(options):
            var = tk.BooleanVar(value=False)
            self.chk_vars[key] = var
            cb = tk.Checkbutton(grid_frame, text=label, variable=var, bg="#f1f8e9",
              activebackground="#f1f8e9",
              font=('Segoe UI', 14),
              selectcolor="#a5d6a7")
            cb.grid(row=(i // 2), column=(i % 2), sticky="w", padx=20, pady=10)
        else:

            def do_start_monitoring():
                selected = [k for k, v in self.chk_vars.items() if v.get()]
                if not selected:
                    if not messagebox.askyesno("تأكيد", "لم تختر أي عنصر متابعة! هل تريد المتابعة فقط (تحضير)؟"):
                        return
                self.current_lesson_criteria = selected
                self.setup_monitoring_table()
                self.show("teacher_monitoring_view")


            tk.Button(frm, text="بدء المتابعة ✅", command=do_start_monitoring, bg="#2e7d32",
              fg="white",
              font=('Segoe UI', 14, 'bold'),
              padx=30,
              pady=10).pack(pady=20)
            return frm

        # ========================================

    # --- page_teacher_criteria_select_do_start_monitoring ---
    def page_teacher_criteria_select_do_start_monitoring():
        # ========================================
        selected = [k for k, v in self.chk_vars.items() if v.get()]
        if not selected:
            if not messagebox.askyesno("تأكيد", "لم تختر أي عنصر متابعة! هل تريد المتابعة فقط (تحضير)؟"):
                return
        self.current_lesson_criteria = selected
        self.setup_monitoring_table()
        self.show("teacher_monitoring_view")

        # ========================================

    # --- setup_monitoring_table ---
    def setup_monitoring_table(self):
        # ========================================


        def norm(x):
            if pd.isna(x):
                return ""
            return normalize_arabic(str(x)).strip()


        target_cls = norm(self.current_lesson_data.get("class", ""))
        target_sec = norm(self.current_lesson_data.get("section", ""))
        self.current_student_list = []
        try:
            all_students = self.m.list_students_simple()
            if not all_students.empty:
                if "الصف" in all_students.columns:
                    if "الشعبة" in all_students.columns:
                        all_students["_norm_class"] = all_students["الصف"].astype(str).apply(normalize_arabic).str.strip()
                        all_students["_norm_section"] = all_students["الشعبة"].astype(str).apply(normalize_arabic).str.strip()
                        mask = (all_students["_norm_class"] == target_cls) & (all_students["_norm_section"] == target_sec)
                        filtered = all_students[mask].to_dict("records")
                        self.current_student_list = sorted(filtered, key=(lambda x: x.get("الاسم", "")))
        except Exception as e:
            try:
                print(f"Error loading students: {e}")
            finally:
                pass

        else:
            if hasattr(self, "monitoring_container"):
                self.build_monitoring_grid(self.monitoring_container)

        # ========================================

    # --- setup_monitoring_table_norm ---
    def setup_monitoring_table_norm(x=None):
        # ========================================
        if pd.isna(x):
            return ""
        return normalize_arabic(str(x)).strip()

        # ========================================

    # --- page_teacher_monitoring_view ---
    def page_teacher_monitoring_view(self, parent=None):
        # ========================================
        frm = tk.Frame(parent, bg="#f1f8e9")
        header = tk.Frame(frm, bg="#2e7d32", padx=20, pady=10)
        header.pack(fill="x")

        def get_stat_text():
            d = self.current_lesson_data
            return f'الصف: {d["class"]} - {d["section"]} | المادة: {d["subject"]} | الحصة: {d["period"]}'


        self.lbl_mon_info = tk.Label(header, text="...", font=('Segoe UI', 12, 'bold'), bg="#2e7d32", fg="white")
        self.lbl_mon_info.pack(side="right")
        tk.Button(header, text="إنهاء الحصة", command=(lambda: self.show("teacher_dashboard")), bg="#d32f2f",
          fg="white").pack(side="left")
        canvas = tk.Canvas(frm, bg="#f1f8e9")
        scrollbar = tk.Scrollbar(frm, orient="vertical", command=(canvas.yview))
        scroll_frame = tk.Frame(canvas, bg="#f1f8e9")
        scroll_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=(canvas.bbox("all"))))
        canvas.create_window((0, 0), window=scroll_frame, anchor="ne")
        canvas.configure(yscrollcommand=(scrollbar.set))
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        self.monitoring_container = scroll_frame
        self.monitoring_widgets = {}
        return frm

        # ========================================

    # --- page_teacher_monitoring_view_get_stat_text ---
    def page_teacher_monitoring_view_get_stat_text():
        # ========================================
        d = self.current_lesson_data
        return f'الصف: {d["class"]} - {d["section"]} | المادة: {d["subject"]} | الحصة: {d["period"]}'

        # ========================================

    # --- refresh_monitoring_view ---
    def refresh_monitoring_view(self):
        # ========================================
        return

        # ========================================

    # --- build_monitoring_grid ---
    def build_monitoring_grid(self, container=None):
        # ========================================
        for widget in container.winfo_children():
            widget.destroy()
        else:
            criteria_keys = self.current_lesson_criteria
            students = self.current_student_list
            headers_row = tk.Frame(container, bg="#c8e6c9")
            headers_row.pack(fill="x", pady=2)
            tk.Label(headers_row, text="#", width=4, bg="#c8e6c9", font=('Segoe UI', 10, 'bold')).pack(side="right", padx=2)
            tk.Label(headers_row, text="اسم الطالب", width=25, bg="#c8e6c9", font=('Segoe UI',
                                                                                   10, 'bold')).pack(side="right", padx=2)
            for k in criteria_keys:
                labels = {
                 'correct': '"جواب صحيح"', 
                 'half': '"نصف جواب"', 
                 'wrong': '"جواب خاطئ"', 
                 'late': '"تأخير"', 
                 'absent': '"غياب"', 
                 'behavior': '"سلوك"', 
                 'participate': '"مشاركة"'}
                lbl = labels.get(k, k)
                tk.Label(headers_row, text=lbl, width=10, bg="#c8e6c9").pack(side="right", padx=2)
            else:
                tk.Label(headers_row, text="متابعة أخرى", width=12, bg="#c8e6c9").pack(side="right", padx=2)
                for idx, student in enumerate(students):
                    row = tk.Frame(container, bg="white")
                    row.pack(fill="x", pady=1)
                    s_name = student.get("name", "Unknown")
                    s_id = student.get("id", "")
                    tk.Label(row, text=(str(idx + 1)), width=4, bg="white").pack(side="right")
                    tk.Label(row, text=s_name, width=25, anchor="e", bg="white").pack(side="right")
                    for crit in criteria_keys:
                        btn = tk.Button(row, text="⚪", width=4, bg="#eceff1", bd=0)
                        btn.configure(command=(lambda b=btn, s=s_name, c=crit: self.toggle_monitoring_btn(b, s, c)))
                        btn.pack(side="right", padx=2)
                    else:
                        tk.Button(row, text="📝", width=4, command=(lambda s=s_name: self.prompt_note(s))).pack(side="right", padx=5)

        # ========================================

    # --- toggle_monitoring_btn ---
    def toggle_monitoring_btn(self, btn=None, student=None, criteria=None):
        # ========================================
        current_bg = btn.cget("bg")
        is_on = current_bg != "#eceff1"
        new_state = not is_on
        positive = [
         "correct", "half", "participate"]
        negative = ['wrong', 'late', 'absent', 'no_book', 'no_hw', 'no_tools', 'sleep', 'play']
        if criteria in positive:
            active_color = "#66bb6a"
        else:
            active_color = "#ef5350"
        if new_state:
            btn.configure(bg=active_color, text="✔️")
        else:
            btn.configure(bg="#eceff1", text="⚪")
        self.save_monitoring_record(student, criteria, new_state)

        # ========================================

    # --- save_monitoring_record ---
    def save_monitoring_record(self, student=None, criteria=None, state=None):
        # ========================================
        try:
            record = {'timestamp':(datetime.now().strftime)("%Y-%m-%d %H:%M:%S"), 
             'date':self.current_lesson_data.get("date"), 
             'teacher':self.current_lesson_data.get("teacher"), 
             'class':self.current_lesson_data.get("class"), 
             'section':self.current_lesson_data.get("section"), 
             'subject':self.current_lesson_data.get("subject"), 
             'period':self.current_lesson_data.get("period"), 
             'student_name':student, 
             'action':criteria, 
             'value':1 if (state is True) else (0 if state is False else str(state)), 
             'details':""}
            self.m.log_monitoring_event(record)
        except Exception as e:
            try:
                print(f"Save error: {e}")
            finally:
                pass

        # ========================================

    # --- prompt_note ---
    def prompt_note(self, student=None):
        # ========================================
        note = simpledialog.askstring("ملاحظة", f"سجل ملاحظة للطالب: {student}")
        if note:
            self.save_monitoring_record(student, "note", note)

        # ========================================
        # Main Module Level 
        # ========================================
        # [DECOMPILATION FAILED]: 



class DataModel:

    # --- __init__ ---
    def __init__(self):
        # ========================================
        self.backup_dir = data_path("backups")
        if not os.path.exists(self.backup_dir):
            try:
                os.makedirs(self.backup_dir)
            except:
                return None

        self.reload_all(silent=True)
        self.STATUS_FINISHED = "منجز"
        self.STATUS_WORKING = "جاري"

        # ========================================

    # --- _read_excel_any ---
    def _read_excel_any(self, path=None, sheet_name=None, silent=None):
        # ========================================

        try:
            if sheet_name:
                return pd.read_excel(path, sheet_name=sheet_name)
            return pd.read_excel(path)
        except Exception as e:
            if not silent: print(f"Excel read error: {e}")
            return pd.DataFrame()






        # ========================================
        pass

    # --- _safe_save_dataframe ---
    def _safe_save_dataframe(self, df=None, path=None, sheet_name=None, beautify=None):
        # ========================================

        try:
            df.to_excel(path, sheet_name=(sheet_name or "Sheet1"), index=False)
            return True
        except Exception as e:
            print(f"Excel save error: {e}")
            return False






        # ========================================
        pass

    # --- _save_formatted_xlsx ---
    def _save_formatted_xlsx(self, df=None, path=None, sheet_name=None):
        # ========================================
        return self._safe_save_dataframe(df, path, sheet_name=sheet_name, beautify=True)

        # ========================================

    # --- reload_all ---
    def reload_all(self, silent=None):
        # ========================================
        ensure_file(FILE_SUBJECTS, pd.DataFrame({"المادة": []}))
        self.df_subjects = self._read_excel_any(FILE_SUBJECTS, silent=silent)
        ensure_file(FILE_STUDENTS, pd.DataFrame({'الاسم':[],  'الصف':[],  'الشعبة':[],  'ولي الأمر':[],  'جوال ولي الأمر':[],  'ملاحظات':[]}))
        self.df_students = self._read_excel_any(FILE_STUDENTS, silent=silent)
        if not os.path.exists(FILE_TIMINGS):
            rows = []
            for p in range(1, 8):
                rows.append({'الحصة':str(p),  'من':"07:00",  'إلى':"07:45"})
            else:
                ensure_file(FILE_TIMINGS, pd.DataFrame(rows))

        else:
            self.df_timings = self._read_excel_any(FILE_TIMINGS, silent=silent)
            self.df_timings = self._normalize_timings(self.df_timings)
            self.max_periods = self._calc_max_periods()
            if not os.path.exists(FILE_CALENDAR):
                df = pd.DataFrame(columns=['الفصل الدراسي', 'الأسبوع', 'اليوم', 'التاريخ الميلادي', 'الملاحظات'])
                ensure_file(FILE_CALENDAR, df, sheet_name="التقويم")
            else:
                self.df_calendar = self.load_calendar(FILE_CALENDAR, silent=True)
                self.teacher_names = []
                if os.path.exists(FILE_MASTER):
                    try:
                        import openpyxl
                        wb = openpyxl.load_workbook(FILE_MASTER, read_only=True)
                        self.teacher_names = list(wb.sheetnames)
                        wb.close()
                    except Exception as e:
                        try:
                            self.teacher_names = []
                        finally:
                            pass

                else:
                    silent or messagebox.showwarning("تحذير", f"الملف غير موجود:\n{FILE_MASTER}\nرجاء وضعه بجانب البرنامج.")

        # ========================================

    # --- merge_attendance ---
    def merge_attendance(self, new_df=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- merge_swaps ---
    def merge_swaps(self, new_df=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- merge_swaps_make_key ---
    def merge_swaps_make_key(row=None):
        # ========================================
        return f'{row.get("طالب المبادلة", "")}_{row.get("المعلم البديل", "")}_{row.get("تاريخ المبادلة", "")}_{row.get("الحصة", "")}'

        # ========================================

    # --- get_all_teachers ---
    def get_all_teachers(self):
        # ========================================
        return sorted(self.teacher_names)

        # ========================================

    # --- get_all_employees ---
    def get_all_employees(self):
        # ========================================
        try:
            if not os.path.exists(FILE_EMPLOYEES_PINS):
                return []
            df = self._read_excel_any(FILE_EMPLOYEES_PINS, silent=True)
            if df.empty:
                return []
            name_col = None
            for c in df.columns:
                if "اسم" in str(c) or "الموظف" in str(c):
                    name_col = c
                    break
            else:
                if not name_col:
                    name_col = df.columns[0]
                names = df[name_col].dropna().astype(str).unique()

            return sorted([tidy(n) for n in names if tidy(n)])
        except Exception:
            return []

        # ========================================

    # --- get_cell ---
    def get_cell(self, teacher_name=None, day=None, period_idx=None):
        # ========================================

        return ""


    # --- _calc_max_periods ---
    def _calc_max_periods(self):
        # ========================================
        try:
            p_nums = []
            for x in self.df_timings["الحصة"].astype(str):
                x = x.strip()
                digits = "".join(filter(str.isdigit, x))
                if digits:
                    p_nums.append(int(digits))
            else:
                if x.isdigit():
                    p_nums.append(int(x))

            return max(p_nums) if p_nums else 7
        except Exception:
            return 7

        # ========================================

    # --- get_period_labels ---
    def get_period_labels(self):
        # ========================================
        return [str(i) for i in range(1, self.max_periods + 1)]

        # ========================================

    # --- get_all_subjects ---
    def get_all_subjects(self):
        # ========================================

        try:
            if "المادة" in self.df_subjects.columns:
                return sorted([str(x).strip() for x in self.df_subjects["المادة"].dropna() if str(x).strip()])
            return []
        except: return []






        # ========================================
        pass

    # --- add_subject ---
    def add_subject(self, name=None):
        # ========================================
        name = tidy(name)
        if not name:
            return (False, 'أدخل اسم مادة.')
        if "المادة" not in self.df_subjects.columns:
            self.df_subjects["المادة"] = []
        if name in list(self.df_subjects["المادة"].astype(str)):
            return (False, 'المادة موجودة مسبقًا.')
        self.df_subjects = pd.concat([self.df_subjects, pd.DataFrame([{"المادة": name}])], ignore_index=True)
        self._save_formatted_xlsx((self.df_subjects), FILE_SUBJECTS, sheet_name="المواد")
        return (True, 'تمت إضافة المادة.')

        # ========================================

    # --- delete_subjects ---
    def delete_subjects(self, names=None):
        # ========================================
        if "المادة" not in self.df_subjects.columns:
            self.df_subjects["المادة"] = []
        names_norm = [tidy(str(n)) for n in names]
        temp_col = self.df_subjects["المادة"].astype(str).apply(tidy)
        before = len(self.df_subjects)
        self.df_subjects = self.df_subjects[~temp_col.isin(names_norm)].copy()
        self._save_formatted_xlsx((self.df_subjects), FILE_SUBJECTS, sheet_name="المواد")
        return before - len(self.df_subjects)

        # ========================================

    # --- import_subjects_from_excel ---
    def import_subjects_from_excel(self, path=None):
        # ========================================

        try:
            df = pd.read_excel(path)
            added = 0
            for c in df.columns:
                for v in df[c].dropna():
                    v_str = str(v).strip()
                    if v_str and v_str not in list(self.df_subjects.get("المادة", [])):
                        self.df_subjects = pd.concat([self.df_subjects, pd.DataFrame([{"المادة": v_str}])], ignore_index=True)
                        added += 1
            if added > 0:
                self._save_formatted_xlsx(self.df_subjects, FILE_SUBJECTS, sheet_name="المواد")
            return added, f"تم استيراد {added} مادة"
        except Exception as e: return 0, f"خطأ: {e}" 






        # ========================================
        pass

    # --- _detect_student_cols ---
    def _detect_student_cols(self):
        # ========================================
        name_col = class_col = section_col = None
        for c in self.df_students.columns:
            c_str = str(c).strip()
            if "اسم" in c_str and ("طالب" in c_str or c_str == "الاسم"):
                name_col = c_str
            if "الاسم" == c_str and not name_col:
                name_col = c_str
            elif "الصف" in c_str:
                class_col = c_str
            elif "شعب" in c_str:
                section_col = c_str
        return name_col or "الاسم", class_col or "الصف", section_col or "الشعبة"



        # ========================================

    # --- get_student_contact_details ---
    def get_student_contact_details(self, student_name=None):
        # ========================================
        if self.df_students.empty:
            return {'mobile':"—", 
             'home':"—",  'work':"—"}
        else:
            n_search = normalize_arabic(str(student_name))
            name_col, _, _ = self._detect_student_cols()
            if not name_col:
                return {'mobile':"Err", 
                 'home':"Err",  'work':"Err"}
            cols = self.df_students.columns
            mob_col = next((c for c in [] if "جوال" in str(c)), None)
            home_col = next((c for c in []), None)
            work_col = next((c for c in [] if "عمل" in str(c)), None)
            mask = self.df_students[name_col].astype(str).apply(normalize_arabic).str.contains(n_search, na=False)
            res = self.df_students[mask]
            ret = {'mobile':"—", 
             'home':"—",  'work':"—"}
            row = res.empty or res.iloc[0]

            def get_val(c):
                if c:
                    v = str(row[c]).strip()
                    if v:
                        if v.lower() != "nan":
                            return v
                    return "—"
                return "— (غير موجود)"


            ret["mobile"] = get_val(mob_col)
            ret["home"] = get_val(home_col)
            ret["work"] = get_val(work_col)
        return ret

        # ========================================

    # --- get_student_contact_details_get_val ---
    def get_student_contact_details_get_val(c=None):
        # ========================================
        if c:
            v = str(row[c]).strip()
            if v:
                if v.lower() != "nan":
                    return v
            return "—"
        return "— (غير موجود)"

        # ========================================

    # --- save_pending_contacts ---
    def save_pending_contacts(self, batch_list=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- load_pending_contacts ---
    def load_pending_contacts(self, date_filter=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- update_contact_status ---
    def update_contact_status(self, student_name=None, date_str=None, new_status=None, notes=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- list_students_simple ---
    def list_students_simple(self):
        # ========================================
        name_col, class_col, section_col = self._detect_student_cols()
        result = pd.DataFrame(index=(self.df_students.index))
        result["الاسم"] = self.df_students[name_col] if name_col in self.df_students.columns else ""
        result["الصف"] = self.df_students[class_col] if class_col in self.df_students.columns else ""
        result["الشعبة"] = self.df_students[section_col] if section_col in self.df_students.columns else ""
        return result.fillna("")

        # ========================================

    # --- get_student_full ---
    def get_student_full(self, idx=None):
        # ========================================
        try:
            idx = int(idx)
            if idx < 0 or idx >= len(self.df_students):
                return {}
            return self.df_students.iloc[idx].to_dict()
        except Exception:
            return {}

        # ========================================

    # --- update_student_full ---
    def update_student_full(self, idx=None, data=None):
        # ========================================
        for k, v in data.items():
            if k not in self.df_students.columns:
                self.df_students[k] = ""
            self.df_students.at[(idx, k)] = v
        else:
            self._save_formatted_xlsx((self.df_students), FILE_STUDENTS, sheet_name="الطلاب")

        # ========================================

    # --- get_available_classes ---
    def get_available_classes(self):
        # ========================================
        _, class_col, _ = self._detect_student_cols()
        if not class_col or class_col not in self.df_students.columns:
            return []
        classes = self.df_students[class_col].dropna().unique()
        classes = [tidy(c) for c in classes if tidy(c)]
        return sorted(classes)

        # ========================================

    # --- get_available_sections ---
    def get_available_sections(self):
        # ========================================
        _, _, section_col = self._detect_student_cols()
        if not section_col or section_col not in self.df_students.columns:
            return []
        sections = self.df_students[section_col].dropna().unique()
        sections = [tidy(str(s)) for s in sections if tidy(str(s))]
        return sorted(sections)

        # ========================================

    # --- get_students_by_class_section ---
    def get_students_by_class_section(self, class_name=None, section=None):
        # ========================================
        results = []
        name_col, class_col, section_col = self._detect_student_cols()
        class_norm = normalize_arabic(class_name)
        section_norm = normalize_arabic(section)
        for idx, row in self.df_students.iterrows():
            row_class = normalize_arabic(row.get(class_col, "")) if class_col else ""
            row_sec = normalize_arabic(row.get(section_col, "")) if section_col else ""
            if row_class == class_norm and row_sec == section_norm:
                results.append({'index':idx,  'name':tidy(row.get(name_col, ""))})
            return results

        # ========================================

    # --- _normalize_timings ---
    def _normalize_timings(self, df=None):
        # ========================================
        if df is None or df.empty:
            return pd.DataFrame(columns=["الحصة", "من", "إلى"])
        cols_map = {}
        for col in df.columns:
            c = str(col).strip()
            cl = c.lower()
            if "حصة" in c or "حدث" in c:
                cols_map[col] = "الحصة"
            elif c == "من" or "من" in c:
                cols_map[col] = "من"
            else:
                if "إلى" in c or "الى" in cl:
                    cols_map[col] = "إلى"
                df = df.rename(columns=cols_map)
                for needed in ('الحصة', 'من', 'إلى'):
                    if needed not in df.columns:
                        df[needed] = ""
                    return df[["الحصة", "من", "إلى"]].copy()

        # ========================================

    # --- get_period_time ---
    def get_period_time(self, period_num_str=None):
        # ========================================
        if self.df_timings is None or self.df_timings.empty:
            return ('', '')
        df = self.df_timings.copy()
        target = str(period_num_str).strip()
        row = df[df["الحصة"].astype(str).str.strip() == target]
        if row.empty:
            for i, r in df.iterrows():
                val = str(r["الحصة"]).strip()
                digits = "".join(filter(str.isdigit, val))
                if digits == target:
                    return (
                     tidy(r.get("من", "")), tidy(r.get("إلى", "")))

        if row.empty:
            return ('', '')
        r = row.iloc[0]
        return (tidy(r.get("من", "")), tidy(r.get("إلى", "")))

        # ========================================

    # --- get_current_active_period ---
    def get_current_active_period(self):
        # ========================================
        from datetime import datetime
        import datetime as dt_module
        try:
            if getattr(self, "df_timings", None) is None or self.df_timings.empty:
                return ("--", "Out", 0)
                
            def parse_t(t_val):
                try:
                    if isinstance(t_val, dt_module.time):
                        return t_val.hour, t_val.minute
                    if isinstance(t_val, dt_module.datetime):
                        return t_val.hour, t_val.minute
                    s = str(t_val).replace("ص", "").replace("م", "").strip()
                    parts = s.split(":")
                    if len(parts) >= 2:
                        return int(parts[0]), int(parts[1])
                    return None
                except Exception:
                    return None
                    
            now = datetime.now()
            t_now_dt = now
            min_dt = None
            max_dt = None
            
            for _, row in self.df_timings.iterrows():
                name = str(row.get("الحصة", ""))
                from_t = parse_t(row.get("من", ""))
                to_t = parse_t(row.get("إلى", ""))
                if from_t and to_t:
                    start_dt = now.replace(hour=from_t[0], minute=from_t[1], second=0, microsecond=0)
                    end_dt = now.replace(hour=to_t[0], minute=to_t[1], second=0, microsecond=0)
                    
                    if min_dt is None or start_dt < min_dt: min_dt = start_dt
                    if max_dt is None or end_dt > max_dt: max_dt = end_dt
                    
                    if start_dt <= t_now_dt <= end_dt:
                        rem_minutes = max(0, int((end_dt - t_now_dt).total_seconds() / 60))
                        return (name, "Active", rem_minutes)
                        
            if min_dt and max_dt:
                if t_now_dt < min_dt or t_now_dt > max_dt:
                    return ("--", "Out", 0)
            return ("--", "Wait", 0)
        except Exception:
            return ("--", "Out", 0)

    # --- load_calendar ---
    def load_calendar(self, path=None, silent=None):
        # ========================================

        df = self._read_excel_any(FILE_CALENDAR, silent=True)
        if df.empty:
            return pd.DataFrame(columns=['الفصل الدراسي', 'الأسبوع', 'اليوم', 'التاريخ الهجري', 'التاريخ الميلادي', 'الملاحظات'])
        return normalize_calendar_columns(df)


    # --- save_calendar ---
    def save_calendar(self):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar)
        self._save_formatted_xlsx(df, FILE_CALENDAR, sheet_name="التقويم")
        self.df_calendar = df

        # ========================================

    # --- import_calendar_from_excel ---
    def import_calendar_from_excel(self, path=None):
        # ========================================
        df = self.load_calendar(path, silent=True)
        if df.empty:
            return (False, 'الملف فارغ أو غير متوافق.')
        self.df_calendar = df
        self.save_calendar()
        return (True, 'تم استيراد التقويم بنجاح.')

        # ========================================

    # --- calendar_summary ---
    def calendar_summary(self):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar)
        if df.empty:
            return {'start':None, 
             'end':None,  'weeks':None,  'terms':[]}
        dates = [safe_parse_date(x) for x in df["التاريخ الميلادي"].tolist()]
        dates = [d for d in dates if d is not None]
        start = min(dates) if dates else None
        end = max(dates) if dates else None
        weeks = None
        try:
            w = pd.to_numeric((df["الأسبوع"]), errors="coerce")
            weeks = int(w.max()) if w.notna().any() else None
        except Exception:
            weeks = None
        else:
            terms = sorted([t for t in df["الفصل الدراسي"].astype(str).unique().tolist() if t if t != "nan"])
            return {'start': start, 'end': end, 'weeks': weeks, 'terms': terms}

        # ========================================

    # --- get_day_status ---
    def get_day_status(self, on_date=None):
        # ========================================
        on_date = on_date or datetime.now().date()
        df = normalize_calendar_columns(self.df_calendar)
        if df.empty:
            return {'date':on_date,  'in_calendar':False,  'is_holiday':False,  'note':"", 
             'term':"",  'week':None,  'day':arabic_day_from_english(on_date.strftime("%A"))}
        df2 = df.copy()
        df2["_d"] = df2["التاريخ الميلادي"].apply(safe_parse_date)
        row = df2[df2["_d"] == on_date]
        day_name_ar = arabic_day_from_english(on_date.strftime("%A"))
        if row.empty:
            is_holiday = day_name_ar in ('الجمعة', 'السبت')
            return {
             'date': on_date, 'in_calendar': False, 'is_holiday': is_holiday, 
             'note': '""', 'term': '""', 'week': None, 'day': day_name_ar, 'hijri': '""'}
        r = row.iloc[0]
        note = tidy(r.get("الملاحظات", ""))
        term = tidy(r.get("الفصل الدراسي", ""))
        week = None
        try:
            week_val = r.get("الأسبوع", None)
            week = int(float(week_val)) if str(week_val).strip() not in ('', 'nan', 'None') else None
        except Exception:
            week = None
        else:
            day = tidy(r.get("اليوم", "")) or arabic_day_from_english(on_date.strftime("%A"))
            note_norm = normalize_arabic(note)
            is_holiday = "اجاز" in note_norm or "عطل" in note_norm or "holiday" in note_norm.lower()
            if day in ('الجمعة', 'السبت'):
                if "دوام" not in note_norm:
                    is_holiday = True
            return {'date':on_date,  'in_calendar':True, 
             'is_holiday':bool(is_holiday), 
             'note':note, 
             'term':term, 
             'week':week, 
             'day':day, 
             'hijri':tidy(r.get("التاريخ الهجري", ""))}

        # ========================================

    # --- get_holiday_range ---
    def get_holiday_range(self, on_date=None):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar).copy()
        df["_d"] = df["التاريخ الميلادي"].apply(safe_parse_date)
        df = df.sort_values("_d").reset_index(drop=True)
        st_start = self.get_day_status(on_date)
        if not st_start.get("is_holiday"):
            return
        note_norm = normalize_arabic(tidy(st_start.get("note", "")))
        if note_norm and len(note_norm) > 3:
            df_h = df[df["الملاحظات"].astype(str).apply(normalize_arabic).str.contains(note_norm, na=False)].copy()
            if not df_h.empty:
                current_s = df_h["_d"].min()
                current_e = df_h["_d"].max()
        else:

            def is_h(dt):
                return self.get_day_status(dt).get("is_holiday")


            current_s = on_date
            while is_h(current_s - timedelta(days=1)):
                current_s -= timedelta(days=1)

            current_e = on_date
            while is_h(current_e + timedelta(days=1)):
                current_e += timedelta(days=1)

            def get_h_str(dt):
                row = df[df["_d"] == dt]
                if not row.empty:
                    if tidy(row.iloc[0].get("التاريخ الهجري", "")):
                        return tidy(row.iloc[0].get("التاريخ الهجري", ""))
                for off in (1, -1, 2, -2):
                    nr = df[df["_d"] == dt + timedelta(days=off)]
                    if not nr.empty:
                        if tidy(nr.iloc[0].get("التاريخ الهجري", "")):
                            return "~" + tidy(nr.iloc[0].get("التاريخ الهجري", ""))
                    return ""


            return {'start':current_s, 
             'end':current_e,  'start_h':get_h_str(current_s), 
             'end_h':get_h_str(current_e)}

        # ========================================

    # --- get_holiday_range_is_h ---
    def get_holiday_range_is_h(dt=None):
        # ========================================
        return self.get_day_status(dt).get("is_holiday")

        # ========================================

    # --- get_holiday_range_get_h_str ---
    def get_holiday_range_get_h_str(dt=None):
        # ========================================
        row = df[df["_d"] == dt]
        if not row.empty:
            if tidy(row.iloc[0].get("التاريخ الهجري", "")):
                return tidy(row.iloc[0].get("التاريخ الهجري", ""))
        for off in (1, -1, 2, -2):
            nr = df[df["_d"] == dt + timedelta(days=off)]
            if not nr.empty:
                if tidy(nr.iloc[0].get("التاريخ الهجري", "")):
                    return "~" + tidy(nr.iloc[0].get("التاريخ الهجري", ""))
            return ""

        # ========================================

    # --- get_holiday_summary_list ---
    def get_holiday_summary_list(self):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar).copy()
        df["_d"] = df["التاريخ الميلادي"].apply(safe_parse_date)
        df = df[df["_d"].notna()].sort_values("_d")
        df = df[df["الأسبوع"].astype(str) != "17"]
        keywords = ['رمضان', 'اضحى', 'فطر', 'تاسيس', 'وطني', 'فصل', 'عطل', 'اجاز', 'عيد', 'يوم']
        
        holidays_list = []
        
        for _, r in df.iterrows():
            note = tidy(r.get("الملاحظات", ""))
            n_norm = normalize_arabic(note)
            is_major = any((k in n_norm for k in keywords))
            if is_major and len(n_norm) > 2:
                dt = r["_d"]
                hij = tidy(r.get("التاريخ الهجري", ""))
                
                merged = False
                for h in holidays_list:
                    if 0 <= (dt - h["end_dt"]).days <= 6:
                        w1 = set(n_norm.split()) - set(['اجازة', 'بداية', 'نهاية'])
                        w2 = set(normalize_arabic(h['name']).split()) - set(['اجازة', 'بداية', 'نهاية'])
                        
                        if w1.intersection(w2) or ('مطول' in n_norm and 'مطول' in normalize_arabic(h['name'])):
                            h["end_dt"] = dt
                            h["end_h"] = hij
                            if 'بداية' in h['name'] and 'بداية' not in note:
                                h['name'] = note
                            merged = True
                            break
                
                if not merged:
                    holidays_list.append({
                        'name': note,
                        'start_dt': dt,
                        'end_dt': dt,
                        'start_h': hij,
                        'end_h': hij
                    })
                    
        for h in holidays_list:
            h["duration"] = (h["end_dt"] - h["start_dt"]).days + 1
            
        return sorted(holidays_list, key=lambda x: x["start_dt"])

        # ========================================

    # --- get_days_until_next_holiday ---
    def get_days_until_next_holiday(self):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar).copy()
        if df.empty:
            target = date(2026, 1, 15)
            today = date.today()
            rem = (target - today).days
            return ("إجازة منتصف العام (تقريبي)", rem if rem > 0 else 0)
        df["_d"] = df["التاريخ الميلادي"].apply(safe_parse_date)
        today = date.today()
        candidates = []
        keywords = "اجاز|عطل|عيد|فطر|اضحى|نهاية|تاسيس|وطني|منتصف|بدء"
        future_h = df[(df["_d"] >= today) & df["الملاحظات"].astype(str).apply(normalize_arabic).str.contains(keywords, na=False)]
        if not future_h.empty:
            next_h_dt = future_h["_d"].min()
            row = future_h[future_h["_d"] == next_h_dt].iloc[0]
            candidates.append((next_h_dt, tidy(row.get("الملاحظات", ""))))
        all_terms = df["الفصل الدراسي"].unique()
        for t_name in all_terms:
            if not isinstance(t_name, str):
                continue
            clean_t = normalize_arabic(t_name)
            if "اجاز" in clean_t:
                continue
            t_rows = df[df["الفصل الدراسي"] == t_name]
            if not t_rows.empty:
                last_day = t_rows["_d"].max()
                if pd.notna(last_day) and last_day >= today:
                    candidates.append((last_day, f"نهاية {t_name}"))
        if not candidates:
            return ('لا توجد أحداث قادمة', 0)
        candidates.sort(key=lambda x: x[0])
        best_date, best_note = candidates[0]
        if "نهاية" in best_note and "الفصل" in best_note:
            best_note = "بداية إجازة " + best_note
        days = (best_date - today).days
        return (best_note, days)

    # --- get_academic_progress ---
    def get_academic_progress(self):
        # ========================================
        from datetime import date, timedelta
        import pandas as pd
        try:
            df = normalize_calendar_columns(self.df_calendar).copy()
            if df.empty:
                return {'year_total': 0, 'year_curr': 0, 'year_pct': 0, 'terms': []}
            
            df["_d"] = df["التاريخ الميلادي"].apply(safe_parse_date)
            df = df[df["_d"].notna()]
            if df.empty:
                return {'year_total': 0, 'year_curr': 0, 'year_pct': 0, 'terms': []}
                
            today = date.today()
            df["composite_week"] = df["الفصل الدراسي"].astype(str) + "_" + df["الأسبوع"].astype(str)
            unique_weeks = df.drop_duplicates(subset=["composite_week"])
            total_year_weeks = len(unique_weeks)
            passed_year_weeks = len(unique_weeks[unique_weeks["_d"] <= today])
            year_pct = int(passed_year_weeks / total_year_weeks * 100) if total_year_weeks > 0 else 0
            
            terms_data = []
            term_names = df["الفصل الدراسي"].unique()
            term_order = []
            for t in term_names:
                first_dt = df[df["الفصل الدراسي"] == t]["_d"].min()
                term_order.append((t, first_dt))
                
            term_order.sort(key=lambda x: x[1])
            current_term_name = self.get_day_status(today).get("term", "")
            
            for t_name, _ in term_order:
                clean_name = normalize_arabic(str(t_name))
                if "اجاز" in clean_name:
                    continue
                    
                t_df = df[df["الفصل الدراسي"] == t_name]
                t_weeks = t_df["الأسبوع"].unique()
                t_total = len(t_weeks)
                t_passed_df = t_df[t_df["_d"] <= today]
                t_passed_count = len(t_passed_df["الأسبوع"].unique())
                t_pct = int(t_passed_count / t_total * 100) if t_total > 0 else 0
                
                t_start_day = t_df["_d"].min()
                t_last_day = t_df["_d"].max()
                t_days_left = (t_last_day - today).days if (pd.notna(t_last_day) and t_last_day >= today) else 0
                
                is_curr = (normalize_arabic(t_name) == normalize_arabic(current_term_name))
                if not is_curr and pd.notna(t_start_day) and pd.notna(t_last_day):
                    extended_end = t_last_day + timedelta(days=25) # Guess current term if close
                    if t_start_day <= today <= extended_end:
                        is_curr = True
                        
                terms_data.append({
                    'name': t_name,
                    'total': t_total,
                    'curr': t_passed_count,
                    'pct': t_pct,
                    'days_left': t_days_left,
                    'is_current': is_curr
                })
                
            return {
                'year_total': total_year_weeks,
                'year_curr': passed_year_weeks,
                'year_pct': year_pct,
                'terms': terms_data
            }
        except Exception:
            return {'year_total': 0, 'year_curr': 0, 'year_pct': 0, 'terms': []}
            
    # --- filter_calendar ---
    def filter_calendar(self, term=None, week=None, text=None):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar)
        if df.empty:
            return df
        out = df.copy()
        if term:
            term_norm = normalize_arabic(term)
            out = out[out["الفصل الدراسي"].astype(str).apply(normalize_arabic) == term_norm]
        if week not in (None, '', 'None'):
            try:
                w = float(week)
                out = out[pd.to_numeric((out["الأسبوع"]), errors="coerce") == w]
            except Exception:
                pass

        if text:
            t_norm = normalize_arabic(text)
            if t_norm:
                out = out[out["الملاحظات"].astype(str).apply(normalize_arabic).str.contains(t_norm, na=False)]
        out["_d"] = out["التاريخ الميلادي"].apply(safe_parse_date)
        out = out.sort_values("_d").drop(columns=["_d"], errors="ignore")
        return out

        # ========================================

    # --- add_calendar_row ---
    def add_calendar_row(self, data=None):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar)
        df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
        self.df_calendar = normalize_calendar_columns(df)
        self.save_calendar()

        # ========================================

    # --- update_calendar_row ---
    def update_calendar_row(self, idx=None, data=None):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar)
        if idx < 0 or idx >= len(df):
            return
        for k, v in data.items():
            if k not in df.columns:
                df[k] = ""
            df.at[(idx, k)] = v
        else:
            self.df_calendar = normalize_calendar_columns(df)
            self.save_calendar()

        # ========================================

    # --- delete_calendar_rows ---
    def delete_calendar_rows(self, indices=None):
        # ========================================
        df = normalize_calendar_columns(self.df_calendar)
        indices = sorted(list(set(int(i) for i in row_indices)), reverse=True)
        for i in indices:
            if 0 <= i < len(df):
                df = df.drop(df.index[i])
            df = df.reset_index(drop=True)
            self.df_calendar = normalize_calendar_columns(df)
            self.save_calendar()

        # ========================================

    # --- get_teacher_grid ---
    def get_teacher_grid(self, teacher_name=None):
        # ========================================
        import openpyxl
        wb = openpyxl.load_workbook(FILE_MASTER)
        if teacher_name not in wb.sheetnames:
            raise RuntimeError(f"لا توجد ورقة باسم المعلم: {teacher_name}")
        ws = wb[teacher_name]
        periods_list = self.get_period_labels()
        header = ["اليوم"] + periods_list
        rows = []
        for i, day in enumerate(DAYS, start=2):
            row = [
             day]
            for p_idx, _ in enumerate(periods_list, start=0):
                val = ws.cell(row=i, column=(p_idx + 2)).value
                row.append(val)
            else:
                rows.append(row)

        else:
            return pd.DataFrame(rows, columns=header)

        # ========================================

    # --- set_teacher_cell ---
    def set_teacher_cell(self, teacher_name=None, day=None, period=None, text=None):
        # ========================================
        import openpyxl
        wb = openpyxl.load_workbook(FILE_MASTER)
        if teacher_name not in wb.sheetnames:
            raise RuntimeError(f"لا توجد ورقة باسم {teacher_name}.")
        ws = wb[teacher_name]
        day_row = {DAYS[i]: i + 2 for i in range(5)}.get(day)
        if not day_row:
            raise RuntimeError("يوم غير صالح.")
        periods = self.get_period_labels()
        try:
            p_idx = periods.index(str(period))
            col = p_idx + 2
        except ValueError:
            raise RuntimeError("حصة غير صالحة.")
        else:
            ws.cell(row=day_row, column=col).value = text if text else "—"
            wb.save(FILE_MASTER)
            threading.Thread(target=(self.beautify_all_internal_databases), daemon=True).start()

        # ========================================

    # --- manager_week_table ---
    def manager_week_table(self):
        # ========================================
        import openpyxl
        import pandas as pd
        if not os.path.exists(FILE_MASTER): return pd.DataFrame()
        result = {d: {p: [] for p in self.get_period_labels()} for d in DAYS}
        try:
            wb = openpyxl.load_workbook(FILE_MASTER)
            periods_list = self.get_period_labels()
            for tname in wb.sheetnames:
                ws = wb[tname]
                for r_i, day in enumerate(DAYS, start=2):
                    for p_idx, p_label in enumerate(periods_list):
                        c_i = p_idx + 2
                        v = ws.cell(row=r_i, column=c_i).value
                        if v and str(v).strip() and str(v).strip() != "—":
                            result[day][p_label].append(f"{tname} – {v}")
            rows = []
            for d in DAYS:
                row = [d]
                for p in periods_list:
                    items = result[d][p]
                    row.append("\n".join(items) if items else "—")
                rows.append(row)
            return pd.DataFrame(rows, columns=(["اليوم"] + [f"الحصة {p}" for p in periods_list]))
        except Exception:
            return pd.DataFrame()

        # ========================================

    # --- find_student_schedule ---
    def find_student_schedule(self, student_name=None, student_class=None, student_section=None):
        # ========================================
        if not os.path.exists(FILE_MASTER):
            return ([], None, None)
        
        student_name_clean = tidy(student_name)
        name_col, class_col, section_col = self._detect_student_cols()
        
        if not student_class or not student_section:
            df_s = self.df_students
            matches = df_s[df_s[name_col].astype(str).str.strip() == student_name_clean]
            if matches.empty:
                matches = df_s[df_s[name_col].astype(str).str.contains(student_name_clean, na=False)]
            if not matches.empty:
                row = matches.iloc[0]
                student_class = tidy(row.get(class_col, ""))
                student_section = tidy(row.get(section_col, ""))
            else:
                return ([], None, None)
                
        import re, openpyxl
        c_nums = re.findall(r"\d+", student_class)
        c_aliases = [student_class, tidy(student_class)] + c_nums
        mapping_cls = [
            ('الأول', '1'), ('الاول', '1'), ('أول', '1'), ('اول', '1'), 
            ('الثاني', '2'), ('الثانى', '2'), ('ثاني', '2'), ('ثانى', '2'), 
            ('الثالث', '3'), ('ثالث', '3'), 
            ('الرابع', '4'), ('رابع', '4'), 
            ('الخامس', '5'), ('خامس', '5'), 
            ('السادس', '6'), ('سادس', '6')
        ]
        for k, v in mapping_cls:
            if k in student_class:
                c_aliases.append(v)
                
        s_nums = re.findall(r"\d+", student_section)
        s_aliases = [student_section, tidy(student_section)] + s_nums
        
        results = []
        periods_list = self.get_period_labels()
        last_period = None
        last_period_time = None
        
        try:
            wb = openpyxl.load_workbook(FILE_MASTER)
            for teacher_name in wb.sheetnames:
                ws = wb[teacher_name]
                for r_i, day in enumerate(DAYS, start=2):
                    for p_idx, p_label in enumerate(periods_list):
                        c_i = p_idx + 2
                        cell_value = ws.cell(row=r_i, column=c_i).value
                        if not cell_value: continue
                        
                        cell_text = str(cell_value).strip()
                        if cell_text == "—" or not cell_text: continue
                        
                        parsed = parse_teacher_cell(cell_text)
                        target = parsed["class_section"] if (parsed and parsed["class_section"]) else cell_text
                        t_norm = tidy(target)
                        
                        c_match = False
                        for ca in c_aliases:
                            if ca and ca in t_norm:
                                c_match = True
                                break
                                
                        s_match = False
                        for sa in s_aliases:
                            if sa and sa in t_norm:
                                s_match = True
                                break
                                
                        if c_match and s_match:
                            t_from, t_to = self.get_period_time(p_label)
                            results.append({
                                'اليوم': day, 
                                'الحصة': str(p_label), 
                                'المعلم': teacher_name, 
                                'المادة/الصف': cell_text, 
                                'من': t_from, 
                                'إلى': t_to
                            })
                            last_period = str(p_label)
                            last_period_time = t_to or ""
                            
            return (results, last_period, last_period_time)
        except Exception:
            return ([], None, None)

        # ========================================

    # --- get_student_current_location ---
    def get_student_current_location(self, student_name=None):
        # ========================================
        try:
            from datetime import datetime
            
            today_idx = datetime.now().weekday()
            days_map = {6: "الأحد", 0: "الاثنين", 1: "الثلاثاء", 2: "الأربعاء", 3: "الخميس"}
            today_name = days_map.get(today_idx)
            
            active_p, status, rem = self.get_current_active_period()
            if status == "Out":
                return ("خارج وقت الدوام 🌙", "", "", "")
                
            if not str(active_p).isdigit():
                return (f"الآن: {active_p} ☕", "", "", str(active_p))
                
            if not today_name:
                return ("خارج أيام الدوام 🌴", "", "", "")

            sched, _, _ = self.find_student_schedule(student_name)
            if not sched:
                return ("الجدول غير متاح 🚫", "", "", str(active_p))
                
            for row in sched:
                if row.get('اليوم') == today_name and str(row.get('الحصة')) == str(active_p):
                    teacher = row.get("المعلم", "")
                    subject = row.get("المادة/الصف", "")
                    loc = f"حصة {active_p} ({subject}) - المعلم: {teacher}"
                    return (loc, subject, teacher, str(active_p))
                    
            return ("لا توجد حصة مسجلة الآن 🤷", "", "", str(active_p))
            
        except Exception as e:
            return (f"خطأ: {e}", "", "", "")

    # --- beautify_all_internal_databases ---
    def beautify_all_internal_databases(self):
        # ========================================
        try:
            if os.path.exists(FILE_SUBJECTS):
                self._save_formatted_xlsx(self.df_subjects, FILE_SUBJECTS, "المواد")
            elif os.path.exists(FILE_STUDENTS):
                self._save_formatted_xlsx(self.df_students, FILE_STUDENTS, "الطلاب")
            if os.path.exists(FILE_CALENDAR):
                self._save_formatted_xlsx(self.df_calendar, FILE_CALENDAR, "التقويم")
            if os.path.exists(FILE_TIMINGS):
                self._save_formatted_xlsx(self.df_timings, FILE_TIMINGS, "التوقيت")
            if os.path.exists(FILE_TEACHERS):
                df_t = self._read_excel_any(FILE_TEACHERS, silent=True)
                if not df_t.empty:
                    self._save_formatted_xlsx(df_t, FILE_TEACHERS, "المعلمين")
            if os.path.exists(FILE_MASTER):
                import openpyxl
                from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
                wb = openpyxl.load_workbook(FILE_MASTER)
                header_fill = PatternFill(start_color="317135", end_color="317135", fill_type="solid")
                cell_fill = PatternFill(start_color="F2F9F2", end_color="F2F9F2", fill_type="solid")
                white_font = Font(color="FFFFFF", bold=True)
                black_font = Font(color="000000")
                center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
                thin_border = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"),
                  bottom=Side(style="thin"))
                for sheet in wb.worksheets:
                    sheet.sheet_view.rightToLeft = True

                if sheet.max_row >= 1:
                    for cell in sheet[1]:
                        cell.fill = header_fill
                        cell.font = white_font
                        cell.alignment = center_align
                        cell.border = thin_border

                else:
                    for row in sheet.iter_rows(min_row=2):
                        for cell in row:
                            cell.fill = cell_fill
                            cell.font = black_font
                            cell.alignment = center_align
                            cell.border = thin_border

                    else:
                        for col in sheet.columns:
                            max_length = 0
                            column = col[0].column_letter
                            for cell in col:
                                try:
                                    if len(str(cell.value)) > max_length:
                                        max_length = len(str(cell.value))
                                except:
                                    pass

                            else:
                                adjusted_width = max_length + 4
                                sheet.column_dimensions[column].width = min(adjusted_width, 30)

                        else:
                            wb.save(FILE_MASTER)

        except Exception as e:
            try:
                print(f"Background beautification failed: {e}")
            finally:
                pass

        # ========================================

    # --- load_employees_pins ---
    def load_employees_pins(self):
        # ========================================
        if not os.path.exists(FILE_EMPLOYEES_PINS): return {}
        try:
            df = pd.read_excel(FILE_EMPLOYEES_PINS)
            if df.empty or len(df.columns) < 2: return {}
            name_col = df.columns[0]
            pin_col = df.columns[1]
            return dict(zip(df[name_col].astype(str).str.strip(), df[pin_col].astype(str).str.strip()))
        except Exception:
            return {}



        # ========================================

    # --- verify_employee_pin ---
    def verify_employee_pin(self, name=None, pin=None):
        # ========================================
        pins = self.load_employees_pins()
        stored_pin = pins.get(name)
        return stored_pin == str(pin).strip()

        # ========================================

    # --- load_task_history ---
    def load_task_history(self):
        # ========================================
        fpath = data_path("employee_task_history.json")
        if not os.path.exists(fpath): return []
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f: return json.load(f)
        except: return []

    # --- save_task_history ---
    def save_task_history(self, history_list=None):
        # ========================================
        fpath = data_path("employee_task_history.json")
        try:
            import json
            with open(fpath, "w", encoding="utf-8") as f: json.dump(history_list, f, ensure_ascii=False)
            return True
        except: return False

    # --- load_employee_achievements ---
    def load_employee_achievements(self):
        fpath = data_path("employee_achievements.json")
        if not os.path.exists(fpath): return []
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f: return json.load(f)
        except: return []

    # --- save_employee_achievements ---
    def save_employee_achievements(self, records=None):
        fpath = data_path("employee_achievements.json")
        try:
            import json
            with open(fpath, "w", encoding="utf-8") as f: json.dump(records, f, ensure_ascii=False)
            return True
        except: return False
        
    # --- log_daily_achievement ---
    def log_daily_achievement(self, employee_name, task_type, items_list):
        records = self.load_employee_achievements()
        records.append({
            "employee": employee_name,
            "task_type": task_type,
            "date": datetime.now().strftime("%Y-%m-%d"),
            "time": datetime.now().strftime("%I:%M:%S %p"),
            "items": items_list,
            "status": "pending_manager_approval" 
        })
        return self.save_employee_achievements(records)

    # --- log_role_assignment ---
    def log_role_assignment(self, emp_name=None, task_name=None, action="add"):
        # ========================================
        history = self.load_task_history()
        today_str = datetime.now().strftime("%Y-%m-%d")
        
        if action == "add":
            active_exists = any((h for h in history if h.get("employee") == emp_name and h.get("task") == task_name and h.get("status") == "تحت العمل"))
            if not active_exists:
                history.append({
                    "employee": emp_name,
                    "task": task_name,
                    "start_date": today_str,
                    "end_date": "",
                    "status": "تحت العمل"
                })
        elif action == "remove":
            for h in history:
                if h.get("employee") == emp_name and h.get("task") == task_name and h.get("status") == "تحت العمل":
                    h["end_date"] = today_str
                    h["status"] = "منتهي"
        self.save_task_history(history)

    # ========================================

    # --- load_attendance_roles ---
    def load_attendance_roles(self):
        # ========================================

        fpath = data_path("attendance_roles.json")
        if not os.path.exists(fpath): return {}
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f: return json.load(f)
        except: return {}


    # --- save_attendance_roles ---
    def save_attendance_roles(self, roles_dict=None):
        # ========================================

        fpath = data_path("attendance_roles.json")
        try:
            import json
            with open(fpath, "w", encoding="utf-8") as f: json.dump(roles_dict, f, ensure_ascii=False)
            return True
        except: return False


    # --- save_assignment ---
    def save_assignment(self, emp_name=None, role_name=None, date_str=None):
        # ========================================
        current = self.load_attendance_roles()
        if role_name not in current:
            current[role_name] = []
        if not isinstance(current[role_name], list):
            val = current[role_name]
            current[role_name] = [str(val)] if val else []
        if str(emp_name).strip() not in [str(x).strip() for x in current[role_name]]:
            current[role_name].append(str(emp_name).strip())
        if self.save_attendance_roles(current):
            self.update_task_status(role_name, date_str, "تحت العمل", "")
            return True
        return False

        # ========================================

    # --- get_employee_role ---
    def get_employee_role(self, emp_name=None):
        # ========================================
        roles = self.load_attendance_roles()
        allowed = []
        target_name = str(emp_name).strip()
        print(f"DEBUG: Checking roles for '{target_name}' among: {roles}")
        for task, assigned_list in roles.items():
            if any((target_name == str(e).strip() for e in assigned_list)):
                allowed.append(task.strip())
            print(f"DEBUG: Found roles: {allowed}")
            return allowed

        # ========================================

    # --- get_employee_message ---
    def get_employee_message(self, emp_name=None):
        # ========================================

        return ""


    # --- save_employee_message ---
    def save_employee_message(self, emp_name=None, msg=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- get_task_status ---
    def get_task_status(self, task=None, date_str=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- get_task_attachment ---
    def get_task_attachment(self, task=None, date_str=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- update_task_status ---
    def update_task_status(self, task=None, date_str=None, status=None, attachment=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- add_custom_assignment ---
    def add_custom_assignment(self, emp_name=None, task_text=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- add_behavior_record ---
    def add_behavior_record(self, teacher=None, student=None, student_class=None, subject=None, b_type=None, points=None, notes=None):
        # ========================================
        fpath = data_path("سجل_السلوك_والمشاركة.xlsx(")
        if os.path.exists(fpath):
            try:
                df = pd.read_excel(fpath)
            except:
                df = pd.DataFrame(columns=['التاريخ', 'المعلم', 'الطالب', 'الصف', 'المادة', 'النوع', 'النقاط', 
                 'ملاحظات'])

        else:
            df = pd.DataFrame(columns=['التاريخ', 'المعلم', 'الطالب', 'الصف', 'المادة', 'النوع', 'النقاط', 'ملاحظات'])
        new_row = {'التاريخ':(datetime.now().strftime)(")%Y-%m-%d %H:%M"), 
         'المعلم':teacher, 
         'الطالب':student, 
         'الصف':student_class, 
         'المادة':subject, 
         'النوع':b_type, 
         'النقاط':points, 
         'ملاحظات':notes}
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        self._save_formatted_xlsx(df, fpath, sheet_name="السلوك والمواظبة")
        return True

        # ========================================

    # --- remove_custom_assignment ---
    def remove_custom_assignment(self, task_name=None, emp_name=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- get_all_assignments ---
    def get_all_assignments(self):
        # ========================================

        fpath = data_path("assignments.json")
        if not os.path.exists(fpath): return []
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f: return json.load(f)
        except: return []


    # --- update_assignment_status ---
    def update_assignment_status(self, assignment_id=None, status=None, attachment=None, reply_msg=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- send_direct_message ---
    def send_direct_message(self, sender=None, receiver=None, msg_text=None, attachment=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- get_my_messages ---
    def get_my_messages(self, receiver_name=None):
        # ========================================
        fpath = data_path("رسائل_مباشرة.json")
        if not os.path.exists(fpath):
            return []
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f:
                all_msgs = json.load(f)
            my_msgs = [m for m in all_msgs if m.get("receiver") == receiver_name]
            return sorted(my_msgs, key=(lambda x: x["id"]), reverse=True)
        except:
            return []

        # ========================================

    # --- mark_message_read ---
    def mark_message_read(self, msg_id=None):
        # ========================================
        fpath = data_path("رسائل_مباشرة.json")
        if not os.path.exists(fpath):
            return
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f:
                msgs = json.load(f)
            for m in msgs:
                if m["id"] == msg_id:
                    m["read"] = True
            else:
                with open(fpath, "w", encoding="utf-8") as f:
                    json.dump(msgs, f, ensure_ascii=False, indent=2)

        except:
            pass

        # ========================================

    # --- send_swap_request ---
    def send_swap_request(self, requester=None, acceptor=None, day=None, period=None, subject=None, comments=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- get_my_swaps ---
    def get_my_swaps(self, user_name=None):
        # ========================================
        fpath = data_path("تبادل_حصص.json")
        if not os.path.exists(fpath):
            return []
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f:
                swaps = json.load(f)
            return [s for s in swaps if not s.get("requester") == user_name if s.get("acceptor") == user_name]
        except:
            return []

        # ========================================

    # --- respond_swap_request ---
    def respond_swap_request(self, swap_id=None, new_status=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- admin_finalize_swap ---
    def admin_finalize_swap(self, swap_id=None, final_decision=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- get_attendance_history ---
    def get_attendance_history(self):
        # ========================================
        if not os.path.exists(FILE_ATTENDANCE):
            return pd.DataFrame()
        try:
            return self._read_excel_any(FILE_ATTENDANCE)
        except:
            return pd.DataFrame()

        # ========================================

    # --- add_attendance_record ---
    def add_attendance_record(self, records_list=None):
        # ========================================
        if not records_list: return False
        try:
            df = self.get_attendance_history()
            new_df = pd.DataFrame(records_list)
            if not df.empty:
                for c in new_df.columns:
                    if c not in df.columns:
                        df[c] = ""
                updated_df = pd.concat([df, new_df], ignore_index=True)
            else:
                updated_df = new_df
            updated_df.to_excel(FILE_ATTENDANCE, index=False)
            return True
        except Exception as e:
            return False

    # --- approve_attendance_records ---
    def approve_attendance_records(self, indices=None):
        # ========================================
        try:
            df = self.get_attendance_history()
            if df.empty or not indices: return False
            if "حالة الاعتماد" not in df.columns: df["حالة الاعتماد"] = "معلق"
            for idx in indices:
                if 0 <= idx < len(df):
                    df.at[idx, "حالة الاعتماد"] = "معتمد"
            df.to_excel(FILE_ATTENDANCE, index=False)
            return True
        except: return False

    # --- reject_attendance_records ---
    def reject_attendance_records(self, indices=None):
        # ========================================
        try:
            df = self.get_attendance_history()
            if df.empty or not indices: return False
            if "حالة الاعتماد" not in df.columns: df["حالة الاعتماد"] = "معلق"
            for idx in indices:
                if 0 <= idx < len(df):
                    df.at[idx, "حالة الاعتماد"] = "مرفوض"
            df.to_excel(FILE_ATTENDANCE, index=False)
            return True
        except: return False

    # --- update_batch_status ---
    def update_batch_status(self, date_str=None, time_str=None, op_type=None, new_status=None):
        # ========================================
        try:
            df = self.get_attendance_history()
            if df.empty: return False
            mask = (df["التاريخ"] == date_str) & (df["الوقت"] == time_str) & (df["نوع العملية"] == op_type)
            df.loc[mask, "حالة الاعتماد"] = new_status
            df.to_excel(FILE_ATTENDANCE, index=False)
            return True
        except: return False
        return None

    # --- get_attendance_stats ---
    def get_attendance_stats(self):
        # ========================================
        df = self.get_attendance_history()
        if df.empty:
            return {
             'total': 0, 'late': 0, 'absent': 0, 'leave': 0}
        return {'total':len(df_today), 
         'late':len(df_today[df_today["نوع العملية"] == "تأخير"]), 
         'absent':len(df_today[df_today["نوع العملية"] == "غياب"]), 
         'leave':len(df_today[df_today["نوع العملية"] == "انصراف"])}

        # ========================================

    # --- get_attendance_records_by_date ---
    def get_attendance_records_by_date(self, target_date=None):
        # ========================================
        df = self.get_attendance_history()
        if df.empty:
            return df
        if target_date is None:
            target_date = datetime.now().strftime("%Y-%m-%d")
        try:
            return df[df["التاريخ"].astype(str) == target_date]
        except:
            return pd.DataFrame()

        # ========================================

    # --- load_swaps ---
    def load_swaps(self):
        # ========================================

        if not os.path.exists(FILE_SWAPS): return pd.DataFrame()
        try: return pd.read_excel(FILE_SWAPS)
        except: return pd.DataFrame()


    # --- save_swaps ---
    def save_swaps(self, df=None):
        # ========================================
        try:
            self._save_formatted_xlsx(df, FILE_SWAPS, sheet_name="سجل المبادلات")
        except Exception as e:
            try:
                messagebox.showerror("خطأ", f"فشل حفظ المبادلات:\n{e}")
            finally:
                pass

        # ========================================

    # --- add_swap_request ---
    def add_swap_request(self, swap_data=None):
        # ========================================
        df = self.load_swaps()
        if df.empty or "رقم المبادلة" not in df.columns:
            next_id = 1
        else:
            try:
                next_id = int(df["رقم المبادلة"].max()) + 1
            except:
                next_id = 1
            else:
                period_num = swap_data.get("رقم الحصة", "")
                time_from, time_to = self.get_period_time(str(period_num))
                now = datetime.now()
                new_record = {'رقم المبادلة':next_id, 
                 'تاريخ الطلب':now.strftime("%Y-%m-%d"), 
                 'وقت الطلب':now.strftime("%H:%M"), 
                 'المعلم المرسل':swap_data.get("المعلم المرسل", ""), 
                 'المعلم المستقبل':swap_data.get("المعلم المستقبل", ""), 
                 'اليوم':swap_data.get("اليوم", ""), 
                 'رقم الحصة':period_num, 
                 'وقت بداية الحصة':time_from, 
                 'وقت نهاية الحصة':time_to, 
                 'المادة':swap_data.get("المادة", ""), 
                 'الصف':swap_data.get("الصف", ""), 
                 'الشعبة':swap_data.get("الشعبة", ""), 
                 'تاريخ بداية المبادلة':swap_data.get("تاريخ بداية المبادلة", ""), 
                 'تاريخ نهاية المبادلة':swap_data.get("تاريخ نهاية المبادلة", ""), 
                 'حالة الاعتماد':"قيد المراجعة", 
                 'تاريخ الاعتماد':"", 
                 'الملاحظات':swap_data.get("الملاحظات", "")}
                df = pd.concat([df, pd.DataFrame([new_record])], ignore_index=True)
                self.save_swaps(df)
                return (
                 True, next_id)

        # ========================================

    # --- update_swap_status ---
    def update_swap_status(self, swap_id=None, new_status=None):
        # ========================================
        df = self.load_swaps()
        if df.empty:
            return (False, 'لا توجد مبادلات')
        else:
            mask = df["رقم المبادلة"] == int(swap_id)
            return mask.any() or (False, 'المبادلة غير موجودة')
        df.loc[(mask, "حالة الاعتماد")] = new_status
        if new_status == "معتمد نهائياً":
            df.loc[(mask, "تاريخ الاعتماد")] = datetime.now().strftime("%Y-%m-%d %H:%M")
        self.save_swaps(df)
        return (True, 'تم تحديث الحالة')

        # ========================================

    # --- get_pending_swaps_for_teacher ---
    def get_pending_swaps_for_teacher(self, teacher_name=None):
        # ========================================
        df = self.load_swaps()
        if df.empty:
            return []
        mask = (df["المعلم المستقبل"] == teacher_name) & (df["حالة الاعتماد"] == "قيد المراجعة")
        pending = df[mask]
        return pending.to_dict("records")

        # ========================================

    # --- get_pending_swaps_for_admin ---
    def get_pending_swaps_for_admin(self):
        # ========================================
        df = self.load_swaps()
        if df.empty:
            return []
        mask = df["حالة الاعتماد"] == "معتمد من المستقبل"
        pending = df[mask]
        return pending.to_dict("records")

        # ========================================

    # --- get_active_swaps ---
    def get_active_swaps(self):
        # ========================================
        df = self.load_swaps()
        if df.empty:
            return []
        today = date.today()
        active = []
        for _, row in df.iterrows():
            if row.get("حالة الاعتماد") != "معتمد نهائياً":
                pass
            else:
                start_date = safe_parse_date(row.get("تاريخ بداية المبادلة"))
                end_date = safe_parse_date(row.get("تاريخ نهاية المبادلة"))
                if start_date and end_date:
                    if start_date <= today <= end_date:
                        active.append(row.to_dict())
                return active

        # ========================================

    # --- check_swap_for_cell ---
    def check_swap_for_cell(self, teacher_name=None, day=None, period_idx=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- log_monitoring_event ---
    def log_monitoring_event(self, record_dict=None):
        # ========================================
        file_path = data_path("سجل_متابعة_الحصص.xlsx")
        try:
            new_df = pd.DataFrame([record_dict])
            if os.path.exists(file_path):
                with pd.ExcelWriter(file_path, mode="a", engine="openpyxl", if_sheet_exists="overlay") as writer:
                    existing_df = pd.read_excel(file_path)
                    combined = pd.concat([existing_df, new_df], ignore_index=True)
                    combined.to_excel(file_path, index=False)
            else:
                new_df.to_excel(file_path, index=False)
        except Exception as e:
            try:
                print(f"Error saving monitoring record: {e}")
            finally:
                pass

        # ========================================

    # --- save_contact_log ---
    def save_contact_log(self, record=None):
        # ========================================
        fpath = data_path("sجلات_التواصل.json")
        data = []
        if os.path.exists(fpath):
            try:
                import json
                with open(fpath, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except:
                pass

        updated = False
        for i, r in enumerate(data):
            if r.get("date") == record["date"] and r.get("employee") == record["employee"]:
                data[i] = record
                updated = True
                break
        else:
            if not updated:
                data.append(record)
            with open(fpath, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

        # ========================================

    # --- get_contact_logs ---
    def get_contact_logs(self):
        # ========================================

        fpath = data_path("سجلات_التواصل.json")
        if not os.path.exists(fpath): return []
        try:
            import json
            with open(fpath, "r", encoding="utf-8") as f: return json.load(f)
        except: return []



class SyncManager:

    # --- generate_export_filename ---
    def generate_export_filename(self, user_name=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- export_work ---
    def export_work(self, user_name=None, target_dir=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None

    # --- import_work ---
    def import_work(self, source_zip=None, data_model=None):
        # ========================================
        # [DECOMPILATION FAILED]: 



        # ========================================
        return None


if __name__ == '__main__':
    app = App(DataModel())
    app.mainloop()
